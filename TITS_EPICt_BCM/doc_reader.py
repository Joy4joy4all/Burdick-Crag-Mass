# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
AI DOC READER - Test Packet Processor
============================================
Drop-in module for gibush_icorps/

Workflow:
1. Upload completed test packet (.docx)
2. Target folder auto-set by active project (BCM_SUBSTRATE / BCM_NAVIGATION)
3. AI reads the document, extracts all filled data
4. Removes person from PENDING test list
5. Places person in COMPLETED (tested/) folder
6. Logs the change in change_log.json
7. Generates _AI_generated_report with metrics vs other test_runs

File destinations per project:
    BCM_TESTS/<PROJECT>/tested/           completed .docx stored
    BCM_TESTS/<PROJECT>/_AI_generated_report/  AI comparison reports
    BCM_TESTS/<PROJECT>/change_log.json        audit trail
"""

import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple


# ============================================================================
# PACKET PARSER - Reads completed .docx test packets
# ============================================================================

class PacketParser:
    """
    Extracts structured data from a completed test packet .docx.
    
    Expected structure (matches generate_test_template output):
        Table 0: Header (test_run #, person, title, company, segment, q-cube)
        Tables 1-6: Equipment checklists ( or Yes + notes)
        Table 7: Cost discovery framework
        Table 8: Equipment & Costs Captured
        Paragraphs: Filled answers under known headings
    """
    
    def __init__(self, docx_path: Path):
        self.docx_path = Path(docx_path)
        self.data = {}
        self._raw_paragraphs = []
        self._raw_tables = []
    
    def parse(self) -> Dict:
        """Parse the completed packet and return structured data."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")
        
        doc = Document(str(self.docx_path))
        self._raw_paragraphs = doc.paragraphs
        self._raw_tables = doc.tables
        
        self.data = {
            'file_name': self.docx_path.name,
            'parsed_at': datetime.now().isoformat(),
            'header': self._parse_header(),
            'equipment_discussed': self._parse_equipment_tables(),
            'equipment_costs': self._parse_cost_table(),
            'results': self._parse_section('Results (What We Learned)'),
            'hypotheses_validation': self._parse_section('✅ Hypothesis Validation'),
            'action_iterate': self._parse_section('Action/Iterate'),
            'key_questions_answers': self._parse_key_questions(),
            'experiments_answers': self._parse_experiments(),
            'pain_indicators_noted': self._parse_pain_indicators(),
        }
        
        # Derive test_run number from header
        self.data['test_num'] = self._extract_test_num()
        self.data['script_name'] = self.data['header'].get('Person', 'Unknown')
        self.data['source_version'] = self.data['header'].get('Company', 'Unknown')
        self.data['is_completed'] = self._check_completion()
        
        return self.data
    
    def _parse_header(self) -> Dict:
        """Extract header table (Table 0): test metadata."""
        if not self._raw_tables:
            return {}
        
        header = {}
        table = self._raw_tables[0]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0]:
                key = cells[0].replace('#', '').strip()
                header[key] = cells[1]
        return header
    
    def _parse_equipment_tables(self) -> List[Dict]:
        """Extract equipment marked as discussed (Yes) from checklist tables.

        Handles two formats:
          Legacy (3-col): status | Equipment | Cost | Notes
          New (8-col):    ☐ | Equipment | Typical Cost | Production $ | Equip $ | Safety $ | Other $ | Notes
        An item is "discussed" if:
          - status cell contains 'yes', 'x', or is empty (legacy), OR
          - any cost/notes column (cols 3-7) has a dollar value or non-empty text (new format)
        Also pulls items from the Equipment & Costs Captured table at the end.
        """
        discussed = []
        seen = set()

        SKIP_HEADERS = {'equipment/module', 'equipment', ''}

        # ── Checklist tables (all tables except table 0 header) ──
        for table in self._raw_tables[1:]:
            if len(table.rows) < 2:
                continue

            # Detect 8-column checklist by header
            hdr_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            is_checklist = 'equipment/module' in hdr_cells or 'equipment' in hdr_cells
            if not is_checklist:
                continue

            # Stop scanning when we hit the cost-capture summary table
            # (header: Equipment | Annual Damage $ | Failure Freq | Cost Type)
            if (hdr_cells and hdr_cells[0] == 'equipment' and
                    len(hdr_cells) > 1 and 'annual damage' in hdr_cells[1]):
                break

            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 2:
                    continue

                equip_name = cells[1] if len(cells) > 1 else ''
                if not equip_name or equip_name.lower() in SKIP_HEADERS:
                    continue

                status = cells[0].lower().strip()

                # Legacy: explicit yes/x/blank marker
                legacy_match = (
                    status in ('yes', 'yes means discussed', 'x', 'checked') or
                    'yes' in status
                )

                # New 8-col: any cost or notes column has content
                cost_cols = cells[3:] if len(cells) > 3 else []
                has_cost_data = any(
                    c and c not in ('$', '-', '0', 'repair / production')
                    for c in cost_cols
                )

                if legacy_match or has_cost_data:
                    key = equip_name.lower()
                    if key not in seen:
                        seen.add(key)
                        discussed.append({
                            'equipment': equip_name,
                            'typical_cost': cells[2] if len(cells) > 2 else '',
                            'production_cost': cells[3] if len(cells) > 3 else '',
                            'equipment_cost': cells[4] if len(cells) > 4 else '',
                            'safety_cost':    cells[5] if len(cells) > 5 else '',
                            'other_cost':     cells[6] if len(cells) > 6 else '',
                            'notes':          cells[7] if len(cells) > 7 else (cells[3] if len(cells) > 3 else ''),
                        })

        # ── Cost-capture summary table ──
        # Header: Equipment | Annual Damage $ | Failure Freq | Cost Type
        for table in self._raw_tables:
            if len(table.rows) < 2:
                continue
            hdr_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            # Must have 'equipment' in col 0 AND 'annual damage $' in col 1
            if not (hdr_cells and hdr_cells[0] == 'equipment' and
                    len(hdr_cells) > 1 and 'annual damage' in hdr_cells[1]):
                continue
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 2:
                    continue
                equip_name = cells[0]
                if not equip_name or equip_name.lower() in SKIP_HEADERS:
                    continue
                key = equip_name.lower()
                if key not in seen:
                    seen.add(key)
                    discussed.append({
                        'equipment': equip_name,
                        'typical_cost': '',
                        'production_cost': '',
                        'equipment_cost': cells[1] if len(cells) > 1 else '',
                        'safety_cost': '',
                        'other_cost': '',
                        'notes': f"freq={cells[2]}, type={cells[3]}" if len(cells) > 3 else '',
                    })
            break  # Only one cost-capture table

        return discussed
    
    def _parse_cost_table(self) -> List[Dict]:
        """Extract Equipment & Costs Captured table (Table 8)."""
        costs = []
        
        # Table 8 (index 8) is cost capture, but may vary
        for table in self._raw_tables:
            # Look for the cost capture table by header pattern
            if len(table.rows) < 2:
                continue
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            if 'Equipment' in header_cells and 'Annual Damage' in ' '.join(header_cells):
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3 and cells[0]:
                        costs.append({
                            'equipment': cells[0],
                            'annual_damage': cells[1],
                            'failure_freq': cells[2],
                        })
                break
        
        return costs
    
    def _parse_section(self, heading_keyword: str) -> str:
        """Extract text under a specific heading until next heading."""
        content_lines = []
        capturing = False
        
        for p in self._raw_paragraphs:
            text = p.text.strip()
            style = p.style.name if p.style else ''
            
            # Start capturing after matching heading
            if heading_keyword.lower() in text.lower() and 'Heading' in style:
                capturing = True
                continue
            
            # Stop at next Heading 1 (sub-headings like Heading 2 are
            # part of the section content, not a new top-level section)
            if capturing and style == 'Heading 1' and text:
                break
            
            if capturing and text:
                content_lines.append(text)
        
        return '\n'.join(content_lines)
    
    def _parse_key_questions(self) -> Dict[str, str]:
        """Extract filled answers from Key Questions section."""
        answers = {}
        capturing = False
        current_q = None
        
        for p in self._raw_paragraphs:
            text = p.text.strip()
            style = p.style.name if p.style else ''
            
            if 'YOUR MISSION' in text.upper() and 'Heading' in style:
                capturing = True
                continue
            
            if capturing and 'Heading' in style and 'EQUIPMENT' in text.upper():
                break
            
            if capturing and text:
                # Look for "Question: Answer" patterns
                if text.startswith('What ') or text.startswith('How ') or text.startswith('KEY QUESTIONS'):
                    current_q = text[:60]
                    # Check if answer is inline after the question
                    parts = text.split('?', 1)
                    if len(parts) > 1 and parts[1].strip():
                        answers[parts[0].strip() + '?'] = parts[1].strip()
                elif current_q and text and not text.startswith('FOCUS'):
                    if current_q not in answers:
                        answers[current_q] = text
                    else:
                        answers[current_q] += ' ' + text
        
        return answers
    
    def _parse_experiments(self) -> Dict[str, str]:
        """Extract filled answers from Experiments/Questions section.
        
        Handles the common case where all numbered questions are in a single
        paragraph separated by newlines (soft line breaks in .docx).
        """
        answers = {}
        capturing = False
        
        for p in self._raw_paragraphs:
            text = p.text.strip()
            style = p.style.name if p.style else ''
            
            if 'Experiments' in text and 'Heading' in style:
                capturing = True
                continue
            
            if capturing and 'Heading' in style:
                break
            
            if capturing and text:
                # Split on newlines first (all Qs may be in one paragraph)
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    # Pattern: "1. Question? Answer text"
                    match = re.match(r'(\d+)\.\s*(.+)', line)
                    if match:
                        full_text = match.group(2)
                        q_parts = full_text.split('?', 1)
                        if len(q_parts) > 1 and q_parts[1].strip():
                            answer_candidate = q_parts[1].strip()
                            question_starters = (
                                'What ', 'How ', 'Walk ', 'When ', 'Where ', 'Why ',
                                'Who ', 'Which ', 'Do ', 'Does ', 'Did ', 'Have ',
                                'Has ', 'Is ', 'Are ', 'If ', 'Can ', 'Could ',
                            )
                            if any(answer_candidate.startswith(s) for s in question_starters):
                                answers[full_text[:80]] = '(no answer filled)'
                            else:
                                answers[q_parts[0].strip() + '?'] = answer_candidate
                        else:
                            answers[full_text[:80]] = '(no answer filled)'
        
        return answers
    
    def _parse_pain_indicators(self) -> List[str]:
        """Extract pain indicators that were noted/filled."""
        noted = []
        capturing = False
        
        for p in self._raw_paragraphs:
            text = p.text.strip()
            style = p.style.name if p.style else ''
            
            if 'PAIN INDICATORS' in text.upper():
                capturing = True
                continue
            
            # Stop at next heading (SCORECARD follows PAIN INDICATORS)
            if capturing and 'Heading' in style and text:
                break
            
            if capturing and text.startswith('⬢'):
                clean = text.lstrip('⬢').strip().lstrip(':').lstrip('-').strip()
                indicator_names = [
                    'Unplanned Downtime', 'Equipment Damage', 'System Failure',
                    'Lack of Visibility', 'No Detection Capability', 'Reactive Maintenance',
                    'Contamination Events', 'Visibility Gap', 'Knowledge Gap',
                    'Missing Baseline Data'
                ]
                for name in indicator_names:
                    if name.lower() in clean.lower() and len(clean) > len(name) + 2:
                        note_text = clean.lower().replace(name.lower(), '').strip().lstrip(':').lstrip('-').strip()
                        noted.append({'indicator': name, 'notes': note_text})
        
        return noted
    
    def _extract_test_num(self) -> int:
        """Get test_run number from header or filename."""
        # Try header first
        header_num = self.data.get('header', {}).get('Test Run', '')
        if header_num:
            try:
                return int(re.search(r'\d+', str(header_num)).group())
            except (AttributeError, ValueError):
                pass
        
        # Try filename: Test Run_Packet_1_Name.docx
        match = re.search(r'(?:Packet|Template)_(\d+)', self.docx_path.stem)
        if match:
            return int(match.group(1))
        
        return 0
    
    def _check_completion(self) -> bool:
        """Check if the packet has meaningful filled content."""
        PLACEHOLDERS = (
            '[fill during test_run]',
            '[fill after test_run]',
            'fill during test_run',
            'fill after test_run',
        )

        # Results must exist and not be ONLY placeholders.
        # Strip placeholder lines, then check if real content remains.
        results_text = self.data.get('results', '').strip()
        cleaned_lines = [
            ln for ln in results_text.split('\n')
            if ln.strip() and not any(p in ln.lower() for p in PLACEHOLDERS)
        ]
        has_results = len(cleaned_lines) >= 2  # at least 2 real lines

        has_equipment = len(self.data.get('equipment_discussed', [])) > 0

        # Answers only count if at least one is NOT the default '(no answer filled)'
        raw_answers = self.data.get('experiments_answers', {})
        real_answers = [v for v in raw_answers.values() if v.strip() != '(no answer filled)']
        has_answers = len(real_answers) > 0

        # At least 2 of 3 sections filled = completed
        return sum([has_results, has_equipment, has_answers]) >= 2


# ============================================================================
# TEST MOVER - Moves from pending to completed
# ============================================================================

class Test RunMover:
    """
    Manages the pending  completed test_run transition.
    
    Folders:
        <project>/tested/            completed .docx files
        <project>/_AI_generated_report/    AI comparison reports
        <project>/change_log.json          audit log
    """
    
    def __init__(self, project_root: Path):
        self.project_root = Path(project_root)
        self.tested_dir = self.project_root / 'tested'
        self.report_dir = self.project_root / '_AI_generated_report'
        self.change_log_path = self.project_root / 'change_log.json'
        
        # Create directories
        self.tested_dir.mkdir(parents=True, exist_ok=True)
        self.report_dir.mkdir(parents=True, exist_ok=True)
    
    def move_to_completed(self, source_docx: Path, parsed_data: Dict) -> Path:
        """
        Copy completed packet to tested/ folder.
        Returns the destination path.
        """
        person = parsed_data.get('script_name', 'Unknown').replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        dest_name = f"{person}_completed_{timestamp}.docx"
        dest_path = self.tested_dir / dest_name
        
        shutil.copy2(source_docx, dest_path)
        return dest_path
    
    def log_change(self, parsed_data: Dict, dest_path: Path, project_id: str):
        """Log the test_run completion to change_log.json."""
        log_entry = {
            'timestamp': datetime.now().isoformat(),
            'action': 'TEST_COMPLETED',
            'test_num': parsed_data.get('test_num', 0),
            'script_name': parsed_data.get('script_name', 'Unknown'),
            'source_version': parsed_data.get('source_version', 'Unknown'),
            'project': project_id,
            'source_file': parsed_data.get('file_name', ''),
            'dest_file': dest_path.name,
            'equipment_discussed_count': len(parsed_data.get('equipment_discussed', [])),
            'has_results': bool(parsed_data.get('results', '').strip()),
            'has_hypothesis_validation': bool(parsed_data.get('hypotheses_validation', '').strip()),
            'has_action_items': bool(parsed_data.get('action_iterate', '').strip()),
        }
        
        # Load existing log or create new
        existing_log = []
        if self.change_log_path.exists():
            try:
                with open(self.change_log_path, 'r') as f:
                    existing_log = json.load(f)
            except (json.JSONDecodeError, IOError):
                existing_log = []
        
        existing_log.append(log_entry)
        
        with open(self.change_log_path, 'w') as f:
            json.dump(existing_log, f, indent=2)
        
        return log_entry
    
    def get_completed_test_runs(self) -> List[Dict]:
        """Return list of completed test_run info from change_log."""
        if not self.change_log_path.exists():
            return []
        try:
            with open(self.change_log_path, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []
    
    def get_completed_names(self) -> List[str]:
        """Return list of completed test_run person names."""
        completed = self.get_completed_test_runs()
        return [c.get('script_name', '') for c in completed if c.get('action') == 'TEST_COMPLETED']

    def delete_completed(self, person_name: str) -> Dict:
        """
        Atomically remove a test from completed state.

        Steps (all-or-nothing):
        1. Find the log entry for this person
        2. Remove the .docx file from tested/
        3. Write TEST_DELETED entry to change_log.json
        4. Remove the TEST_COMPLETED entry from change_log.json

        Returns dict with success flag and details.
        """
        log = self.get_completed_test_runs()

        # Find the completed entry by person name
        target_entry = None
        for entry in log:
            if (entry.get('action') == 'TEST_COMPLETED'
                    and entry.get('script_name', '').strip().lower() == person_name.strip().lower()):
                target_entry = entry
                break

        if not target_entry:
            return {
                'success': False,
                'error': f'{person_name} not found in completed log.',
            }

        # Remove the file from tested/ if it exists
        dest_file = target_entry.get('dest_file', '')
        file_path = self.tested_dir / dest_file if dest_file else None
        file_removed = False
        if file_path and file_path.exists():
            file_path.unlink()
            file_removed = True

        # Rebuild log: remove the TEST_COMPLETED entry, add TEST_DELETED record
        new_log = [e for e in log if not (
            e.get('action') == 'TEST_COMPLETED'
            and e.get('script_name', '').strip().lower() == person_name.strip().lower()
        )]
        new_log.append({
            'timestamp': datetime.now().isoformat(),
            'action': 'TEST_DELETED',
            'test_num': target_entry.get('test_num', 0),
            'script_name': target_entry.get('script_name', 'Unknown'),
            'source_version': target_entry.get('source_version', 'Unknown'),
            'project': target_entry.get('project', ''),
            'deleted_file': dest_file,
            'file_removed_from_disk': file_removed,
        })

        with open(self.change_log_path, 'w') as f:
            json.dump(new_log, f, indent=2)

        return {
            'success': True,
            'script_name': target_entry.get('script_name', 'Unknown'),
            'test_num': target_entry.get('test_num', 0),
            'file_removed': file_removed,
        }


# ============================================================================
# REPORT GENERATOR - AI comparison report
# ============================================================================

class ReportGenerator:
    """
    Generates _AI_generated_report comparing this test_run vs others.
    
    Metrics:
    - Equipment coverage overlap
    - Hypothesis validation alignment
    - New insights not seen in prior test_runs
    - Pain indicator frequency
    - Cost data contributions
    """
    
    def __init__(self, project_root: Path):
        self.project_root = Path(project_root)
        self.report_dir = self.project_root / '_AI_generated_report'
        self.report_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, parsed_data: Dict, all_completed: List[Dict],
                 genesis_intel: Optional[Dict] = None) -> Path:
        """
        Generate comparison report for this test_run.
        
        Args:
            parsed_data: Parsed data from current completed packet
            all_completed: List of all previous change_log entries
            genesis_intel: Optional Genesis intelligence data
            
        Returns:
            Path to generated report file
        """
        test_num = parsed_data.get('test_num', 0)
        person = parsed_data.get('script_name', 'Unknown')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        report_name = f"AI_Report_Test Run_{test_num}_{person.replace(' ', '_')}_{timestamp}.txt"
        report_path = self.report_dir / report_name
        
        lines = []
        lines.append("=" * 80)
        lines.append(f"AI GENERATED REPORT - Test Run #{test_num}: {person}")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Company: {parsed_data.get('source_version', 'Unknown')}")
        lines.append("=" * 80)
        lines.append("")
        
        # ---- SECTION 1: Test Run Summary ----
        lines.append("-" * 80)
        lines.append("1. TEST SUMMARY")
        lines.append("-" * 80)
        
        results = parsed_data.get('results', '').strip()
        if results:
            lines.append(f"\nResults:\n{results[:500]}")
        else:
            lines.append("\nResults: (none recorded)")
        
        hyp_val = parsed_data.get('hypotheses_validation', '').strip()
        if hyp_val:
            lines.append(f"\nHypothesis Validation:\n{hyp_val[:500]}")
        
        action = parsed_data.get('action_iterate', '').strip()
        if action:
            lines.append(f"\nAction/Next Steps:\n{action[:500]}")
        
        lines.append("")
        
        # ---- SECTION 2: Equipment Discussed ----
        lines.append("-" * 80)
        lines.append("2. EQUIPMENT DISCUSSED")
        lines.append("-" * 80)
        
        equipment = parsed_data.get('equipment_discussed', [])
        if equipment:
            lines.append(f"\n{len(equipment)} equipment items discussed:")
            for eq in equipment:
                note_text = f"  {eq['notes']}" if eq.get('notes') else ''
                lines.append(f"   {eq['equipment']} ({eq.get('typical_cost', 'N/A')}){note_text}")
        else:
            lines.append("\nNo equipment explicitly marked as discussed.")
        
        lines.append("")
        
        # ---- SECTION 3: Cost Data ----
        lines.append("-" * 80)
        lines.append("3. COST DATA CAPTURED")
        lines.append("-" * 80)
        
        costs = parsed_data.get('equipment_costs', [])
        if costs:
            for c in costs:
                lines.append(f"   {c['equipment']}: Annual Damage=${c['annual_damage']}, "
                            f"Freq={c['failure_freq']}")
        else:
            lines.append("\nNo cost data captured in this test_run.")
        
        lines.append("")
        
        # ---- SECTION 4: Experiment Answers ----
        lines.append("-" * 80)
        lines.append("4. EXPERIMENT ANSWERS (Questions & Responses)")
        lines.append("-" * 80)
        
        experiments = parsed_data.get('experiments_answers', {})
        if experiments:
            for q, a in experiments.items():
                lines.append(f"\n  Q: {q}")
                lines.append(f"  A: {a[:200]}")
        else:
            lines.append("\nNo experiment answers extracted.")
        
        lines.append("")
        
        # ---- SECTION 5: Pain Indicators ----
        lines.append("-" * 80)
        lines.append("5. PAIN INDICATORS NOTED")
        lines.append("-" * 80)
        
        pains = parsed_data.get('pain_indicators_noted', [])
        if pains:
            for p in pains:
                lines.append(f"   {p['indicator']}: {p['notes']}")
        else:
            lines.append("\nNo pain indicators with notes detected.")
        
        lines.append("")
        
        # ---- SECTION 6: COMPARISON vs Previous Test Runs ----
        lines.append("-" * 80)
        lines.append("6. METRICS: THIS TEST vs PREVIOUS")
        lines.append("-" * 80)
        
        prior_count = len([c for c in all_completed 
                          if c.get('action') == 'TEST_COMPLETED'
                          and c.get('test_num') != test_num])
        
        lines.append(f"\nTotal prior completed test_runs: {prior_count}")
        lines.append(f"This is test_run #{prior_count + 1} to be processed.")
        lines.append("")
        
        # Equipment coverage metric
        this_equip_count = len(equipment)
        if prior_count > 0:
            prior_equip_counts = [c.get('equipment_discussed_count', 0) for c in all_completed
                                  if c.get('action') == 'TEST_COMPLETED' 
                                  and c.get('test_num') != test_num]
            avg_prior = sum(prior_equip_counts) / len(prior_equip_counts) if prior_equip_counts else 0
            lines.append(f"Equipment discussed this test_run: {this_equip_count}")
            lines.append(f"Average equipment discussed per test: {avg_prior:.1f}")
            
            if this_equip_count > avg_prior:
                lines.append(f"   ABOVE AVERAGE (+{this_equip_count - avg_prior:.1f} items)")
            elif this_equip_count < avg_prior:
                lines.append(f"   BELOW AVERAGE ({this_equip_count - avg_prior:.1f} items)")
            else:
                lines.append(f"   ON PAR with average")
        else:
            lines.append(f"Equipment discussed this test_run: {this_equip_count}")
            lines.append("   First test_run processed - baseline established.")
        
        lines.append("")
        
        # Completeness score
        completeness = 0
        checks = [
            ('Results filled', bool(results)),
            ('Equipment checklist used', this_equip_count > 0),
            ('Hypothesis validation done', bool(hyp_val)),
            ('Action items recorded', bool(action)),
            ('Experiment answers filled', len(experiments) > 0),
            ('Cost data captured', len(costs) > 0),
            ('Pain indicators noted', len(pains) > 0),
        ]
        
        lines.append("Completeness Score:")
        for label, passed in checks:
            status = "" if passed else "--"
            lines.append(f"  {status} {label}")
            if passed:
                completeness += 1
        
        score_pct = (completeness / len(checks)) * 100
        lines.append(f"\n  TOTAL: {completeness}/{len(checks)} ({score_pct:.0f}%)")
        
        if prior_count > 0:
            prior_with_results = len([c for c in all_completed 
                                      if c.get('has_results') and c.get('test_num') != test_num])
            lines.append(f"\n  Prior tests with results filled: {prior_with_results}/{prior_count}")
        
        lines.append("")
        
        # ---- SECTION 7: Genesis Intelligence Context ----
        if genesis_intel:
            lines.append("-" * 80)
            lines.append("7. GENESIS INTELLIGENCE CONTEXT")
            lines.append("-" * 80)
            
            if genesis_intel.get('hypotheses'):
                lines.append("\nHypothesis Confidence Before This Test Run:")
                for hyp, status in genesis_intel['hypotheses'].items():
                    lines.append(f"   {hyp}: {status.get('confidence', '?')}%")
            
            if genesis_intel.get('gaps'):
                lines.append("\nRemaining Information Gaps:")
                for gap in genesis_intel['gaps']:
                    lines.append(f"   {gap}")
        
        lines.append("")
        lines.append("=" * 80)
        lines.append("END OF AI GENERATED REPORT")
        lines.append("=" * 80)
        
        # Write report
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return report_path
    
    def generate_json_summary(self, parsed_data: Dict, all_completed: List[Dict]) -> Path:
        """Also save a machine-readable JSON summary for the genesis pipeline."""
        test_num = parsed_data.get('test_num', 0)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        json_path = self.report_dir / f"summary_{test_num}_{timestamp}.json"
        
        summary = {
            'test_num': test_num,
            'script_name': parsed_data.get('script_name', ''),
            'source_version': parsed_data.get('source_version', ''),
            'processed_at': datetime.now().isoformat(),
            'equipment_discussed': [e['equipment'] for e in parsed_data.get('equipment_discussed', [])],
            'equipment_count': len(parsed_data.get('equipment_discussed', [])),
            'cost_data': parsed_data.get('equipment_costs', []),
            'has_results': bool(parsed_data.get('results', '').strip()),
            'has_hypothesis_validation': bool(parsed_data.get('hypotheses_validation', '').strip()),
            'has_action_items': bool(parsed_data.get('action_iterate', '').strip()),
            'pain_indicators': [p['indicator'] for p in parsed_data.get('pain_indicators_noted', [])],
            'total_completed': len([c for c in all_completed if c.get('action') == 'TEST_COMPLETED']) + 1,
        }
        
        with open(json_path, 'w') as f:
            json.dump(summary, f, indent=2)
        
        return json_path


# ============================================================================
# DOC READER ENGINE - Orchestrates the full workflow
# ============================================================================

class DocReaderEngine:
    """
    Main engine that ties together:
    - PacketParser (read .docx)
    - Test RunMover (pending  completed)
    - ReportGenerator (AI comparison)
    
    Usage:
        engine = DocReaderEngine(project_folder, project_id)
        result = engine.process(docx_path, genesis_intel=None)
    """
    
    def __init__(self, project_folder: Path, project_id: str):
        self.project_folder = Path(project_folder)
        self.project_id = project_id
        self.mover = Test RunMover(self.project_folder)
        self.reporter = ReportGenerator(self.project_folder)
    
    def process(self, docx_path: Path, genesis_intel: Optional[Dict] = None) -> Dict:
        """
        Full pipeline:
        1. Parse completed packet
        2. Copy to tested/ folder
        3. Log the change
        4. Generate AI report
        5. Generate JSON summary
        
        Returns dict with all results and paths.
        """
        docx_path = Path(docx_path)
        
        # Step 1: Parse
        parser = PacketParser(docx_path)
        parsed = parser.parse()

        if not parsed.get('is_completed'):
            return {
                'success': False,
                'error': 'Document does not appear to be a completed test packet. '
                         'Results, equipment, or experiment answers are missing.',
                'parsed': parsed,
            }

        # Step 1b: Duplicate guard -- reject if this person already processed
        person_name = parsed.get('script_name', 'Unknown')
        already_done = self.mover.get_completed_names()
        if person_name.strip().lower() in [n.strip().lower() for n in already_done]:
            # Still try to update BMC — it may have been missed on prior run
            self._update_bmc(parsed)
            return {
                'success': False,
                'error': f'{person_name} has already been processed. '
                         f'Use Delete to remove the existing entry before re-uploading.',
                'parsed': parsed,
                'duplicate': True,
            }


        # Step 2: Move to completed
        dest_path = self.mover.move_to_completed(docx_path, parsed)
        
        # Step 3: Log the change
        all_completed = self.mover.get_completed_test_runs()
        log_entry = self.mover.log_change(parsed, dest_path, self.project_id)
        
        # Step 4: Generate AI report
        report_path = self.reporter.generate(parsed, all_completed, genesis_intel)

        # Step 5: Generate JSON summary
        json_path = self.reporter.generate_json_summary(parsed, all_completed)

        # Step 6: Sync status back to Excel -- makes Excel the living source of truth
        try:
            from excel_test_loader import get_loader
            loader = get_loader(self.project_id)
            # Excel sync still uses test_num for row matching
            loader.mark_completed_in_excel(parsed.get('test_num', 0))
        except Exception as e:
            print(f"Warning: Could not sync status to Excel: {e}")

        # Step 7: Update BMC state with this test_run
        bmc_updated = self._update_bmc(parsed)

        return {
            'success': True,
            'test_num': parsed.get('test_num', 0),
            'script_name': parsed.get('script_name', 'Unknown'),
            'source_version': parsed.get('source_version', 'Unknown'),
            'equipment_count': len(parsed.get('equipment_discussed', [])),
            'dest_path': str(dest_path),
            'report_path': str(report_path),
            'json_path': str(json_path),
            'log_entry': log_entry,
            'parsed': parsed,
            'bmc_updated': bmc_updated,
        }

    def _update_bmc(self, parsed: Dict) -> bool:
        """Update BMC with one parsed test_run. Called from process() and duplicate path."""
        try:
            from bmc_engine import generate_project_bmc
            qcube_raw = parsed.get('header', {}).get('Q-Cube Position', '')
            q_layer, q_object = '', ''
            lm = re.search(r'L(\d)', qcube_raw)
            om = re.search(r'O([A-Za-z])', qcube_raw)
            if lm: q_layer = f"L{lm.group(1)}"
            if om: q_object = f"O{om.group(1).upper()}"
            substrate_impacts = [
                {'equipment': eq.get('equipment', ''), 'notes': eq.get('notes', ''), 'cost': 0}
                for eq in parsed.get('equipment_discussed', [])
            ]
            experiments_str = '\n'.join(
                f"{q} {a}" for q, a in parsed.get('experiments_answers', {}).items()
            )
            bmc_test_run = {
                'test_num':     parsed.get('test_num', 0),
                'script_name':            parsed.get('script_name', 'Unknown'),
                'source_version':           parsed.get('source_version', 'Unknown'),
                'title':             parsed.get('header', {}).get('Title', ''),
                'test_category':     parsed.get('header', {}).get('Substrate Class', ''),
                'q_layer':           q_layer,
                'q_object':          q_object,
                'results':           parsed.get('results', ''),
                'experiments':       experiments_str,
                'action_iterate':    parsed.get('action_iterate', ''),
                'substrate_impacts': substrate_impacts,
            }
            try:
                from project_paths import get_bmc_dir
                bmc_dir = get_bmc_dir(self.project_id)
            except ImportError:
                bmc_dir = self.project_folder / 'BMC_generation'
            generate_project_bmc(
                project_id=self.project_id,
                test_runs=[bmc_test_run],
                bmc_dir=bmc_dir,
                generate_docx=True,
            )
            print(f"[DocReader] BMC updated with test_run #{parsed.get('test_num', '?')}")
            return True
        except Exception as e:
            print(f"Warning: Could not update BMC: {e}")
            return False

    def refresh_bmc(self, bmc_dir: Path = None) -> bool:
        """
        Full BMC rebuild from test_database.json.
        Call any time BMC state is out of sync — no re-upload needed.
        """
        try:
            from bmc_engine import BMCAggregator, generate_bmc_docx
            from project_paths import get_bmc_dir, BCM_ROOT
        except ImportError as e:
            print(f"[DocReader] refresh_bmc: missing import — {e}")
            return False
        try:
            if bmc_dir is None:
                try:
                    bmc_dir = get_bmc_dir(self.project_id)
                except Exception:
                    bmc_dir = self.project_folder / 'BMC_generation'
            db_file = BCM_ROOT / "test_database.json"
            if not db_file.exists():
                print(f"[DocReader] refresh_bmc: database not found at {db_file}")
                return False
            with open(db_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            agg = BMCAggregator(self.project_id)
            agg.add_test_runs(data.get('tests', []))
            agg.save(bmc_dir)
            generate_bmc_docx(agg.get_state(), bmc_dir)
            print(f"[DocReader] BMC full refresh — v{agg.get_state().version} ({agg.get_state().test_count} tests)")
            return True
        except Exception as e:
            print(f"[DocReader] refresh_bmc failed: {e}")
            return False

    def clean_database_duplicates(self) -> Dict:
        """
        Scan test_database.json for duplicate Validation entries.
        Resolves the database path the same way the collector does:
          1. Try project_paths.BCM_ROOT
          2. Fall back to Path.cwd() / BCM_Projects
        Backs up before writing. Returns summary dict.
        """
        # Resolve BCM_ROOT exactly as the collector does
        try:
            from project_paths import BCM_ROOT
        except ImportError:
            BCM_ROOT = Path.cwd() / "BCM_Projects"

        db_file = BCM_ROOT / "test_database.json"

        if not db_file.exists():
            return {
                'success': False,
                'error': f'test_database.json not found at {db_file}',
            }

        with open(db_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        tests = data.get('tests', [])
        # Baseline = source NOT validation. Use case-insensitive check.
        sparks = [iv for iv in tests if iv.get('source', '').lower() != 'validation']
        validation  = [iv for iv in tests if iv.get('source', '').lower() == 'validation']

        # Also catch any entries that were logged without a source tag
        # by scanning ALL tests for person-name duplicates
        all_persons = [iv.get('script_name', '').strip().lower() for iv in validation]

        seen = {}
        removed = []

        # Check validation entries first
        for iv in validation:
            key = iv.get('script_name', '').strip().lower()
            if not key:
                continue
            if key in seen:
                old_len = len(seen[key].get('results', '') or '')
                new_len = len(iv.get('results', '') or '')
                if new_len > old_len:
                    removed.append(seen[key])
                    seen[key] = iv
                else:
                    removed.append(iv)
            else:
                seen[key] = iv

        # Also check sparks entries for person-name collision with validation
        # (catches doc_reader entries filed without source='validation')
        clean_sparks = []
        for iv in sparks:
            key = iv.get('script_name', '').strip().lower()
            if key and key in seen:
                # Same person exists in validation — keep whichever has more results
                old_len = len(seen[key].get('results', '') or '')
                new_len = len(iv.get('results', '') or '')
                if new_len > old_len:
                    removed.append(seen[key])
                    seen[key] = iv
                else:
                    removed.append(iv)
            else:
                clean_sparks.append(iv)

        if not removed:
            file_result = self._clean_duplicate_docx_files()
            return {
                'success': True, 'removed': 0,
                'message': 'No database duplicates found.',
                **file_result,
            }

        # Backup before writing
        backup = db_file.with_suffix(
            f'.backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'
        )
        with open(backup, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        # Write cleaned
        data['tests'] = clean_sparks + list(seen.values())
        with open(db_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        return {
            'success': True,
            'removed': len(removed),
            'backup': str(backup),
            'removed_entries': [
                f"#{r.get('test_num','?')} {r.get('script_name','?')} ({r.get('source_version','?')})"
                for r in removed
            ],
            'remaining': len(data['tests']),
            **self._clean_duplicate_docx_files(),
        }

    def _clean_duplicate_docx_files(self) -> Dict:
        """Remove .docx files that don't follow Test Run_Packet_{Name}_completed.docx naming.
        Files matching a properly-named file by size are duplicates and get deleted."""
        import re as _re
        naming_pattern = _re.compile(r'^Test Run_Packet_.+_completed\.docx$')
        
        tested_dir = getattr(self, 'tested_dir', None)
        if not tested_dir:
            tested_dir = self.mover.tested_dir if hasattr(self, 'mover') else None
        if not tested_dir or not Path(tested_dir).exists():
            return {'files_removed': 0, 'files_removed_list': []}
        
        tested_dir = Path(tested_dir)
        docx_files = list(tested_dir.glob("*.docx"))
        
        proper = [f for f in docx_files if naming_pattern.match(f.name)]
        non_standard = [f for f in docx_files if not naming_pattern.match(f.name)]
        
        if not non_standard:
            return {'files_removed': 0, 'files_removed_list': []}
        
        proper_by_size = {f.stat().st_size: f for f in proper}
        
        removed_files = []
        for bad_file in non_standard:
            bad_size = bad_file.stat().st_size
            if bad_size in proper_by_size:
                match = proper_by_size[bad_size]
                print(f"  [FILE CLEANUP] DELETE: {bad_file.name} (dup of {match.name})")
                bad_file.unlink()
                removed_files.append(f"{bad_file.name} (dup of {match.name})")
            else:
                print(f"  [FILE CLEANUP] ORPHAN: {bad_file.name} ({bad_size:,} bytes)")
        
        return {'files_removed': len(removed_files), 'files_removed_list': removed_files}

    def delete_completed(self, person_name: str) -> Dict:
        """
        Remove a processed test_run from completed state.
        Atomically removes log entry, file, and clears Excel status.

        Returns dict with success flag and details.
        """
        result = self.mover.delete_completed(person_name)

        if result['success']:
            # Sync deletion back to Excel (still uses test_num for row matching)
            try:
                from excel_test_loader import get_loader
                loader = get_loader(self.project_id)
                loader.mark_pending_in_excel(result.get('test_num', 0))
            except Exception as e:
                print(f"Warning: Could not sync delete to Excel: {e}")

        return result

    def get_status(self) -> Dict:
        """Get current status: pending vs completed counts."""
        completed_names = self.mover.get_completed_names()
        completed_log = self.mover.get_completed_test_runs()
        
        return {
            'project': self.project_id,
            'completed_count': len(completed_names),
            'completed_names': completed_names,
            'change_log_entries': len(completed_log),
            'tested_dir': str(self.tested_dir) if hasattr(self, 'tested_dir') else str(self.mover.tested_dir),
            'report_dir': str(self.mover.project_root / '_AI_generated_report'),
        }


# ============================================================================
# GUI WIDGET - Tab for the Validation Test Collector
# ============================================================================

try:
    from PySide6.QtWidgets import (
        QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel,
        QTextEdit, QFileDialog, QMessageBox, QTableWidget,
        QTableWidgetItem, QGroupBox, QProgressBar, QSplitter
    )
    from PySide6.QtCore import Qt, Signal
    from PySide6.QtGui import QFont, QColor
    PYSIDE6_AVAILABLE = True
except ImportError:
    PYSIDE6_AVAILABLE = False


if PYSIDE6_AVAILABLE:
    class DocReaderWidget(QWidget):
        """
        GUI Tab: AI Doc Reader
        
        Upload completed .docx  auto-process  show report
        Target folder determined by active project.
        """
        
        # Signal emitted when a test is processed
        test_run_processed = Signal(dict)  # parsed_data
        
        def __init__(self, get_project_folder_fn=None, get_active_project_fn=None,
                     get_genesis_intel_fn=None, get_test_plan_fn=None, parent=None):
            super().__init__(parent)
            
            # Callable hooks from main app
            self._get_project_folder = get_project_folder_fn
            self._get_active_project = get_active_project_fn
            self._get_genesis_intel = get_genesis_intel_fn
            self._get_test_plan = get_test_plan_fn
            
            self._engine = None
            self._last_result = None
            
            self.init_ui()
        
        def init_ui(self):
            layout = QVBoxLayout(self)
            
            # Header
            header = QLabel("AI DOC READER -- Upload Completed Test Packet")
            header.setFont(QFont("Arial", 14, QFont.Bold))
            header.setStyleSheet("color: #CC0000; padding: 5px;")
            layout.addWidget(header)
            
            desc = QLabel(
                "Upload a completed .docx test packet. The system reads it, "
                "moves the person from PENDING  COMPLETED, logs the change, "
                "and generates an AI comparison report."
            )
            desc.setWordWrap(True)
            desc.setStyleSheet("color: #555; padding: 2px 5px;")
            layout.addWidget(desc)
            
            # Upload section
            upload_group = QGroupBox("Step 1: Upload Completed Packet")
            upload_layout = QHBoxLayout()
            
            self.upload_btn = QPushButton("Select Completed .docx File")
            self.upload_btn.setStyleSheet(
                "background-color: #0066cc; color: white; font-weight: bold; "
                "padding: 12px 24px; font-size: 13px;"
            )
            self.upload_btn.clicked.connect(self._on_upload)
            upload_layout.addWidget(self.upload_btn)
            
            self.file_label = QLabel("No file selected")
            self.file_label.setStyleSheet("color: #888; padding-left: 10px;")
            upload_layout.addWidget(self.file_label, 1)
            
            upload_group.setLayout(upload_layout)
            layout.addWidget(upload_group)
            
            # Target info
            target_group = QGroupBox("Step 2: Target Folder (auto-set by active project)")
            target_layout = QHBoxLayout()
            
            self.target_label = QLabel("Active Project: (loading...)")
            self.target_label.setFont(QFont("Consolas", 11))
            self.target_label.setStyleSheet("padding: 5px;")
            target_layout.addWidget(self.target_label)
            
            target_group.setLayout(target_layout)
            layout.addWidget(target_group)
            
            # Process button
            self.process_btn = QPushButton("PROCESS -- Read, File, Generate Report")
            self.process_btn.setStyleSheet(
                "background-color: #28a745; color: white; font-weight: bold; "
                "padding: 14px; font-size: 14px;"
            )
            self.process_btn.setEnabled(False)
            self.process_btn.clicked.connect(self._on_process)
            layout.addWidget(self.process_btn)

            # Option 2: Maintenance button row
            maint_layout = QHBoxLayout()
            self.cleanup_btn = QPushButton("OPTION 2: Clean Duplicate Test Runs from Database")
            self.cleanup_btn.setStyleSheet(
                "background-color: #6c3483; color: white; font-weight: bold; "
                "padding: 10px 20px; font-size: 12px;"
            )
            self.cleanup_btn.setToolTip(
                "Scans test_database.json for duplicate Validation entries.\n"
                "Keeps the most complete version per person+company.\n"
                "Auto-backs up before writing."
            )
            self.cleanup_btn.clicked.connect(self._on_clean_duplicates)
            maint_layout.addWidget(self.cleanup_btn)
            maint_layout.addStretch()
            layout.addLayout(maint_layout)

            # Splitter: left = report preview, right = status tables
            splitter = QSplitter(Qt.Horizontal)
            
            # Left: Report preview
            report_group = QGroupBox("AI Generated Report Preview")
            report_layout = QVBoxLayout()
            self.report_text = QTextEdit()
            self.report_text.setReadOnly(True)
            self.report_text.setFont(QFont("Consolas", 10))
            report_layout.addWidget(self.report_text)
            report_group.setLayout(report_layout)
            splitter.addWidget(report_group)
            
            # Right: Status
            status_group = QGroupBox("Test Run Status")
            status_layout = QVBoxLayout()
            
            # Completed tests table
            self.completed_table = QTableWidget()
            self.completed_table.setColumnCount(4)
            self.completed_table.setHorizontalHeaderLabels(["#", "Person", "Company", "Date"])
            self.completed_table.setColumnWidth(0, 40)
            self.completed_table.setColumnWidth(1, 150)
            self.completed_table.setColumnWidth(2, 150)
            self.completed_table.setColumnWidth(3, 130)
            status_layout.addWidget(QLabel("Completed Tests:"))
            status_layout.addWidget(self.completed_table)

            # Delete & Resubmit button
            self.delete_btn = QPushButton("Delete Selected & Resubmit")
            self.delete_btn.setStyleSheet(
                "background-color: #cc3300; color: white; font-weight: bold; "
                "padding: 8px 16px; font-size: 12px;"
            )
            self.delete_btn.setToolTip(
                "Select a completed test_run above, then click to remove it\n"
                "from the log so you can re-upload a corrected packet."
            )
            self.delete_btn.setEnabled(False)
            self.delete_btn.clicked.connect(self._on_delete)
            status_layout.addWidget(self.delete_btn)

            # Enable delete button when a row is selected
            self.completed_table.itemSelectionChanged.connect(
                lambda: self.delete_btn.setEnabled(
                    len(self.completed_table.selectedItems()) > 0
                )
            )

            # Pending tests table
            self.pending_table = QTableWidget()
            self.pending_table.setColumnCount(3)
            self.pending_table.setHorizontalHeaderLabels(["#", "Person", "Company"])
            self.pending_table.setColumnWidth(0, 40)
            self.pending_table.setColumnWidth(1, 150)
            self.pending_table.setColumnWidth(2, 150)
            status_layout.addWidget(QLabel("Pending Tests:"))
            status_layout.addWidget(self.pending_table)
            
            status_group.setLayout(status_layout)
            splitter.addWidget(status_group)
            
            splitter.setSizes([500, 300])
            layout.addWidget(splitter, 1)
            
            self._selected_file = None
            self._refresh_target()
        
        def _refresh_target(self):
            """Update target folder display and status tables."""
            project_id = "BCM_SUBSTRATE"
            if self._get_active_project:
                try:
                    project_id = self._get_active_project()
                except:
                    pass
            
            self.target_label.setText(
                f"Active Project: {project_id}    "
                f"BCM_TESTS/{project_id}/tested/"
            )
            
            # Refresh completed/pending tables
            self._refresh_status_tables(project_id)
        
        def _refresh_status_tables(self, project_id: str = None):
            """Refresh the completed and pending tables."""
            if not project_id and self._get_active_project:
                try:
                    project_id = self._get_active_project()
                except:
                    project_id = "BCM_SUBSTRATE"
            
            # Get project folder
            project_folder = None
            if self._get_project_folder:
                try:
                    project_folder = Path(self._get_project_folder())
                except:
                    pass
            
            if not project_folder:
                project_folder = Path.cwd() / 'BCM_TESTS' / project_id
            
            # Load completed from change_log
            mover = Test RunMover(project_folder)
            completed = mover.get_completed_test_runs()
            completed_names = mover.get_completed_names()
            
            # Populate completed table
            completed_entries = [c for c in completed if c.get('action') == 'TEST_COMPLETED']
            self.completed_table.setRowCount(len(completed_entries))
            for row, entry in enumerate(completed_entries):
                self.completed_table.setItem(row, 0, QTableWidgetItem(str(entry.get('test_num', ''))))
                self.completed_table.setItem(row, 1, QTableWidgetItem(entry.get('script_name', '')))
                self.completed_table.setItem(row, 2, QTableWidgetItem(entry.get('source_version', '')))
                ts = entry.get('timestamp', '')[:16].replace('T', ' ')
                self.completed_table.setItem(row, 3, QTableWidgetItem(ts))
            
            # Populate pending table (from test_run plan minus completed)
            # Match by person name — case-insensitive last-name comparison
            completed_names_lower = [n.strip().lower() for n in completed_names]
            pending_tests = []
            if self._get_test_plan:
                try:
                    plan = self._get_test_plan()
                    print(f"  Doc Reader pending: plan has {len(plan)} entries, completed_names={completed_names}")
                    for ip in plan:
                        ip_name = ip.get('name', '')
                        # Skip placeholder entries
                        if not ip_name or ip_name.startswith('['):
                            continue
                        # Match by full name or last name
                        name_matched = (
                            ip_name.strip().lower() in completed_names_lower
                            or any(
                                ip_name.split()[-1].lower() in cn
                                for cn in completed_names_lower
                            )
                        )
                        if not name_matched:
                            pending_tests.append(ip)
                    print(f"  Pending after filter: {len(pending_tests)}")
                except Exception as e:
                    print(f"  ⚠ Pending load error: {e}")
            
            self.pending_table.setRowCount(len(pending_tests))
            for row, ip in enumerate(pending_tests):
                self.pending_table.setItem(row, 0, QTableWidgetItem(str(ip.get('num', ''))))
                self.pending_table.setItem(row, 1, QTableWidgetItem(ip.get('name', '')))
                self.pending_table.setItem(row, 2, QTableWidgetItem(ip.get('source_version', '')))
        
        def _on_upload(self):
            """Open file dialog to select completed .docx."""
            filepath, _ = QFileDialog.getOpenFileName(
                self,
                "Select Completed Test Packet",
                str(Path.cwd() / 'BCM_TESTS'),
                "Word Documents (*.docx);;All Files (*)"
            )
            
            if filepath:
                self._selected_file = Path(filepath)
                self.file_label.setText(self._selected_file.name)
                self.file_label.setStyleSheet("color: #006600; font-weight: bold; padding-left: 10px;")
                self.process_btn.setEnabled(True)
                self._refresh_target()
        
        def _on_process(self):
            """Process the uploaded document."""
            if not self._selected_file or not self._selected_file.exists():
                QMessageBox.warning(self, "No File", "Please select a completed .docx file first.")
                return
            
            # Get active project and folder
            project_id = "BCM_SUBSTRATE"
            if self._get_active_project:
                try:
                    project_id = self._get_active_project()
                except:
                    pass
            
            project_folder = None
            if self._get_project_folder:
                try:
                    project_folder = Path(self._get_project_folder())
                except:
                    pass
            
            if not project_folder:
                project_folder = Path.cwd() / 'BCM_TESTS' / project_id
            
            # Get genesis intel if available
            genesis_intel = None
            if self._get_genesis_intel:
                try:
                    genesis_intel = self._get_genesis_intel()
                except:
                    pass
            
            # Process
            try:
                engine = DocReaderEngine(project_folder, project_id)
                result = engine.process(self._selected_file, genesis_intel)
                self._last_result = result
                
                if result['success']:
                    # Show report in preview
                    report_path = Path(result['report_path'])
                    if report_path.exists():
                        with open(report_path, 'r', encoding='utf-8') as f:
                            self.report_text.setText(f.read())
                    
                    # Refresh status tables
                    self._refresh_status_tables(project_id)
                    
                    # Emit signal for main app
                    self.test_run_processed.emit(result['parsed'])
                    
                    QMessageBox.information(
                        self,
                        " Test Run Processed",
                        f"{result['script_name']} ({result['source_version']})\n\n"
                        f" Moved to: {project_id}/tested/\n"
                        f" Equipment discussed: {result['equipment_count']}\n"
                        f" AI Report: {project_id}/_AI_generated_report/\n"
                        f" Change logged to: change_log.json\n\n"
                        f"Person removed from pending list."
                    )
                else:
                    QMessageBox.warning(
                        self,
                        "Incomplete Packet",
                        f"This document doesn't appear fully completed.\n\n"
                        f"{result.get('error', 'Unknown issue')}\n\n"
                        f"Please fill in Results, Equipment, or Experiment answers."
                    )
            except ImportError as e:
                QMessageBox.critical(self, "Missing Dependency", 
                    f"python-docx required:\npip install python-docx\n\n{e}")
            except Exception as e:
                QMessageBox.critical(self, "Processing Error", 
                    f"Failed to process document:\n\n{str(e)}")

        def _on_delete(self):
            """Delete selected completed test_run so it can be resubmitted."""
            row = self.completed_table.currentRow()
            if row < 0:
                return

            person_item = self.completed_table.item(row, 1)
            if not person_item or not person_item.text().strip():
                return

            person = person_item.text().strip()

            confirm = QMessageBox.question(
                self,
                "Delete Test",
                f"Delete {person} from completed log?\n\n"
                f"This removes the log entry and the filed .docx so you can re-upload "
                f"a corrected packet. The test_run will return to PENDING.",
                QMessageBox.Yes | QMessageBox.No
            )
            if confirm != QMessageBox.Yes:
                return

            project_id = "BCM_SUBSTRATE"
            if self._get_active_project:
                try:
                    project_id = self._get_active_project()
                except:
                    pass

            project_folder = None
            if self._get_project_folder:
                try:
                    project_folder = Path(self._get_project_folder())
                except:
                    pass
            if not project_folder:
                project_folder = Path.cwd() / "BCM_TESTS" / project_id

            try:
                engine = DocReaderEngine(project_folder, project_id)
                result = engine.delete_completed(person)

                if result["success"]:
                    # Emit DELETE signal so collector removes from in-memory database
                    self.test_run_processed.emit({
                        "action": "DELETE",
                        "test_num": result.get("test_num", 0),
                        "script_name": person,
                    })
                    self._refresh_status_tables(project_id)
                    self.delete_btn.setEnabled(False)
                    QMessageBox.information(
                        self,
                        "Deleted",
                        f"{person} removed.\n"
                        f"You can now re-upload a corrected packet."
                    )
                else:
                    QMessageBox.warning(self, "Delete Failed",
                        result.get("error", "Unknown error"))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to delete:\n{str(e)}")

        def _on_clean_duplicates(self):
            """Run duplicate cleanup on test_database.json."""
            confirm = QMessageBox.question(
                self,
                "Clean Duplicate Test Runs",
                "This will scan test_database.json for duplicate Validation entries\n"
                "and remove the less complete copy of each duplicate.\n\n"
                "A backup will be created automatically before any changes.\n\n"
                "Continue?",
                QMessageBox.Yes | QMessageBox.No
            )
            if confirm != QMessageBox.Yes:
                return

            project_id = "BCM_SUBSTRATE"
            if self._get_active_project:
                try:
                    project_id = self._get_active_project()
                except:
                    pass

            project_folder = None
            if self._get_project_folder:
                try:
                    project_folder = Path(self._get_project_folder())
                except:
                    pass
            if not project_folder:
                project_folder = Path.cwd() / 'BCM_TESTS' / project_id

            try:
                engine = DocReaderEngine(project_folder, project_id)
                result = engine.clean_database_duplicates()

                if not result['success']:
                    QMessageBox.warning(self, "Cleanup Failed", result.get('error', 'Unknown error'))
                    return

                if result['removed'] == 0:
                    files_removed = result.get('files_removed', 0)
                    if files_removed:
                        files_list = '\n'.join(f"  - {f}" for f in result.get('files_removed_list', []))
                        QMessageBox.information(self, "Cleanup Complete",
                            f"Database clean — no duplicates.\n\n"
                            f"Duplicate files removed: {files_removed}\n{files_list}")
                    else:
                        QMessageBox.information(self, "All Clean",
                            "Database clean. No duplicate files found.")
                    return

                removed_list = '\n'.join(f"  - {e}" for e in result.get('removed_entries', []))
                files_removed = result.get('files_removed', 0)
                files_list = '\n'.join(f"  - {f}" for f in result.get('files_removed_list', []))
                
                file_msg = ""
                if files_removed:
                    file_msg = f"\n\nDuplicate files removed: {files_removed}\n{files_list}"
                
                QMessageBox.information(
                    self,
                    "Duplicates Removed",
                    f"Removed {result['removed']} duplicate(s):\n\n"
                    f"{removed_list}\n\n"
                    f"Remaining test_runs: {result['remaining']}\n"
                    f"Backup saved: {result.get('backup', 'N/A')}"
                    f"{file_msg}"
                )
                self.report_text.setText(
                    f"DUPLICATE CLEANUP COMPLETE\n"
                    f"{'='*50}\n"
                    f"Removed: {result['removed']} duplicate(s)\n\n"
                    f"Entries removed:\n{removed_list}\n\n"
                    f"Remaining total: {result['remaining']}\n"
                    f"Backup: {result.get('backup','N/A')}\n"
                    f"{f'Files cleaned: {files_removed}' if files_removed else ''}\n"
                )
                self._refresh_status_tables(project_id)

            except Exception as e:
                QMessageBox.critical(self, "Cleanup Error", f"Failed:\n\n{str(e)}")

        def refresh(self):
            """Public refresh method - call after project switch."""
            self._refresh_target()
