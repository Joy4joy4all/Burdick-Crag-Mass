# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
BMC ENGINE — Accumulative KNN Business Model Canvas
=====================================================
Location: TITS_GIBUSH_AISOS_SPINE_ICORPS/bmc_engine.py

RULE: Every BMC claim must trace to >= 1 test_run number.
      Claims with < 3 tests are flagged "NEEDS VALIDATION"
      BMC state ACCUMULATES — each test adds, never regenerates.

PIPELINE:
  Test Run → Q-Cube Position → BMC Field Scoring → Aggregation → KNN Weighting

KNN LOGIC:
  Test Runs in the same Q-Cube neighborhood reinforce each other.
  L2+OC tests saying "$300K damage" weight higher when 3 neighbors agree.
  Isolated claims from single tests get "NEEDS VALIDATION" flag.

HIERARCHY: THE PERSON > TITS > GIBUSH > AISOS/SPINE > CaaS > SaaS
© 2025-2026 Stephen J. Burdick Sr. — All Rights Reserved.
For all the industrial workers lost to preventable acts.
"""

import json
import re
import math
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass, field, asdict


# ══════════════════════════════════════════════════════════════
# BMC FIELD DEFINITIONS — 10 fields per Strategyzer/BCM
# ══════════════════════════════════════════════════════════════

BMC_FIELDS = [
    "substrate_classs",
    "value_propositions",
    "channels",
    "customer_relationships",
    "revenue_streams",
    "key_resources",
    "key_activities",
    "key_partnerships",
    "cost_structure",
    "unfair_advantage",
]

# Display names for UI
BMC_DISPLAY_NAMES = {
    "substrate_classs":      "Substrate Classs",
    "value_propositions":     "Value Propositions",
    "channels":               "Channels",
    "customer_relationships": "Customer Relationships",
    "revenue_streams":        "Revenue Streams",
    "key_resources":          "Key Resources",
    "key_activities":         "Key Activities",
    "key_partnerships":       "Key Partnerships",
    "cost_structure":         "Cost Structure",
    "unfair_advantage":       "Unfair Advantage",
}


# ══════════════════════════════════════════════════════════════
# KEYWORD EXTRACTORS — per BMC field
# ══════════════════════════════════════════════════════════════

BMC_EXTRACTORS = {
    "substrate_classs": {
        "source_fields": ["test_category", "title", "source_version", "results"],
        "cube_weight": {"L2": 1.0, "L3": 0.8, "L1": 0.5},
        "themes": {
            "mill_maintenance_manager": {
                "keywords": ["maintenance", "manager", "superintendent", "supt",
                             "reliability", "planner", "coordinator"],
                "label": "Mill Maintenance/Reliability Managers",
            },
            "mill_operations": {
                "keywords": ["operator", "operations", "shift", "foreman",
                             "process", "board operator"],
                "label": "Mill Operations Personnel",
            },
            "procurement_exec": {
                "keywords": ["procurement", "buyer", "purchasing", "regional",
                             "VP", "executive", "director"],
                "label": "Procurement/Executive Decision Makers",
            },
            "oem_service": {
                "keywords": ["voith", "andritz", "valmet", "OEM", "service",
                             "consultant", "independent"],
                "label": "OEM / Service Partners",
            },
            "wood_supply": {
                "keywords": ["hauler", "trucking", "wood", "log", "yard",
                             "supplier", "chip supplier", "secondary"],
                "label": "Wood Supply Chain / Haulers",
            },
        },
        "blocker_keywords": ["no money", "won't invest", "budget constrained",
                             "no budget", "not justified"],
        "buyer_keywords": ["would buy", "worth it", "immediately", "will try",
                           "AMO budget", "authorize"],
    },

    "value_propositions": {
        "source_fields": ["results", "experiments", "substrate_impacts"],
        "cube_weight": {"OC": 1.0, "OB": 0.8, "OA": 0.6},
        "themes": {
            "damage_prevention": {
                "keywords": ["damage", "$100k", "rock event", "shutdown",
                             "digester", "screen basket", "rock", "tramp"],
                "label": "Damage Prevention ($100K-$1M events)",
            },
            "planned_maintenance": {
                "keywords": ["$300k", "reactive", "accepted damage", "planned",
                             "emergency", "PM", "repair cost", "annual"],
                "label": "Reactive → Planned Maintenance ($300K/yr)",
            },
            "operator_visibility": {
                "keywords": ["will try", "would adopt", "readings", "Go look",
                             "real-time", "visibility", "monitoring"],
                "label": "Real-Time Operator Visibility",
            },
            "chemical_optimization": {
                "keywords": ["$500k", "chemical savings", "dosing", "off-quality",
                             "consistency", "calcium", "scale"],
                "label": "Chemical Optimization ($500K/yr potential)",
            },
            "supplier_accountability": {
                "keywords": ["vendor", "supplier", "docking", "certified clean",
                             "verification", "certified", "quality spec"],
                "label": "Supplier Accountability / Verification",
            },
            "bark_line_safety": {
                "keywords": ["bark blow", "bark line", "hog fuel", "boiler grate",
                             "boiler tube", "bark knife"],
                "label": "Bark/Hog Fuel Line Safety Monitoring",
            },
            "commissioning_baseline": {
                "keywords": ["baseline", "commissioning", "new equipment",
                             "day one", "$1B", "investment protection"],
                "label": "New Equipment Commissioning Baseline Data",
            },
        },
        "validation_threshold": 3,
    },

    "channels": {
        "source_fields": ["action_iterate", "results", "source_version"],
        "cube_weight": {"L3": 1.0, "L2": 0.8, "L1": 0.5},
        "themes": {
            "expert_referral": {
                "keywords": ["contact", "referral", "introduced", "recommended",
                             "network", "knows"],
                "label": "Expert/Industry Referral Network",
            },
            "oem_bundling": {
                "keywords": ["voith", "andritz", "valmet", "oem", "bundled",
                             "service agreement"],
                "label": "OEM Service Bundling",
            },
            "direct_sales": {
                "keywords": ["cavanaugh", "calibration", "service partner",
                             "direct", "install"],
                "label": "Direct Sales / Service Partner",
            },
            "conference": {
                "keywords": ["tappi", "conference", "publication", "trade show",
                             "PIMA"],
                "label": "Industry Conference / Publication",
            },
        },
    },

    "customer_relationships": {
        "source_fields": ["results", "experiments", "action_iterate"],
        "cube_weight": {"L2": 1.0, "L1": 0.8, "L3": 0.6},
        "themes": {
            "diagnostic_first": {
                "keywords": ["find the source", "diagnostic", "where", "upstream",
                             "root cause", "pinpoint"],
                "label": "Diagnostic-First Engagement",
            },
            "proof_of_value": {
                "keywords": ["proof", "PV", "verify", "pilot", "trial", "demo",
                             "test install"],
                "label": "Proof-of-Value / Pilot Program",
            },
            "operator_training": {
                "keywords": ["training", "adopt", "certification", "empower",
                             "MT would adopt", "usability"],
                "label": "Operator Training & Adoption",
            },
            "amo_path": {
                "keywords": ["AMO", "maintenance budget", "not CapEx", "authorize",
                             "operating budget", "annual maintenance"],
                "label": "AMO Budget Path (not CapEx)",
            },
        },
    },

    "revenue_streams": {
        "source_fields": ["results", "substrate_impacts"],
        "cube_weight": {"L2": 1.0, "L3": 0.9, "L1": 0.3},
        "themes": {
            "annual_damage_prevention": {
                "keywords": ["per year", "/yr", "annual", "yearly", "$300k",
                             "$500k", "recurring"],
                "label": "Annual Damage Prevention Revenue",
            },
            "event_prevention": {
                "keywords": ["per event", "each", "per incident", "$100k",
                             "one-time", "shutdown cost"],
                "label": "Event Prevention Value",
            },
            "willingness_to_pay": {
                "keywords": ["would buy", "worth it", "immediately",
                             "pay for itself", "ROI", "<2yr", "payback"],
                "label": "Buyer Willingness Signals",
            },
            "price_sensitivity": {
                "keywords": ["won't invest", "too expensive", "no budget",
                             "constrained", "severely constrained"],
                "label": "Price Sensitivity / Budget Blockers",
            },
        },
    },

    "key_resources": {
        "source_fields": ["results", "action_iterate"],
        "cube_weight": {"OB": 1.0, "OC": 0.9, "OA": 0.7},
        "themes": {
            "detection_gap": {
                "keywords": ["nothing", "no detection", "don't have",
                             "doesn't exist", "no way to", "can't detect",
                             "Go look", "nothing on the piles"],
                "label": "Detection Technology Gap (nothing exists)",
            },
            "algorithm_ip": {
                "keywords": ["algorithm", "acoustic", "sensor", "classification",
                             "triangulation", "real-time", "patent"],
                "label": "Algorithm / Patent IP",
            },
            "domain_expertise": {
                "keywords": ["30 year", "kraft", "mill experience", "foreman",
                             "operator knowledge", "tribal"],
                "label": "Domain Expertise (30yr kraft mill)",
            },
        },
    },

    "key_activities": {
        "source_fields": ["action_iterate", "experiments"],
        "cube_weight": {"L1": 1.0, "L2": 1.0, "L3": 1.0},
        "themes": {
            "pilot_install": {
                "keywords": ["next step", "pilot", "install", "test",
                             "deploy", "prototype"],
                "label": "Pilot Installation / Testing",
            },
            "test_run_validate": {
                "keywords": ["test_run", "validate", "customer discovery",
                             "BCM", "validation"],
                "label": "Customer Discovery / Validation",
            },
            "develop_build": {
                "keywords": ["develop", "build", "design", "engineer",
                             "code", "sensor array"],
                "label": "Product Development / Engineering",
            },
            "quantify_damage": {
                "keywords": ["quantify", "measure", "baseline", "data",
                             "track", "log"],
                "label": "Damage Quantification & Data Collection",
            },
        },
    },

    "key_partnerships": {
        "source_fields": ["source_version", "results", "action_iterate"],
        "cube_weight": {"L3": 1.0, "L2": 0.8, "L1": 0.5},
        "themes": {
            "mill_partner": {
                "keywords": ["woodland", "IP", "international paper",
                             "georgia pacific", "GP", "sappi", "verso"],
                "label": "Mill Operational Partners",
            },
            "oem_partner": {
                "keywords": ["voith", "andritz", "valmet", "OEM",
                             "kadant", "AFT"],
                "label": "OEM Equipment Partners",
            },
            "academic_partner": {
                "keywords": ["MIT", "UMaine", "university", "NSF",
                             "BCM", "research"],
                "label": "Academic / Research Partners",
            },
            "service_partner": {
                "keywords": ["cavanaugh", "service", "calibration",
                             "installation", "maintenance provider"],
                "label": "Service / Installation Partners",
            },
        },
    },

    "cost_structure": {
        "source_fields": ["results", "substrate_impacts"],
        "cube_weight": {"L2": 1.0, "L3": 0.9, "L1": 0.5},
        "themes": {
            "budget_blockers": {
                "keywords": ["no money", "won't invest", "budget constrained",
                             "severely constrained", "no CapEx"],
                "label": "Budget Blockers Identified",
            },
            "budget_path": {
                "keywords": ["AMO", "maintenance budget", "operating budget",
                             "not CapEx", "annual maintenance outage"],
                "label": "AMO/Operating Budget Path",
            },
            "price_anchor": {
                "keywords": ["$75K", "system price", "payback", "ROI",
                             "4 month", "return"],
                "label": "Price Anchor: $75K system vs $300K+/yr damage",
            },
        },
    },

    "unfair_advantage": {
        "source_fields": ["results", "action_iterate", "experiments"],
        "cube_weight": {"L1": 1.0, "L2": 1.0, "L3": 1.0},
        "themes": {
            "operator_inventor": {
                "keywords": ["30 year", "operator", "invented", "foreman",
                             "kraft", "built from floor"],
                "label": "Operator-Inventor: 30yr kraft mill foreman",
            },
            "no_competition": {
                "keywords": ["nothing exists", "no detection", "first",
                             "nobody", "only", "unique"],
                "label": "No Existing Competition in Detection Space",
            },
            "patent_ip": {
                "keywords": ["patent", "patented", "IP", "proprietary",
                             "algorithm", "physics-based"],
                "label": "Patented Physics-Based Detection Technology",
            },
        },
    },
}


# ══════════════════════════════════════════════════════════════
# DATA CLASSES
# ══════════════════════════════════════════════════════════════

@dataclass
class ThemeEvidence:
    """Evidence for a single theme within a BMC field."""
    theme_id: str
    label: str
    test_nums: List[int] = field(default_factory=list)
    quotes: List[str] = field(default_factory=list)
    score: float = 0.0
    equipment_costs: List[float] = field(default_factory=list)
    status: str = "NEEDS_VALIDATION"  # NEEDS_VALIDATION | EMERGING | VALIDATED


@dataclass
class BMCFieldState:
    """State of a single BMC field after aggregation."""
    field_id: str
    display_name: str
    themes: Dict[str, ThemeEvidence] = field(default_factory=dict)
    total_test_runs: int = 0
    confidence: float = 0.0  # 0.0 - 1.0
    summary: str = ""


@dataclass
class BMCState:
    """Complete BMC state — all 10 fields."""
    project_id: str
    version: int = 1
    generated_at: str = ""
    test_count: int = 0
    fields: Dict[str, BMCFieldState] = field(default_factory=dict)
    
    # Diff tracking
    previous_version: int = 0
    changes_since_last: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict:
        """Serialize for JSON storage."""
        d = {
            "project_id": self.project_id,
            "version": self.version,
            "generated_at": self.generated_at,
            "test_count": self.test_count,
            "previous_version": self.previous_version,
            "changes_since_last": self.changes_since_last,
            "fields": {},
        }
        for fid, fstate in self.fields.items():
            fd = {
                "field_id": fstate.field_id,
                "display_name": fstate.display_name,
                "total_test_runs": fstate.total_tests,
                "confidence": round(fstate.confidence, 3),
                "summary": fstate.summary,
                "themes": {},
            }
            for tid, te in fstate.themes.items():
                fd["themes"][tid] = {
                    "theme_id": te.theme_id,
                    "label": te.label,
                    "test_nums": te.test_nums,
                    "quotes": te.quotes,
                    "score": round(te.score, 3),
                    "equipment_costs": te.equipment_costs,
                    "status": te.status,
                }
            d["fields"][fid] = fd
        return d

    @classmethod
    def from_dict(cls, data: Dict) -> 'BMCState':
        """Deserialize from JSON."""
        state = cls(
            project_id=data.get("project_id", ""),
            version=data.get("version", 1),
            generated_at=data.get("generated_at", ""),
            test_count=data.get("test_count", 0),
            previous_version=data.get("previous_version", 0),
            changes_since_last=data.get("changes_since_last", []),
        )
        for fid, fd in data.get("fields", {}).items():
            fstate = BMCFieldState(
                field_id=fd["field_id"],
                display_name=fd["display_name"],
                total_test_runs=fd.get("total_test_runs", 0),
                confidence=fd.get("confidence", 0.0),
                summary=fd.get("summary", ""),
            )
            for tid, td in fd.get("themes", {}).items():
                te = ThemeEvidence(
                    theme_id=td["theme_id"],
                    label=td["label"],
                    test_nums=td.get("test_nums", []),
                    quotes=td.get("quotes", []),
                    score=td.get("score", 0.0),
                    equipment_costs=td.get("equipment_costs", []),
                    status=td.get("status", "NEEDS_VALIDATION"),
                )
                fstate.themes[tid] = te
            state.fields[fid] = fstate
        return state


# ══════════════════════════════════════════════════════════════
# BMC FIELD SCORER — scores one test_run against all fields
# ══════════════════════════════════════════════════════════════

class BMCFieldScorer:
    """
    Score a single test_run against all 10 BMC fields.
    
    Returns dict: {field_id: {theme_id: (score, quotes)}}
    
    KNN weight comes from Q-Cube position matching the field's
    preferred cube weight. L2 tests score higher on substrate_classs
    because managers are buyers.
    """

    def __init__(self, extractors: Dict = None):
        self.extractors = extractors or BMC_EXTRACTORS

    def score_test_run(self, test_run: dict) -> Dict[str, Dict[str, Tuple[float, List[str]]]]:
        """
        Score one test_run against all BMC fields.
        
        Args:
            test_run: dict with keys matching Test Run dataclass
                       (results, experiments, action_iterate, q_layer, q_object,
                        substrate_impacts, company, test_category, title, person,
                        test_num)
        
        Returns:
            {field_id: {theme_id: (weighted_score, [evidence_quotes])}}
        """
        scores = {}
        
        for field_id, extractor in self.extractors.items():
            field_scores = {}
            
            # Build searchable text from source fields
            text_parts = []
            for src_field in extractor.get("source_fields", []):
                val = test_entry.get(src_field, "")
                if isinstance(val, str):
                    text_parts.append(val)
                elif isinstance(val, list):
                    # substrate_impacts is a list of dicts
                    for item in val:
                        if isinstance(item, dict):
                            text_parts.append(str(item.get("equipment", "")))
                            text_parts.append(str(item.get("notes", "")))
                        else:
                            text_parts.append(str(item))
            
            search_text = " ".join(text_parts).lower()
            if not search_text.strip():
                scores[field_id] = field_scores
                continue
            
            # Strip placeholder and template sections from search text
            for ph in ('[pending', '[fill', '[synthesized', '(no answer',
                       '[ ] validated', '[ ] invalidated', '[ ] unclear',
                       'targeted questions:'):
                while ph in search_text:
                    idx = search_text.find(ph)
                    # Remove from placeholder to next sentence boundary or end
                    end = search_text.find('.', idx + 20)
                    if end == -1:
                        end = min(idx + 100, len(search_text))
                    search_text = search_text[:idx] + search_text[end:]
            
            # Calculate cube weight for this test_run + field
            cube_weight = self._get_cube_weight(test_run, extractor)
            
            # Score each theme
            for theme_id, theme_def in extractor.get("themes", {}).items():
                keywords = theme_def.get("keywords", [])
                hit_count, evidence = self._match_keywords(search_text, keywords,
                                                            test_run)
                
                if hit_count > 0:
                    # Score = hits * cube_weight
                    # More keyword hits = stronger evidence
                    # Cube position matching = higher weight
                    raw_score = hit_count * cube_weight
                    
                    # Boost for component metric data
                    eq_costs = self._extract_equipment_costs(test_run)
                    if eq_costs:
                        # Having $ evidence is a strong signal
                        raw_score *= 1.5
                    
                    field_scores[theme_id] = (raw_score, evidence)
            
            scores[field_id] = field_scores
        
        return scores

    def _get_cube_weight(self, test_run: dict, extractor: dict) -> float:
        """Get cube position weight for this test_run/field combo."""
        weights = extractor.get("cube_weight", {})
        
        # Try layer weight first (L1/L2/L3)
        layer = test_entry.get("q_layer", "")
        layer_w = weights.get(layer, 0.5)
        
        # Try object weight (OA/OB/OC)
        obj = test_entry.get("q_object", "")
        obj_w = weights.get(obj, 0.5)
        
        # Use whichever is defined in the extractor
        # If both layer and object weights exist, average them
        has_layer = any(k.startswith("L") for k in weights)
        has_obj = any(k.startswith("O") for k in weights)
        
        if has_layer and has_obj:
            return (layer_w + obj_w) / 2
        elif has_layer:
            return layer_w
        elif has_obj:
            return obj_w
        return 0.5

    def _match_keywords(self, text: str, keywords: List[str],
                        test_run: dict) -> Tuple[int, List[str]]:
        """Match keywords in text, return (hit_count, evidence_quotes)."""
        hits = 0
        evidence = []
        seen_snippets = set()  # Within-theme dedup
        
        # Patterns that indicate template structure, not real evidence
        skip_patterns = (
            '[fill', '[pending', '[synthesized', '(no answer',
            '[ ] validated', '[ ] invalidated', '[ ] unclear',
            'targeted questions:', 'q: what', 'q: how', 'q: do ',
            'q: when', 'q: where', 'q: is ',
            '\na: .', '\na: ,', 'a: . ',
        )
        
        for kw in keywords:
            kw_lower = kw.lower()
            if kw_lower in text:
                hits += 1
                # Extract surrounding context as quote (±50 chars)
                idx = text.find(kw_lower)
                start = max(0, idx - 50)
                end = min(len(text), idx + len(kw_lower) + 50)
                snippet = text[start:end].strip()
                
                if not snippet:
                    continue
                # Skip template/placeholder text
                if any(pat in snippet for pat in skip_patterns):
                    continue
                # Skip if snippet is too short to be meaningful (<20 chars of real content)
                if len(snippet.replace('.', '').replace(',', '').strip()) < 20:
                    continue
                # Within-theme dedup by first 60 chars
                snippet_key = snippet[:60]
                if snippet_key in seen_snippets:
                    continue
                seen_snippets.add(snippet_key)
                
                num = test_entry.get("test_num", "?")
                person = test_entry.get("script_name", "?")
                evidence.append(f"[#{num} {person}] ...{snippet}...")
        
        return hits, evidence

    def _extract_equipment_costs(self, test_run: dict) -> List[float]:
        """Pull numeric costs from substrate_impacts."""
        costs = []
        for eq in test_entry.get("substrate_impacts", []):
            if isinstance(eq, dict):
                cost = eq.get("cost", 0)
                if isinstance(cost, (int, float)) and cost > 0:
                    costs.append(float(cost))
        return costs


# ══════════════════════════════════════════════════════════════
# BMC AGGREGATOR — combines all tests into BMC state
# ══════════════════════════════════════════════════════════════

class BMCAggregator:
    """
    Aggregate all tests into a complete BMC state.
    
    ACCUMULATIVE: Each call builds on previous state.
    KNN: Test Runs in the same Q-Cube neighborhood reinforce each other.
    
    Usage:
        agg = BMCAggregator(active_project_id)
        agg.load_previous()  # loads last saved state
        agg.add_test_runs(new_tests)  # accumulates
        state = agg.get_state()
        agg.save(bmc_dir)
    """

    def __init__(self, project_id: str):
        self.project_id = project_id
        self.scorer = BMCFieldScorer()
        self.state = BMCState(project_id=project_id)
        self._processed_nums: set = set()
        self._processed_names: set = set()
        self._all_test_runs: List[dict] = []
        
        # Initialize empty fields
        self._init_fields()

    def _init_fields(self):
        """Initialize all 10 BMC fields with empty themes."""
        for field_id in BMC_FIELDS:
            extractor = BMC_EXTRACTORS.get(field_id, {})
            fstate = BMCFieldState(
                field_id=field_id,
                display_name=BMC_DISPLAY_NAMES.get(field_id, field_id),
            )
            # Pre-populate theme slots
            for theme_id, theme_def in extractor.get("themes", {}).items():
                fstate.themes[theme_id] = ThemeEvidence(
                    theme_id=theme_id,
                    label=theme_def.get("label", theme_id),
                )
            self.state.fields[field_id] = fstate

    def load_previous(self, bmc_dir: Path) -> bool:
        """
        Load previous BMC state from disk.
        Returns True if loaded, False if starting fresh.
        """
        state_file = bmc_dir / "bmc_state.json"
        if state_file.exists():
            try:
                with open(state_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                self.state = BMCState.from_dict(data)
                # Track which tests are already scored
                for fstate in self.state.fields.values():
                    for te in fstate.themes.values():
                        self._processed_nums.update(te.test_nums)
                print(f"[BMC] Loaded v{self.state.version} "
                      f"({self.state.test_count} tests)")
                return True
            except Exception as e:
                print(f"[BMC] Error loading state: {e}")
        return False

    @staticmethod
    def _has_real_content(iv: dict) -> bool:
        """Check if test_run has real results vs only placeholders.
        Requires the 'results' field specifically to have real content."""
        placeholders = ('[pending', '[fill', '[synthesized')
        
        # The results field is the primary evidence — must have real content
        results = str(iv.get('results', '') or '').strip()
        if not results:
            return False
        
        results_lower = results.lower()
        # Reject if results starts with a placeholder
        if any(results_lower.startswith(ph) for ph in placeholders):
            return False
        
        # Reject if only empty bullets
        stripped = results.replace('⬢', '').replace('•', '').replace('\n', '').strip()
        if not stripped or stripped.lower().startswith('[fill'):
            return False
        
        # Strip placeholder text and check remaining length
        clean = results
        for ph in ('[PENDING', '[FILL', '[fill', '[Pending', '[SYNTHESIZED'):
            idx = clean.find(ph)
            if idx >= 0:
                clean = clean[:idx]
        
        return len(clean.strip()) > 50

    def add_test_runs(self, test_runs: List[dict]):
        """
        ACCUMULATE new tests into BMC state.
        Skips already-processed tests (by test_num).
        This is the KNN accumulation — each new test adds evidence.
        """
        new_count = 0
        changes = []
        
        for iv in test_runs:
            num = iv.get("test_num", 0)
            person = iv.get("script_name", "").strip().lower()
            
            # Skip already-processed tests (by person name)
            if person and person in self._processed_names:
                continue
            # Legacy fallback: also skip by test_num
            if num in self._processed_nums:
                continue
            
            # Score this test_run against all fields
            scores = self.scorer.score_test_run(iv)
            
            # Accumulate scores into state
            for field_id, field_scores in scores.items():
                fstate = self.state.fields.get(field_id)
                if not fstate:
                    continue
                
                for theme_id, (score, evidence) in field_scores.items():
                    te = fstate.themes.get(theme_id)
                    if not te:
                        # New theme discovered
                        extractor = BMC_EXTRACTORS.get(field_id, {})
                        theme_def = extractor.get("themes", {}).get(theme_id, {})
                        te = ThemeEvidence(
                            theme_id=theme_id,
                            label=theme_def.get("label", theme_id),
                        )
                        fstate.themes[theme_id] = te
                    
                    # ACCUMULATE — add test evidence (dedup by first 60 chars)
                    if num not in te.test_nums:
                        te.test_nums.append(num)
                    existing_keys = {q[:60] for q in te.quotes}
                    for eq in evidence:
                        if eq[:60] not in existing_keys:
                            te.quotes.append(eq)
                            existing_keys.add(eq[:60])
                    
                    # Keep quotes manageable (max 5 per theme)
                    if len(te.quotes) > 5:
                        te.quotes = te.quotes[-5:]
                    
                    # Accumulate component metrics
                    eq_costs = self.scorer._extract_equipment_costs(iv)
                    te.equipment_costs.extend(eq_costs)
                    
                    # KNN-weighted score: old_score + new_score * neighbor_boost
                    neighbor_boost = self._knn_boost(iv, te.test_nums)
                    te.score += score * neighbor_boost
                    
                    # Update status thresholds
                    n_tests = len(te.test_nums)
                    if n_tests >= 3:
                        old_status = te.status
                        te.status = "VALIDATED"
                        if old_status != "VALIDATED":
                            changes.append(
                                f"{BMC_DISPLAY_NAMES.get(field_id, field_id)}: "
                                f"'{te.label}' → VALIDATED ({n_test_runs} tests)")
                    elif n_tests >= 2:
                        te.status = "EMERGING"
                    else:
                        te.status = "NEEDS_VALIDATION"
            
            self._processed_nums.add(num)
            if person:
                self._processed_names.add(person)
            self._all_tests.append(iv)
            new_count += 1
        
        if new_count > 0:
            # Update state metadata
            self.state.previous_version = self.state.version
            self.state.version += 1
            self.state.generated_at = datetime.now().isoformat()
            self.state.test_count = len(self._processed_nums)
            self.state.changes_since_last = changes
            
            # Recalculate field-level confidence and summaries
            self._recalculate_fields()
            
            print(f"[BMC] +{new_count} tests → v{self.state.version} "
                  f"({self.state.test_count} total)")

    def _knn_boost(self, new_test_run: dict,
                   existing_nums: List[int]) -> float:
        """
        KNN neighborhood boost.
        If the new test is in the same Q-Cube area as existing
        evidence tests, the score gets boosted (neighbors agree).
        
        boost = 1.0 + 0.2 * (number_of_same_cube_neighbors)
        Max boost = 2.0 (capped)
        """
        if not self._all_test_runs:
            return 1.0
        
        new_layer = new_test_entry.get("q_layer", "")
        new_obj = new_test_entry.get("q_object", "")
        
        neighbor_count = 0
        for iv in self._all_test_runs:
            if iv.get("test_num") in existing_nums:
                if (iv.get("q_layer") == new_layer or
                        iv.get("q_object") == new_obj):
                    neighbor_count += 1
        
        boost = 1.0 + 0.2 * neighbor_count
        return min(boost, 2.0)

    def _recalculate_fields(self):
        """Recalculate confidence and summary for each field."""
        for field_id, fstate in self.state.fields.items():
            # Count unique tests touching this field
            all_nums = set()
            validated_themes = 0
            total_themes = 0
            
            for te in fstate.themes.values():
                if te.test_nums:
                    total_themes += 1
                    all_nums.update(te.test_nums)
                    if te.status == "VALIDATED":
                        validated_themes += 1
            
            fstate.total_tests = len(all_nums)
            
            # Confidence = validated_themes / total_possible_themes
            extractor = BMC_EXTRACTORS.get(field_id, {})
            possible = len(extractor.get("themes", {}))
            if possible > 0:
                fstate.confidence = validated_themes / possible
            else:
                fstate.confidence = 0.0
            
            # Generate summary text
            fstate.summary = self._generate_field_summary(field_id, fstate)

    def _generate_field_summary(self, field_id: str,
                                fstate: BMCFieldState) -> str:
        """Generate human-readable summary for one BMC field."""
        lines = []
        
        # Sort themes: VALIDATED first, then EMERGING, then NEEDS_VALIDATION
        status_order = {"VALIDATED": 0, "EMERGING": 1, "NEEDS_VALIDATION": 2}
        sorted_themes = sorted(
            fstate.themes.values(),
            key=lambda t: (status_order.get(t.status, 9), -t.score)
        )
        
        for te in sorted_themes:
            if not te.test_nums:
                continue
            
            nums_str = ", ".join(f"#{n}" for n in sorted(te.test_nums))
            cost_str = ""
            if te.equipment_costs:
                total = sum(te.equipment_costs)
                cost_str = f" [${total:,.0f} documented]"
            
            status_icon = {
                "VALIDATED": "✓",
                "EMERGING": "◐",
                "NEEDS_VALIDATION": "○",
            }.get(te.status, "?")
            
            lines.append(
                f"{status_icon} {te.label}\n"
                f"  Evidence: {nums_str}{cost_str}\n"
                f"  Status: {te.status} ({len(te.test_nums)} tests)"
            )
        
        return "\n\n".join(lines) if lines else "(No evidence yet)"

    def get_state(self) -> BMCState:
        """Get current BMC state."""
        return self.state

    def save(self, bmc_dir: Path):
        """Save BMC state to project's BMC_generation folder."""
        bmc_dir.mkdir(parents=True, exist_ok=True)
        
        state_file = bmc_dir / "bmc_state.json"
        with open(state_file, 'w', encoding='utf-8') as f:
            json.dump(self.state.to_dict(), f, indent=2, ensure_ascii=False)
        
        # Also save versioned snapshot
        snapshot = bmc_dir / f"bmc_state_v{self.state.version}.json"
        with open(snapshot, 'w', encoding='utf-8') as f:
            json.dump(self.state.to_dict(), f, indent=2, ensure_ascii=False)
        
        print(f"[BMC] Saved v{self.state.version} → {state_file}")


# ══════════════════════════════════════════════════════════════
# BMC DIFF ENGINE — what changed between versions
# ══════════════════════════════════════════════════════════════

class BMCDiffEngine:
    """Compare two BMC states and report what changed."""

    @staticmethod
    def diff(old_state: BMCState, new_state: BMCState) -> List[str]:
        """
        Compare old and new BMC states.
        Returns list of human-readable change descriptions.
        """
        changes = []
        
        if old_state.test_count != new_state.test_count:
            delta = new_state.test_count - old_state.test_count
            changes.append(f"+{delta} tests added "
                          f"({old_state.test_count} → {new_state.test_count})")
        
        for field_id in BMC_FIELDS:
            old_f = old_state.fields.get(field_id)
            new_f = new_state.fields.get(field_id)
            if not old_f or not new_f:
                continue
            
            fname = BMC_DISPLAY_NAMES.get(field_id, field_id)
            
            # Confidence change
            old_conf = old_f.confidence
            new_conf = new_f.confidence
            if abs(new_conf - old_conf) > 0.01:
                direction = "↑" if new_conf > old_conf else "↓"
                changes.append(
                    f"{fname}: confidence {direction} "
                    f"({old_conf:.0%} → {new_conf:.0%})")
            
            # Theme status changes
            for tid in set(list(old_f.themes.keys()) + list(new_f.themes.keys())):
                old_t = old_f.themes.get(tid)
                new_t = new_f.themes.get(tid)
                
                if new_t and not old_t:
                    changes.append(f"{fname}: NEW theme '{new_t.label}'")
                elif old_t and new_t and old_t.status != new_t.status:
                    changes.append(
                        f"{fname}: '{new_t.label}' "
                        f"{old_t.status} → {new_t.status}")
        
        return changes


# ══════════════════════════════════════════════════════════════
# DOCX GENERATOR — proper BMC document output
# ══════════════════════════════════════════════════════════════

def generate_bmc_docx(state: BMCState, output_path: Path) -> Path:
    """
    Generate a proper Business Model Canvas DOCX from BMC state.
    All 10 fields. Every claim cited. Accumulative version tracking.
    
    Args:
        state: BMCState from aggregator
        output_path: Directory to save into
    
    Returns:
        Path to generated docx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[BMC] python-docx not installed. pip install python-docx")
        return None
    
    doc = Document()
    
    # ── Title ──
    title = doc.add_heading(
        f'Business Model Canvas — {state.project_id}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Version {state.version} | "
        f"{state.test_count} Test Runs | "
        f"Generated {state.generated_at[:10] if state.generated_at else 'N/A'}\n"
    ).font.size = Pt(10)
    meta.add_run(
        "RULE: Every claim traces to ≥1 test_run. "
        "< 3 tests = NEEDS VALIDATION\n"
    ).font.size = Pt(8)
    meta.add_run(
        "© 2025-2026 Stephen J. Burdick Sr. — All Rights Reserved."
    ).font.size = Pt(8)
    
    # ── Changes since last version ──
    if state.changes_since_last:
        doc.add_heading("Changes Since Last Version", level=2)
        for change in state.changes_since_last:
            doc.add_paragraph(f"• {change}", style='List Bullet')
    
    # ── Each BMC field ──
    used_quotes = set()  # Track quotes across all sections to prevent duplicates
    
    for field_id in BMC_FIELDS:
        fstate = state.fields.get(field_id)
        if not fstate:
            continue
        
        # Field heading with confidence
        conf_pct = f"{fstate.confidence:.0%}"
        heading = doc.add_heading(
            f"{fstate.display_name} [{conf_pct} confident]", level=1)
        
        # Test Run count
        doc.add_paragraph(
            f"Evidence from {fstate.total_test_runs} test_runs",
            style='Intense Quote'
        )
        
        # Themes
        status_order = {"VALIDATED": 0, "EMERGING": 1, "NEEDS_VALIDATION": 2}
        sorted_themes = sorted(
            fstate.themes.values(),
            key=lambda t: (status_order.get(t.status, 9), -t.score)
        )
        
        for te in sorted_themes:
            if not te.test_nums:
                continue
            
            # Theme label with status
            status_icon = {"VALIDATED": "✓", "EMERGING": "◐",
                           "NEEDS_VALIDATION": "○"}.get(te.status, "?")
            
            p = doc.add_paragraph()
            run = p.add_run(f"{status_icon} {te.label}")
            run.bold = True
            if te.status == "VALIDATED":
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif te.status == "EMERGING":
                run.font.color.rgb = RGBColor(204, 153, 0)  # Amber
            else:
                run.font.color.rgb = RGBColor(180, 0, 0)  # Red
            
            # Test Run citations
            nums_str = ", ".join(f"#{n}" for n in sorted(te.test_nums))
            doc.add_paragraph(f"  Tests: {nums_str}")
            
            # Equipment costs
            if te.equipment_costs:
                total = sum(te.equipment_costs)
                doc.add_paragraph(f"  Documented costs: ${total:,.0f}")
            
            # Evidence quotes (max 3, no duplicates across document)
            quote_count = 0
            for quote in te.quotes:
                if quote_count >= 3:
                    break
                # Deduplicate: extract person tag + first 40 chars of content
                # Quote format: "[#num person] ...snippet..."
                m = re.match(r'(\[#\d+\s+[^\]]+\])', quote)
                person_tag = m.group(1) if m else ''
                content_start = quote[len(person_tag):].strip(' .') if person_tag else quote
                quote_key = f"{person_tag}|{content_start[:40]}"
                if quote_key in used_quotes:
                    continue
                used_quotes.add(quote_key)
                q_para = doc.add_paragraph(f"  \"{quote}\"")
                quote_count += 1
                q_para.style = 'Quote'
        
        # Gap warning if no themes have evidence
        active_themes = [t for t in fstate.themes.values() if t.test_nums]
        if not active_themes:
            gap = doc.add_paragraph("⚠ NO EVIDENCE — needs test_run coverage")
            gap.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    
    # ── Save ──
    output_path.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"BMC_Generated_{ts}.docx"
    filepath = output_path / filename
    doc.save(str(filepath))
    
    print(f"[BMC] DOCX → {filepath}")
    return filepath


# ══════════════════════════════════════════════════════════════
# CONVENIENCE: One-shot BMC generation for a project
# ══════════════════════════════════════════════════════════════

def generate_project_bmc(project_id: str,
                         test_runs: List[dict],
                         bmc_dir: Path,
                         generate_docx: bool = True) -> BMCState:
    """
    One-call BMC generation.
    
    1. Loads previous state (accumulative)
    2. Adds new tests
    3. Saves state
    4. Optionally generates DOCX
    
    Args:
        project_id: Active project ID (from collector's project selector)
        test_runs: list of test_run dicts
        bmc_dir: project's BMC_generation/ folder
        generate_docx: whether to also create DOCX
    
    Returns:
        Updated BMCState
    """
    agg = BMCAggregator(project_id)
    agg.load_previous(bmc_dir)
    agg.add_test_runs(tests)
    agg.save(bmc_dir)
    
    if generate_docx:
        generate_bmc_docx(agg.get_state(), bmc_dir)
    
    return agg.get_state()


# ══════════════════════════════════════════════════════════════
# STANDALONE TEST
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys
    
    print("=" * 60)
    print("  BMC ENGINE — Standalone Test")
    print("  The Person is the Highest Responsibility")
    print("=" * 60)
    
    # Import project_paths for discovery — zero hardcoding
    try:
        from project_paths import (
            list_deployed_projects, get_project_dir,
            get_bmc_dir, get_baseline_dir
        )
    except ImportError:
        print("  ERROR: project_paths.py not found in same directory.")
        print("  Place bmc_engine.py alongside project_paths.py.")
        sys.exit(1)
    
    # Discover ALL deployed projects
    projects = list_deployed_projects()
    if not projects:
        print("  No deployed projects found.")
        sys.exit(0)
    
    print(f"  Found {len(projects)} project(s): {projects}")
    
    for project_id in projects:
        print(f"\n  ── {project_id} ──")
        
        # Find Baseline deck (or any deck) for this project
        baseline_dir = get_baseline_dir(project_id)
        tests = []
        
        if baseline_dir.exists():
            for deck_file in baseline_dir.glob("*deck*.json"):
                try:
                    with open(deck_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    if isinstance(data, dict) and 'tests' in data:
                        tests = data['tests']
                    elif isinstance(data, list):
                        tests = data
                    print(f"  Loaded {len(tests)} tests from {deck_file.name}")
                except Exception as e:
                    print(f"  Error reading {deck_file.name}: {e}")
                break
        
        if not test_runs:
            print(f"  No test_run deck found for {project_id}, skipping BMC.")
            continue
        
        # Generate BMC into the project's own BMC_generation folder
        bmc_dir = get_bmc_dir(project_id)
        state = generate_project_bmc(project_id, tests, bmc_dir,
                                     generate_docx=True)
        
        print(f"  BMC v{state.version} — {state.test_count} test_runs")
        for field_id in BMC_FIELDS:
            fstate = state.fields.get(field_id)
            if fstate:
                active = sum(1 for t in fstate.themes.values() if t.test_nums)
                validated = sum(1 for t in fstate.themes.values()
                              if t.status == "VALIDATED")
                print(f"    {fstate.display_name:<25} "
                      f"conf={fstate.confidence:.0%}  "
                      f"themes={active}  validated={validated}  "
                      f"test_runs={fstate.total_test_runs}")
    
    print(f"\n  Done.")
