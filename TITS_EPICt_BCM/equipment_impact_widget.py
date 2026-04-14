# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
Substrate Impact Assessment - Auto-Accumulating Widget
=======================================================
Drop-in replacement for SubstrateImpactWidget in validation_test_collector.py

SPARKS vs FUSION:
  Baseline = discovery phase baselines from initial deck JSON
  Validation = validation phase costs from completed test packets
  Both show in the table, tracked separately, sourced to test_run #

PROJECT-AWARE:
  Reads config from config/{project_keyword}/ via keyword matching
  Loads Baseline deck from project's Baseline_deck_tests_final/ folder
  Scans completed tests filtered by active project

COLUMNS:
  ✓ | Equipment | Config Range | Baseline | Validation | Events/yr | Annual Impact $ | Sources

Emerald Entities LLC — GIBUSH Systems
"""

import json
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

try:
    from PySide6.QtWidgets import (
        QWidget, QVBoxLayout, QHBoxLayout, QLabel, QScrollArea,
        QCheckBox, QLineEdit, QPushButton, QGroupBox, QMessageBox,
        QTableWidget, QTableWidgetItem, QHeaderView, QFrame, QSplitter,
        QFileDialog, QTableView, QAbstractItemView, QSpinBox, QMenu, QApplication
    )
    from PySide6.QtCore import Qt, Signal, QPoint
    from PySide6.QtGui import QColor, QFont, QBrush, QKeySequence, QShortcut
    PYSIDE_AVAILABLE = True
except ImportError:
    PYSIDE_AVAILABLE = False


# ============================================================================
# CONFIG LOADER - Project-aware, keyword-matched
# ============================================================================

def find_substrate_config(project_type: str) -> Optional[Path]:
    """
    Find equipment config JSON using keyword matching.
    AISOS_SPINE → keywords 'aisos','spine' → finds config/bcm_navigation/
    BCM_SUBSTRATE → keywords 'chip','blow','line' → finds config/bcm_substrate/
    """
    subfolder = project_type.lower()
    keywords = [w for w in subfolder.split("_") if len(w) > 2]
    
    script_dir = Path(__file__).parent if '__file__' in dir() else Path.cwd()
    
    anchors = [script_dir, Path.cwd()]
    if sys.platform == "win32":
        anchors.append(Path("C:/TITS/TITS_GIBUSH_AISOS_SPINE/TITS_GIBUSH_AISOS_SPINE_ICORPS"))
    
    for anchor in anchors:
        config_dir = anchor / "config"
        if not config_dir.exists():
            continue
        
        # Exact match first
        exact = config_dir / subfolder
        if exact.exists():
            for f in exact.glob("substrate_config*.json"):
                return f
        
        # Keyword match against subfolder names
        for sub in config_dir.iterdir():
            if sub.is_dir() and any(kw in sub.name.lower() for kw in keywords):
                for f in sub.glob("substrate_config*.json"):
                    return f
        
        # Keyword match in filename
        for f in config_dir.glob("substrate_config*.json"):
            if any(kw in f.stem.lower() for kw in keywords):
                return f
    
    return None


def load_equipment_from_config(project_type: str = "BCM_SUBSTRATE") -> Tuple[Dict, Dict]:
    """
    Load equipment categories from config JSON.
    Includes baseline_metric and baseline_sources if present.
    
    Returns (components_by_category, pain_indicators)
    """
    config_file = find_substrate_config(project_type)
    
    components_by_category = {}
    pain_indicators = {}
    
    if not config_file or not config_file.exists():
        print(f"  ✗ No equipment config found for {project_type}")
        return components_by_category, pain_indicators
    
    print(f"  ✓ Equipment config: {config_file}")
    
    with open(config_file, 'r', encoding='utf-8') as f:
        config = json.load(f)
    
    if 'module_categories' in config:
        categories = config.get('module_categories', {})
        for cat_id, cat_data in categories.items():
            cat_name = cat_data.get('name', cat_id)
            components_by_category[cat_name] = []
            for mod_id, mod_data in cat_data.get('modules', {}).items():
                components_by_category[cat_name].append({
                    'id': mod_id,
                    'name': mod_data.get('display_name', mod_id),
                    'cost': mod_data.get('typical_cost', ''),
                    'keywords': mod_data.get('keywords', []),
                    'baseline_metric': mod_data.get('baseline_metric', 0),
                    'baseline_sources': mod_data.get('baseline_sources', []),
                })
    else:
        categories = config.get('equipment_categories', {})
        for cat_id, cat_data in categories.items():
            cat_name = cat_data.get('name', cat_id)
            components_by_category[cat_name] = []
            for equip_id, equip_data in cat_data.get('equipment', {}).items():
                components_by_category[cat_name].append({
                    'id': equip_id,
                    'name': equip_data.get('display_name', equip_id),
                    'cost': equip_data.get('typical_cost_range', ''),
                    'keywords': equip_data.get('keywords', []),
                    'notes': equip_data.get('notes', ''),
                    'baseline_metric': equip_data.get('baseline_metric', 0),
                    'baseline_sources': equip_data.get('baseline_sources', []),
                })
    
    pain_indicators = config.get('pain_indicators', {})
    return components_by_category, pain_indicators


# ============================================================================
# SPARKS DECK LOADER - Reads from project's Baseline folder
# ============================================================================

def load_baseline_equipment(project_root: Path, project_id: str) -> Dict[str, Dict]:
    """
    Load component metric data from Baseline deck JSON.
    Searches: project_root/{project_id}/Baseline_deck_tests_final/*_deck.json
    
    Returns: { "component_name": { 
        "cost": float, "sources": ["#5 Rocky Smith: $70K..."], "test_nums": [5]
    }}
    """
    baseline_dir = project_root / project_id / "Baseline_deck_tests_final"
    baseline_equipment = {}
    
    if not baseline_dir.exists():
        return baseline_equipment
    
    for deck_file in baseline_dir.glob("*_deck.json"):
        try:
            with open(deck_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            tests = data.get('tests', [])
            for iv in test_runs:
                script_key = iv.get("script_name", "Unknown")
                person = iv.get('script_name', '')
                
                # Source 1: structured substrate_impacts
                for impact in iv.get('substrate_impacts', []):
                    equip_name = impact.get('equipment', '').strip()
                    cost = impact.get('cost', 0)
                    if equip_name and cost:
                        if isinstance(cost, str):
                            cost = _parse_cost(cost)
                        if equip_name not in baseline_equipment:
                            baseline_equipment[equip_name] = {'costs': [], 'sources': [], 'persons': []}
                        baseline_equipment[equip_name]['costs'].append(cost)
                        baseline_equipment[equip_name]['sources'].append(f"#{num} {person}: ${cost:,.0f}")
                        baseline_equipment[equip_name]['persons'].append(script_key)
                
                # Source 2: dollar amounts in text near equipment keywords
                text = f"{iv.get('results', '')} {iv.get('action_iterate', '')}"
                if not text.strip():
                    continue
                
                equip_patterns = {
                    'CTS Rollers': ['cts', 'chip thickness screen', 'roller'],
                    'Screen Baskets': ['screen basket', 'basket'],
                    'Blow Line': ['blow line', 'chip line repair'],
                    'Brown Stock Washers': ['brown stock washer', 'washer'],
                    'Recovery Boiler': ['recovery boiler', 'pre-tube', 'tube pipe', 'grate'],
                    'Digester': ['digester', 'chemical savings'],
                    'Debarker': ['debarker'],
                    'Chipper Knives': ['chipper kniv', 'knife', 'knives'],
                }
                
                text_lower = text.lower()
                for equip_name, kws in equip_patterns.items():
                    for kw in kws:
                        kw_pos = text_lower.find(kw)
                        if kw_pos >= 0:
                            # Extract $ near keyword
                            window = text[max(0, kw_pos-80):kw_pos+len(kw)+120]
                            costs = re.findall(r'\$[\d,]+(?:\.\d+)?(?:\s*[kKmM])?', window)
                            for c in costs:
                                val = _parse_cost(c)
                                if val > 0:
                                    if equip_name not in baseline_equipment:
                                        baseline_equipment[equip_name] = {'costs': [], 'sources': [], 'persons': []}
                                    if script_key not in baseline_equipment[equip_name]['persons']:
                                        baseline_equipment[equip_name]['costs'].append(val)
                                        baseline_equipment[equip_name]['sources'].append(f"#{num} {person}: {c}")
                                        baseline_equipment[equip_name]['persons'].append(script_key)
                            break
        except Exception as e:
            print(f"  ⚠ Baseline deck error ({deck_file.name}): {e}")
    
    return baseline_equipment


# ============================================================================
# FUZZY MATCHER
# ============================================================================

def fuzzy_match_equipment(test_run_name: str, config_name: str, config_keywords: list) -> bool:
    """Check if a test equipment name matches a config equipment entry."""
    i_lower = test_run_name.lower().strip()
    c_lower = config_name.lower().strip()
    
    if i_lower in c_lower or c_lower in i_lower:
        return True
    
    if len(i_lower) >= 10 and len(c_lower) >= 10:
        if i_lower[:20] == c_lower[:20]:
            return True
    
    for kw in config_keywords:
        if kw.lower() in i_lower:
            return True
    
    i_words = set(re.findall(r'\w+', i_lower))
    c_words = set(re.findall(r'\w+', c_lower))
    if i_words and c_words:
        overlap = len(i_words & c_words)
        if overlap >= 2 or (overlap >= 1 and min(len(i_words), len(c_words)) <= 2):
            return True
    
    return False


# ============================================================================
# COST PARSER
# ============================================================================

def _parse_cost(cost_str) -> float:
    """Parse '$300K', '$45,000', '$1.2M', '$15,000 - $80,000' → number."""
    if not cost_str:
        return 0
    if isinstance(cost_str, (int, float)):
        return float(cost_str)
    
    clean = str(cost_str).strip().replace('$', '').replace(',', '')
    
    mult_match = re.match(r'^([\d.]+)\s*([kKmM])', clean)
    if mult_match:
        val = float(mult_match.group(1))
        mult = mult_match.group(2).upper()
        return val * 1000 if mult == 'K' else val * 1000000
    
    nums = re.findall(r'[\d.]+', clean)
    if len(nums) >= 2:
        try:
            return (float(nums[0]) + float(nums[1])) / 2
        except:
            pass
    elif len(nums) == 1:
        try:
            return float(nums[0])
        except:
            pass
    return 0


# ============================================================================
# TEST SCANNER - Separates Baseline from Validation
# ============================================================================

def scan_tests_for_equipment(database, source_filter: str = None) -> Dict[str, Dict]:
    """
    Sca tests for equipment mentions.
    source_filter: 'sparks', 'validation', or None (all)
    """
    equipment_data = {}
    
    if not database or not hasattr(database, 'tests'):
        return equipment_data
    
    for test_entry in database.test_runs:
        # Filter by source if requested
        source = getattr(test_run, 'source', '')
        if source_filter and source != source_filter:
            continue
        
        int_person = getattr(test_run, 'script_name', '')
        
        # Source 1: structured substrate_impacts
        if hasattr(test_run, 'substrate_impacts') and test_run.substrate_impacts:
            for impact in test_run.substrate_impacts:
                equip_name = impact.get('equipment', '').strip()
                if not equip_name or equip_name == 'Equipment/Module':
                    continue
                
                if equip_name not in equipment_data:
                    equipment_data[equip_name] = {
                        'tests': [], 'total_mentions': 0,
                        'costs_reported': [], 'notes': [],
                        'discussed_count': 0, 'costed_count': 0,
                    }
                
                already = any(i.get('script_name', '').strip().lower() == int_person.strip().lower()
                              for i in equipment_data[equip_name]['tests'])
                if not already:
                    cost = impact.get('cost', '')
                    notes = impact.get('notes', '')
                    
                    equipment_data[equip_name]['total_mentions'] += 1
                    equipment_data[equip_name]['discussed_count'] += 1
                    equipment_data[equip_name]['tests'].append({
                        'num': int_num, 'script_name': int_person,
                        'cost': cost, 'notes': notes, 'source': source or 'unknown',
                    })
                    
                    if cost and str(cost) != '0':
                        equipment_data[equip_name]['costs_reported'].append(str(cost))
                        equipment_data[equip_name]['costed_count'] += 1
                    if notes:
                        equipment_data[equip_name]['notes'].append(notes)
        
        # Source 2: text scanning with cost extraction
        all_text = ''
        for field in ('results', 'experiments', 'hypotheses', 'action_iterate'):
            val = getattr(test_run, field, None)
            if val:
                all_text += ' ' + val
        
        if all_text.strip():
            text_lower = all_text.lower()
            equip_keywords = {
                'CTS': ['cts', 'chip thickness screen'],
                'Blow Line': ['blow line', 'blowline'],
                'Screen Baskets': ['screen basket'],
                'Chipper Knives': ['chipper kniv', 'knife', 'knives'],
                'Debarker': ['debarker', 'debarking'],
                'Digester': ['digester'],
                'Recovery Boiler': ['recovery boiler'],
                'Log Truck': ['log truck', 'hauling truck'],
                'Chip Van': ['chip van', 'walking floor'],
                'Brown Stock Washers': ['brown stock washer'],
                'Tub Pulper': ['tub pulper', 'pulper'],
            }
            
            for equip_name, kws in equip_keywords.items():
                for kw in kws:
                    kw_pos = text_lower.find(kw)
                    if kw_pos >= 0:
                        if equip_name not in equipment_data:
                            equipment_data[equip_name] = {
                                'tests': [], 'total_mentions': 0,
                                'costs_reported': [], 'notes': [],
                                'discussed_count': 0, 'costed_count': 0,
                            }
                        
                        already = any(i['num'] == int_num for i in equipment_data[equip_name]['tests'])
                        if not already:
                            # Extract dollar amount near keyword
                            window = all_text[max(0, kw_pos-100):kw_pos+len(kw)+150]
                            cost_matches = re.findall(
                                r'\$[\d,]+(?:\.\d+)?(?:\s*[kKmM])?', window
                            )
                            cost_found = cost_matches[0].strip() if cost_matches else ''
                            
                            equipment_data[equip_name]['total_mentions'] += 1
                            equipment_data[equip_name]['discussed_count'] += 1
                            equipment_data[equip_name]['tests'].append({
                                'num': int_num, 'script_name': int_person,
                                'cost': cost_found, 'notes': f'(text: "{kw}")',
                                'source': source or 'unknown',
                            })
                            
                            if cost_found:
                                equipment_data[equip_name]['costs_reported'].append(cost_found)
                                equipment_data[equip_name]['costed_count'] += 1
                        break
    
    return equipment_data


# ============================================================================
# COMPLETED TEST FOLDER SCANNER
# ============================================================================

def scan_completed_folder(tested_dir: Path, config_names: list) -> Dict[str, Dict]:
    """
    Scan a completed tests folder for .docx files.
    Parses each using the same table structure as doc_reader:
      - Tables 1-6: equipment checklist rows (cell[0]=yes, cell[1]=name, cell[3]=notes)
      - Table 8:    cost capture table (Equipment, Annual Damage $, Failure Freq)
      - Table 0:    test metadata (Test Run #, Person, Company)

    Returns:
      {
        "Equipment Name": {
            "costs":        [float, ...],
            "sources":      ["#7 Ferrel Fowler: $15,000", ...],
            "test_nums": [7, ...],
            "notes":        ["note text", ...],
        }
      }
    Keys use exact config display_names so the widget slots them directly.
    """
    try:
        from docx import Document
    except ImportError:
        print("  ⚠ python-docx not installed — cannot scan completed test_runs")
        return {}

    if not tested_dir or not tested_dir.exists():
        return {}

    # Build lowercase lookup for fuzzy→exact config name resolution
    config_lower = {n.lower(): n for n in config_names}

    def resolve_name(raw: str) -> Optional[str]:
        """Map a raw equipment name from the docx to a config display_name."""
        raw = raw.strip()
        if not raw or raw in ('Equipment/Module', 'Equipment'):
            return None
        # Exact match
        if raw in config_names:
            return raw
        # Case-insensitive
        if raw.lower() in config_lower:
            return config_lower[raw.lower()]
        # Substring / word overlap
        r_words = set(re.findall(r'\w+', raw.lower()))
        for cn in config_names:
            c_words = set(re.findall(r'\w+', cn.lower()))
            overlap = len(r_words & c_words)
            if overlap >= 2 or (overlap >= 1 and min(len(r_words), len(c_words)) <= 2):
                return cn
        return None

    folder_data: Dict[str, Dict] = {}

    docx_files = sorted(tested_dir.glob("*.docx"))
    for docx_path in docx_files:
        if docx_path.name.startswith('~'):
            continue  # skip temp files
        try:
            doc = Document(str(docx_path))
        except Exception as e:
            print(f"  ⚠ Could not open {docx_path.name}: {e}")
            continue

        tables = doc.tables

        # ── Test Run metadata ──────────────────────────────────────────────
        test_num = 0
        person = docx_path.stem
        company = ''
        if tables:
            for row in tables[0].rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0].lower()
                    val = cells[1]
                    if 'test_run' in key and '#' in key or key == 'test_run #':
                        try:
                            test_num = int(re.findall(r'\d+', val)[0])
                        except:
                            pass
                    elif key == 'script_name':
                        person = val
                    elif key == 'source_version':
                        company = val

        # ── Equipment checklists (Tables 1-6) ──────────────────────────────
        discussed = {}  # config_name → notes
        for table in tables[1:7]:
            if len(table.rows) < 2:
                continue
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                status = cells[0].lower()
                if 'yes' in status or status in ('x', 'checked', '✓', '✔'):
                    raw_name = cells[1] if len(cells) > 1 else ''
                    notes = cells[3] if len(cells) > 3 else ''
                    config_name = resolve_name(raw_name)
                    if config_name:
                        discussed[config_name] = notes

        # ── Cost capture table (Table 8 pattern) ───────────────────────────
        cost_rows = []  # [(config_name, cost_float, freq)]
        for table in tables:
            if len(table.rows) < 2:
                continue
            header = [c.text.strip() for c in table.rows[0].cells]
            if 'Equipment' in header and any('Annual' in h or 'Damage' in h for h in header):
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2 and cells[0]:
                        config_name = resolve_name(cells[0])
                        cost_val = _parse_cost(cells[1]) if len(cells) > 1 else 0
                        freq_val = cells[2] if len(cells) > 2 else '1'
                        if config_name:
                            cost_rows.append((config_name, cost_val, freq_val))
                break

        # ── Merge into folder_data ──────────────────────────────────────────
        source_tag = f"{person}"

        # Equipment discussed (may have no cost — still record mention)
        for config_name, notes in discussed.items():
            if config_name not in folder_data:
                folder_data[config_name] = {
                    'costs': [], 'sources': [], 'persons': [], 'notes': []
                }
            if person not in folder_data[config_name]['persons']:
                folder_data[config_name]['persons'].append(person)
                folder_data[config_name]['notes'].append(notes)
                folder_data[config_name]['sources'].append(
                    f"{source_tag} ({company[:20]}): discussed"
                )

        # Cost rows override/add costs
        for config_name, cost_val, freq_val in cost_rows:
            if config_name not in folder_data:
                folder_data[config_name] = {
                    'costs': [], 'sources': [], 'persons': [], 'notes': []
                }
            folder_data[config_name]['costs'].append(cost_val)
            # Update source to show cost
            src = f"{source_tag} ({company[:20]}): ${cost_val:,.0f}"
            # Replace the 'discussed' source if already added
            existing_srcs = folder_data[config_name]['sources']
            replaced = False
            for i, s in enumerate(existing_srcs):
                if source_tag in s and 'discussed' in s:
                    existing_srcs[i] = src
                    replaced = True
                    break
            if not replaced:
                existing_srcs.append(src)
                if person not in folder_data[config_name]['persons']:
                    folder_data[config_name]['persons'].append(person)

    return folder_data


# ============================================================================
# EQUIPMENT IMPACT WIDGET
# ============================================================================

if PYSIDE_AVAILABLE:
    class SubstrateImpactWidget(QWidget):
        """
        Auto-accumulating equipment impact assessment.
        
        Shows SPARKS baselines and FUSION test_run costs side by side.
        Reads config from project-specific JSON.
        Scans Baseline deck + completed tests for cost data.
        """
        
        impact_updated = Signal(float)
        
        def __init__(self, database=None, parent=None):
            super().__init__(parent)
            self.database = database
            self._get_active_project = None
            self._get_project_root = None
            self.component_rows = {}
            self.test_run_equipment = {}
            self.baseline_equipment = {}
            
            self.init_ui()
        
        def set_project_callback(self, fn):
            self._get_active_project = fn
        
        def set_project_getter(self, fn):
            """Alias for collector compatibility."""
            self._get_active_project = fn
        
        def set_project_root_callback(self, fn):
            self._get_project_root = fn
        
        def set_database(self, database):
            """Set/update database reference."""
            self.database = database
        
        def _get_project_type(self):
            if self._get_active_project:
                try:
                    return self._get_active_project()
                except:
                    pass
            return "BCM_SUBSTRATE"
        
        def _get_root(self):
            if self._get_project_root:
                try:
                    return Path(self._get_project_root())
                except:
                    pass
            # Default
            return Path.cwd() / "BCM_Projects" / "project_caas_deployed"
        
        def init_ui(self):
            layout = QVBoxLayout(self)
            
            # Header
            title = QLabel("<h2>Substrate Impact Assessment</h2>")
            title.setStyleSheet("color: #c41e3a;")
            layout.addWidget(title)
            
            # Refresh button
            btn_layout = QHBoxLayout()
            
            self.info_label = QLabel("Loading config...")
            self.info_label.setWordWrap(True)
            self.info_label.setStyleSheet("color: #555; padding: 4px;")
            btn_layout.addWidget(self.info_label)
            
            btn_layout.addStretch()

            # ── Table controls: font size + row height ──
            font_lbl = QLabel("Font:")
            font_lbl.setStyleSheet("color: #888; font-size: 10px;")
            btn_layout.addWidget(font_lbl)
            self._font_spin = QSpinBox()
            self._font_spin.setRange(8, 18)
            self._font_spin.setValue(11)
            self._font_spin.setSuffix("px")
            self._font_spin.setFixedWidth(65)
            self._font_spin.setStyleSheet(
                "QSpinBox { background: #222; color: #ddd; border: 1px solid #555; padding: 2px; }"
            )
            self._font_spin.valueChanged.connect(self._apply_font_size)
            btn_layout.addWidget(self._font_spin)

            row_lbl = QLabel("Row:")
            row_lbl.setStyleSheet("color: #888; font-size: 10px;")
            btn_layout.addWidget(row_lbl)
            self._row_spin = QSpinBox()
            self._row_spin.setRange(20, 80)
            self._row_spin.setValue(36)
            self._row_spin.setSuffix("px")
            self._row_spin.setFixedWidth(65)
            self._row_spin.setStyleSheet(
                "QSpinBox { background: #222; color: #ddd; border: 1px solid #555; padding: 2px; }"
            )
            self._row_spin.valueChanged.connect(self._apply_row_height)
            btn_layout.addWidget(self._row_spin)

            btn_layout.addSpacing(8)
            
            self.refresh_btn = QPushButton("🔄 Refresh from Test Runs")
            self.refresh_btn.setStyleSheet(
                "background-color: #111; color: #00FF88; font-weight: bold; "
                "padding: 6px 16px; border: 1px solid #00FF88;"
            )
            self.refresh_btn.clicked.connect(self.refresh)
            btn_layout.addWidget(self.refresh_btn)
            
            layout.addLayout(btn_layout)
            
            # Pivot table — columns are built dynamically on refresh
            self.table = QTableWidget()
            self.table.setColumnCount(2)
            self.table.setHorizontalHeaderLabels(["Equipment", "Sector"])

            header = self.table.horizontalHeader()
            header.setSectionResizeMode(0, QHeaderView.Stretch)
            header.setSectionResizeMode(1, QHeaderView.Fixed)
            self.table.setColumnWidth(1, 100)
            self.table.horizontalHeader().setDefaultSectionSize(130)
            self.table.verticalHeader().setVisible(False)
            self.table.setColumnWidth(2, 140)
            self.table.setColumnWidth(3, 100)
            self.table.setColumnWidth(4, 100)
            self.table.setColumnWidth(5, 70)
            self.table.setColumnWidth(6, 120)
            header.setSectionResizeMode(7, QHeaderView.Stretch)
            
            self.table.setAlternatingRowColors(False)
            self.table.verticalHeader().setVisible(True)
            self.table.verticalHeader().setSectionResizeMode(QHeaderView.Interactive)
            self.table.verticalHeader().setDefaultSectionSize(36)
            self.table.verticalHeader().setMinimumSectionSize(20)
            self.table.verticalHeader().setStyleSheet(
                "QHeaderView::section { background: #1a1a1a; color: #888; "
                "border: 1px solid #333; font-size: 9px; padding: 1px; }"
            )
            self.table.setWordWrap(True)
            self.table.setTextElideMode(Qt.ElideNone)
            self.table.setStyleSheet(
                "QTableWidget { background-color: #111; color: #ddd; gridline-color: #333; "
                "font-size: 11px; }"
                "QHeaderView::section { background-color: #1a1a1a; color: #FFD700; "
                "font-weight: bold; border: 1px solid #333; padding: 4px; }"
            )
            self.table.horizontalHeader().setMinimumSectionSize(60)
            self.table.horizontalHeader().setDefaultSectionSize(130)
            self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)

            # ── Interactive features ──
            self.table.setSelectionMode(QAbstractItemView.ExtendedSelection)
            self.table.setSelectionBehavior(QAbstractItemView.SelectItems)
            self.table.setContextMenuPolicy(Qt.CustomContextMenu)
            self.table.customContextMenuRequested.connect(self._context_menu)
            self.table.itemSelectionChanged.connect(self._on_selection_changed)

            # Header click to highlight full column
            self.table.horizontalHeader().sectionClicked.connect(self._highlight_column)
            self.table.verticalHeader().sectionClicked.connect(self._highlight_row)

            # Ctrl+C copy shortcut
            copy_shortcut = QShortcut(QKeySequence.Copy, self.table)
            copy_shortcut.activated.connect(self._copy_selection)

            layout.addWidget(self.table)

            # ── Formula bar — shows SUM/AVG/COUNT/MIN/MAX of selected cells ──
            self._formula_bar = QLabel("Select cells to see SUM | AVG | MIN | MAX")
            self._formula_bar.setStyleSheet(
                "background-color: #0a0a1a; color: #00ccff; padding: 4px 8px; "
                "font-family: Consolas, monospace; font-size: 11px; "
                "border: 1px solid #333; border-top: none;"
            )
            self._formula_bar.setFixedHeight(24)
            layout.addWidget(self._formula_bar)

            # ── Frozen column overlay (Equipment + Sector stay visible) ──
            self.frozen_cols = 2  # columns 0 and 1 are frozen
            self.frozen_view = QTableView(self.table)
            self.frozen_view.setModel(self.table.model())
            self.frozen_view.setSelectionModel(self.table.selectionModel())
            self.frozen_view.verticalHeader().hide()
            self.frozen_view.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            self.frozen_view.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            self.frozen_view.setWordWrap(True)
            self.frozen_view.setTextElideMode(Qt.ElideNone)
            self.frozen_view.setStyleSheet(
                "QTableView { background-color: #111; color: #ddd; "
                "gridline-color: #333; border: none; border-right: 2px solid #9333EA; "
                "font-size: 11px; }"
                "QHeaderView::section { background-color: #1a1a1a; color: #FFD700; "
                "font-weight: bold; border: 1px solid #333; padding: 4px; }"
            )
            self.frozen_view.setFocusPolicy(Qt.NoFocus)

            # Sync vertical scrolling both directions
            self.table.verticalScrollBar().valueChanged.connect(
                self.frozen_view.verticalScrollBar().setValue)
            self.frozen_view.verticalScrollBar().valueChanged.connect(
                self.table.verticalScrollBar().setValue)

            # Sync row heights when main table changes
            self.table.verticalHeader().sectionResized.connect(
                lambda idx, old, new: self.frozen_view.setRowHeight(idx, new))

            # Reposition frozen overlay when columns resize or table resizes
            self.table.horizontalHeader().sectionResized.connect(
                lambda: self._update_frozen_geometry())

            # Use viewport event filter to catch table resize
            self.table.viewport().installEventFilter(self)
            
            # Summary bar
            summary_frame = QFrame()
            summary_frame.setStyleSheet(
                "background-color: #111; border: 1px solid #9333EA; padding: 8px;"
            )
            summary_layout = QHBoxLayout()
            
            self.total_label = QLabel("<b>Total Annual Impact: $0</b>")
            self.total_label.setStyleSheet("font-size: 14pt; color: #c41e3a;")
            summary_layout.addWidget(self.total_label)
            
            self.baseline_total_label = QLabel("Baseline Repair: $0")
            self.production_total_label = QLabel("Production Impact: $0")
            summary_layout.addWidget(self.production_total_label)
            self.production_total_label.setStyleSheet("color: #FF8C00; font-weight: bold; padding: 4px 10px; background: #111; border-radius: 4px;")
            self.baseline_total_label.setStyleSheet("color: #FFD700;")
            summary_layout.addWidget(self.baseline_total_label)
            
            self.validation_total_label = QLabel("Validation: $0")
            self.validation_total_label.setStyleSheet("color: #00FF88;")
            summary_layout.addWidget(self.validation_total_label)
            
            self.stats_label = QLabel("")
            self.stats_label.setStyleSheet("color: #888;")
            summary_layout.addWidget(self.stats_label)
            
            summary_layout.addStretch()
            
            export_btn = QPushButton("📊 Export Equipment Analysis")
            export_btn.setStyleSheet(
                "background-color: #111; color: #FFD700; font-weight: bold; "
                "padding: 6px 16px; border: 1px solid #9333EA;"
            )
            export_btn.clicked.connect(self.export_analysis)
            summary_layout.addWidget(export_btn)
            
            summary_frame.setLayout(summary_layout)
            layout.addWidget(summary_frame)
            
            self.setLayout(layout)
            self._build_table()
        
        def _build_table(self):
            """Placeholder — pivot table is built entirely in refresh()."""
            pass

        def _update_frozen_columns(self):
            """
            Sync the frozen overlay after the main table is rebuilt.
            Hides all columns except the first self.frozen_cols in the overlay,
            then sizes and positions the overlay to cover exactly those columns.
            """
            n_cols = self.table.columnCount()
            if n_cols < self.frozen_cols:
                self.frozen_view.hide()
                return

            # Show frozen columns, hide the rest in the overlay
            for col in range(n_cols):
                if col < self.frozen_cols:
                    self.frozen_view.setColumnHidden(col, False)
                    self.frozen_view.setColumnWidth(col, self.table.columnWidth(col))
                else:
                    self.frozen_view.setColumnHidden(col, True)

            # Match row heights
            for row in range(self.table.rowCount()):
                self.frozen_view.setRowHeight(row, self.table.rowHeight(row))

            # Size the overlay to cover exactly the frozen columns
            self._update_frozen_geometry()
            self.frozen_view.show()
            self.frozen_view.raise_()

        def _update_frozen_geometry(self):
            """Position and size the frozen overlay to cover the frozen columns."""
            frozen_width = sum(
                self.table.columnWidth(c) for c in range(self.frozen_cols)
            )
            # Account for the vertical header width (hidden, but still 0)
            vh_width = self.table.verticalHeader().width() if self.table.verticalHeader().isVisible() else 0
            # Account for the frame width
            fw = self.table.frameWidth()

            self.frozen_view.setGeometry(
                fw + vh_width,
                fw,
                frozen_width + 1,
                self.table.viewport().height() + self.table.horizontalHeader().height()
            )

        def eventFilter(self, obj, event):
            """Reposition frozen overlay when the table viewport resizes."""
            if obj == self.table.viewport() and event.type() == event.Type.Resize:
                self._update_frozen_geometry()
            return super().eventFilter(obj, event)

        def _apply_font_size(self, size: int):
            """Apply font size to the main table and frozen view."""
            self.table.setStyleSheet(
                f"QTableWidget {{ background-color: #111; color: #ddd; gridline-color: #333; "
                f"font-size: {size}px; }}"
                f"QHeaderView::section {{ background-color: #1a1a1a; color: #FFD700; "
                f"font-weight: bold; border: 1px solid #333; padding: 4px; font-size: {max(size - 1, 8)}px; }}"
            )
            self.frozen_view.setStyleSheet(
                f"QTableView {{ background-color: #111; color: #ddd; "
                f"gridline-color: #333; border: none; border-right: 2px solid #9333EA; "
                f"font-size: {size}px; }}"
                f"QHeaderView::section {{ background-color: #1a1a1a; color: #FFD700; "
                f"font-weight: bold; border: 1px solid #333; padding: 4px; font-size: {max(size - 1, 8)}px; }}"
            )

        def _apply_row_height(self, height: int):
            """Apply row height to all rows in the main table and frozen view."""
            self.table.verticalHeader().setDefaultSectionSize(height)
            for row in range(self.table.rowCount()):
                self.table.setRowHeight(row, height)
                self.frozen_view.setRowHeight(row, height)

        # ── Interactive features ──────────────────────────────────

        def _on_selection_changed(self):
            """Update formula bar with SUM/AVG/COUNT/MIN/MAX of selected numeric cells."""
            selected = self.table.selectedItems()
            if not selected:
                self._formula_bar.setText("Select cells to see SUM | AVG | MIN | MAX")
                return

            values = []
            for item in selected:
                text = item.text().strip()
                num = self._extract_number(text)
                if num is not None and num != 0:
                    values.append(num)

            count = len(selected)
            if values:
                total = sum(values)
                avg = total / len(values)
                lo = min(values)
                hi = max(values)
                self._formula_bar.setText(
                    f"CELLS: {count}  |  NUMERIC: {len(values)}  |  "
                    f"SUM: ${total:,.0f}  |  AVG: ${avg:,.0f}  |  "
                    f"MIN: ${lo:,.0f}  |  MAX: ${hi:,.0f}"
                )
            else:
                self._formula_bar.setText(f"CELLS: {count}  |  No numeric values in selection")

        def _extract_number(self, text: str) -> float:
            """Pull a number from cell text like 'P:$197,100' or '$300K' or '✓'."""
            if not text:
                return None
            import re
            # Strip prefixes like P: D: 
            for prefix in ('P:', 'D:', 'R:', 'S:'):
                if text.startswith(prefix):
                    text = text[len(prefix):]
            text = text.strip()
            if not text or text in ('✓', '-', '.', ''):
                return None
            return _parse_cost(text) if any(c.isdigit() for c in text) else None

        def _highlight_column(self, col_index: int):
            """Click column header to select entire column."""
            self.table.clearSelection()
            for row in range(self.table.rowCount()):
                item = self.table.item(row, col_index)
                if item:
                    item.setSelected(True)

        def _highlight_row(self, row_index: int):
            """Click row header to select entire row."""
            self.table.clearSelection()
            for col in range(self.table.columnCount()):
                item = self.table.item(row_index, col)
                if item:
                    item.setSelected(True)

        def _copy_selection(self):
            """Ctrl+C — copy selected cells as tab-delimited text."""
            selected = self.table.selectedItems()
            if not selected:
                return

            # Find bounds of selection
            rows = sorted(set(item.row() for item in selected))
            cols = sorted(set(item.column() for item in selected))

            lines = []
            for r in rows:
                row_data = []
                for c in cols:
                    item = self.table.item(r, c)
                    row_data.append(item.text() if item else '')
                lines.append('\t'.join(row_data))

            text = '\n'.join(lines)
            clipboard = QApplication.clipboard()
            clipboard.setText(text)
            self._formula_bar.setText(f"Copied {len(rows)} rows × {len(cols)} cols to clipboard")

        def _context_menu(self, pos: QPoint):
            """Right-click context menu."""
            menu = QMenu(self.table)
            menu.setStyleSheet(
                "QMenu { background: #222; color: #ddd; border: 1px solid #555; }"
                "QMenu::item:selected { background: #9333EA; }"
            )

            copy_action = menu.addAction("📋 Copy Selection (Ctrl+C)")
            copy_action.triggered.connect(self._copy_selection)

            menu.addSeparator()

            sum_action = menu.addAction("Σ Sum Selected")
            sum_action.triggered.connect(self._sum_selected_to_bar)

            col_action = menu.addAction("▐ Select Column")
            col_action.triggered.connect(lambda: self._highlight_column(
                self.table.currentColumn()))

            row_action = menu.addAction("▬ Select Row")
            row_action.triggered.connect(lambda: self._highlight_row(
                self.table.currentRow()))

            menu.exec_(self.table.viewport().mapToGlobal(pos))

        def _sum_selected_to_bar(self):
            """Sum numeric values in selected cells and show in formula bar."""
            self._on_selection_changed()
        
        def _get_tested_dir(self) -> Optional[Path]:
            """Resolve the completed tests folder for the active project."""
            root = self._get_root()
            project_type = self._get_project_type()
            candidate = root / project_type / "tested"
            if candidate.exists():
                return candidate
            for anchor in [Path.cwd(), Path(__file__).parent if "__file__" in dir() else Path.cwd()]:
                candidate = anchor / "BCM_Projects" / "project_caas_deployed" / project_type / "tested"
                if candidate.exists():
                    return candidate
            if sys.platform == "win32":
                candidate = Path("C:/TITS/TITS_GIBUSH_AISOS_SPINE/TITS_GIBUSH_AISOS_SPINE_ICORPS") /                             "BCM_Projects" / "project_caas_deployed" / project_type / "tested"
                if candidate.exists():
                    return candidate
            return None

        def _load_baseline_deck_live(self) -> Dict[str, Dict]:
            """
            Read Baseline deck JSON from disk right now.
            Returns { component_name: { costs, sources, test_nums } }
            No caching — always reads current file.
            """
            project_type = self._get_project_type()
            root = self._get_root()

            # Search path: project_caas_deployed/{project}/Baseline_deck_tests_final/
            search_dirs = [
                root / project_type / "Baseline_deck_tests_final",
                root / project_type,
            ]
            if sys.platform == "win32":
                base = Path("C:/TITS/TITS_GIBUSH_AISOS_SPINE/TITS_GIBUSH_AISOS_SPINE_ICORPS")
                search_dirs.append(
                    base / "BCM_Projects" / "project_caas_deployed" / project_type / "Baseline_deck_tests_final"
                )

            baseline_data = {}
            for search_dir in search_dirs:
                if not search_dir.exists():
                    continue
                for deck_file in search_dir.glob("*deck*.json"):
                    try:
                        with open(deck_file, 'r', encoding='utf-8') as f:
                            deck = json.load(f)
                        for iv in deck.get('tests', []):
                            script_key = iv.get("script_name", "Unknown")
                            person = iv.get('script_name', '')
                            for e in iv.get('substrate_impacts', []):
                                name = e.get('equipment', '').strip()
                                cost = float(e.get('cost', 0))
                                cost_type = e.get('cost_type', 'repair')
                                if name and cost > 0 and cost_type == 'repair':
                                    if name not in baseline_data:
                                        baseline_data[name] = {'costs': [], 'sources': [], 'persons': []}
                                    baseline_data[name]['costs'].append(cost)
                                    baseline_data[name]['persons'].append(script_key)
                                    baseline_data[name]['sources'].append(
                                        f"{person}: ${cost:,.0f}(S)"
                                    )
                        print(f"  ✓ Baseline deck loaded: {deck_file.name}")
                        return baseline_data  # first found wins
                    except Exception as ex:
                        print(f"  ⚠ Baseline deck read error: {ex}")
            return baseline_data

        def _load_config_live(self) -> Dict:
            """
            Re-read equipment config from disk right now.
            Returns flat dict: { display_name: eq_data_dict }
            No caching — always reads current file.
            """
            project_type = self._get_project_type()
            config_file = find_substrate_config(project_type)
            live_config = {}
            if not config_file or not config_file.exists():
                return live_config
            try:
                with open(config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                cats = config.get('equipment_categories', config.get('module_categories', {}))
                for cat_id, cat_data in cats.items():
                    for eq_id, eq_data in cat_data.get('equipment', cat_data.get('modules', {})).items():
                        display = eq_data.get('display_name', eq_id)
                        live_config[display] = eq_data
            except Exception as ex:
                print(f"  ⚠ Config reload error: {ex}")
            return live_config

        def refresh(self):
            """
            Pivot table refresh — reads config, Baseline deck, and tested/ from disk.
            Rows = equipment items grouped by category.
            Columns = fixed (Equipment, Sector) + one column per test (Baseline then Validation).
            Add a docx to tested/ → new column on next refresh.
            Delete one → gone.
            """
            print("  ↻ Substrate Impact: pivot rebuild...")

            # ── 1. Load live config ──────────────────────────────────────────
            live_config = self._load_config_live()
            project_type = self._get_project_type()
            config_file = find_substrate_config(project_type)
            equipment_by_cat = {}
            if config_file and config_file.exists():
                with open(config_file, 'r', encoding='utf-8') as f:
                    raw_cfg = json.load(f)
                cats = raw_cfg.get('equipment_categories', {})
                for cat_id, cat_data in cats.items():
                    cat_name = cat_data.get('name', cat_id)
                    items = []
                    for eq_id, eq_data in cat_data.get('equipment', {}).items():
                        items.append({
                            'id': eq_id,
                            'name': eq_data.get('display_name', eq_id),
                            'sector': eq_data.get('sector', ''),
                            'cost_type': eq_data.get('cost_type', 'repair'),
                            'baseline_metric': eq_data.get('baseline_metric', 0),
                            'baseline_production_impact': eq_data.get('baseline_production_impact', 0),
                            'baseline_sources': eq_data.get('baseline_sources', []),
                            'keywords': eq_data.get('keywords', []),
                        })
                    if items:
                        equipment_by_cat[cat_name] = items

            if not equipment_by_cat:
                self.info_label.setText("⚠️ No config found. Check config/ folder.")
                self.info_label.setStyleSheet("color: #FF4444; padding: 6px;")
                return

            # ── 2. Load Baseline deck tests ──────────────────────────────
            # Returns { iv_num: { 'script_name', 'source_version', 'source', impacts: {equip_name: {cost, cost_type, notes}} } }
            baseline_tests = {}
            baseline_deck_data = {}  # equip_name → list of (num, cost)
            for search_dir in [
                self._get_root() / project_type / "Baseline_deck_tests_final",
                self._get_root() / project_type,
            ] + ([Path("C:/TITS/TITS_GIBUSH_AISOS_SPINE/TITS_GIBUSH_AISOS_SPINE_ICORPS") /
                  "BCM_Projects" / "project_caas_deployed" / project_type / "Baseline_deck_tests_final"]
                 if sys.platform == "win32" else []):
                if not Path(search_dir).exists():
                    continue
                for deck_file in Path(search_dir).glob("*deck*.json"):
                    try:
                        with open(deck_file, 'r', encoding='utf-8') as f:
                            deck = json.load(f)
                        for iv in deck.get('tests', []):
                            script_key = iv.get("script_name", "Unknown")
                            baseline_test_runs[script_key] = {
                                'script_name': script_key,
                                'source_version': iv.get('source_version', ''),
                                'source': 'sparks',
                                'impacts': {},
                            }
                            for e in iv.get('substrate_impacts', []):
                                name = e.get('equipment', '').strip()
                                cost = float(e.get('cost', 0))
                                ct = e.get('cost_type', 'repair')
                                notes = e.get('notes', '')
                                if name:
                                    baseline_test_runs[script_key]['impacts'][name] = {
                                        'cost': cost, 'cost_type': ct, 'notes': notes
                                    }
                        break
                    except Exception as ex:
                        print(f"  ⚠ Baseline deck: {ex}")
                break

            # ── 3. Load Validation tests from tested/ folder ───────────
            validation_tests = {}
            tested_dir = self._get_tested_dir()
            if tested_dir and tested_dir.exists():
                try:
                    from docx import Document as DocxDoc
                    SECTOR_CFG = {eq['name']: eq for items in equipment_by_cat.values() for eq in items}
                    config_names_list = list(SECTOR_CFG.keys())

                    for docx_path in sorted(tested_dir.glob("*.docx")):
                        if docx_path.name.startswith('~'):
                            continue
                        try:
                            doc = DocxDoc(str(docx_path))
                            tables = doc.tables
                            num = 0; person = docx_path.stem; company = ''
                            if tables:
                                for row in tables[0].rows:
                                    cells = [c.text.strip() for c in row.cells]
                                    if len(cells) >= 2:
                                        k = cells[0].lower().replace('#','').strip()
                                        if k == 'test_run':
                                            try: num = int(re.findall(r'\d+', cells[1])[0])
                                            except: pass
                                        elif k == 'script_name': person = cells[1]
                                        elif k == 'source_version': company = cells[1]
                            # Fallback: parse number from filename if header blank
                            if num == 0:
                                m = re.search(r'(?:Packet_|Test Run_)(\d+)', docx_path.stem)
                                if m: num = int(m.group(1))
                            # Validation tests must always be above SPARKS_MAX (50)
                            # If filename gave us a number in Baseline range, reassign
                            # person is the key now
                            # test_num assignment no longer needed for keying
                                pass  # person name is the key
                                
                                
                            if not person or person == docx_path.stem:
                                # Try to get person from filename
                                stem = docx_path.stem
                                for prefix in ['Test Run_Packet_','Test Run_Packet','Test Run_']:
                                    stem = stem.replace(prefix,'')
                                parts = stem.split('_')
                                # Strip leading numbers
                                while parts and parts[0].isdigit(): parts.pop(0)
                                person = ' '.join(parts[:3]).replace('_',' ').split('completed')[0].strip()

                            validation_test_runs[person] = {
                                'script_name': person, 'source_version': company,
                                'source': 'validation', 'impacts': {},
                            }
                            # Equipment checklists — scan ALL tables, stop at cost capture
                            for table in tables[1:]:
                                if len(table.rows) < 2: continue
                                hdr = [c.text.strip() for c in table.rows[0].cells]
                                if 'Equipment' in hdr and any('Annual' in h or 'Damage' in h for h in hdr):
                                    break  # hit cost capture table
                                for row in table.rows[1:]:
                                    cells = [c.text.strip() for c in row.cells]
                                    if len(cells) < 2: continue
                                    if 'yes' in cells[0].lower() or cells[0].lower() in ('x','checked','✓','✔'):
                                        raw = cells[1]
                                        ncols = len(cells)
                                        # 8-col: Production $, Equipment $, Safety $, Other $, Notes
                                        if ncols >= 8:
                                            prod_cost = _parse_cost(cells[3])
                                            dmg_cost  = _parse_cost(cells[4])
                                            safe_cost = _parse_cost(cells[5])
                                            other_cost= _parse_cost(cells[6])
                                            notes     = cells[7]
                                            best_cost = prod_cost or dmg_cost or safe_cost or other_cost
                                            cost_type = 'production' if prod_cost > 0 else ('safety' if safe_cost > 0 else 'repair')
                                        else:
                                            best_cost = 0
                                            cost_type = 'repair'
                                            notes = cells[3] if ncols > 3 else ''
                                            prod_cost = dmg_cost = safe_cost = other_cost = 0
                                        matched = raw if raw in SECTOR_CFG else None
                                        if not matched:
                                            for cn in config_names_list:
                                                if fuzzy_match_equipment(raw, cn, SECTOR_CFG.get(cn, {}).get('keywords', [])):
                                                    matched = cn; break
                                        if matched:
                                            validation_test_runs[person]['impacts'][matched] = {
                                                'cost': best_cost, 'cost_type': cost_type,
                                                'production_cost': prod_cost,
                                                'damage_cost': dmg_cost,
                                                'safety_cost': safe_cost,
                                                'other_cost': other_cost,
                                                'notes': notes,
                                            }
                            # Cost capture table — update costs from summary
                            for table in tables:
                                if len(table.rows) < 2: continue
                                header = [c.text.strip() for c in table.rows[0].cells]
                                if 'Equipment' in header and any('Annual' in h or 'Damage' in h for h in header):
                                    for row in table.rows[1:]:
                                        cells = [c.text.strip() for c in row.cells]
                                        if len(cells) >= 2 and cells[0]:
                                            raw = cells[0]
                                            cost_val = _parse_cost(cells[1]) if len(cells) > 1 else 0
                                            ct = cells[3].lower().strip() if len(cells) > 3 else 'repair'
                                            matched = raw if raw in SECTOR_CFG else None
                                            if not matched:
                                                for cn in config_names_list:
                                                    if fuzzy_match_equipment(raw, cn, SECTOR_CFG.get(cn, {}).get('keywords', [])):
                                                        matched = cn; break
                                            if matched:
                                                if matched in validation_test_runs[person]['impacts']:
                                                    validation_test_runs[person]['impacts'][matched]['cost'] = cost_val
                                                    validation_test_runs[person]['impacts'][matched]['cost_type'] = ct
                                                else:
                                                    validation_test_runs[person]['impacts'][matched] = {
                                                        'cost': cost_val, 'cost_type': ct,
                                                        'production_cost': cost_val if ct == 'production' else 0,
                                                        'damage_cost': cost_val if ct == 'repair' else 0,
                                                        'safety_cost': cost_val if ct == 'safety' else 0,
                                                        'other_cost': 0, 'notes': '',
                                                    }
                                    break
                        except Exception as ex:
                            print(f"  ⚠ {docx_path.name}: {ex}")
                except ImportError:
                    print("  ⚠ python-docx not installed")

            # ── 4. Build column order: Baseline tests sorted, then Validation sorted ──
            all_tests = {}
            all_tests.update(baseline_tests)
            all_tests.update(validation_tests)
            baseline_cols = sorted(baseline_tests.keys())
            validation_cols  = sorted(validation_tests.keys())
            test_run_cols = baseline_cols + validation_cols  # ordered

            # ── 5. Location color map ────────────────────────────────────────
            LOCATION_COLORS = {
                'Woodland Pulp':       '#7B6000',
                'Saint Croix Chipping':'#1a5c1a',
                'St Croix Tissue':     '#1a1a6a',
                'LumberScan':          '#4a2a6a',
            }
            def loc_color(company: str, source: str) -> str:
                for k, v in LOCATION_COLORS.items():
                    if k.lower() in company.lower():
                        return v
                return '#003030' if source == 'validation' else '#3a2800'

            SECTOR_COLORS = {
                'fiber_supply': '#8B6914', 'pulp_mill': '#1a4a1a',
                'tissue_machine': '#1a1a4a', 'paper_machine': '#2a1a4a',
                'occ_recycling': '#3a2a0a', 'safety': '#4a1a1a',
            }

            # ── 6. Build the pivot table ─────────────────────────────────────
            try:
                self.table.cellChanged.disconnect()
            except RuntimeError:
                pass
            self.table.blockSignals(True)
            self.table.clear()
            self.table.setRowCount(0)
            self.component_rows = {}

            # Fixed cols: Equipment(0), Sector(1), then one per test
            n_fixed = 2
            n_cols = n_fixed + len(test_run_cols)
            self.table.setColumnCount(n_cols)

            # Set headers
            headers = ["Equipment", "Sector"]
            for iv_key in test_run_cols:
                iv = all_test_runs[iv_key]
                src_tag = "(S)" if iv['source'] == 'sparks' else "(F)"
                company_short = iv['source_version'][:18] if iv['source_version'] else ''
                person_short = iv['script_name'][:20]
                headers.append(f"{person_short}\n{company_short} {src_tag}")
            self.table.setHorizontalHeaderLabels(headers)

            # Column widths — Interactive so user can drag to resize
            hdr = self.table.horizontalHeader()
            hdr.setSectionResizeMode(0, QHeaderView.Interactive)
            self.table.setColumnWidth(0, 260)
            hdr.setSectionResizeMode(1, QHeaderView.Interactive)
            self.table.setColumnWidth(1, 90)
            for i in range(n_fixed, n_cols):
                hdr.setSectionResizeMode(i, QHeaderView.Interactive)
                self.table.setColumnWidth(i, 130)
            hdr.setStretchLastSection(False)

            # Color test_run column headers by phase + location
            for col_i, iv_num in enumerate(test_run_cols):
                iv = all_test_runs[iv_num]
                color = loc_color(iv['source_version'], iv['source'])
                item = self.table.horizontalHeaderItem(n_fixed + col_i)
                if item:
                    item.setBackground(QBrush(QColor(color)))
                    txt_color = "#FFD700" if iv['source'] == 'sparks' else "#00FF88"
                    item.setForeground(QBrush(QColor(txt_color)))
                    # Add phase border hint to tooltip
                    phase = "SPARKS" if iv['source'] == 'sparks' else "FUSION"
                    item.setToolTip(f"{phase} | {iv['script_name']} | {iv['source_version']}")

            # ── Spanning phase header row (row 0) ───────────────────────
            # SPARKS block | FUSION block — uses setSpan for visual grouping
            self.table.setRowCount(1)
            self.table.setRowHeight(0, 28)

            # Empty fixed cols
            for col in range(n_fixed):
                ph = QTableWidgetItem("")
                ph.setFlags(Qt.ItemIsEnabled)
                ph.setBackground(QBrush(QColor("#111")))
                self.table.setItem(0, col, ph)

            # Baseline span
            if baseline_cols:
                sp_start = n_fixed
                sp_end   = n_fixed + len(baseline_cols) - 1
                sp_item  = QTableWidgetItem(f"◆  SPARKS  ({len(baseline_cols)} tests)")
                sp_item.setFlags(Qt.ItemIsEnabled)
                sp_item.setTextAlignment(Qt.AlignCenter)
                sp_item.setBackground(QBrush(QColor("#5C4000")))
                sp_item.setForeground(QBrush(QColor("#FFD700")))
                font = QFont(); font.setBold(True); sp_item.setFont(font)
                self.table.setItem(0, sp_start, sp_item)
                if len(baseline_cols) > 1:
                    self.table.setSpan(0, sp_start, 1, len(baseline_cols))
                # Remaining sparks cols need placeholder or setSpan covers them
                for c in range(sp_start + 1, sp_end + 1):
                    ph = QTableWidgetItem("")
                    ph.setFlags(Qt.ItemIsEnabled)
                    ph.setBackground(QBrush(QColor("#5C4000")))
                    self.table.setItem(0, c, ph)

            # Validation span
            if validation_cols:
                fu_start = n_fixed + len(baseline_cols)
                fu_item  = QTableWidgetItem(f"◈  FUSION  ({len(validation_cols)} tests)")
                fu_item.setFlags(Qt.ItemIsEnabled)
                fu_item.setTextAlignment(Qt.AlignCenter)
                fu_item.setBackground(QBrush(QColor("#003C3C")))
                fu_item.setForeground(QBrush(QColor("#00FF88")))
                font = QFont(); font.setBold(True); fu_item.setFont(font)
                self.table.setItem(0, fu_start, fu_item)
                if len(validation_cols) > 1:
                    self.table.setSpan(0, fu_start, 1, len(validation_cols))
                for c in range(fu_start + 1, fu_start + len(validation_cols)):
                    ph = QTableWidgetItem("")
                    ph.setFlags(Qt.ItemIsEnabled)
                    ph.setBackground(QBrush(QColor("#003C3C")))
                    self.table.setItem(0, c, ph)

            # ── Data rows start at row 1 ─────────────────────────────────────
            row_idx = 1
            total_rows = sum(1 + len(items) for items in equipment_by_cat.values())
            self.table.setRowCount(1 + total_rows + len(equipment_by_cat))  # +phase row +headers

            for cat_name, items in equipment_by_cat.items():
                # Category header row
                for col in range(n_cols):
                    filler = QTableWidgetItem(cat_name.upper() if col == 0 else "")
                    filler.setFlags(Qt.ItemIsEnabled)
                    filler.setBackground(QBrush(QColor("#28143C")))
                    if col == 0:
                        font = QFont(); font.setBold(True)
                        filler.setFont(font)
                        filler.setForeground(QBrush(QColor("#FFD700")))
                    self.table.setItem(row_idx, col, filler)
                row_idx += 1

                for equip in items:
                    eq_name = equip['name']
                    sector = equip['sector']
                    cost_type = equip['cost_type']

                    # Col 0: Equipment name
                    name_item = QTableWidgetItem(eq_name)
                    name_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
                    sc = SECTOR_COLORS.get(sector, '#222')
                    name_item.setBackground(QBrush(QColor(sc)))
                    name_item.setForeground(QBrush(QColor("#ddd")))
                    self.table.setItem(row_idx, 0, name_item)

                    # Col 1: Sector
                    sect_item = QTableWidgetItem(sector.replace('_',' '))
                    sect_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
                    sect_item.setBackground(QBrush(QColor(sc)))
                    sect_item.setForeground(QBrush(QColor("#aaa")))
                    self.table.setItem(row_idx, 1, sect_item)

                    # Test Run columns
                    for col_i, iv_num in enumerate(test_run_cols):
                        iv = all_test_runs[iv_num]
                        impact = iv['impacts'].get(eq_name)
                        col = n_fixed + col_i
                        color = loc_color(iv['source_version'], iv['source'])

                        if impact is not None:
                            prod  = float(impact.get('production_cost', 0) or 0)
                            dmg   = float(impact.get('damage_cost',     0) or 0)
                            safe  = float(impact.get('safety_cost',     0) or 0)
                            other = float(impact.get('other_cost',      0) or 0)
                            cost  = float(impact.get('cost', 0) or 0) or prod or dmg or safe or other
                            ct    = impact.get('cost_type', 'repair')
                            notes = impact.get('notes', '')
                            eq_cost_type = equip.get('cost_type', 'repair')
                            effective_ct = ct if ct != 'repair' else eq_cost_type
                            if cost > 0:
                                # Show breakdown if multiple cost types filled
                                parts = []
                                if prod  > 0: parts.append(f"P:${prod:,.0f}")
                                if dmg   > 0: parts.append(f"D:${dmg:,.0f}")
                                if safe  > 0: parts.append(f"S:${safe:,.0f}")
                                if other > 0: parts.append(f"O:${other:,.0f}")
                                display = " | ".join(parts) if parts else f"${cost:,.0f}"
                                if prod > 0:
                                    text_color = "#FF8C00"   # orange = production
                                elif safe > 0:
                                    text_color = "#FF4444"   # red = safety
                                elif iv['source'] == 'sparks':
                                    text_color = "#FFD700"   # gold = sparks
                                else:
                                    text_color = "#00FF88"   # green = validation repair
                            else:
                                display = "✓" if notes else "·"
                                text_color = "#FF8C00" if effective_ct == 'production' else "#888"
                            cell = QTableWidgetItem(display)
                            cell.setTextAlignment(Qt.AlignCenter)
                            cell.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
                            cell.setBackground(QBrush(QColor(color)))
                            cell.setForeground(QBrush(QColor(text_color)))
                            if notes:
                                cell.setToolTip(notes[:200])
                        else:
                            cell = QTableWidgetItem("")
                            cell.setFlags(Qt.ItemIsEnabled)
                            cell.setBackground(QBrush(QColor("#111")))
                        self.table.setItem(row_idx, col, cell)

                    self.component_rows[eq_name] = {'row': row_idx, 'config': equip}
                    row_idx += 1

            self.table.setRowCount(row_idx)
            self.table.blockSignals(False)

            # ── 7. Update totals ─────────────────────────────────────────────
            baseline_repair = sum(
                eq['baseline_metric'] for eq in live_config.values()
                if eq.get('baseline_metric', 0) > 0
            )
            baseline_prod = sum(
                eq['baseline_production_impact'] for eq in live_config.values()
                if eq.get('baseline_production_impact', 0) > 0
            )
            validation_repair = 0
            for iv in validation_tests.values():
                for name, impact in iv['impacts'].items():
                    if impact['cost'] > 0 and impact['cost_type'] == 'repair':
                        validation_repair += impact['cost']

            self.total_label.setText(
                f"<b>Baseline Repair: ${baseline_repair:,.0f} | "
                f"Production Impact: ${baseline_prod:,.0f} | "
                f"Validation Repair: ${validation_repair:,.0f}</b>"
            )
            if hasattr(self, 'baseline_total_label'):
                self.baseline_total_label.setText(f"Baseline Repair: ${baseline_repair:,.0f}")
            if hasattr(self, 'validation_total_label'):
                self.validation_total_label.setText(f"Validation Repair: ${validation_repair:,.0f}")
            if hasattr(self, 'production_total_label'):
                self.production_total_label.setText(f"Production: ${baseline_prod:,.0f}")

            n_iv = len(test_run_cols)
            self.info_label.setText(
                f"✓ {len(self.component_rows)} equipment rows | "
                f"{len(baseline_cols)} Baseline tests | "
                f"{len(validation_cols)} Validation tests | "
                f"{n_iv} total columns"
            )
            self.info_label.setStyleSheet("color: #00FF88; padding: 6px;")
            self.impact_updated.emit(baseline_repair + validation_repair)

            # ── 8. Update frozen column overlay ──────────────────────────────
            self._update_frozen_columns()

        def _on_cell_changed(self, row, col):
            """Recalc annual impact when cost or frequency changes."""
            if col in (4, 6):  # Validation metric cost or events/yr
                self._recalc_row(row)
                self._recalc_total()
        
        def _recalc_row(self, row):
            """Recalculate annual impact for one row."""
            try:
                cost_text = self.table.item(row, 4).text().replace('$', '').replace(',', '')
                cost = float(cost_text) if cost_text else 0

                # If no validation repair cost, fall back to sparks repair
                if cost == 0:
                    baseline_text = self.table.item(row, 3).text().replace('$', '').replace(',', '')
                    cost = float(baseline_text) if baseline_text else 0

                freq_text = self.table.item(row, 6).text()
                freq = float(freq_text) if freq_text else 1

                annual = cost * freq
                self.table.item(row, 7).setText(f"${annual:,.0f}")
            except:
                pass
        
        def _recalc_total(self):
            """Recalculate total annual impact."""
            total = 0
            for row in range(self.table.rowCount()):
                item = self.table.item(row, 7)
                if item:
                    val = _parse_cost(item.text())
                    if val > 0:
                        total += val
            
            self.total_label.setText(f"<b>Total Annual Impact: ${total:,.0f}</b>")
            self.impact_updated.emit(total)
        
        def rebuild(self):
            """Full rebuild — call when project changes."""
            self.table.cellChanged.disconnect()
            self.table.setRowCount(0)
            self.component_rows = {}
            self._build_table()
        
        def get_equipment_summary(self) -> Dict:
            """Get summary from the live pivot table for export."""
            summary = {'items': [], 'total': 0, 'baseline_total': 0, 'validation_total': 0, 'production_total': 0}
            
            n_fixed = 2  # Equipment, Sector
            n_cols = self.table.columnCount()
            
            # Determine which columns are Baseline vs Validation from headers
            col_source = {}  # col_index -> 'sparks' or 'validation'
            col_person = {}  # col_index -> person name
            for col in range(n_fixed, n_cols):
                header_item = self.table.horizontalHeaderItem(col)
                if header_item:
                    htext = header_item.text()
                    col_source[col] = 'sparks' if '(S)' in htext else 'validation'
                    col_person[col] = htext.split('\n')[0].strip()
            
            # Scan each equipment row
            for eq_name, row_data in self.component_rows.items():
                row = row_data['row']
                baseline_cost = 0.0
                validation_cost = 0.0
                production_cost = 0.0
                sources = []
                
                for col in range(n_fixed, n_cols):
                    cell = self.table.item(row, col)
                    if not cell:
                        continue
                    text = cell.text().strip()
                    if not text or text in ('', '·', '✓'):
                        if text == '✓':
                            sources.append(col_person.get(col, ''))
                        continue
                    
                    # Parse cost parts: "P:$197,100 | D:$225,000" or "$424,320"
                    total_cell = 0.0
                    prod_cell = 0.0
                    for part in text.split('|'):
                        part = part.strip()
                        if part.startswith('P:'):
                            val = _parse_cost(part[2:])
                            prod_cell += val
                            total_cell += val
                        elif part.startswith('D:') or part.startswith('S:') or part.startswith('O:'):
                            val = _parse_cost(part[2:])
                            total_cell += val
                        else:
                            val = _parse_cost(part)
                            total_cell += val
                    
                    if total_cell > 0:
                        source = col_source.get(col, 'validation')
                        if source == 'sparks':
                            baseline_cost += total_cell
                        else:
                            validation_cost += total_cell
                        production_cost += prod_cell
                        sources.append(col_person.get(col, ''))
                
                annual = baseline_cost + validation_cost
                if annual > 0 or sources:
                    summary['items'].append({
                        'equipment': eq_name,
                        'baseline_cost': baseline_cost,
                        'validation_cost': validation_cost,
                        'production_impact': production_cost,
                        'annual_impact': annual,
                        'sources': ', '.join(s for s in sources if s),
                    })
                    summary['baseline_total'] += baseline_cost
                    summary['validation_total'] += validation_cost
                    summary['production_total'] += production_cost
                    summary['total'] += annual
            
            return summary
        
        def export_analysis(self):
            """Export equipment impact analysis — auto-saves to equipment_analysis/ folder."""
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"substrate_impact_{timestamp}.txt"

            # Default path: system equipment_analysis folder
            default_dir = None
            if sys.platform == "win32":
                system_path = Path(
                    "C:/TITS/TITS_GIBUSH_AISOS_SPINE/"
                    "TITS_GIBUSH_AISOS_SPINE_ICORPS/equipment_analysis"
                )
                if not system_path.exists():
                    try:
                        system_path.mkdir(parents=True, exist_ok=True)
                    except Exception:
                        pass
                if system_path.exists():
                    default_dir = system_path

            if default_dir is None:
                # Fallback: next to script
                default_dir = Path(__file__).parent if "__file__" in dir() else Path.cwd()
                fallback = default_dir / "equipment_analysis"
                fallback.mkdir(exist_ok=True)
                default_dir = fallback

            default_path = str(default_dir / filename)

            filepath, _ = QFileDialog.getSaveFileName(
                self, "Export Equipment Analysis",
                default_path,
                "Text Files (*.txt);;All Files (*)"
            )
            if not filepath:
                return
            
            summary = self.get_equipment_summary()
            lines = [
                "EQUIPMENT IMPACT ANALYSIS",
                f"Project: {self._get_project_type()}",
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                f"Baseline Established: ${summary['baseline_total']:,.0f}",
                f"Validation Confirmed: ${summary['validation_total']:,.0f}",
                f"Total Annual Impact: ${summary['total']:,.0f}",
                "",
                f"{'Equipment':<40} {'Baseline Repair':>13} {'Validation Repair':>13} {'Production':>12} {'Annual $':>12}  Sources",
                "-" * 110,
            ]
            for item in summary['items']:
                lines.append(
                    f"{item['equipment']:<40} "
                    f"${item['baseline_cost']:>10,.0f} "
                    f"${item['validation_cost']:>10,.0f} "
                    f"${item.get('production_impact',0):>10,.0f} "
                    f"${item['annual_impact']:>10,.0f}  "
                    f"{item['sources']}"
                )
            
            with open(filepath, 'w') as f:
                f.write("\n".join(lines))
            
            QMessageBox.information(self, "Exported", f"Equipment analysis exported to:\n{filepath}")
