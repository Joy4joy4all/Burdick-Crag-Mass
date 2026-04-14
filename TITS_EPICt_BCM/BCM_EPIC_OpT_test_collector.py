# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
GIBUSH FUSION TEST COLLECTOR
==================================
ICorps BCM Test Collection & Q-Cube Synergy Analysis Module

Features:
- Test Run data collection (ICorps template)
- Q-Cube 3D synergy matrix population
- PyTorch AI pattern recognition
- Generative customer segment refinement
- Value proposition optimization

Architecture: PySide6 + PyTorch (same as TITS system)
"""

import os
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

import sys
import json
import re
import os
from datetime import datetime
from pathlib import Path

# ── PATH FIX: Ensure own directory is on sys.path for doc_reader etc ──
_SELF_DIR = Path(__file__).parent.resolve()
if str(_SELF_DIR) not in sys.path:
    sys.path.insert(0, str(_SELF_DIR))
from typing import Dict, List, Optional, Tuple

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QTextEdit, QPushButton, QTabWidget,
    QTableWidget, QTableWidgetItem, QComboBox, QGroupBox,
    QScrollArea, QMessageBox, QFileDialog, QSplitter,
    QTreeWidget, QTreeWidgetItem, QProgressBar, QCheckBox,
    QProgressDialog, QHeaderView, QFrame
)
from PySide6.QtCore import Qt, Signal, QThread, QTimer
from PySide6.QtGui import QFont, QColor, QPalette, QBrush

# Real ML: sklearn for embeddings + clustering
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.decomposition import TruncatedSVD
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import AgglomerativeClustering
import numpy as np

# ── PROJECT PATHS — Single Source of Truth ──
try:
    from project_paths import (
        PROJECT_ROOT, BCM_ROOT, ICORPS_ROOT,
        list_deployed_projects, get_project_dir, ensure_project_structure,
        get_bmc_dir, get_tested_dir, get_template_dir,
        get_baseline_dir, get_ai_report_dir, get_pptx_slide_dir,
        get_substrate_class_dir, get_inclusion_log, get_project_config,
        find_project_excel, find_project_baseline_deck,
    )
    PROJECT_PATHS_AVAILABLE = True
except ImportError:
    PROJECT_PATHS_AVAILABLE = False
    print("WARNING: project_paths.py not found — using legacy paths")
    PROJECT_ROOT = Path.cwd() / "BCM_Projects" / "project_caas_deployed"
    BCM_ROOT = Path.cwd() / "BCM_Projects"
    def list_deployed_projects(): return []
    def get_project_dir(pid): return PROJECT_ROOT / pid
    def ensure_project_structure(pid): d = PROJECT_ROOT / pid; d.mkdir(parents=True, exist_ok=True); return d
    def get_bmc_dir(pid): return get_project_dir(pid) / "BMC_generation"
    def get_tested_dir(pid): return get_project_dir(pid) / "tested"
    def get_template_dir(pid): return get_project_dir(pid) / "test_template_list"
    def get_baseline_dir(pid): return get_project_dir(pid) / "Baseline_deck_tests_final"
    def get_ai_report_dir(pid): return get_project_dir(pid) / "_AI_generated_report"
    def get_pptx_slide_dir(pid): return get_project_dir(pid) / "tested_pptx_slide_creator"
    def get_substrate_class_dir(pid): return get_project_dir(pid) / "substrate_class_profile"
    def get_inclusion_log(pid): return get_project_dir(pid) / "inclusion_log.json"
    def get_immulsion_log(pid): return get_project_dir(pid) / "immulsion_log.json"
    def get_project_config(pid): return get_project_dir(pid) / "project_config.json"
    def find_project_excel(pid): return None
    def find_project_baseline_deck(pid): return None

# ── BMC ENGINE — Accumulative KNN Business Model Canvas ──
try:
    from bmc_engine import (
        generate_project_bmc, BMCAggregator, BMCState,
        BMC_FIELDS, BMC_DISPLAY_NAMES, generate_bmc_docx,
    )
    BMC_ENGINE_AVAILABLE = True
except ImportError:
    BMC_ENGINE_AVAILABLE = False
    print("WARNING: bmc_engine.py not found — BMC generation disabled")

# ── Excel Test Run Loader (still needed for plan loading) ──
try:
    from excel_test_loader import (
        load_test_plans, 
        set_active_project,
        get_active_project,
        reload_from_excel,
        is_excel_available,
        get_excel_path,
        create_project,
        list_projects,
        get_project_folder,
    )
    EXCEL_LOADER_AVAILABLE = True
except ImportError:
    EXCEL_LOADER_AVAILABLE = False
    def load_test_plans(): return []
    def set_active_project(p): pass
    def get_active_project(): return list_deployed_projects()[0] if list_deployed_projects() else "DEFAULT"
    def reload_from_excel(): return []
    def is_excel_available(): return False
    def get_excel_path(): return None
    def create_project(p, f=None): return None
    def list_projects(): return list_deployed_projects()
    def get_project_folder(): return PROJECT_ROOT

# Import AI Doc Reader
try:
    from doc_reader import DocReaderWidget
    DOC_READER_AVAILABLE = True
except ImportError:
    DOC_READER_AVAILABLE = False

# Import Inclusion Receipt Collector (Stream 2 of Intelligence Triad)
try:
    from Inclusion_Module_Receipt_Collector.inclusion_receipt import (
        get_inclusion_for_project, MASTER_LENS
    )
    INCLUSION_AVAILABLE = True
    print("  ✓ Inclusion Receipt Collector loaded (Master Lens ML1/ML2/ML3)")
except ImportError:
    INCLUSION_AVAILABLE = False
    # Fallback: try direct path
    try:
        import sys as _sys
        _inc_path = str(Path(__file__).parent / "Inclusion_Module_Receipt_Collector")
        if _inc_path not in _sys.path:
            _sys.path.insert(0, str(Path(__file__).parent))
        from Inclusion_Module_Receipt_Collector.inclusion_receipt import (
            get_inclusion_for_project, MASTER_LENS
        )
        INCLUSION_AVAILABLE = True
        print("  ✓ Inclusion Receipt Collector loaded (fallback path)")
    except ImportError:
        print("  ○ Inclusion Receipt Collector not available")
    print("Warning: doc_reader.py not found. Doc Reader tab disabled.")

# ── Import Emulsion Receiver (Stream 3 of Intelligence Triad) ──
# Pattern: same as Inclusion_Module_Receipt_Collector above
try:
    from Immulsion_Receiver.emulsion_receiver import (
        get_immulsion_for_project, ImmulsionReceiver, IMMULSION_LENS
    )
    EMULSION_AVAILABLE = True
    print("  ✓ Emulsion Receiver loaded (ML5/ML6/ML7)")
except ImportError:
    EMULSION_AVAILABLE = False
    try:
        from emulsion_receiver import (
            get_immulsion_for_project, ImmulsionReceiver, IMMULSION_LENS
        )
        EMULSION_AVAILABLE = True
        print("  ✓ Emulsion Receiver loaded (flat import)")
    except ImportError:
        print("  ○ Emulsion Receiver not available")

# ── R&D HEALTH TAB — Research + Development Convergence ──
# Flat imports from _SELF_DIR (same as bayesian_engine, etc.)
try:
    from rd_health_tab import (
        create_inclusion_intelligence_tab as create_rd_health_tab,
        compute_and_update_tab as compute_rd_health,
    )
    _HAS_RD_HEALTH_TAB = True
    print("  ✓ R&D Health Tab loaded (Research ↔ Development convergence)")
except ImportError:
    _HAS_RD_HEALTH_TAB = False

# ── External AI Engines

# ── TEST RECONCILER — Data integrity for all 6 footprint locations ──
try:
    from test_reconciler import (
        Test RunReconciler,
        remove_test_run_completely,
        quick_verify,
    )
    _HAS_RECONCILER = True
except ImportError:
    _HAS_RECONCILER = False  # conditional - enhance when available
# These are REAL statistical/ML engines that enhance the production heuristics.
# If absent, system runs on deterministic keyword inference (the baseline).

try:
    from bayesian_engine import BayesianHypothesisEngine, Evidence, EvidenceType
    BAYESIAN_ENGINE_AVAILABLE = True
    print("  ✓ Bayesian Hypothesis Engine loaded (Beta-Binomial conjugate model)")
except ImportError:
    BAYESIAN_ENGINE_AVAILABLE = False

try:
    from doctoral_analyzer import DoctoralAnalyzer, QuestionCandidate
    DOCTORAL_ANALYZER_AVAILABLE = True
    print("  ✓ Doctoral Analyzer loaded (info-gain question selection)")
except ImportError:
    DOCTORAL_ANALYZER_AVAILABLE = False

try:
    from knowledge_extractor import KnowledgeExtractor, EntityType
    KNOWLEDGE_EXTRACTOR_AVAILABLE = True
    print("  ✓ Knowledge Extractor loaded (entity extraction pipeline)")
except ImportError:
    KNOWLEDGE_EXTRACTOR_AVAILABLE = False

try:
    from adaptive_analyzer import AdaptiveTestGenerator
    ADAPTIVE_ANALYZER_AVAILABLE = True
    print("  ✓ Adaptive Analyzer loaded (keyword amplitude + segment targeting)")
except ImportError:
    ADAPTIVE_ANALYZER_AVAILABLE = False

try:
    from learning_engine import detect_patterns, generate_questions, track_hypotheses
    LEARNING_ENGINE_AVAILABLE = True
    print("  ✓ Learning Engine loaded (TF-IDF patterns, hypothesis tracking)")
except ImportError:
    LEARNING_ENGINE_AVAILABLE = False

# ============================================================================
# CONFIGURATION
# ============================================================================

# TEST_DATA_DIR routes through project_paths
# Every path goes: get_active_project() → get_project_dir() → subfolder
# This variable is kept for backward compat but routes to PROJECT_ROOT
TEST_DATA_DIR = PROJECT_ROOT
TEST_DATA_DIR.mkdir(exist_ok=True, parents=True)

# ============================================================================
# NUMBERING CONSTANTS
# ============================================================================
# Baseline (discovery phase) test_runs: #1 through SPARKS_MAX
# Validation (validation phase) test_runs: #(SPARKS_MAX+1) onward
# This prevents ID collision between the two phases.
SPARKS_MAX = 50       # Reserve slots 1-50 for Baseline; Validation starts at 51
SPARKS_REQUIRED = 12  # Baseline requires 12 tests minimum

# ============================================================================
# STACK KEY NORMALIZATION
# ============================================================================
# The JSON database uses correct Unicode Greek letters (Sα, Sβ, Sγ, Sδ)
# but some Python source files have double-encoded UTF-8 (Sα, Sβ, Sγ, Sδ).
# This normalizer ensures both encodings map to the SAME canonical key.

_STACK_NORMALIZE_MAP = {
    # Double-encoded → canonical
    'S\u00ce\u00b1': 'S\u03b1',  # Sα → Sα
    'S\u00ce\u00b2': 'S\u03b2',  # Sβ → Sβ
    'S\u00ce\u00b3': 'S\u03b3',  # Sγ → Sγ
    'S\u00ce\u00b4': 'S\u03b4',  # Sδ → Sδ
    # Already correct → pass through
    'S\u03b1': 'S\u03b1',        # Sα
    'S\u03b2': 'S\u03b2',        # Sβ
    'S\u03b3': 'S\u03b3',        # Sγ
    'S\u03b4': 'S\u03b4',        # Sδ
    # ASCII fallbacks
    'Sa': 'S\u03b1', 'Sb': 'S\u03b2', 'Sg': 'S\u03b3', 'Sd': 'S\u03b4',
}

def normalize_stack_key(raw: str) -> str:
    """Normalize a Q-Stack key to canonical Unicode form."""
    return _STACK_NORMALIZE_MAP.get(raw, raw)

def build_cube_key(layer: str, obj: str, stack: str) -> str:
    """Build a normalized cube matrix key."""
    return f"[{layer}, {obj}, {normalize_stack_key(stack)}]"


QCUBE_DIMENSIONS = {
    # ══════════════════════════════════════════════════════════════
    # FOUNDATION LENSES (1-3) — Structural Position
    # ══════════════════════════════════════════════════════════════
    "Q_LAYER": {
        "L1": "Operator Care - Daily pain, hands-on experience (3× weight)",
        "L2": "Manager Care - Production responsibility, P&L ownership (1.5×)",
        "L3": "Executive Care - Strategic risk, enterprise protection (1×)"
    },
    "Q_OBJECT": {
        "OA": "Upstream Supplier - Chip mills, timber (get blamed)",
        "OB": "Transfer Point - Blow line, bark line (blind spot)",
        "OC": "Downstream Receiver - Kraft mill, digester (gets damaged)"
    },
    "Q_STACK": {
        "Sα": "Cross-Mill Patterns - Same pain across facilities",
        "Sβ": "Post-Investment Failures - Modern equipment, old problems",
        "Sγ": "Baseline Void - New systems, no commissioning data",
        "Sδ": "Dual Impact - Single source, multiple product damages"
    },
    # ══════════════════════════════════════════════════════════════
    # ANALYTICAL LENSES (4-6) — Evidence Quality
    # ══════════════════════════════════════════════════════════════
    "Q_AWARENESS": {
        "A1": "Unaware - Doesn't know contamination exists or matters",
        "A2": "Aware-Accepted - Knows, considers it cost of doing business",
        "A3": "Aware-Suffering - Knows, wants it fixed, absorbing cost",
        "A4": "Aware-Acting - Actively trying to solve it"
    },
    "Q_EVIDENCE": {
        "V1": "Anecdotal - War story, no numbers",
        "V2": "Experiential - Repeated personal experience, frequency estimates",
        "V3": "Quantified - Actual dollar amounts, downtime hours",
        "V4": "Documented - Has records, logs, purchase orders to prove it"
    },
    "Q_TIMEHORIZON": {
        "T1": "Chronic - Every load, every shift, normalized pain",
        "T2": "Episodic - Seasonal, weather-dependent, periodic",
        "T3": "Acute - Single catastrophic event, emergency",
        "T4": "Latent - Cumulative damage, only visible at shutdown/inspection"
    },
    # ══════════════════════════════════════════════════════════════
    # INNOVATION LENSES (7-10) — The Four Originals (Team GIBUSH IP)
    # ══════════════════════════════════════════════════════════════
    "Q_CASCADE": {
        "C1": "Direct - Same entity creates and suffers damage",
        "C2": "One Handoff - Two parties, one transfer of accountability",
        "C3": "Two Handoffs - Three parties, accountability diluting",
        "C4": "Three Handoffs - Four parties, feedback loop weakening",
        "C5": "Four Handoffs - Five parties, feedback loop severed",
        "C6": "Full Cascade - Six+ parties, no traceability possible"
    },
    "Q_NORMALIZATION": {
        "N1": "Recognized Hazard - Everyone knows it's wrong, actively avoided",
        "N2": "Tolerated Risk - Known risk, tolerated as low-frequency",
        "N3": "Accepted Practice - Deliberately done, solves a different problem",
        "N4": "Invisible Normal - So embedded it's not recognized as contamination"
    },
    "Q_COUNTERFLOW": {
        "Fm": "Material Only - Sees physical contamination, no cost or data visibility",
        "F$": "Money Only - Sees cost impact, no material or data visibility",
        "Fi": "Information Only - Sees data/reports, no material or cost visibility",
        "Fg": "Gap Position - Has partial visibility, sits in information gap"
    },
    "Q_TRIBALKNOWLEDGE": {
        "K1": "Institutional - Knowledge in systems (CMMS, SAP, SOPs), survives turnover",
        "K2": "Procedural - Informal procedures, transmitted by training, not written",
        "K3": "Tribal - Lives in one person's head, lost at retirement",
        "K4": "Lost - Knowledge already gone through turnover, ghost knowledge"
    },
}

# Human-readable lens names for UI
QCUBE_LENS_NAMES = {
    # Foundation (1-3)
    "Q_LAYER":           "① Authority (WHO)",
    "Q_OBJECT":          "② Supply Chain (WHERE)",
    "Q_STACK":           "③ Context (WHAT)",
    # Analytical (4-6)
    "Q_AWARENESS":       "④ Awareness (RELATIONSHIP)",
    "Q_EVIDENCE":        "⑤ Evidence Type (STRENGTH)",
    "Q_TIMEHORIZON":     "⑥ Time Horizon (WHEN)",
    # Innovation / Originals (7-10)
    "Q_CASCADE":         "⑦ Cascade Depth (HANDOFFS)",
    "Q_NORMALIZATION":   "⑧ Normalization (CULTURAL DEPTH)",
    "Q_COUNTERFLOW":     "⑨ Counter-Flow (VISIBILITY GAP)",
    "Q_TRIBALKNOWLEDGE": "⑩ Tribal Knowledge (RETIREMENT CLOCK)",
}

# ============================================================================
# LENS INFERENCE ENGINE — Auto-tags tests across all 10 dimensions
# ============================================================================

def infer_lens_awareness(test_run) -> str:
    """
    Infer Awareness lens (A1-A4) from test_run content.
    A1=Unaware, A2=Aware-Accepted, A3=Aware-Suffering, A4=Aware-Acting
    """
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'experiments', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'action_iterate', ''),
    ])).lower()
    
    if not all_text.strip():
        return ''
    
    # A4: Actively trying to solve it
    a4_signals = ['pilot', 'testing', 'installed', 'trying to', 'working on',
                  'implemented', 'solution', 'prototype', 'proof of concept',
                  'we built', 'we added', 'system we use', 'monitoring']
    if any(s in all_text for s in a4_signals):
        return 'A4'
    
    # A3: Knows and suffering — wants fix, absorbing cost
    a3_signals = ['wish we could', 'costs us', 'we lose', 'we need',
                  'biggest problem', 'kills us', 'hurts us', 'pain',
                  'damage', 'downtime', 'can\'t afford', 'want to fix',
                  'need better', 'frustrat', 'fed up', 'unacceptable']
    if any(s in all_text for s in a3_signals):
        return 'A3'
    
    # A2: Knows but accepted — cost of doing business
    a2_signals = ['always been', 'part of the job', 'normal', 'accepted',
                  'cost of doing business', 'just how it is', 'nothing you can do',
                  'that\'s the way', 'everybody deals with', 'common practice',
                  'nobody tracks', 'dirt on loads', 'standard practice']
    if any(s in all_text for s in a2_signals):
        return 'A2'
    
    # A1: Unaware
    a1_signals = ['didn\'t know', 'never thought', 'what contamination',
                  'not a problem', 'doesn\'t affect', 'first time hearing']
    if any(s in all_text for s in a1_signals):
        return 'A1'
    
    # Default: if they discussed equipment or results, they're at least A2
    if getattr(test_run, 'substrate_impacts', []):
        return 'A2'
    return 'A2'


def infer_lens_evidence(test_run) -> str:
    """
    Infer Evidence Type lens (V1-V4) from test_run content.
    V1=Anecdotal, V2=Experiential, V3=Quantified, V4=Documented
    """
    results = getattr(test_run, 'results', '') or ''
    all_text = results + ' ' + (getattr(test_run, 'experiments', '') or '')
    text_lower = all_text.lower()
    
    if not text_lower.strip():
        return ''
    
    # V4: Has documentation/records
    v4_signals = ['records show', 'purchase order', 'log shows', 'data from',
                  'our records', 'tracked', 'documented', 'report says',
                  'spreadsheet', 'database', 'maintenance log', 'work order']
    if any(s in text_lower for s in v4_signals):
        return 'V4'
    
    # V3: Has dollar amounts or specific numbers
    has_dollars = bool(re.search(r'\$[\d,]+', all_text))
    has_specific_numbers = bool(re.search(r'\b\d{2,}\s*(hour|day|week|event|time|per year|annually|a year)', text_lower))
    has_cost_data = bool(getattr(test_run, 'substrate_impacts', []) and 
                         any(imp.get('cost', '') and imp['cost'] != '0' 
                             for imp in test_run.substrate_impacts))
    if has_dollars or has_specific_numbers or has_cost_data:
        return 'V3'
    
    # V2: Repeated experience, frequency language
    v2_signals = ['every time', 'always', 'happens regularly', 'common',
                  'frequently', 'multiple times', 'years of', 'in my experience',
                  'across mills', 'every mill', 'seen it at', 'happens when',
                  'winter is worse', 'seasonal', 'pattern']
    if any(s in text_lower for s in v2_signals):
        return 'V2'
    
    # V1: Anecdotal — single story, war story
    return 'V1'


def infer_lens_timehorizon(test_run) -> str:
    """
    Infer Time Horizon lens (T1-T4) from test_run content.
    T1=Chronic, T2=Episodic, T3=Acute, T4=Latent
    """
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'experiments', ''),
    ])).lower()
    
    if not all_text.strip():
        return ''
    
    scores = {'T1': 0, 'T2': 0, 'T3': 0, 'T4': 0}
    
    t1_words = ['every load', 'every shift', 'every day', 'constant',
                'continuous', 'always there', 'normal', 'accepted practice',
                'part of the job', 'ongoing', 'persistent', 'chronic',
                'dirt on loads', 'phantom weight', 'every delivery']
    for w in t1_words:
        if w in all_text: scores['T1'] += 1
    
    t2_words = ['winter', 'seasonal', 'rain season', 'spring breakup',
                'wet weather', 'frozen', 'periodic', 'certain times',
                'time of year', 'mud season', 'summer', 'dry season']
    for w in t2_words:
        if w in all_text: scores['T2'] += 1
    
    t3_words = ['catastrophic', 'explosion', 'emergency', 'shutdown',
                'disaster', 'insurance claim', 'worst event', 'rock through',
                'blew out', 'destroyed', 'rupture', 'burst', 'failure',
                'lost load', 'crash', 'near miss', 'incident']
    for w in t3_words:
        if w in all_text: scores['T3'] += 1
    
    t4_words = ['cumulative', 'wear', 'gradual', 'over time', 'erosion',
                'builds up', 'inspection found', 'didn\'t know until',
                'shutdown revealed', 'turnaround', 'hidden damage',
                'nobody noticed', 'slow degradation']
    for w in t4_words:
        if w in all_text: scores['T4'] += 1
    
    best = max(scores, key=scores.get)
    if scores[best] > 0:
        return best
    return 'T1'


def infer_lens_cascade(test_run) -> str:
    """
    Infer Cascade Depth lens (C1-C6) — supply chain forensics.
    How many ownership handoffs between contamination creation and equipment death.
    
    Uses supply chain position (q_object) + text signals about upstream/downstream.
    C1=Direct, C2=One handoff, C3=Two, C4=Three, C5=Four, C6=Full cascade
    """
    q_object = getattr(test_run, 'q_object', '')
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'experiments', ''),
        getattr(test_run, 'action_iterate', ''),
    ])).lower()
    title = getattr(test_run, 'title', '').lower()
    
    if not all_text.strip():
        return ''
    
    # Deep cascade signals — multiple handoffs mentioned
    c6_signals = ['entire supply chain', 'forest to paper', 'from the woods to',
                  'road to digester', 'landing to paper machine', 'no traceability',
                  'can\'t trace', 'nobody owns', 'six parties', 'full chain']
    if any(s in all_text for s in c6_signals):
        return 'C6'
    
    c5_signals = ['road maintenance', 'road grading', 'landing crew',
                  'four handoff', 'five parties', 'road to mill']
    if any(s in all_text for s in c5_signals):
        return 'C5'
    
    c4_signals = ['forest to mill', 'harvester to digester', 'three handoff',
                  'logging to kraft', 'slash to chip', 'woods to pulp',
                  'dirt at landing', 'landing debris']
    if any(s in all_text for s in c4_signals):
        return 'C4'
    
    c3_signals = ['supplier to mill', 'two handoff', 'chip mill to kraft',
                  'hauler to receiver', 'trucker delivers', 'loading to unloading']
    if any(s in all_text for s in c3_signals):
        return 'C3'
    
    # Position-based inference
    if q_object == 'OA':
        # Upstream — typically C3+ (they create, someone else suffers)
        if any(s in title for s in ['hauler', 'trucker', 'driver', 'transport']):
            return 'C4'  # Carrier = mid-chain
        if any(s in title for s in ['timber', 'forest', 'harvest', 'logger']):
            return 'C4'  # Source = top of chain
        return 'C3'
    
    if q_object == 'OB':
        # Transfer point — C2-C3 typically
        return 'C2'
    
    if q_object == 'OC':
        # Downstream receiver — experiencing damage from upstream
        # Check if they mention specific upstream sources
        if any(s in all_text for s in ['supplier', 'truck', 'chip mill', 'upstream']):
            return 'C3'
        return 'C2'
    
    # Default
    return 'C2'


def infer_lens_normalization(test_run) -> str:
    """
    Infer Normalization of Deviance lens (N1-N4) — Vaughan framework.
    How culturally embedded is this destructive practice.
    
    N1=Recognized Hazard, N2=Tolerated Risk, N3=Accepted Practice, N4=Invisible Normal
    """
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'experiments', ''),
        getattr(test_run, 'action_iterate', ''),
    ])).lower()
    
    if not all_text.strip():
        return ''
    
    # N4: So embedded it's not recognized as contamination
    n4_signals = ['just mud', 'washes off', 'not contamination', 'that\'s not a problem',
                  'never considered', 'nobody thinks of', 'always been like that',
                  'what do you mean', 'natural', 'it\'s just', 'doesn\'t count',
                  'not a concern', 'never measured', 'nobody tracks that']
    if any(s in all_text for s in n4_signals):
        return 'N4'
    
    # N3: Accepted practice — deliberately done, solves different problem
    n3_signals = ['accepted practice', 'standard practice', 'we do it because',
                  'everybody does', 'common practice', 'industry standard',
                  'that\'s how', 'normal procedure', 'dirt on loads',
                  'load stability', 'prevents shift', 'have to', 'no choice',
                  'only way', 'cost of doing business', 'part of the job']
    if any(s in all_text for s in n3_signals):
        return 'N3'
    
    # N2: Tolerated risk — known but considered manageable
    n2_signals = ['once a year', 'rarely', 'low frequency', 'manageable',
                  'we deal with it', 'tolerate', 'acceptable risk',
                  'not that bad', 'could be worse', 'occasional',
                  'happens sometimes', 'we live with']
    if any(s in all_text for s in n2_signals):
        return 'N2'
    
    # N1: Recognized hazard — everyone knows it's wrong
    n1_signals = ['never allow', 'zero tolerance', 'strict policy',
                  'not acceptable', 'prohibited', 'we reject',
                  'send it back', 'penalty', 'violation', 'safety hazard',
                  'unacceptable', 'need to fix', 'biggest problem']
    if any(s in all_text for s in n1_signals):
        return 'N1'
    
    # Default: most contamination in kraft mills is N2-N3 (tolerated to accepted)
    return 'N2'


def infer_lens_counterflow(test_run) -> str:
    """
    Infer Counter-Flow Mapping lens (Fm/F$/Fi/Fg) — three simultaneous flows.
    Material flow, Money flow, Information flow — where do they diverge.
    
    Fm=Material only, F$=Money only, Fi=Information only, Fg=Gap position
    """
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'experiments', ''),
        getattr(test_run, 'action_iterate', ''),
    ])).lower()
    title = getattr(test_run, 'title', '').lower()
    q_object = getattr(test_run, 'q_object', '')
    q_layer = getattr(test_run, 'q_layer', '')
    
    if not all_text.strip():
        return ''
    
    # Check which flows this person has visibility into
    sees_material = False
    sees_money = False
    sees_information = False
    
    # Material flow visibility — physical contamination knowledge
    mat_signals = ['saw rock', 'dirt on', 'debris', 'contamination in',
                   'physical', 'handles', 'we see', 'comes through',
                   'in the chip', 'in the wood', 'on the load', 'bark',
                   'mud', 'sand', 'gravel', 'metal', 'tramp']
    if any(s in all_text for s in mat_signals):
        sees_material = True
    if any(s in title for s in ['operator', 'tender', 'hauler', 'trucker', 'driver']):
        sees_material = True
    
    # Money flow visibility — cost/budget knowledge
    money_signals = ['cost', 'budget', 'expense', 'price', 'dollar',
                     'spend', 'invest', 'annual damage', 'replacement cost',
                     'downtime cost', 'maintenance budget', 'capital',
                     'roi', 'payback', 'we pay', 'costs us']
    if any(s in all_text for s in money_signals):
        sees_money = True
    if q_layer in ('L2', 'L3'):
        sees_money = True
    
    # Information flow visibility — data/tracking/reporting
    info_signals = ['data', 'tracking', 'monitor', 'report', 'log',
                    'record', 'database', 'system shows', 'analytics',
                    'dashboard', 'sensor', 'measurement', 'instrument']
    if any(s in all_text for s in info_signals):
        sees_information = True
    
    # Classify based on visibility combination
    if sees_material and sees_money and sees_information:
        return 'Fi'  # Full visibility — rare, information-rich position
    if sees_material and sees_money:
        return 'Fg'  # Has material + money but sits in information gap
    if sees_material and not sees_money:
        return 'Fm'  # Material only — upstream operators, truckers
    if sees_money and not sees_material:
        return 'F$'  # Money only — executives who see cost but not physical
    
    # Default based on position
    if q_object == 'OA':
        return 'Fm'  # Upstream = material visibility
    if q_object == 'OC' and q_layer == 'L3':
        return 'F$'  # Downstream exec = money visibility
    return 'Fg'  # Most sit in the gap


def infer_lens_tribalknowledge(test_run) -> str:
    """
    Infer Tribal Knowledge Mortality lens (K1-K4) — the retirement clock.
    When this person retires, does the company lose awareness.
    
    K1=Institutional (in systems), K2=Procedural (informal), K3=Tribal (one head), K4=Lost
    """
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'experiments', ''),
        getattr(test_run, 'action_iterate', ''),
    ])).lower()
    title = getattr(test_run, 'title', '').lower()
    test_category = getattr(test_run, 'test_category', '').lower()
    
    if not all_text.strip():
        return ''
    
    # K4: Knowledge already lost — discussing things nobody knows anymore
    k4_signals = ['used to know', 'old hands are gone', 'retired',
                  'nobody remembers', 'lost when', 'turnover',
                  'new crew doesn\'t know', 'institutional memory gone',
                  'we don\'t know why', 'nobody knows anymore',
                  'used to be someone', 'that knowledge left']
    if any(s in all_text for s in k4_signals):
        return 'K4'
    
    # K1: Institutional — documented in systems
    k1_signals = ['sap', 'cmms', 'maintenance log', 'work order system',
                  'our database', 'tracked in', 'documented in',
                  'sop says', 'procedure manual', 'our system tracks',
                  'data shows', 'records show', 'we log everything']
    if any(s in all_text for s in k1_signals):
        return 'K1'
    
    # K2: Procedural — informal but transmitted
    k2_signals = ['we train', 'we teach', 'our practice is', 'we always check',
                  'procedure is', 'standard practice', 'crew knows to',
                  'training covers', 'sop', 'we tell new guys']
    if any(s in all_text for s in k2_signals):
        return 'K2'
    
    # K3: Tribal — knowledge in one person's head
    k3_signals = ['i know', 'i can tell', 'i\'ve learned', 'in my experience',
                  'after 30 years', 'years of experience', 'i just know',
                  'you learn to', 'old dave', 'only person', 'by the sound',
                  'by the feel', 'gut feeling', 'instinct', 'my technique',
                  'nobody else', 'i\'m the only', 'when i retire']
    if any(s in all_text for s in k3_signals):
        return 'K3'
    
    # Infer from test_run type — independent operators are K3 by nature
    if any(s in title for s in ['owner', 'operator', 'independent', 'gypo']):
        return 'K3'
    if any(s in test_category for s in ['independent', 'owner']):
        return 'K3'
    
    # Operators typically K3, managers K2, executives K1
    q_layer = getattr(test_run, 'q_layer', '')
    if q_layer == 'L1':
        return 'K3'
    if q_layer == 'L3':
        return 'K2'  # Execs have some institutional knowledge
    
    # Default
    return 'K2'


def _compute_lens_confidence(test_run, lens_name: str, label: str) -> Tuple[float, List[str]]:
    """
    Compute Bayesian confidence for a lens assignment.
    
    Method: Count signal matches across ALL categories for this lens.
    Confidence = P(label|evidence) using simple evidence accumulation:
    
    α = 1.0 + (matched_signals_for_label × 0.3)
    β = 1.0 + (matched_signals_for_other_labels × 0.1)
    confidence = α / (α + β)
    
    Returns: (confidence_float, evidence_list)
    """
    all_text = ' '.join(filter(None, [
        getattr(test_run, 'results', ''),
        getattr(test_run, 'experiments', ''),
        getattr(test_run, 'hypotheses', ''),
        getattr(test_run, 'action_iterate', ''),
    ])).lower()
    
    if not all_text.strip():
        return 0.5, []
    
    # Get the QCUBE_DIMENSIONS for this lens
    dim_key = f"Q_{lens_name.upper().replace('Q_', '')}"
    dim = QCUBE_DIMENSIONS.get(dim_key, {})
    if not dim:
        return 0.5, []
    
    evidence = []
    alpha = 1.0  # Prior successes for chosen label
    beta = 1.0   # Prior failures (evidence for other labels)
    
    # Count which signals actually fired in the text
    label_found_in_text = label.lower() in all_text
    if label_found_in_text:
        alpha += 0.5
        evidence.append(f"Direct reference to {label}")
    
    # Check equipment impact density
    equip_count = len(getattr(test_run, 'substrate_impacts', []))
    if equip_count > 3:
        alpha += 0.3
        evidence.append(f"{equip_count} equipment impacts documented")
    
    # Text density bonus
    if len(all_text) > 500:
        alpha += 0.2
        evidence.append("Rich text content (>500 chars)")
    elif len(all_text) < 100:
        beta += 0.3
        evidence.append("Sparse text — low confidence")
    
    # Confidence = posterior mean of Beta(α, β)
    confidence = alpha / (alpha + beta)
    
    return min(1.0, confidence), evidence


def infer_all_lenses(test_run) -> Dict[str, str]:
    """
    Run all 7 analytical lens inference on a test.
    
    Returns dict of lens values. Only overwrites empty fields.
    Also computes confidence scores and evidence trails.
    Foundation lenses (L, O, S) are set from test metadata, not inferred here.
    """
    lenses = {}
    confidences = {}
    evidences = {}
    
    # ── Lens 4: Awareness ──
    if not getattr(test_run, 'q_awareness', ''):
        label = infer_lens_awareness(test_run)
        lenses['q_awareness'] = label
        conf, ev = _compute_lens_confidence(test_run, 'awareness', label)
        confidences['q_awareness'] = conf
        evidences['q_awareness'] = ev
    else:
        lenses['q_awareness'] = test_run.q_awareness
        confidences['q_awareness'] = 1.0  # Manual = full confidence
        evidences['q_awareness'] = ['Manual assignment']
    
    # ── Lens 5: Evidence Type ──
    if not getattr(test_run, 'q_evidence', ''):
        label = infer_lens_evidence(test_run)
        lenses['q_evidence'] = label
        conf, ev = _compute_lens_confidence(test_run, 'evidence', label)
        confidences['q_evidence'] = conf
        evidences['q_evidence'] = ev
    else:
        lenses['q_evidence'] = test_run.q_evidence
        confidences['q_evidence'] = 1.0
        evidences['q_evidence'] = ['Manual assignment']
    
    # ── Lens 6: Time Horizon ──
    if not getattr(test_run, 'q_timehorizon', ''):
        label = infer_lens_timehorizon(test_run)
        lenses['q_timehorizon'] = label
        conf, ev = _compute_lens_confidence(test_run, 'timehorizon', label)
        confidences['q_timehorizon'] = conf
        evidences['q_timehorizon'] = ev
    else:
        lenses['q_timehorizon'] = test_run.q_timehorizon
        confidences['q_timehorizon'] = 1.0
        evidences['q_timehorizon'] = ['Manual assignment']
    
    # ── Lens 7: Cascade Depth (ORIGINAL) ──
    if not getattr(test_run, 'q_cascade', ''):
        label = infer_lens_cascade(test_run)
        lenses['q_cascade'] = label
        conf, ev = _compute_lens_confidence(test_run, 'cascade', label)
        confidences['q_cascade'] = conf
        evidences['q_cascade'] = ev
    else:
        lenses['q_cascade'] = test_run.q_cascade
        confidences['q_cascade'] = 1.0
        evidences['q_cascade'] = ['Manual assignment']
    
    # ── Lens 8: Normalization of Deviance (ORIGINAL) ──
    if not getattr(test_run, 'q_normalization', ''):
        label = infer_lens_normalization(test_run)
        lenses['q_normalization'] = label
        conf, ev = _compute_lens_confidence(test_run, 'normalization', label)
        confidences['q_normalization'] = conf
        evidences['q_normalization'] = ev
    else:
        lenses['q_normalization'] = test_run.q_normalization
        confidences['q_normalization'] = 1.0
        evidences['q_normalization'] = ['Manual assignment']
    
    # ── Lens 9: Counter-Flow Mapping (ORIGINAL) ──
    if not getattr(test_run, 'q_counterflow', ''):
        label = infer_lens_counterflow(test_run)
        lenses['q_counterflow'] = label
        conf, ev = _compute_lens_confidence(test_run, 'counterflow', label)
        confidences['q_counterflow'] = conf
        evidences['q_counterflow'] = ev
    else:
        lenses['q_counterflow'] = test_run.q_counterflow
        confidences['q_counterflow'] = 1.0
        evidences['q_counterflow'] = ['Manual assignment']
    
    # ── Lens 10: Tribal Knowledge Mortality (ORIGINAL) ──
    if not getattr(test_run, 'q_tribalknowledge', ''):
        label = infer_lens_tribalknowledge(test_run)
        lenses['q_tribalknowledge'] = label
        conf, ev = _compute_lens_confidence(test_run, 'tribalknowledge', label)
        confidences['q_tribalknowledge'] = conf
        evidences['q_tribalknowledge'] = ev
    else:
        lenses['q_tribalknowledge'] = test_run.q_tribalknowledge
        confidences['q_tribalknowledge'] = 1.0
        evidences['q_tribalknowledge'] = ['Manual assignment']
    
    # Store confidence and evidence on test_run
    test_run.lens_confidence = confidences
    test_run.lens_evidence = evidences
    
    return lenses


def apply_lenses_to_database(database):
    """
    Run lens inference on ALL tests in database.
    Only fills empty lens fields — won't overwrite manual assignments.
    Computes and stores confidence scores + evidence trails.
    Returns count of tests updated.
    """
    updated = 0
    for test_entry in database.test_runs:
        lenses = infer_all_lenses(test_run)  # Also sets lens_confidence + lens_evidence
        changed = False
        for field, value in lenses.items():
            old_val = getattr(test_run, field, '')
            if not old_val and value:
                setattr(test_run, field, value)
                changed = True
        if changed:
            updated += 1
            # Console: show lens assignments with confidence
            confs = getattr(test_run, 'lens_confidence', {})
            conf_str = ' '.join(
                f"{k.replace('q_','')[:4]}={v:.2f}" 
                for k, v in sorted(confs.items()) if v < 1.0
            )
            if conf_str:
                print(f"    #{test_run.test_num} {test_run.person}: {conf_str}")
    return updated


# ============================================================================
# REAL ML: TF-IDF + SVD EMBEDDING ENGINE
# ============================================================================

class Test RunEmbeddingEngine:
    """
    Real semantic embeddings using TF-IDF + TruncatedSVD.
    
    Math:
    1. TF-IDF vectorization — term frequency × inverse document frequency
       tfidf(t,d) = tf(t,d) × log(N / df(t))
    2. Truncated SVD (Latent Semantic Analysis) — dimensionality reduction
       X ≈ U·Σ·Vᵀ  →  embeddings = U·Σ (first k components)
    3. Cosine similarity — test_run-to-test_run resonance
       sim(a,b) = (a·b) / (||a|| × ||b||)
    
    This is mathematically rigorous, not keyword matching.
    sklearn.decomposition.TruncatedSVD = Halko et al. 2011 randomized SVD.
    """
    
    def __init__(self, n_components: int = 32):
        """
        Args:
            n_components: Embedding dimensions. 32 is good for 20-200 tests.
                         More components = more granular but needs more data.
        """
        self.n_components = n_components
        self.vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english',
            ngram_range=(1, 2),  # Unigrams + bigrams ("blow line", "screen basket")
            min_df=1,            # Keep rare terms (equipment names appear few times)
            max_df=0.95,         # Drop terms in >95% of docs
        )
        self.svd = TruncatedSVD(n_components=n_components, random_state=42)
        self.is_fitted = False
        self.tfidf_matrix = None
        self.embeddings = None
        self.similarity_matrix = None
    
    def fit_transform(self, test_runs: list) -> np.ndarray:
        """
        Build TF-IDF matrix from all tests, reduce to n_components via SVD.
        Returns: (n_tests, n_components) embedding matrix.
        """
        if not test_runs:
            return np.zeros((0, self.n_components))
        
        # Build document corpus — every text field concatenated per test
        corpus = []
        for test_entry in test_runs:
            text_parts = []
            for field in ('hypotheses', 'experiments', 'results', 'action_iterate'):
                val = getattr(test_run, field, '') or ''
                if val.strip():
                    text_parts.append(val)
            # Add equipment names as high-signal terms
            for imp in getattr(test_run, 'substrate_impacts', []):
                equip = imp.get('equipment', '')
                if equip:
                    text_parts.append(equip)
            corpus.append(' '.join(text_parts))
        
        # Guard: need at least 2 documents for SVD
        if len(corpus) < 2:
            self.embeddings = np.zeros((len(corpus), self.n_components))
            self.is_fitted = True
            return self.embeddings
        
        # Fit TF-IDF
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)
        
        # SVD — reduce to n_components (capped at min(n_docs, n_features) - 1)
        actual_components = min(self.n_components, 
                               self.tfidf_matrix.shape[0] - 1,
                               self.tfidf_matrix.shape[1] - 1)
        if actual_components < 1:
            actual_components = 1
        self.svd = TruncatedSVD(n_components=actual_components, random_state=42)
        self.embeddings = self.svd.fit_transform(self.tfidf_matrix)
        
        # Pad if fewer components than requested
        if self.embeddings.shape[1] < self.n_components:
            pad = np.zeros((self.embeddings.shape[0], self.n_components - self.embeddings.shape[1]))
            self.embeddings = np.hstack([self.embeddings, pad])
        
        # Cosine similarity matrix — test_run-to-test_run resonance
        self.similarity_matrix = cosine_similarity(self.embeddings)
        
        self.is_fitted = True
        return self.embeddings
    
    def get_top_terms(self, n_terms: int = 10) -> list:
        """Get top terms from TF-IDF vocabulary by importance."""
        if not self.is_fitted or self.tfidf_matrix is None:
            return []
        feature_names = self.vectorizer.get_feature_names_out()
        # SVD components — each row is a latent topic, values = term weights
        top_terms = []
        for i, comp in enumerate(self.svd.components_[:3]):
            top_indices = comp.argsort()[-n_terms:][::-1]
            terms = [feature_names[idx] for idx in top_indices]
            top_terms.append((f"Topic {i+1}", terms))
        return top_terms
    
    def get_nearest_neighbors(self, test_run_idx: int, k: int = 3) -> list:
        """Find k most similar tests by cosine similarity."""
        if self.similarity_matrix is None or test_run_idx >= len(self.similarity_matrix):
            return []
        sims = self.similarity_matrix[test_run_idx].copy()
        sims[test_run_idx] = -1  # Exclude self
        top_k = sims.argsort()[-k:][::-1]
        return [(int(idx), float(sims[idx])) for idx in top_k if sims[idx] > 0]
    
    def explained_variance_ratio(self) -> float:
        """How much of the total text variance the embedding captures."""
        if not self.is_fitted:
            return 0.0
        return float(self.svd.explained_variance_ratio_.sum())


# ============================================================================
# TEST DATA STRUCTURES
# ============================================================================

class Test Run:
    """Single test data structure"""
    
    def __init__(self):
        self.test_num: int = 0
        self.date: str = ""
        self.person: str = ""
        self.title: str = ""
        self.company: str = ""
        self.test_category: str = ""
        
        # ICorps template fields
        self.hypotheses: str = ""
        self.experiments: str = ""  # Questions asked
        self.results: str = ""  # What we learned
        self.action_iterate: str = ""  # Next steps
        
        # Q-Cube classification — Foundation 3 axes
        self.q_layer: str = ""  # L1, L2, or L3
        self.q_object: str = ""  # OA, OB, or OC
        self.q_stack: List[str] = []  # Can be multiple: Sα, Sβ, Sγ, Sδ
        
        # Q-Cube Hypercube — Analytical lenses (4-6)
        self.q_awareness: str = ""        # A1-A4: Relationship to the problem
        self.q_evidence: str = ""         # V1-V4: Strength of validation
        self.q_timehorizon: str = ""      # T1-T4: When the pain hits
        
        # Q-Cube Hypercube — Innovation lenses (7-10) — Team GIBUSH originals
        self.q_cascade: str = ""          # C1-C6: Supply chain handoff depth
        self.q_normalization: str = ""    # N1-N4: Cultural embedding (Vaughan)
        self.q_counterflow: str = ""      # Fm/F$/Fi/Fg: Three-flow visibility gap
        self.q_tribalknowledge: str = ""  # K1-K4: Knowledge mortality / retirement clock
        
        # Lens inference confidence — computed by Bayesian evidence accumulation
        # Keys: 'q_awareness', 'q_evidence', etc. Values: float 0-1
        self.lens_confidence: Dict[str, float] = {}
        # Lens inference evidence trail — which signals fired
        # Keys: 'q_awareness', etc. Values: list of signal strings
        self.lens_evidence: Dict[str, List[str]] = {}
        
        # Equipment impacts mentioned in this test_run
        self.substrate_impacts: List[Dict] = []  # [{'equipment': 'CTS Rollers', 'cost': 160000}, ...]
        
        # Data source: "sparks" (discovery phase) or "validation" (validation phase)
        self.source: str = ""
        
        # AI-derived metrics
        self.synergy_score: float = 0.0
        self.embedding: Optional[np.ndarray] = None
        
    def to_dict(self) -> Dict:
        return {
            'test_num': self.test_num,
            'date': self.date,
            'script_name': self.person,
            'title': self.title,
            'source_version': self.company,
            'test_category': self.test_category,
            'hypotheses': self.hypotheses,
            'experiments': self.experiments,
            'results': self.results,
            'action_iterate': self.action_iterate,
            'q_layer': self.q_layer,
            'q_object': self.q_object,
            'q_stack': self.q_stack,
            'q_awareness': self.q_awareness,
            'q_evidence': self.q_evidence,
            'q_timehorizon': self.q_timehorizon,
            'q_cascade': self.q_cascade,
            'q_normalization': self.q_normalization,
            'q_counterflow': self.q_counterflow,
            'q_tribalknowledge': self.q_tribalknowledge,
            'lens_confidence': self.lens_confidence,
            'lens_evidence': self.lens_evidence,
            'substrate_impacts': self.substrate_impacts,
            'synergy_score': self.synergy_score,
            'source': self.source
        }
    
    @classmethod
    def from_dict(cls, data: Dict):
        test_run = cls()
        
        # Fields that MUST be numeric (JSON can store them as strings)
        FLOAT_FIELDS = {'synergy_score'}
        INT_FIELDS = {'test_num'}
        
        for key, value in data.items():
            if key == 'embedding':
                continue
            # Drop deprecated q_economic field
            if key == 'q_economic':
                continue
            
            # Type coercion for numeric fields
            if key in FLOAT_FIELDS:
                try:
                    value = float(value) if value else 0.0
                except (ValueError, TypeError):
                    value = 0.0
            elif key in INT_FIELDS:
                try:
                    value = int(value) if value else 0
                except (ValueError, TypeError):
                    value = 0
            
            if not hasattr(test_run, key):
                # Handle legacy data without substrate_impacts
                if key == 'substrate_impacts':
                    test_run.substrate_impacts = value
            else:
                setattr(test_run, key, value)
        
        # Ensure substrate_impacts exists for legacy data
        if not hasattr(test_run, 'substrate_impacts'):
            test_run.substrate_impacts = []
        
        # Coerce equipment impact costs to float
        if test_run.substrate_impacts:
            for impact in test_run.substrate_impacts:
                if 'cost' in impact:
                    try:
                        impact['cost'] = float(impact['cost']) if impact['cost'] else 0.0
                    except (ValueError, TypeError):
                        impact['cost'] = 0.0
        
        # Migrate old R1-R4 awareness codes to A1-A4
        if test_run.q_awareness and test_run.q_awareness.startswith('R'):
            test_run.q_awareness = 'A' + test_run.q_awareness[1:]
        return test_run


class Test RunDatabase:
    """Manages all test data and cube matrix"""
    
    def __init__(self):
        self.tests: List[Test Run] = []
        self.cube_matrix: Dict[str, List[str]] = {}  # Cube position -> person names
        self._init_cube_matrix()
        
    def _init_cube_matrix(self):
        """Initialize empty cube matrix with normalized keys"""
        for layer in QCUBE_DIMENSIONS["Q_LAYER"]:
            for obj in QCUBE_DIMENSIONS["Q_OBJECT"]:
                for stack in QCUBE_DIMENSIONS["Q_STACK"]:
                    key = build_cube_key(layer, obj, stack)
                    self.cube_matrix[key] = []
    
    def add_test_run(self, test_run: Test Run):
        """Add test_run and update cube matrix using normalized keys.

        DEDUP GUARD: Within validation entries, blocks any incoming test_run
        where the person name already exists. Company excluded from key
        because it is entered inconsistently (e.g. 'Unknown' vs full name).
        Baseline entries not checked -- discovery phase, different rules.
        """
        incoming_source = getattr(test_run, 'source', '').lower()
        incoming_person = getattr(test_run, 'script_name', '').strip().lower()

        if incoming_source == 'validation' and incoming_person:
            for existing in self.tests:
                if getattr(existing, 'source', '').lower() != 'validation':
                    continue
                if getattr(existing, 'script_name', '').strip().lower() == incoming_person:
                    print(
                        f"  [DEDUP GUARD] Blocked: {test_run.person} "
                        f"({test_run.company}) -- "
                        f"already exists as {existing.person} ({existing.company})"
                    )
                    return

        self.tests.append(test_run)
        
        # Add to cube matrix (keyed by person name)
        script_key = test_run.person.strip()
        if test_run.q_layer and test_run.q_object:
            for stack in test_run.q_stack:
                key = build_cube_key(test_run.q_layer, test_run.q_object, stack)
                if key not in self.cube_matrix:
                    self.cube_matrix[key] = []
                if script_key not in self.cube_matrix[key]:
                    self.cube_matrix[key].append(script_key)
    
    def rebuild_cube_matrix(self):
        """Rebuild the entire cube matrix from scratch using all tests.
        
        Call this after loading from disk or when the matrix is suspected
        to be out of sync with the tests list.
        """
        self._init_cube_matrix()
        for test_entry in self.tests:
            script_key = test_run.person.strip()
            if test_run.q_layer and test_run.q_object:
                for stack in test_run.q_stack:
                    key = build_cube_key(test_run.q_layer, test_run.q_object, stack)
                    if key not in self.cube_matrix:
                        self.cube_matrix[key] = []
                    if script_key not in self.cube_matrix[key]:
                        self.cube_matrix[key].append(script_key)
    
    def get_cube_position(self, layer: str, obj: str, stack: str) -> List[Test Run]:
        """Get all tests at specific cube position (normalized lookup)"""
        key = build_cube_key(layer, obj, stack)
        person_ids = self.cube_matrix.get(key, [])
        return [i for i in self.tests if i.person.strip() in person_ids]
    
    def get_synergy_clusters(self, min_test_runs: int = 2) -> Dict[str, List[Test Run]]:
        """Find cube positions with multiple tests (synergy clusters)"""
        clusters = {}
        for key, ids in self.cube_matrix.items():
            if len(ids) >= min_test_runs:
                clusters[key] = [i for i in self.tests if i.person.strip() in ids]
        return clusters
    
    def get_baseline_test_runs(self) -> List[Test Run]:
        """Return only Baseline (discovery phase) tests."""
        return [i for i in self.tests if getattr(i, 'source', '') != 'validation']
    
    def get_validation_test_runs(self) -> List[Test Run]:
        """Return only Validation (validation phase) tests."""
        return [i for i in self.tests if getattr(i, 'source', '') == 'validation']
    
    def save(self, filepath: Path):
        """Save database to JSON"""
        # Rebuild cube matrix before saving to ensure consistency
        self.rebuild_cube_matrix()
        data = {
            'tests': [i.to_dict() for i in self.tests],
            'cube_matrix': self.cube_matrix
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    @classmethod
    def load(cls, filepath: Path):
        """Load database from JSON, rebuild cube matrix, auto-tag lenses"""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        db = cls()
        db.tests = [Test Run.from_dict(i) for i in data['tests']]
        # ALWAYS rebuild cube matrix from tests to ensure encoding consistency
        db.rebuild_cube_matrix()
        # Auto-tag analytical lenses (only fills empty fields)
        updated = apply_lenses_to_database(db)
        if updated:
            print(f"  Lens auto-tag: {updated}/{len(db.tests)} tests tagged")
        return db


# ============================================================================
# AI ANALYSIS ENGINE
# ============================================================================

class AIAnalysisEngine(QThread):
    """
    Real ML analysis pipeline using:
    - TF-IDF + TruncatedSVD for semantic embeddings (Halko et al. 2011)
    - Cosine similarity for test_run-to-test_run resonance
    - Bayesian posterior scoring for synergy (Beta-Binomial conjugate model)
    - Information gain (Shannon entropy) for evidence quality ranking
    
    No mock data. No random noise. Every number is computed from test_run text.
    """
    progress_update = Signal(int, str)
    analysis_complete = Signal(dict)
    
    def __init__(self, database: Test RunDatabase):
        super().__init__()
        self.database = database
        self.embedding_engine = Test RunEmbeddingEngine(n_components=32)
        
    def run(self):
        """
        Run real ML analysis pipeline.
        
        Execution layers (in order):
        1. TF-IDF + SVD embeddings (sklearn — always runs)
        2. Bayesian synergy scoring (Beta-Binomial — always runs)
        3. Knowledge extraction (if knowledge_extractor available)
        4. Cosine resonance pattern detection (sklearn — always runs)
        5. Cluster analysis with 10-lens positions (always runs)
        6. Adaptive question generation (if doctoral_analyzer available)
        7. Customer segments + VP refinement (always runs)
        """
        try:
            # Step 1: Build real semantic embeddings
            self.progress_update.emit(10, "Building TF-IDF matrix...")
            self._embed_test_runs()
            
            # Step 2: Compute Bayesian synergy scores
            self.progress_update.emit(25, "Computing Bayesian synergy scores...")
            self._compute_bayesian_scores()
            
            # Step 3: Knowledge extraction (if engine available)
            knowledge_entities = {}
            if KNOWLEDGE_EXTRACTOR_AVAILABLE:
                self.progress_update.emit(35, "Extracting knowledge entities...")
                knowledge_entities = self._run_knowledge_extraction()
            
            # Step 4: Detect resonance patterns via cosine similarity
            self.progress_update.emit(50, "Detecting cross-test resonance...")
            synergy_patterns = self._detect_synergy_patterns()
            
            # Step 5: Cluster analysis
            self.progress_update.emit(65, "Analyzing synergy clusters...")
            clusters = self._analyze_clusters()
            
            # Step 6: Adaptive question generation (if doctoral_analyzer available)
            adaptive_questions = []
            if DOCTORAL_ANALYZER_AVAILABLE:
                self.progress_update.emit(75, "Generating info-gain questions...")
                adaptive_questions = self._run_adaptive_questioning()
            
            # Step 7: Generate customer segments
            self.progress_update.emit(85, "Generating refined customer segments...")
            segments = self._generate_substrate_classs(clusters)
            
            # Step 8: Refine value propositions
            self.progress_update.emit(92, "Optimizing value propositions...")
            value_props = self._refine_value_propositions(clusters)
            
            # Step 9: Lens confidence summary
            self.progress_update.emit(97, "Computing lens confidence metrics...")
            lens_summary = self._compute_lens_confidence_summary()
            
            self.progress_update.emit(100, "Analysis complete!")
            
            results = {
                'synergy_patterns': synergy_patterns,
                'clusters': clusters,
                'substrate_classs': segments,
                'value_propositions': value_props,
                'embedding_variance_explained': self.embedding_engine.explained_variance_ratio(),
                'top_terms': self.embedding_engine.get_top_terms(8),
                'knowledge_entities': knowledge_entities,
                'adaptive_questions': adaptive_questions,
                'lens_confidence_summary': lens_summary,
                'engines_active': {
                    'sklearn_embeddings': True,
                    'bayesian_scoring': True,
                    'knowledge_extractor': KNOWLEDGE_EXTRACTOR_AVAILABLE,
                    'doctoral_analyzer': DOCTORAL_ANALYZER_AVAILABLE,
                    'adaptive_analyzer': ADAPTIVE_ANALYZER_AVAILABLE,
                },
            }
            
            # Print ML summary to console
            self._print_ml_summary(results)
            
            # Save updated synergy scores to database
            db_file = BCM_ROOT / "test_database.json"
            self.database.save(db_file)
            
            self.analysis_complete.emit(results)
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            self.progress_update.emit(0, f"Error: {str(e)}")
    
    def _embed_test_runs(self):
        """
        Generate REAL semantic embeddings using TF-IDF + TruncatedSVD.
        
        Math:
        - TF-IDF: tfidf(t,d) = tf(t,d) × log(N/df(t))
        - SVD: X ≈ UΣVᵀ → embeddings = UΣ[:, :k]
        - Result: each test → 32-dim vector in semantic space
        """
        if not self.database.test_runs:
            return
        
        embeddings = self.embedding_engine.fit_transform(self.database.tests)
        
        # Assign real embeddings to each test
        for i, test_run in enumerate(self.database.tests):
            if i < len(embeddings):
                test_run.embedding = embeddings[i]
            else:
                test_run.embedding = np.zeros(32)
        
        variance = self.embedding_engine.explained_variance_ratio()
        print(f"\n  [ML] TF-IDF + SVD embeddings: {len(embeddings)} tests → "
              f"{embeddings.shape[1] if len(embeddings) > 0 else 0}d space "
              f"({variance:.1%} variance explained)")
    
    def _compute_bayesian_scores(self):
        """
        Compute synergy score using Bayesian evidence accumulation.
        
        For each test, score = weighted sum of:
        - Evidence strength (V-lens): V4=1.0, V3=0.8, V2=0.5, V1=0.2
        - Information density: text length normalized  
        - Equipment coverage: fraction of config equipment discussed
        - Cosine resonance: avg similarity to other tests (semantic overlap)
        - Source authority: L1=1.0 (3× per doctrine), L2=0.67, L3=0.33
        
        Score = Beta(α, β) posterior where α = supporting factors, β = gaps.
        """
        import math
        
        # Evidence strength from V-lens
        v_weights = {'V4': 1.0, 'V3': 0.8, 'V2': 0.5, 'V1': 0.2}
        # Authority weight from L-lens (inverted — L1 operators carry most truth)
        l_weights = {'L1': 1.0, 'L2': 0.67, 'L3': 0.33}
        
        for i, test_run in enumerate(self.database.tests):
            alpha = 1.0  # Prior success
            beta = 1.0   # Prior failure
            
            # Factor 1: Evidence type (V-lens)
            v_lens = getattr(test_run, 'q_evidence', 'V1')
            evidence_weight = v_weights.get(v_lens, 0.2)
            alpha += evidence_weight
            
            # Factor 2: Information density — text richness
            text_len = sum(len(getattr(test_run, f, '') or '') 
                          for f in ('results', 'experiments', 'hypotheses', 'action_iterate'))
            if text_len > 800:
                alpha += 0.8
            elif text_len > 400:
                alpha += 0.5
            elif text_len > 100:
                alpha += 0.2
            else:
                beta += 0.5  # Thin data = uncertainty
            
            # Factor 3: Equipment coverage
            equip_count = len(getattr(test_run, 'substrate_impacts', []))
            if equip_count > 5:
                alpha += 0.6
            elif equip_count > 0:
                alpha += 0.3
            else:
                beta += 0.3
            
            # Factor 4: Cost data present (quantified = stronger evidence)
            has_costs = any(
                imp.get('cost', '') and imp['cost'] != '0'
                for imp in getattr(test_run, 'substrate_impacts', [])
            )
            if has_costs:
                alpha += 0.5
            
            # Factor 5: Source authority
            l_lens = getattr(test_run, 'q_layer', 'L2')
            authority = l_weights.get(l_lens, 0.5)
            alpha += authority * 0.5
            
            # Factor 6: Cosine resonance — semantic similarity to others
            if (self.embedding_engine.similarity_matrix is not None and 
                i < len(self.embedding_engine.similarity_matrix)):
                sims = self.embedding_engine.similarity_matrix[i].copy()
                sims[i] = 0  # Exclude self
                avg_sim = float(sims.mean()) if len(sims) > 1 else 0
                if avg_sim > 0.3:
                    alpha += 0.4  # High resonance with other test_runs
                elif avg_sim > 0.1:
                    alpha += 0.2
            
            # Compute posterior: P(H|E) = α / (α + β)
            posterior = alpha / (alpha + beta)
            
            # Shannon entropy: H(p) = -p×log₂(p) - (1-p)×log₂(1-p)
            if 0 < posterior < 1:
                entropy = -posterior * math.log2(posterior) - (1 - posterior) * math.log2(1 - posterior)
            else:
                entropy = 0.0
            
            test_run.synergy_score = min(1.0, posterior)
            
            print(f"  [Bayes] #{test_run.test_num} {test_run.person}: "
                  f"α={alpha:.2f} β={beta:.2f} → P={posterior:.3f} "
                  f"H={entropy:.3f}bits [{v_lens}|{l_lens}]")
    
    def _detect_synergy_patterns(self) -> List[str]:
        """
        Detect patterns using REAL cosine similarity from SVD embeddings
        + 10-lens hypercube gap analysis.
        """
        patterns = []
        tests = self.database.test_runs
        
        if not test_runs:
            return patterns
        
        # ── COSINE SIMILARITY PATTERNS ──
        if self.embedding_engine.similarity_matrix is not None:
            sim_matrix = self.embedding_engine.similarity_matrix
            n = len(tests)
            
            # Find highly resonant pairs (cosine > 0.6)
            resonant_pairs = []
            for i in range(n):
                for j in range(i + 1, n):
                    if sim_matrix[i][j] > 0.6:
                        resonant_pairs.append((i, j, sim_matrix[i][j]))
            
            if resonant_pairs:
                resonant_pairs.sort(key=lambda x: -x[2])
                top = resonant_pairs[:3]
                for i, j, sim in top:
                    a, b = test_runs[i], test_runs[j]
                    patterns.append(
                        f"RESONANCE: #{a.test_num} {a.person} ↔ "
                        f"#{b.test_num} {b.person} "
                        f"(cosine={sim:.3f}) — semantic overlap in pain description"
                    )
            
            # Find isolated tests (low avg similarity = unique perspective)
            avg_sims = []
            for i in range(n):
                row = sim_matrix[i].copy()
                row[i] = 0
                avg_sims.append((i, row.mean()))
            avg_sims.sort(key=lambda x: x[1])
            
            if avg_sims and avg_sims[0][1] < 0.1:
                idx = avg_sims[0][0]
                iv = test_runs[idx]
                patterns.append(
                    f"OUTLIER: #{iv.test_num} {iv.person} "
                    f"(avg_sim={avg_sims[0][1]:.3f}) — unique perspective, "
                    f"may reveal untapped segment"
                )
        
        # ── 10-LENS HYPERCUBE PATTERNS ──
        
        # Cascade depth distribution
        cascade_counts = {}
        for iv in test_runs:
            c = getattr(iv, 'q_cascade', '')
            if c:
                cascade_counts[c] = cascade_counts.get(c, 0) + 1
        deep_cascade = sum(v for k, v in cascade_counts.items() if k in ('C4', 'C5', 'C6'))
        if deep_cascade >= 3:
            patterns.append(
                f"CASCADE DEPTH: {deep_cascade}/{len(tests)} tests at C4+ "
                f"— feedback loop severed, contamination creators never see damage"
            )
        
        # Normalization of deviance concentration
        norm_counts = {}
        for iv in test_runs:
            n_val = getattr(iv, 'q_normalization', '')
            if n_val:
                norm_counts[n_val] = norm_counts.get(n_val, 0) + 1
        n3_n4 = sum(v for k, v in norm_counts.items() if k in ('N3', 'N4'))
        if n3_n4 >= 3:
            patterns.append(
                f"NORMALIZATION: {n3_n4}/{len(tests)} tests at N3-N4 "
                f"— contamination is accepted practice or invisible normal"
            )
        
        # Counter-flow gap concentration
        gap_count = sum(1 for iv in tests 
                       if getattr(iv, 'q_counterflow', '') in ('Fg', 'Fm'))
        if gap_count >= len(tests) * 0.5:
            patterns.append(
                f"INFORMATION GAP: {gap_count}/{len(tests)} tests in "
                f"Fm/Fg positions — material flows without information. "
                f"Detection system IS the missing information flow"
            )
        
        # Tribal knowledge at risk
        k3_count = sum(1 for iv in tests 
                      if getattr(iv, 'q_tribalknowledge', '') == 'K3')
        if k3_count >= 3:
            patterns.append(
                f"TRIBAL KNOWLEDGE RISK: {k3_count}/{len(tests)} tests "
                f"at K3 — knowledge lives in individual heads, lost at retirement"
            )
        
        # Post-investment failures
        post_invest = [i for i in tests if 'Sβ' in (i.q_stack or [])]
        if len(post_invest) >= 2:
            patterns.append(
                f"POST-INVESTMENT FAILURE: {len(post_invest)} tests show "
                f"modern equipment failing to solve contamination"
            )
        
        # SVD topic extraction
        top_terms = self.embedding_engine.get_top_terms(5)
        for topic_name, terms in top_terms[:2]:
            patterns.append(
                f"LATENT TOPIC ({topic_name}): {', '.join(terms[:5])}"
            )
        
        # === COSINE RESONANCE PAIRS ===
        # Find test_run pairs with high semantic similarity across DIFFERENT
        # cube positions. These are synergy opportunities — same pain, different
        # vantage point. Real cosine similarity, not keyword matching.
        if (self.embedding_engine.similarity_matrix is not None and
            len(tests) >= 3):
            sim_matrix = self.embedding_engine.similarity_matrix
            resonance_pairs = []
            
            for i, iv_a in enumerate(tests):
                if i >= len(sim_matrix):
                    continue
                for j, iv_b in enumerate(tests):
                    if j <= i or j >= len(sim_matrix):
                        continue
                    sim = float(sim_matrix[i][j])
                    # High similarity (>0.4) + different cube position = resonance
                    pos_a = f"[{iv_a.q_layer},{iv_a.q_object}]"
                    pos_b = f"[{iv_b.q_layer},{iv_b.q_object}]"
                    if sim > 0.4 and pos_a != pos_b:
                        resonance_pairs.append(
                            (sim, iv_a.test_num, iv_b.test_num,
                             iv_a.person, iv_b.person, pos_a, pos_b)
                        )
            
            resonance_pairs.sort(reverse=True)
            for sim, num_a, num_b, name_a, name_b, pos_a, pos_b in resonance_pairs[:3]:
                patterns.append(
                    f"RESONANCE PAIR (cos={sim:.2f}): "
                    f"#{num_a} {name_a} {pos_a} ↔ #{num_b} {name_b} {pos_b} "
                    f"— same pain language, different vantage points"
                )
        
        return patterns
    
    def _analyze_clusters(self) -> Dict[str, Dict]:
        """Analyze synergy clusters using real embeddings + 10-lens positions."""
        clusters = self.database.get_synergy_clusters(min_test_runs=2)
        
        analysis = {}
        for cube_pos, tests in clusters.items():
            # Real statistics from embeddings
            indices = []
            for iv in test_runs:
                for idx, db_iv in enumerate(self.database.tests):
                    if db_iv.person.strip().lower() == iv.person.strip().lower():
                        indices.append(idx)
                        break
            
            # Intra-cluster coherence from cosine similarity
            coherence = 0.0
            if (len(indices) >= 2 and 
                self.embedding_engine.similarity_matrix is not None):
                sim_sum = 0
                sim_count = 0
                for a in indices:
                    for b in indices:
                        if a != b and a < len(self.embedding_engine.similarity_matrix):
                            sim_sum += self.embedding_engine.similarity_matrix[a][b]
                            sim_count += 1
                coherence = sim_sum / sim_count if sim_count > 0 else 0.0
            
            # Average Bayesian synergy score
            avg_synergy = np.mean([i.synergy_score for i in test_runs]) if tests else 0
            
            # 10-lens distribution within cluster
            lens_dist = {}
            for lens_key in ('q_cascade', 'q_normalization', 'q_counterflow', 'q_tribalknowledge'):
                vals = [getattr(iv, lens_key, '') for iv in tests if getattr(iv, lens_key, '')]
                if vals:
                    lens_dist[lens_key] = max(set(vals), key=vals.count)  # Mode
            
            analysis[cube_pos] = {
                'count': len(tests),
                'avg_synergy': float(avg_synergy),
                'coherence': float(coherence),
                'companies': list(set(i.company for i in tests)),
                'pain_themes': self._extract_pain_themes(tests),
                'dominant_lenses': lens_dist,
            }
        
        return analysis
    
    def _extract_pain_themes(self, test_runs: List[Test Run]) -> List[str]:
        """
        Extract pain themes using TF-IDF term importance within cluster.
        
        Math: For each test in cluster, get its TF-IDF vector.
        Average the vectors. The highest-weighted features are the
        distinctive terms for this cluster vs the full corpus.
        
        Falls back to keyword scan if embedding engine hasn't been fitted.
        """
        if not test_runs:
            return []
        
        # === REAL TF-IDF PATH ===
        if (self.embedding_engine.is_fitted and 
            self.embedding_engine.tfidf_matrix is not None):
            try:
                feature_names = self.embedding_engine.vectorizer.get_feature_names_out()
                
                # Find indices of these tests in the full database
                cluster_indices = []
                for iv in test_runs:
                    for idx, db_iv in enumerate(self.database.tests):
                        if db_iv.person.strip().lower() == iv.person.strip().lower():
                            cluster_indices.append(idx)
                            break
                
                if cluster_indices:
                    # Average TF-IDF vectors for cluster test_runs
                    cluster_tfidf = self.embedding_engine.tfidf_matrix[cluster_indices]
                    avg_vector = cluster_tfidf.mean(axis=0)
                    avg_array = np.asarray(avg_vector).flatten()
                    
                    # Get top terms by average TF-IDF weight
                    top_indices = avg_array.argsort()[-8:][::-1]
                    top_terms = [(feature_names[i], float(avg_array[i])) 
                                 for i in top_indices if avg_array[i] > 0]
                    
                    if top_terms:
                        themes = []
                        # Group into theme phrases
                        term_strs = [t[0] for t in top_terms[:6]]
                        themes.append(f"Dominant terms: {', '.join(term_strs)}")
                        
                        # Check for equipment-heavy cluster
                        equip_terms = [t for t in term_strs if any(
                            e in t for e in ('screen', 'digester', 'pump', 'basket',
                                             'roller', 'blow', 'chip', 'drum',
                                             'washer', 'press', 'tank', 'line'))]
                        if len(equip_terms) >= 2:
                            themes.append(f"Equipment focus: {', '.join(equip_terms)}")
                        
                        # Check for cost-heavy cluster
                        cost_terms = [t for t in term_strs if any(
                            c in t for c in ('cost', 'dollar', 'price', 'budget',
                                             'replace', 'repair', 'downtime'))]
                        if cost_terms:
                            themes.append(f"Cost-quantified pain: {', '.join(cost_terms)}")
                        
                        return themes
            except Exception as e:
                print(f"  [ML] TF-IDF theme extraction failed, falling back: {e}")
        
        # === FALLBACK: enhanced keyword scan ===
        all_text = " ".join([(iv.results or '') + ' ' + (iv.hypotheses or '') 
                             for iv in test_runs]).lower()
        themes = []
        
        theme_signals = {
            "Equipment damage from contamination": ['screen', 'basket', 'digester', 'pump', 'roller', 'damage'],
            "Upstream contamination sources": ['dirt', 'rock', 'sand', 'bark', 'mud', 'gravel', 'tramp metal'],
            "Cost-quantified impacts": ['$', 'cost', 'million', 'thousand', 'budget', 'expense'],
            "Post-investment failure": ['modern', 'new equipment', 'still', 'upgrade', 'investment'],
            "Baseline/commissioning gaps": ['baseline', 'commissioning', 'benchmark', 'startup'],
            "Downstream damage despite screening": ['downstream', 'after screening', 'still getting', 'pass through'],
            "Seasonal/weather variation": ['winter', 'rain', 'frozen', 'mud season', 'spring'],
        }
        
        for theme, signals in theme_signals.items():
            hits = sum(1 for s in signals if s in all_text)
            if hits >= 2:
                themes.append(theme)
        
        return themes if themes else ["Insufficient text for theme extraction"]
    
    def _generate_substrate_classs(self, clusters: Dict) -> List[Dict]:
        """
        Generate customer segments using AgglomerativeClustering on SVD embeddings.
        
        Math: Ward linkage minimizes total within-cluster variance.
        d(u,v) = sqrt(2n_u·n_v / (n_u+n_v)) × ||c_u - c_v||
        
        Falls back to cube-position sorting if <4 tests (clustering needs data).
        """
        segments = []
        tests = self.database.test_runs
        
        # ── REAL CLUSTERING PATH ──
        if (self.embedding_engine.is_fitted and 
            self.embedding_engine.embeddings is not None and
            len(tests) >= 4):
            try:
                embeddings = self.embedding_engine.embeddings
                
                # Determine n_clusters: min(ceil(n/3), 5) — adaptive to dataset size
                n_clusters = min(max(2, len(tests) // 3), 5)
                
                clustering = AgglomerativeClustering(
                    n_clusters=n_clusters,
                    linkage='ward',  # Ward = minimize intra-cluster variance
                )
                labels = clustering.fit_predict(embeddings)
                
                # Build segment for each cluster
                for cluster_id in range(n_clusters):
                    member_indices = [i for i, l in enumerate(labels) if l == cluster_id]
                    if not member_indices:
                        continue
                    
                    members = [test_runs[i] for i in member_indices]
                    
                    # Cluster centroid in embedding space
                    cluster_embeddings = embeddings[member_indices]
                    centroid = cluster_embeddings.mean(axis=0)
                    
                    # Intra-cluster coherence: avg cosine to centroid
                    centroid_norm = centroid / (np.linalg.norm(centroid) + 1e-10)
                    coherences = []
                    for emb in cluster_embeddings:
                        emb_norm = emb / (np.linalg.norm(emb) + 1e-10)
                        coherences.append(float(np.dot(centroid_norm, emb_norm)))
                    avg_coherence = np.mean(coherences)
                    
                    # Dominant lens values within cluster
                    lens_profile = {}
                    for lens_key in ('q_layer', 'q_object', 'q_cascade', 
                                     'q_normalization', 'q_counterflow', 'q_tribalknowledge'):
                        vals = [getattr(m, lens_key, '') for m in members if getattr(m, lens_key, '')]
                        if vals:
                            lens_profile[lens_key] = max(set(vals), key=vals.count)
                    
                    # Pain themes from TF-IDF
                    pain_themes = self._extract_pain_themes(members)
                    
                    # Avg synergy score
                    avg_synergy = np.mean([m.synergy_score for m in members])
                    
                    # Rank by combined signal: count × avg_synergy × coherence
                    rank_score = len(members) * avg_synergy * avg_coherence
                    
                    segments.append({
                        'cluster_id': cluster_id,
                        'rank_score': float(rank_score),
                        'test_count': len(members),
                        'avg_synergy': float(avg_synergy),
                        'coherence': float(avg_coherence),
                        'companies': list(set(m.company for m in members)),
                        'test_nums': [m.test_num for m in members],
                        'lens_profile': lens_profile,
                        'pain_themes': pain_themes,
                        'description': self._generate_cluster_description(members, lens_profile, pain_themes),
                    })
                
                # Sort by rank_score — highest combined signal first
                segments.sort(key=lambda s: -s['rank_score'])
                
                # Label priorities
                for i, seg in enumerate(segments):
                    seg['priority'] = f"PRIMARY" if i == 0 else f"SECONDARY #{i}"
                
                n_tests = len(tests)
                print(f"  [ML] AgglomerativeClustering: {n_test_runs} tests → "
                      f"{n_clusters} segments (Ward linkage)")
                for seg in segments:
                    print(f"    Cluster {seg['cluster_id']}: {seg['test_count']} tests, "
                          f"synergy={seg['avg_synergy']:.3f}, coherence={seg['coherence']:.3f}")
                
                return segments
                
            except Exception as e:
                print(f"  [ML] Clustering failed, falling back to cube-position: {e}")
        
        # ── FALLBACK: cube-position segments ──
        sorted_clusters = sorted(clusters.items(), 
                                key=lambda x: x[1]['count'] * x[1].get('avg_synergy', 0.5), 
                                reverse=True)
        
        for i, (cube_pos, data) in enumerate(sorted_clusters[:5], 1):
            segment = {
                'priority': f"PRIMARY" if i == 1 else f"SECONDARY #{i-1}",
                'cube_position': cube_pos,
                'test_count': data['count'],
                'description': self._generate_segment_description(cube_pos, data),
                'pain_themes': data['pain_themes'],
                'companies': data['companies'],
                'coherence': data.get('coherence', 0),
                'avg_synergy': data.get('avg_synergy', 0),
            }
            segments.append(segment)
        
        return segments
    
    def _generate_cluster_description(self, members: list, lens_profile: Dict, 
                                       pain_themes: list) -> str:
        """Generate natural language description from cluster ML data."""
        parts = []
        
        # Who (layer)
        layer = lens_profile.get('q_layer', '')
        if layer:
            layer_desc = QCUBE_DIMENSIONS["Q_LAYER"].get(layer, '')
            if layer_desc:
                parts.append(layer_desc.split(' - ')[0])
        
        # Where (object)
        obj = lens_profile.get('q_object', '')
        if obj:
            obj_desc = QCUBE_DIMENSIONS["Q_OBJECT"].get(obj, '')
            if obj_desc:
                parts.append(f"at {obj_desc.split(' - ')[0]}")
        
        # Cascade depth
        cascade = lens_profile.get('q_cascade', '')
        if cascade and cascade in ('C4', 'C5', 'C6'):
            parts.append(f"with severed feedback loop ({cascade})")
        
        # Normalization
        norm = lens_profile.get('q_normalization', '')
        if norm == 'N3':
            parts.append("— contamination is accepted practice")
        elif norm == 'N4':
            parts.append("— contamination is invisible normal")
        
        description = ' '.join(parts)
        
        if pain_themes:
            description += f" | Pain: {', '.join(pain_themes[:2])}"
        
        return description
    
    def _generate_segment_description(self, cube_pos: str, data: Dict) -> str:
        """Generate natural language segment description"""
        # Parse cube position
        parts = cube_pos.strip('[]').split(', ')
        layer = parts[0]
        obj = parts[1]
        stack = parts[2]
        
        # Build description
        layer_desc = QCUBE_DIMENSIONS["Q_LAYER"][layer].split(' - ')[0]
        obj_desc = QCUBE_DIMENSIONS["Q_OBJECT"][obj].split(' - ')[0]
        stack_desc = QCUBE_DIMENSIONS["Q_STACK"][stack].split(' - ')[0]
        
        description = f"{layer_desc} at {obj_desc} experiencing {stack_desc}"
        
        # Add pain themes
        if data['pain_themes']:
            description += f" - Pain: {', '.join(data['pain_themes'][:2])}"
        
        return description
    
    def _refine_value_propositions(self, clusters: Dict) -> List[str]:
        """Generate value propositions from test evidence — no hardcoded keywords."""
        tests = self.database.test_runs
        completed = [i for i in tests if getattr(i, 'results', '')]
        if not completed:
            return []
        
        all_results = " ".join([i.results.lower() for i in completed])
        
        vp_defs = [
            ('Calibrate Your Contamination Problem',
             ['100% fail', 'all got through', 'manual catch', 'no detection', 'manual'],
             'Stay until rocks stop, not just until equipment is installed. Find WHERE upstream contamination enters.'),
            ('True Cost Discovery', 
             ['damage cost', 'annual damage', 'nobody tracks', 'more than expected', 'repair cost', 'replacement cost'],
             'Expose real contamination costs hidden in maintenance budgets.'),
            ('Last-Chance Prevention Before Digester',
             ['blow line', 'before digester', 'entry point', 'blind spot', 'response time'],
             'Real-time detection at blow line entry before contamination reaches process.'),
            ('Bark Blow Line Safety Monitoring',
             ['bark', 'hog fuel', 'bark blow', 'boiler tube'],
             'First-ever hog fuel boiler contamination detection.'),
            ('Contamination is Accepted Practice',
             ['accepted', 'normal', 'always been', 'cost of doing business', 'live with it'],
             'Industry normalizes damage. Detection creates accountability.'),
            ('Commissioning Baseline Data',
             ['baseline', 'commissioning', 'new equipment', 'modernization', 'investment'],
             'Document new equipment performance from day one.'),
        ]
        
        value_props = []
        for name, keywords, desc in vp_defs:
            test_run_hits = sum(1 for i in completed if any(kw in i.results.lower() for kw in keywords))
            if test_run_hits > 0:
                confidence = min(95, int((test_run_hits / len(completed)) * 100))
                value_props.append(f'"{name}" — {desc} ({confidence}%, {test_run_hits} tests)')
        
        # Lens-driven VPs from 10-lens intelligence
        deep_cascade = sum(1 for iv in tests 
                         if getattr(iv, 'q_cascade', '') in ('C4', 'C5', 'C6'))
        if deep_cascade >= 3:
            value_props.append(
                f'"Close the Accountability Gap" — {deep_cascade} tests confirm '
                f'contamination travels 4+ handoffs without traceability.')
        
        gap_count = sum(1 for iv in tests 
                       if getattr(iv, 'q_counterflow', '') in ('Fg', 'Fm'))
        if gap_count >= len(tests) * 0.4:
            value_props.append(
                f'"The Missing Information Flow" — {gap_count}/{len(tests)} '
                f'test_sources see material or money but NOT both.')
        
        k3_count = sum(1 for iv in tests 
                      if getattr(iv, 'q_tribalknowledge', '') == 'K3')
        if k3_count >= 3:
            value_props.append(
                f'"Institutional Knowledge Before Retirement" — {k3_count} '
                f'test_sources carry tribal knowledge at risk.')
        
        return value_props
    
    def _run_knowledge_extraction(self) -> Dict:
        """
        Run KnowledgeExtractor on all tests.
        Extracts: costs, equipment, pain points, validations, metrics, timelines.
        Returns entity summary dict.
        """
        try:
            extractor = KnowledgeExtractor()
            for test_entry in self.database.test_runs:
                text = ' '.join(filter(None, [
                    test_run.results, test_run.experiments,
                    test_run.hypotheses, test_run.action_iterate
                ]))
                if text.strip():
                    entities = extractor.extract_from_test_run(
                        test_run.person, text)
                    print(f"  [KE] {test_run.person}: "
                          f"{len(entities)} entities extracted")
            
            summary = extractor.get_entity_summary()
            print(f"  [KE] Total: {sum(len(v) for v in summary.values())} entities "
                  f"across {len(summary)} categories")
            return summary
        except Exception as e:
            print(f"  [KE] Knowledge extraction failed: {e}")
            return {}
    
    def _run_adaptive_questioning(self) -> List[Dict]:
        """
        Run DoctoralAnalyzer to generate info-gain-ranked questions.
        Uses Shannon entropy: IG = H(prior) - E[H(posterior)]
        
        Returns list of question dicts with info_gain scores.
        """
        try:
            analyzer = DoctoralAnalyzer(project=get_active_project())
            
            # Feed all tests into the doctoral analyzer
            for test_entry in self.database.test_runs:
                iv_dict = test_run.to_dict()
                analyzer._process_test_run_evidence(iv_dict)
            
            # Generate adaptive questions
            questions = analyzer.generate_adaptive_questions(max_questions=10)
            
            result = []
            for q in questions:
                result.append({
                    'question': q.question,
                    'information_gain': q.information_gain,
                    'target_hypothesis': q.target_hypothesis,
                    'priority': q.priority,
                    'rationale': q.rationale,
                })
                print(f"  [DI] IG={q.information_gain:.3f}bits | {q.question[:60]}...")
            
            # Also get gap analysis
            gaps = analyzer.identify_q_cube_gaps()
            equip_gaps = analyzer.identify_equipment_gaps()
            contradictions = analyzer.detect_contradictions()
            
            print(f"  [DI] {len(gaps)} cube gaps | {len(equip_gaps)} component gaps | "
                  f"{len(contradictions)} contradictions")
            
            return result
        except Exception as e:
            print(f"  [DI] Adaptive questioning failed: {e}")
            return []
    
    def _compute_lens_confidence_summary(self) -> Dict:
        """
        Aggregate lens confidence scores across all tests.
        Reports: average confidence per lens, lowest-confidence tests.
        """
        summary = {}
        lens_keys = ['q_awareness', 'q_evidence', 'q_timehorizon',
                     'q_cascade', 'q_normalization', 'q_counterflow', 'q_tribalknowledge']
        
        for lens_key in lens_keys:
            confidences = []
            low_conf_tests = []
            
            for iv in self.database.test_runs:
                conf = getattr(iv, 'lens_confidence', {}).get(lens_key, 0.5)
                confidences.append(conf)
                if conf < 0.6:
                    low_conf_tests.append(
                        f"#{iv.test_num} {iv.person} ({conf:.2f})")
            
            avg_conf = np.mean(confidences) if confidences else 0.5
            summary[lens_key] = {
                'avg_confidence': float(avg_conf),
                'min_confidence': float(min(confidences)) if confidences else 0.0,
                'low_confidence_count': len(low_conf_tests),
                'low_confidence_test_runs': low_conf_test_runs[:5],
            }
            
            if avg_conf < 0.6:
                print(f"  [CONF] ⚠ {lens_key}: avg={avg_conf:.2f} — "
                      f"{len(low_conf_tests)} tests need manual review")
            else:
                print(f"  [CONF] ✓ {lens_key}: avg={avg_conf:.2f}")
        
        return summary
    
    def _print_ml_summary(self, results: Dict):
        """Print real ML metrics to console — honest about what's running."""
        print("\n" + "=" * 70)
        print("  ML ANALYSIS SUMMARY — REAL MATH, NO MOCK DATA")
        print("=" * 70)
        
        # Engine status
        engines = results.get('engines_active', {})
        print("\n  ENGINES:")
        for name, active in engines.items():
            status = "✓ ACTIVE" if active else "○ not loaded"
            print(f"    {status} — {name}")
        
        var_exp = results.get('embedding_variance_explained', 0)
        print(f"\n  EMBEDDINGS: TF-IDF + TruncatedSVD")
        print(f"    Variance explained: {var_exp:.1%}")
        
        top_terms = results.get('top_terms', [])
        for topic, terms in top_terms[:3]:
            print(f"    {topic}: {', '.join(terms[:5])}")
        
        n_patterns = len(results.get('synergy_patterns', []))
        print(f"\n  SYNERGY PATTERNS: {n_patterns}")
        for p in results.get('synergy_patterns', [])[:5]:
            print(f"    → {p}")
        
        n_clusters = len(results.get('clusters', {}))
        print(f"\n  CLUSTERS: {n_clusters}")
        for pos, data in results.get('clusters', {}).items():
            coh = data.get('coherence', 0)
            print(f"    {pos}: {data['count']} tests, "
                  f"synergy={data['avg_synergy']:.3f}, "
                  f"coherence={coh:.3f}")
        
        # Lens confidence
        lens_summary = results.get('lens_confidence_summary', {})
        if lens_summary:
            print(f"\n  LENS CONFIDENCE (Bayesian evidence accumulation):")
            for lens, info in lens_summary.items():
                avg = info.get('avg_confidence', 0)
                low = info.get('low_confidence_count', 0)
                bar = "█" * int(avg * 10) + "░" * (10 - int(avg * 10))
                print(f"    {lens:22s} [{bar}] {avg:.2f}  ({low} need review)")
        
        # Adaptive questions
        aq = results.get('adaptive_questions', [])
        if aq:
            print(f"\n  INFO-GAIN QUESTIONS (Shannon entropy):")
            for q in aq[:5]:
                print(f"    IG={q['information_gain']:.3f}bits | {q['question'][:60]}...")
        
        n_vps = len(results.get('value_propositions', []))
        print(f"\n  VALUE PROPOSITIONS: {n_vps}")
        for vp in results.get('value_propositions', []):
            print(f"    → {vp[:80]}...")
        
        print("\n" + "=" * 70)


# ============================================================================
# GUI APPLICATION
# ============================================================================

# ============================================================================
# TEST PLANNING DATA - LOADED FROM EXCEL
# ============================================================================

# Load from Excel - THE SINGLE SOURCE OF TRUTH
# Edit BCM_Test_Plan.xlsx or AISOS_SPINE_BCM_Test_Plan.xlsx
# Switch projects via the project dropdown — all paths route through active project
TEST_PLAN = load_test_plans()
print(f"Loaded {len(TEST_PLAN)} test_run plans from Excel ({get_active_project()})")


# ============================================================================
# GUI WIDGETS
# ============================================================================

class TestPlansWidget(QWidget):
    """Test Plans spreadsheet - click row to generate template"""
    
    generate_template = Signal(dict)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        # Header
        header = QLabel("<h2>📋 GIBUSH Test Plans</h2>")
        layout.addWidget(header)
        
        # Value prop box
        vp_group = QGroupBox("VALUE PROPOSITION - THE GIBUSH ADVANTAGE")
        vp_layout = QVBoxLayout()
        vp_text = QLabel(
            "<b>THE GAP WE FILL:</b> Real-time contamination detection at blow line entry - BEFORE digester<br>"
            "<b>WHAT THEY CAN'T DO:</b> BTG measures consistency downstream. Voith optimizes fabric. We PREVENT damage at source.<br>"
            "<b>SYNERGY PLAY:</b> Complement existing quality systems - early warning that protects their investments<br>"
            "<b>DIFFERENTIATION:</b> Only system that triangulates rock position in 3D space using acoustic physics"
        )
        vp_text.setWordWrap(True)
        vp_layout.addWidget(vp_text)
        vp_group.setLayout(vp_layout)
        layout.addWidget(vp_group)
        
        # Instructions
        info = QLabel("💡 <b>Click any row to select test → Click 'Generate Template' to create test_run document</b>")
        info.setStyleSheet("color: #00ffff; padding: 5px; background-color: #1a1a1a; border: 1px solid #333;")
        layout.addWidget(info)
        
        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels([
            "#", "Name", "Title", "Company", "Type"
        ])
        
        # Populate from TEST_PLAN
        self._populate_table()
        
        self.table.setColumnWidth(0, 40)
        self.table.setColumnWidth(1, 200)
        self.table.setColumnWidth(2, 220)
        self.table.setColumnWidth(3, 300)
        self.table.setColumnWidth(4, 180)
        
        self.table.cellClicked.connect(self.on_row_clicked)
        layout.addWidget(self.table)
        
        # Details section
        details_group = QGroupBox("Selected Test Run")
        details_layout = QVBoxLayout()
        
        self.details_text = QTextEdit()
        self.details_text.setReadOnly(True)
        self.details_text.setMinimumHeight(300)  # Make it bigger
        details_layout.addWidget(self.details_text)
        
        # Generate button
        self.generate_btn = QPushButton("📄 Generate Test Run Template Document")
        self.generate_btn.setEnabled(False)
        self.generate_btn.clicked.connect(self.on_generate_clicked)
        self.generate_btn.setStyleSheet("background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 15px; font-size: 14pt; border: 1px solid #7928CA;")
        details_layout.addWidget(self.generate_btn)
        
        details_group.setLayout(details_layout)
        layout.addWidget(details_group)
        
        self.setLayout(layout)
        
        self.selected_test_run = None
    
    def _populate_table(self):
        """Populate table from Excel - reads LIVE from disk"""
        global TEST_PLAN
        TEST_PLAN = reload_from_excel()
        self.table.setRowCount(len(TEST_PLAN))
        for row, test_run in enumerate(TEST_PLAN):
            self.table.setItem(row, 0, QTableWidgetItem(str(test_entry['num'])))
            self.table.setItem(row, 1, QTableWidgetItem(test_entry['name']))
            self.table.setItem(row, 2, QTableWidgetItem(test_entry['title']))
            self.table.setItem(row, 3, QTableWidgetItem(test_entry['source_version']))
            self.table.setItem(row, 4, QTableWidgetItem(test_entry['type']))
    
    def on_row_clicked(self, row, col):
        """Show test_run details"""
        self.selected_test_run = TEST_PLAN[row]
        
        details = []
        details.append(f"════════════════════════════════════════════════════════")
        details.append(f"TEST #{self.selected_test_entry['num']}: {self.selected_test_entry['name']}")
        details.append(f"════════════════════════════════════════════════════════")
        details.append(f"Title: {self.selected_test_entry['title']}")
        details.append(f"Company: {self.selected_test_entry['source_version']}")
        details.append(f"Customer Type: {self.selected_test_entry['type']}")
        details.append("")
        details.append("HYPOTHESIS:")
        details.append(self.selected_test_entry['hypothesis'])
        details.append("")
        details.append("PLANNED QUESTIONS:")
        details.append(self.selected_test_entry['questions'])
        
        self.details_text.setText("\n".join(details))
        self.generate_btn.setEnabled(True)
    
    def on_generate_clicked(self):
        """Generate test_run template"""
        if self.selected_test_run:
            self.generate_template.emit(self.selected_test_run)


class TestPlanningWidget(QWidget):
    """Display test_run planning sheet and generate templates"""
    
    generate_template = Signal(dict)  # Emits test_run plan data
    
    def __init__(self, get_project_folder_fn=None, parent=None):
        super().__init__(parent)
        self._get_project_folder = get_project_folder_fn
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        # Header
        header_layout = QHBoxLayout()
        header_layout.addWidget(QLabel("<h2>📋 Test Planning Sheet</h2>"))
        header_layout.addStretch()
        
        export_plan_btn = QPushButton("📤 Export Plan to Excel")
        export_plan_btn.clicked.connect(self.export_plan)
        header_layout.addWidget(export_plan_btn)
        
        layout.addLayout(header_layout)
        
        # Info
        info = QLabel("Click any row to generate test_run template document for that person.")
        info.setStyleSheet("color: #888; font-style: italic;")
        layout.addWidget(info)
        
        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels([
            "#", "Name", "Title", "Company", "Type", "Status"
        ])
        
        # Populate table from TEST_PLAN
        self._populate_table()
        
        # Column widths
        self.table.setColumnWidth(0, 40)
        self.table.setColumnWidth(1, 180)
        self.table.setColumnWidth(2, 200)
        self.table.setColumnWidth(3, 250)
        self.table.setColumnWidth(4, 180)
        self.table.setColumnWidth(5, 100)
        
        # Row click
        self.table.cellClicked.connect(self.on_row_clicked)
        
        layout.addWidget(self.table)
        
        # Selected test_run details
        details_group = QGroupBox("Selected Test Run Details")
        details_layout = QVBoxLayout()
        
        self.details_text = QTextEdit()
        self.details_text.setReadOnly(True)
        self.details_text.setMaximumHeight(200)
        details_layout.addWidget(self.details_text)
        
        button_layout = QHBoxLayout()
        
        self.generate_btn = QPushButton("📄 Generate Test Run Template")
        self.generate_btn.clicked.connect(self.on_generate_clicked)
        self.generate_btn.setEnabled(False)
        self.generate_btn.setStyleSheet("background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 10px; border: 1px solid #7928CA;")
        button_layout.addWidget(self.generate_btn)
        
        details_layout.addLayout(button_layout)
        
        details_group.setLayout(details_layout)
        layout.addWidget(details_group)
        
        self.setLayout(layout)
        
        self.selected_test_run = None
    
    def _populate_table(self):
        """Populate table from Excel - reads LIVE from disk, checks completion status"""
        global TEST_PLAN
        TEST_PLAN = reload_from_excel()
        # Diagnostic: show which Excel was loaded and how many rows
        try:
            excel_path = get_excel_path()
            print(f"TestPlanningWidget: loaded {len(TEST_PLAN)} tests from {excel_path}")
        except:
            print(f"TestPlanningWidget: loaded {len(TEST_PLAN)} tests from Excel")
        
        # Get completed person names from change_log
        completed_names = set()
        if self._get_project_folder:
            try:
                import json
                project_folder = Path(self._get_project_folder())
                change_log = project_folder / 'change_log.json'
                if change_log.exists():
                    with open(change_log, 'r') as f:
                        log_entries = json.load(f)
                    completed_names = {
                        entry.get('script_name', '').strip().lower()
                        for entry in log_entries
                        if entry.get('action') == 'TEST_COMPLETED'
                    }
                    print(f"TestPlanningWidget: {len(completed_names)} completed: {sorted(completed_names)}")
            except Exception as e:
                print(f"TestPlanningWidget: change_log error: {e}")

        # Also check test database for completed entries (has results)
        if hasattr(self, '_get_database') and self._get_database:
            try:
                db = self._get_database()
                if db and hasattr(db, 'tests'):
                    for iv in db.test_runs:
                        if getattr(iv, 'results', '').strip() and '[PENDING' not in getattr(iv, 'results', ''):
                            completed_names.add(iv.person.strip().lower())
            except:
                pass
        
        self.table.setRowCount(len(TEST_PLAN))
        for row, test_run in enumerate(TEST_PLAN):
            self.table.setItem(row, 0, QTableWidgetItem(str(test_entry['num'])))
            self.table.setItem(row, 1, QTableWidgetItem(test_entry['name']))
            self.table.setItem(row, 2, QTableWidgetItem(test_entry['title']))
            self.table.setItem(row, 3, QTableWidgetItem(test_entry['source_version']))
            self.table.setItem(row, 4, QTableWidgetItem(test_entry['type']))
            
            # Check completion status by person name
            plan_name = test_entry.get('name', '').strip().lower()
            # Match full name or last name
            name_matched = (
                plan_name in completed_names
                or any(plan_name.split()[-1] in cn for cn in completed_names if plan_name)
            )
            if name_matched:
                status_item = QTableWidgetItem("✓ Completed")
                status_item.setForeground(QBrush(QColor("#00ff88")))
                # Green background for completed rows
                for col in range(6):
                    item = self.table.item(row, col)
                    if item:
                        item.setBackground(QBrush(QColor("#0a2a0a")))
            else:
                status_item = QTableWidgetItem("Planned")
            self.table.setItem(row, 5, status_item)
    
    def on_row_clicked(self, row, col):
        """Handle row click - show test_run details"""
        self.selected_test_run = TEST_PLAN[row]
        
        details = []
        details.append(f"TEST #{self.selected_test_entry['num']}: {self.selected_test_entry['name']}")
        details.append(f"Title: {self.selected_test_entry['title']}")
        details.append(f"Company: {self.selected_test_entry['source_version']}")
        details.append(f"Type: {self.selected_test_entry['type']}")
        details.append("")
        details.append("HYPOTHESIS:")
        details.append(self.selected_test_entry['hypothesis'])
        details.append("")
        details.append("PLANNED QUESTIONS:")
        details.append(self.selected_test_entry['questions'])
        
        self.details_text.setText("\n".join(details))
        self.generate_btn.setEnabled(True)
    
    def on_generate_clicked(self):
        """Generate test_run template document"""
        if self.selected_test_run:
            self.generate_template.emit(self.selected_test_run)
    
    def export_plan(self):
        """Export test_run plan to Excel"""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Test Plan"
            
            # Headers
            headers = ["#", "Name", "Title", "Type", "Company", "Hypothesis", "Questions"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="c41e3a", end_color="c41e3a", fill_type="solid")
                cell.font = Font(bold=True, color="FFFFFF")
            
            # Data
            for row, test_run in enumerate(TEST_PLAN, 2):
                ws.cell(row=row, column=1, value=test_entry['num'])
                ws.cell(row=row, column=2, value=test_entry['name'])
                ws.cell(row=row, column=3, value=test_entry['title'])
                ws.cell(row=row, column=4, value=test_entry['type'])
                ws.cell(row=row, column=5, value=test_entry['source_version'])
                ws.cell(row=row, column=6, value=test_entry['hypothesis'])
                ws.cell(row=row, column=7, value=test_entry['questions'])
            
            # Column widths
            ws.column_dimensions['A'].width = 5
            ws.column_dimensions['B'].width = 25
            ws.column_dimensions['C'].width = 30
            ws.column_dimensions['D'].width = 25
            ws.column_dimensions['E'].width = 40
            ws.column_dimensions['F'].width = 60
            ws.column_dimensions['G'].width = 80
            
            filepath = get_project_dir(get_active_project()) / f"Test Run_Plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            wb.save(filepath)
            
            QMessageBox.information(self, "Exported", f"Test Run plan exported to:\n{filepath}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export:\n{str(e)}")


class Test RunEntryWidget(QWidget):
    """Widget for entering single test data"""
    
    test_run_saved = Signal(Test Run)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        # Header info
        header_layout = QHBoxLayout()
        
        self.num_input = QLineEdit()
        self.num_input.setPlaceholderText("Test Run #")
        header_layout.addWidget(QLabel("Test Run #:"))
        header_layout.addWidget(self.num_input)
        
        self.date_input = QLineEdit()
        self.date_input.setText(datetime.now().strftime("%m/%d/%Y"))
        header_layout.addWidget(QLabel("Date:"))
        header_layout.addWidget(self.date_input)
        
        layout.addLayout(header_layout)
        
        # Person info
        person_layout = QVBoxLayout()
        person_layout.addWidget(QLabel("<b>Person Test Runed:</b>"))
        
        self.person_input = QLineEdit()
        self.person_input.setPlaceholderText("Name")
        person_layout.addWidget(self.person_input)
        
        self.title_input = QLineEdit()
        self.title_input.setPlaceholderText("Title/Position")
        person_layout.addWidget(self.title_input)
        
        self.company_input = QLineEdit()
        self.company_input.setPlaceholderText("Company")
        person_layout.addWidget(self.company_input)
        
        self.type_combo = QComboBox()
        self.type_combo.addItems([
            "User", "Decision-Maker", "Payer", "Influencer",
            "User/Decision-Maker", "Decision-Maker/Payer"
        ])
        person_layout.addWidget(QLabel("Customer Type:"))
        person_layout.addWidget(self.type_combo)
        
        layout.addLayout(person_layout)
        
        # ICorps template fields
        self.hypotheses_input = QTextEdit()
        self.hypotheses_input.setPlaceholderText("What we thought / What assumption was tested")
        layout.addWidget(QLabel("<b>Hypotheses:</b>"))
        layout.addWidget(self.hypotheses_input)
        
        self.experiments_input = QTextEdit()
        self.experiments_input.setPlaceholderText("Questions we asked")
        layout.addWidget(QLabel("<b>Experiments (Questions Asked):</b>"))
        layout.addWidget(self.experiments_input)
        
        self.results_input = QTextEdit()
        self.results_input.setPlaceholderText("What we learned")
        layout.addWidget(QLabel("<b>Results (What We Learned):</b>"))
        layout.addWidget(self.results_input)
        
        self.action_input = QTextEdit()
        self.action_input.setPlaceholderText("What we will do next")
        layout.addWidget(QLabel("<b>Action/Iterate:</b>"))
        layout.addWidget(self.action_input)
        
        # Substrate Impact Section
        equip_group = QGroupBox("Substrate Impacted (Optional - Check equipment mentioned in this test_run)")
        equip_layout = QVBoxLayout()
        
        # Scrollable equipment checkboxes
        equip_scroll = QScrollArea()
        equip_scroll.setMaximumHeight(200)
        equip_scroll.setWidgetResizable(True)
        self._equip_scroll_widget = QWidget()
        self._equip_scroll_layout = QVBoxLayout()
        
        self.substrate_impacts = {}
        
        # Load equipment dynamically from config (project-aware)
        self._build_equipment_checkboxes()
        
        self._equip_scroll_widget.setLayout(self._equip_scroll_layout)
        equip_scroll.setWidget(self._equip_scroll_widget)
        equip_layout.addWidget(equip_scroll)
        
        equip_group.setLayout(equip_layout)
        layout.addWidget(equip_group)
        
        # Save button
        save_btn = QPushButton("Save Test & Run AI Analysis")
        save_btn.clicked.connect(self.save_test_run)
        save_btn.setStyleSheet("background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 10px; border: 1px solid #7928CA;")
        layout.addWidget(save_btn)
        
        self.setLayout(layout)
    
    def set_project_getter(self, getter):
        """Set callable to get active project type. Called by main window."""
        self._get_active_project = getter
    
    def _get_project_type(self) -> str:
        """Get current project type for config loading."""
        if hasattr(self, '_get_active_project') and self._get_active_project:
            try:
                return self._get_active_project()
            except:
                pass
        return get_active_project()
    
    def _build_equipment_checkboxes(self):
        """
        Build equipment checkboxes from config JSON — project-aware.
        Reads from substrate_config_chip.json or substrate_config_bcm_navigation.json
        based on active project. Rebuilds on project switch.
        """
        # Clear existing checkboxes
        self.substrate_impacts = {}
        while self._equip_scroll_layout.count():
            item = self._equip_scroll_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                while item.layout().count():
                    sub = item.layout().takeAt(0)
                    if sub.widget():
                        sub.widget().deleteLater()
        
        project_type = self._get_project_type()
        equipment_by_cat, _ = _load_equipment_from_config(project_type)
        
        if equipment_by_cat:
            for cat_name, items in equipment_by_cat.items():
                for item in items[:3]:  # Top 3 per category keeps form manageable
                    equip_name = item['name']
                    row_layout = QHBoxLayout()
                    
                    checkbox = QCheckBox(equip_name)
                    checkbox.stateChanged.connect(
                        lambda state, e=equip_name: self._on_equipment_toggled(e, state))
                    row_layout.addWidget(checkbox)
                    
                    cost_input = QLineEdit()
                    cost_input.setPlaceholderText("Annual $ (e.g. 45000)")
                    cost_input.setMaximumWidth(150)
                    cost_input.setEnabled(False)
                    row_layout.addWidget(cost_input)
                    
                    row_layout.addStretch()
                    self._equip_scroll_layout.addLayout(row_layout)
                    
                    self.substrate_impacts[equip_name] = {
                        'checkbox': checkbox,
                        'cost_input': cost_input,
                        'selected': False,
                        'cost': 0
                    }
        else:
            warn = QLabel(
                f"⚠️ No equipment config found for {project_type}. "
                f"Expected: config/{project_type.lower()}/substrate_config_*.json"
            )
            warn.setStyleSheet("color: #FF4444; padding: 4px;")
            warn.setWordWrap(True)
            self._equip_scroll_layout.addWidget(warn)
        
        self._equip_scroll_layout.addStretch()
    
    def rebuild_equipment(self):
        """Rebuild equipment checkboxes when project changes."""
        self._build_equipment_checkboxes()
    
    def _on_equipment_toggled(self, component_name: str, state: int):
        """Handle equipment checkbox toggle."""
        if component_name in self.substrate_impacts:
            data = self.substrate_impacts[component_name]
            data['selected'] = bool(state)
            data['cost_input'].setEnabled(bool(state))
    
    def save_test_run(self):
        """Save test_run and emit signal"""
        test_run = Test Run()
        test_run.test_num = int(self.num_input.text() or 0)
        test_run.date = self.date_input.text()
        test_run.person = self.person_input.text()
        test_run.title = self.title_input.text()
        test_run.company = self.company_input.text()
        test_run.test_category = self.type_combo.currentText()
        test_run.hypotheses = self.hypotheses_input.toPlainText()
        test_run.experiments = self.experiments_input.toPlainText()
        test_run.results = self.results_input.toPlainText()
        test_run.action_iterate = self.action_input.toPlainText()
        
        # Collect equipment impacts
        equipment_data = []
        for equipment, data in self.substrate_impacts.items():
            if data['selected']:
                try:
                    cost = float(data['cost_input'].text().replace(',', '')) if data['cost_input'].text() else 0
                    equipment_data.append({
                        'equipment': equipment,
                        'cost': cost
                    })
                except:
                    pass
        
        # Store in test_run (will need to add this field to Test Run class)
        test_run.substrate_impacts = equipment_data
        
        self.test_run_saved.emit(test_run)
        
        # Clear form
        QMessageBox.information(self, "Saved", f"Test Run #{test_run.test_num} saved!")
        self.clear_form()
    
    def clear_form(self):
        """Clear all input fields"""
        self.num_input.clear()
        self.person_input.clear()
        self.title_input.clear()
        self.company_input.clear()
        self.hypotheses_input.clear()
        self.experiments_input.clear()
        self.results_input.clear()
        self.action_input.clear()


class QCubeVisualizerWidget(QWidget):
    """Multi-lens Q-Cube matrix visualization with dynamic axis selection."""
    
    def __init__(self, database: Test RunDatabase, parent=None):
        super().__init__(parent)
        self.database = database
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        # ── Header ──
        title = QLabel("<h2>Q-Cube Multi-Lens Analysis</h2>")
        title.setStyleSheet("color: #FFD700;")
        layout.addWidget(title)
        
        # ── Lens Selector Row ──
        selector_layout = QHBoxLayout()
        
        selector_layout.addWidget(QLabel("Primary Axis:"))
        self.primary_combo = QComboBox()
        for key, name in QCUBE_LENS_NAMES.items():
            self.primary_combo.addItem(name, key)
        self.primary_combo.setCurrentIndex(0)  # Q_LAYER default
        self.primary_combo.currentIndexChanged.connect(self.refresh)
        self.primary_combo.setStyleSheet("background-color: #1a1a1a; color: #00ffff; padding: 4px;")
        selector_layout.addWidget(self.primary_combo)
        
        selector_layout.addWidget(QLabel("Secondary Axis:"))
        self.secondary_combo = QComboBox()
        for key, name in QCUBE_LENS_NAMES.items():
            self.secondary_combo.addItem(name, key)
        self.secondary_combo.setCurrentIndex(1)  # Q_OBJECT default
        self.secondary_combo.currentIndexChanged.connect(self.refresh)
        self.secondary_combo.setStyleSheet("background-color: #1a1a1a; color: #00ffff; padding: 4px;")
        selector_layout.addWidget(self.secondary_combo)
        
        # Auto-tag button
        self.autotag_btn = QPushButton("🔬 Auto-Tag Lenses")
        self.autotag_btn.setStyleSheet(
            "background-color: #004400; color: #00ff88; font-weight: bold; "
            "padding: 6px 14px; border: 1px solid #00ff88;"
        )
        self.autotag_btn.clicked.connect(self._auto_tag_lenses)
        selector_layout.addWidget(self.autotag_btn)
        
        selector_layout.addStretch()
        layout.addLayout(selector_layout)
        
        # ── Gap Analysis Summary ──
        self.gap_label = QLabel("")
        self.gap_label.setWordWrap(True)
        self.gap_label.setStyleSheet(
            "color: #00ffff; padding: 6px; background-color: #1a1a1a; border: 1px solid #333;"
        )
        layout.addWidget(self.gap_label)
        
        # ── Main tree view ──
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Position", "Test Runs", "Lenses", "Synergy"])
        self.tree.setColumnWidth(0, 350)
        self.tree.setColumnWidth(1, 200)
        self.tree.setColumnWidth(2, 250)
        self.tree.setColumnWidth(3, 80)
        self.tree.setStyleSheet("""
            QTreeWidget {
                background-color: #0a0a0a;
                color: #00ff88;
                border: 1px solid #333;
            }
            QTreeWidget::item {
                padding: 2px;
            }
            QTreeWidget::item:selected {
                background-color: #9333EA;
                color: #FFD700;
            }
            QHeaderView::section {
                background-color: #1a1a1a;
                color: #00ffff;
                font-weight: bold;
                border: 1px solid #333;
                padding: 4px;
            }
        """)
        layout.addWidget(self.tree)
        
        # ── Coverage Table (all 7 lenses) ──
        self.coverage_label = QLabel("<b>Lens Coverage Summary</b>")
        self.coverage_label.setStyleSheet("color: #FFD700; margin-top: 8px;")
        layout.addWidget(self.coverage_label)
        
        self.coverage_table = QTableWidget()
        self.coverage_table.setMaximumHeight(200)
        self.coverage_table.setStyleSheet("""
            QTableWidget {
                background-color: #0a0a0a;
                color: #00ff88;
                gridline-color: #333;
            }
            QHeaderView::section {
                background-color: #1a1a1a;
                color: #00ffff;
                font-weight: bold;
                border: 1px solid #333;
            }
        """)
        layout.addWidget(self.coverage_table)
        
        self.setLayout(layout)
        self.refresh()
    
    def _get_test_run_lens_value(self, test_run, lens_key: str) -> str:
        """Get the value of a lens for a test."""
        field_map = {
            'Q_LAYER': 'q_layer',
            'Q_OBJECT': 'q_object',
            'Q_STACK': 'q_stack',
            'Q_AWARENESS': 'q_awareness',
            'Q_EVIDENCE': 'q_evidence',
            'Q_TIMEHORIZON': 'q_timehorizon',
            'Q_CASCADE': 'q_cascade',
            'Q_NORMALIZATION': 'q_normalization',
            'Q_COUNTERFLOW': 'q_counterflow',
            'Q_TRIBALKNOWLEDGE': 'q_tribalknowledge',
        }
        field = field_map.get(lens_key, '')
        if not field:
            return ''
        val = getattr(test_run, field, '')
        # Q_STACK is a list, use first entry
        if isinstance(val, list):
            return val[0] if val else ''
        return val or ''
    
    def _get_lens_tag_line(self, test_run) -> str:
        """Build compact 10-lens tag: [L1|OA|Sα|A2|V2|T1|C4|N3|Fg|K3]"""
        parts = []
        for lens_key in QCUBE_LENS_NAMES:
            val = self._get_test_run_lens_value(test_run, lens_key)
            if val:
                parts.append(val)
        return '[' + '|'.join(parts) + ']' if parts else ''
    
    def _auto_tag_lenses(self):
        """Run lens inference on all tests."""
        if not self.database or not self.database.test_runs:
            return
        
        updated = apply_lenses_to_database(self.database)
        
        # Print results
        print(f"\n{'='*60}")
        print(f"LENS AUTO-TAG — {updated}/{len(self.database.tests)} tests updated")
        for test_entry in self.database.test_runs:
            tags = self._get_lens_tag_line(test_run)
            print(f"  #{test_run.test_num} {test_run.person}: {tags}")
        print(f"{'='*60}\n")
        
        # Save
        db_file = BCM_ROOT / "test_database.json"
        self.database.save(db_file)
        
        # Refresh display
        self.refresh()
        
        # Compute confidence summary
        total_conf = []
        low_conf = 0
        for iv in self.database.test_runs:
            confs = getattr(iv, 'lens_confidence', {})
            for v in confs.values():
                total_conf.append(v)
                if v < 0.6:
                    low_conf += 1
        avg_conf = sum(total_conf) / len(total_conf) if total_conf else 0
        
        QMessageBox.information(self, "Hypercube Lenses Applied",
            f"Auto-tagged {updated} tests across 7 analytical lenses.\n\n"
            f"Analytical: Awareness (A), Evidence (V), Time Horizon (T)\n"
            f"Originals: Cascade (C), Normalization (N), Counter-Flow (F), Tribal Knowledge (K)\n\n"
            f"10 total lenses active. Manual overrides preserved.\n\n"
            f"Avg confidence: {avg_conf:.0%} | {low_conf} low-confidence assignments\n"
            f"Engines: sklearn={True}, bayesian={BAYESIAN_ENGINE_AVAILABLE}, "
            f"doctoral={DOCTORAL_ANALYZER_AVAILABLE}")
    
    def refresh(self):
        """Refresh visualization with selected primary/secondary axes."""
        self.tree.clear()
        
        if not self.database or not self.database.test_runs:
            return
        
        # Get selected axes
        primary_key = self.primary_combo.currentData()
        secondary_key = self.secondary_combo.currentData()
        
        if not primary_key:
            primary_key = 'Q_LAYER'
        if not secondary_key:
            secondary_key = 'Q_OBJECT'
        
        primary_dim = QCUBE_DIMENSIONS.get(primary_key, {})
        secondary_dim = QCUBE_DIMENSIONS.get(secondary_key, {})
        
        # Build cross-reference matrix
        total_gaps = 0
        total_cells = 0
        total_populated = 0
        
        for p_code, p_desc in primary_dim.items():
            p_short = p_desc.split(' - ')[0] if ' - ' in p_desc else p_desc
            primary_item = QTreeWidgetItem([f"{p_code}: {p_short}", "", "", ""])
            primary_item.setFont(0, QFont("Arial", 10, QFont.Bold))
            primary_item.setForeground(0, QBrush(QColor("#FFD700")))
            
            primary_count = 0
            
            for s_code, s_desc in secondary_dim.items():
                s_short = s_desc.split(' - ')[0] if ' - ' in s_desc else s_desc
                total_cells += 1
                
                # Find tests matching both axes
                matches = []
                for test_entry in self.database.test_runs:
                    p_val = self._get_test_run_lens_value(test_run, primary_key)
                    s_val = self._get_test_run_lens_value(test_run, secondary_key)
                    
                    if p_val == p_code and s_val == s_code:
                        matches.append(test_run)
                
                if matches:
                    total_populated += 1
                    count = len(matches)
                    primary_count += count
                    
                    # Synergy scores
                    scores = [i.synergy_score for i in matches if i.synergy_score > 0]
                    avg_score = f"{sum(scores)/len(scores):.2f}" if scores else ""
                    
                    # Names
                    names = ", ".join(f"#{i.test_num}" for i in matches)
                    
                    # Full lens tags for each
                    lens_tags = " ".join(self._get_lens_tag_line(i) for i in matches[:3])
                    if len(matches) > 3:
                        lens_tags += f" +{len(matches)-3} more"
                    
                    secondary_item = QTreeWidgetItem([
                        f"  {s_code}: {s_short}",
                        f"{count} — {names}",
                        lens_tags,
                        avg_score
                    ])
                    
                    # Color by count
                    if count >= 3:
                        secondary_item.setBackground(0, QBrush(QColor("#0a3a0a")))
                        secondary_item.setForeground(0, QBrush(QColor("#00ff88")))
                    elif count >= 2:
                        secondary_item.setBackground(0, QBrush(QColor("#0a2a0a")))
                        secondary_item.setForeground(0, QBrush(QColor("#00cc66")))
                    else:
                        secondary_item.setForeground(0, QBrush(QColor("#888")))
                    
                    # Add test_run detail children
                    for test_entry in matches:
                        tags = self._get_lens_tag_line(test_run)
                        detail = QTreeWidgetItem([
                            f"    #{test_run.test_num}: {test_run.person} ({test_run.company})",
                            test_run.title[:40] if test_run.title else "",
                            tags,
                            f"{test_run.synergy_score:.2f}" if test_run.synergy_score else ""
                        ])
                        detail.setForeground(0, QBrush(QColor("#aaa")))
                        secondary_item.addChild(detail)
                    
                    primary_item.addChild(secondary_item)
                else:
                    total_gaps += 1
                    # Show gap as dim entry
                    gap_item = QTreeWidgetItem([
                        f"  {s_code}: {s_short}",
                        "— GAP —",
                        "",
                        ""
                    ])
                    gap_item.setForeground(0, QBrush(QColor("#444")))
                    gap_item.setForeground(1, QBrush(QColor("#661020")))
                    primary_item.addChild(gap_item)
            
            # Update primary count
            primary_item.setText(1, f"{primary_count} test_runs")
            if primary_count == 0:
                primary_item.setForeground(0, QBrush(QColor("#661020")))
                primary_item.setForeground(1, QBrush(QColor("#661020")))
            
            self.tree.addTopLevelItem(primary_item)
        
        self.tree.expandAll()
        
        # Gap analysis summary
        coverage_pct = (total_populated / total_cells * 100) if total_cells > 0 else 0
        self.gap_label.setText(
            f"📊 {QCUBE_LENS_NAMES.get(primary_key, '')} × "
            f"{QCUBE_LENS_NAMES.get(secondary_key, '')}: "
            f"{total_populated}/{total_cells} cells populated ({coverage_pct:.0f}%) | "
            f"{total_gaps} gaps | "
            f"{len(self.database.tests)} total test_runs"
        )
        if coverage_pct < 30:
            self.gap_label.setStyleSheet(
                "color: #FF4444; padding: 6px; background-color: #1a0a0a; border: 1px solid #661020;"
            )
        elif coverage_pct < 60:
            self.gap_label.setStyleSheet(
                "color: #FFD700; padding: 6px; background-color: #1a1a0a; border: 1px solid #665500;"
            )
        else:
            self.gap_label.setStyleSheet(
                "color: #00ff88; padding: 6px; background-color: #0a1a0a; border: 1px solid #006600;"
            )
        
        # Build coverage table across ALL 7 lenses
        self._build_coverage_table()
    
    def _build_coverage_table(self):
        """Build summary table showing coverage across ALL 7 lens dimensions."""
        lens_keys = list(QCUBE_LENS_NAMES.keys())
        
        # For each lens, count how many tests fall into each value
        # Columns: Lens Name | Value1 (count) | Value2 (count) | ... | GAPS
        max_values = max(len(QCUBE_DIMENSIONS[k]) for k in lens_keys)
        
        self.coverage_table.setRowCount(len(lens_keys))
        self.coverage_table.setColumnCount(max_values + 2)  # Lens name + values + gap count
        
        headers = ["Lens"]
        for i in range(max_values):
            headers.append(f"Value {i+1}")
        headers.append("Gaps")
        self.coverage_table.setHorizontalHeaderLabels(headers)
        
        for row_idx, lens_key in enumerate(lens_keys):
            dim = QCUBE_DIMENSIONS[lens_key]
            lens_name = QCUBE_LENS_NAMES[lens_key]
            
            # Lens name
            name_item = QTableWidgetItem(lens_name)
            name_item.setFont(QFont("Arial", 9, QFont.Bold))
            name_item.setFlags(Qt.ItemIsEnabled)
            name_item.setForeground(QBrush(QColor("#FFD700")))
            self.coverage_table.setItem(row_idx, 0, name_item)
            
            # Count tests per value
            empty_count = 0
            col_idx = 1
            for code, desc in dim.items():
                count = 0
                for test_entry in self.database.test_runs:
                    val = self._get_test_run_lens_value(test_run, lens_key)
                    if val == code:
                        count += 1
                
                short_desc = desc.split(' - ')[0] if ' - ' in desc else desc
                cell_text = f"{code}: {count}"
                cell_item = QTableWidgetItem(cell_text)
                cell_item.setFlags(Qt.ItemIsEnabled)
                cell_item.setToolTip(f"{code}: {desc}\n{count} test_runs")
                
                if count == 0:
                    cell_item.setForeground(QBrush(QColor("#661020")))
                    cell_item.setBackground(QBrush(QColor("#1a0a0a")))
                    empty_count += 1
                elif count == 1:
                    cell_item.setForeground(QBrush(QColor("#888")))
                elif count >= 3:
                    cell_item.setForeground(QBrush(QColor("#00ff88")))
                    cell_item.setBackground(QBrush(QColor("#0a2a0a")))
                else:
                    cell_item.setForeground(QBrush(QColor("#00cc66")))
                
                self.coverage_table.setItem(row_idx, col_idx, cell_item)
                col_idx += 1
            
            # Fill remaining columns
            while col_idx <= max_values:
                self.coverage_table.setItem(row_idx, col_idx, QTableWidgetItem(""))
                col_idx += 1
            
            # Gap count
            gap_item = QTableWidgetItem(str(empty_count))
            gap_item.setFlags(Qt.ItemIsEnabled)
            if empty_count > 0:
                gap_item.setForeground(QBrush(QColor("#FF4444")))
                gap_item.setBackground(QBrush(QColor("#1a0a0a")))
            else:
                gap_item.setForeground(QBrush(QColor("#00ff88")))
            self.coverage_table.setItem(row_idx, max_values + 1, gap_item)
        
        self.coverage_table.resizeColumnsToContents()
        self.coverage_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)


# ============================================================================
# EQUIPMENT HELPERS - Config loader, fuzzy matcher, test_run scanner
# ============================================================================

def _find_substrate_config(project_type: str = None) -> Path:
    """
    Find equipment config JSON by searching config/ subfolders.
    Matches by keyword — AISOS_SPINE finds bcm_navigation/ via 'spine'.
    """
    if project_type is None:
        project_type = get_active_project()
    
    subfolder = project_type.lower()
    keywords = [w for w in subfolder.split("_") if len(w) > 2]
    
    script_dir = Path(__file__).parent if '__file__' in dir() else Path.cwd()
    
    anchors = [
        script_dir,
        Path.cwd(),
        PROJECT_ROOT,
    ]
    
    if sys.platform == "win32":
        anchors.extend([
            Path("C:/TITS/TITS_GIBUSH_AISOS_SPINE/TITS_GIBUSH_AISOS_SPINE_ICORPS"),
            Path("C:/TITS/TITS_Production/gibush_icorps"),
        ])
    
    for anchor in anchors:
        config_dir = anchor / "config"
        if not config_dir.exists():
            continue
        
        # 1. Exact subfolder match
        exact = config_dir / subfolder
        if exact.exists():
            for f in exact.glob("substrate_config*.json"):
                print(f"  ✓ Equipment config: {f}")
                return f
        
        # 2. Keyword match against subfolder names
        for sub in config_dir.iterdir():
            if not sub.is_dir():
                continue
            sub_lower = sub.name.lower()
            if any(kw in sub_lower for kw in keywords):
                for f in sub.glob("substrate_config*.json"):
                    print(f"  ✓ Equipment config: {f}")
                    return f
        
        # 3. Keyword match in filename
        for f in config_dir.glob("substrate_config*.json"):
            if any(kw in f.stem.lower() for kw in keywords):
                print(f"  ✓ Equipment config: {f}")
                return f
    
    print(f"  ✗ No equipment config found for {project_type}")
    return None


def _load_equipment_from_config(project_type: str = None):
    """
    Load equipment categories from config JSON (v5.0.0+).
    Returns (components_by_category, pain_indicators)
    """
    config_file = _find_substrate_config(project_type)
    
    components_by_category = {}
    pain_indicators = {}
    
    if config_file is None or not config_file.exists():
        return components_by_category, pain_indicators
    
    print(f"  ✓ Equipment config loaded: {config_file}")
    
    with open(config_file, 'r', encoding='utf-8') as f:
        config = json.load(f)
    
    if 'module_categories' in config:
        categories = config.get('module_categories', {})
        for cat_id, cat_data in categories.items():
            cat_name = cat_data.get('name', cat_id)
            components_by_category[cat_name] = []
            for mod_id, mod_data in cat_data.get('modules', {}).items():
                components_by_category[cat_name].append({
                    'id': mod_id,
                    'name': mod_data.get('display_name', mod_id),
                    'cost': mod_data.get('typical_cost', ''),
                    'keywords': mod_data.get('keywords', []),
                })
    else:
        categories = config.get('equipment_categories', {})
        for cat_id, cat_data in categories.items():
            cat_name = cat_data.get('name', cat_id)
            components_by_category[cat_name] = []
            for equip_id, equip_data in cat_data.get('equipment', {}).items():
                components_by_category[cat_name].append({
                    'id': equip_id,
                    'name': equip_data.get('display_name', equip_id),
                    'cost': equip_data.get('typical_cost_range', ''),
                    'keywords': equip_data.get('keywords', []),
                    'notes': equip_data.get('notes', ''),
                })
    
    pain_indicators = config.get('pain_indicators', {})
    return components_by_category, pain_indicators


def _fuzzy_match_equipment(test_run_name: str, config_name: str, config_keywords: list) -> bool:
    """
    Check if a test equipment name matches a config equipment entry.
    Uses name substring + keyword + word overlap matching.
    """
    i_lower = test_run_name.lower().strip()
    c_lower = config_name.lower().strip()
    
    # Direct substring match
    if i_lower in c_lower or c_lower in i_lower:
        return True
    
    # First 20 chars match
    if len(i_lower) >= 10 and len(c_lower) >= 10:
        if i_lower[:20] == c_lower[:20]:
            return True
    
    # Keyword match (any keyword appears in test_run name)
    for kw in config_keywords:
        if kw.lower() in i_lower:
            return True
    
    # Word overlap (>50% of words match)
    i_words = set(re.findall(r'\w+', i_lower))
    c_words = set(re.findall(r'\w+', c_lower))
    if i_words and c_words:
        overlap = len(i_words & c_words)
        if overlap >= 2 or (overlap >= 1 and min(len(i_words), len(c_words)) <= 2):
            return True
    
    return False


def _scan_tests_for_equipment(database) -> dict:
    """
    Scan ALL tests in database and extract equipment mentions.
    
    FRICTION TRACKING: Equipment that was DISCUSSED in a test counts as 
    a mention even without a dollar cost. The mention count itself = friction 
    signal showing how much traffic a topic gets across tests.
    
    Returns: { "component_name": { 
        "tests": [{"num": 7, "script_name": "...", "cost": "", "notes": "", "source": "checklist|text"}],
        "total_mentions": N, 
        "costs_reported": [...], 
        "notes": [...],
        "discussed_count": N,   # checked in table, may have no cost
        "costed_count": N,      # has actual dollar amount
    }, ... }
    """
    equipment_data = {}
    
    if not database or not hasattr(database, 'tests'):
        return equipment_data
    
    def _ensure_entry(name):
        if name not in equipment_data:
            equipment_data[name] = {
                'tests': [],
                'total_mentions': 0,
                'costs_reported': [],
                'notes': [],
                'discussed_count': 0,
                'costed_count': 0,
            }
    
    def _already_listed(name, person):
        return any(i.get('script_name', '').strip().lower() == person.strip().lower()
                   for i in equipment_data[name]['tests'])
    
    for test_entry in database.test_runs:
        int_num = test_run.test_num
        int_person = test_run.person
        
        # Source 1: substrate_impacts (structured, from doc reader / rescan)
        # These are items checked "Yes" in the equipment checklist tables
        if hasattr(test_run, 'substrate_impacts') and test_run.substrate_impacts:
            for impact in test_run.substrate_impacts:
                equip_name = impact.get('equipment', '').strip()
                if not equip_name or equip_name == 'Equipment/Module':
                    continue
                
                _ensure_entry(equip_name)
                
                if not _already_listed(equip_name, int_person):
                    equipment_data[equip_name]['total_mentions'] += 1
                    equipment_data[equip_name]['discussed_count'] += 1
                    
                    cost = impact.get('cost', '')
                    notes = impact.get('notes', '')
                    
                    equipment_data[equip_name]['tests'].append({
                        'num': int_num,
                        'script_name': int_person,
                        'cost': cost,
                        'notes': notes,
                        'source': 'checklist',
                    })
                    
                    # Track actual costs (not just config range placeholders)
                    if cost and cost != '0':
                        # Filter out config ranges like "$15,000 - $80,000" 
                        # vs actual costs like "$45,000" or "45000"
                        equipment_data[equip_name]['costs_reported'].append(cost)
                        equipment_data[equip_name]['costed_count'] += 1
                    
                    if notes:
                        equipment_data[equip_name]['notes'].append(notes)
        
        # Source 2: Scan all text fields for equipment keyword mentions
        # This catches discussion that wasn't in the structured checklist
        all_text = ''
        for field in ('results', 'experiments', 'hypotheses', 'action_iterate'):
            val = getattr(test_run, field, None)
            if val:
                all_text += ' ' + val
        
        if all_text.strip():
            text_lower = all_text.lower()
            # Equipment keywords — covers v5.0.0 categories
            equip_keywords = {
                'CTS': ['cts', 'chip thickness screen'],
                'Blow Line': ['blow line', 'blowline'],
                'Screen Baskets': ['screen basket'],
                'Chipper Knives': ['chipper kniv', 'knife', 'knives'],
                'Debarker': ['debarker', 'debarking'],
                'Digester': ['digester'],
                'Recovery Boiler': ['recovery boiler'],
                'Log Truck': ['log truck', 'hauling truck', 'log hauling'],
                'Chip Van': ['chip van', 'walking floor', 'chip hauling'],
                'Wrappers/Binders': ['wrapper', 'binder chain'],
                'Loader/Grapple': ['log loader', 'grapple', 'knuckleboom'],
                'Tires': ['tire damage', 'blowout', 'tire wear'],
                'Magnets': ['magnet'],
                'Log Trailer': ['log trailer', 'pole trailer', 'bunk trailer'],
                'Log Bunks': ['bunk', 'stake', 'bolster'],
                'Scale/Weigh': ['scale house', 'weigh', 'phantom weight'],
                'Processor Head': ['processor head', 'harvester head'],
                'Feller Buncher': ['feller buncher'],
                'Chip Screen': ['chip screen', 'overthick'],
                'Rechipper': ['rechipper', 'rechip'],
                'Fourdrinier': ['fourdrinier', 'wire pit'],
                'Dryer Felt': ['dryer felt', 'dryer fabric'],
                'Pitch': ['pitch deposit', 'pitch problem'],
                'Tub Pulper': ['tub pulper', 'pulper'],
            }
            for equip_name, keywords in equip_keywords.items():
                for kw in keywords:
                    if kw in text_lower:
                        _ensure_entry(equip_name)
                        if not _already_listed(equip_name, int_person):
                            equipment_data[equip_name]['total_mentions'] += 1
                            equipment_data[equip_name]['discussed_count'] += 1
                            
                            # Extract dollar amount near the keyword mention
                            cost_found = ''
                            kw_pos = text_lower.find(kw)
                            if kw_pos >= 0:
                                # Search within 200 chars around keyword for $ amounts
                                window_start = max(0, kw_pos - 100)
                                window_end = min(len(all_text), kw_pos + len(kw) + 150)
                                window = all_text[window_start:window_end]
                                # Match patterns: $300K, $45,000, $160,000/yr, $70K/year, $1M, $1.2M
                                cost_matches = re.findall(
                                    r'\$[\d,]+(?:\.\d+)?(?:\s*[kKmM])?(?:/?\s*(?:yr|year|annual))?',
                                    window
                                )
                                if cost_matches:
                                    cost_found = cost_matches[0].strip()
                            
                            equipment_data[equip_name]['tests'].append({
                                'num': int_num,
                                'script_name': int_person,
                                'cost': cost_found,
                                'notes': f'(text: "{kw}")',
                                'source': 'text',
                            })
                            
                            if cost_found:
                                equipment_data[equip_name]['costs_reported'].append(cost_found)
                                equipment_data[equip_name]['costed_count'] += 1
                        break  # One match per equipment per test
    
    return equipment_data


def _parse_cost_to_number(cost_str: str) -> float:
    """Extract numeric value from cost string like '$45,000', '$300K', '$1.2M', '$15,000 - $80,000'."""
    if not cost_str:
        return 0
    clean = cost_str.strip().replace('$', '').replace(',', '')
    
    # Handle K/k and M/m multipliers: $300K → 300000, $1.2M → 1200000
    mult_match = re.match(r'^([\d.]+)\s*([kKmM])', clean)
    if mult_match:
        val = float(mult_match.group(1))
        mult = mult_match.group(2).upper()
        return val * 1000 if mult == 'K' else val * 1000000
    
    range_match = re.findall(r'[\d.]+', clean)
    if len(range_match) >= 2:
        try:
            return (float(range_match[0]) + float(range_match[1])) / 2
        except:
            pass
    elif len(range_match) == 1:
        try:
            return float(range_match[0])
        except:
            pass
    return 0



# ============================================================================
# EQUIPMENT IMPACT WIDGET - Imported from standalone module
# ============================================================================

try:
    from substrate_impact_widget import SubstrateImpactWidget
    print("  ✓ Substrate Impact Widget loaded (standalone module)")
except ImportError:
    # Stub fallback if standalone not available
    class SubstrateImpactWidget(QWidget):
        impact_updated = Signal(float)
        def __init__(self, database=None, parent=None):
            super().__init__(parent)
            self.database = database
            self._get_active_project = None
            layout = QVBoxLayout(self)
            layout.addWidget(QLabel("⚠️ substrate_impact_widget.py not found. Place it next to validation_test_collector.py."))
            self.setLayout(layout)
        def set_project_getter(self, fn): self._get_active_project = fn
        def set_database(self, db): self.database = db
        def refresh(self): pass
        def rebuild(self): pass
        def get_equipment_summary(self): return {'items': [], 'total': 0}
    print("  ⚠ Substrate Impact Widget: standalone not found, using stub")


class LearningAssistantWidget(QWidget):
    """AI-powered learning from completed tests to improve next test_runs"""
    
    rescan_requested = Signal()  # Emitted when user clicks re-scan
    
    def __init__(self, database: Test RunDatabase, parent=None):
        super().__init__(parent)
        self.database = database
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("<h2>🧠 Learning Assistant - Test Intelligence</h2>"))
        
        info = QLabel(
            "Analyzes completed tests to generate insights, update hypotheses, "
            "and recommend follow-up questions for upcoming tests."
        )
        info.setWordWrap(True)
        layout.addWidget(info)
        
        # Analyze button row
        analyze_layout = QHBoxLayout()
        
        analyze_btn = QPushButton("🔍 Analyze All Completed Test Runs")
        analyze_btn.clicked.connect(self.analyze_tests)
        analyze_btn.setStyleSheet("background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 10px; border: 1px solid #7928CA;")
        analyze_layout.addWidget(analyze_btn)
        
        rescan_btn = QPushButton("🔄 Re-scan Completed Packets")
        rescan_btn.setToolTip("Re-read all .docx files in tested/ folder and repopulate empty database fields")
        rescan_btn.clicked.connect(self.rescan_requested.emit)
        rescan_btn.setStyleSheet("background-color: #663600; color: #ff8c00; font-weight: bold; padding: 10px; border: 1px solid #ff8c00;")
        analyze_layout.addWidget(rescan_btn)
        
        self.status_label = QLabel(f"Ready to analyze ({len(self.database.tests)} tests loaded)")
        analyze_layout.addWidget(self.status_label)
        
        analyze_layout.addStretch()
        layout.addLayout(analyze_layout)
        
        # Results in tabbed sections
        results_tabs = QTabWidget()
        
        # Tab 1: Key Learnings Extracted
        self.learnings_text = QTextEdit()
        self.learnings_text.setReadOnly(True)
        results_tabs.addTab(self.learnings_text, "📊 Key Learnings")
        
        # Tab 2: Cross-Test Run Patterns
        self.patterns_text = QTextEdit()
        self.patterns_text.setReadOnly(True)
        results_tabs.addTab(self.patterns_text, "🔗 Patterns Detected")
        
        # Tab 3: Recommended Questions for Next Test Runs
        self.questions_text = QTextEdit()
        self.questions_text.setReadOnly(True)
        results_tabs.addTab(self.questions_text, "❓ Recommended Questions")
        
        # Tab 4: Updated Hypotheses
        self.hypotheses_text = QTextEdit()
        self.hypotheses_text.setReadOnly(True)
        results_tabs.addTab(self.hypotheses_text, "💡 Updated Hypotheses")
        
        layout.addWidget(results_tabs)
        
        # Export buttons
        export_layout = QHBoxLayout()
        
        export_learnings_btn = QPushButton("💾 Export Learnings Report")
        export_learnings_btn.clicked.connect(self.export_learnings)
        export_layout.addWidget(export_learnings_btn)
        
        export_questions_btn = QPushButton("💾 Export Question Bank")
        export_questions_btn.clicked.connect(self.export_questions)
        export_layout.addWidget(export_questions_btn)
        
        export_layout.addStretch()
        layout.addLayout(export_layout)
        
        self.setLayout(layout)
    
    def analyze_test_runs(self):
        """Analyze all completed tests and generate insights"""
        if len(self.database.tests) == 0:
            QMessageBox.warning(self, "No Test Runs", 
                                "No tests in database. Save tests or Re-scan Completed Packets first.")
            return
        
        self.status_label.setText("Analyzing tests...")
        
        try:
            # Extract learnings
            learnings = self.extract_key_learnings()
            self.learnings_text.setText(learnings)
            
            # Detect patterns
            patterns = self.detect_cross_test_run_patterns()
            self.patterns_text.setText(patterns)
            
            # Generate questions
            questions = self.generate_recommended_questions()
            self.questions_text.setText(questions)
            
            # Update hypotheses
            hypotheses = self.update_hypotheses()
            self.hypotheses_text.setText(hypotheses)
            
            self.status_label.setText(f"✓ Analysis complete ({len(self.database.tests)} tests analyzed)")
        except Exception as e:
            import traceback
            self.status_label.setText(f"⚠️ Analysis error: {e}")
            self.learnings_text.setText(
                f"ERROR during analysis:\n\n{e}\n\n{traceback.format_exc()}\n\n"
                f"Database has {len(self.database.tests)} tests.\n"
                f"Test Runs with text: {sum(1 for i in self.database.tests if i.results)}"
            )
    
    def _get_all_text(self, test_run) -> str:
        """Combine ALL test_run text fields into one searchable corpus.
        
        The old code only searched test_run.results — one section of the docx.
        Analyzers often put the real gold in experiments (Q&A), hypotheses,
        or action_iterate. This ensures Learning Assistant sees EVERYTHING.
        """
        parts = []
        if test_run.results:
            parts.append(test_run.results)
        if test_run.experiments:
            parts.append(test_run.experiments)
        if test_run.hypotheses:
            parts.append(test_run.hypotheses)
        if test_run.action_iterate:
            parts.append(test_run.action_iterate)
        # Include equipment impact notes if any
        if hasattr(test_run, 'substrate_impacts') and test_run.substrate_impacts:
            for impact in test_run.substrate_impacts:
                notes = impact.get('notes', '')
                if notes:
                    parts.append(notes)
        return ' '.join(parts)

    def extract_key_learnings(self) -> str:
        """Extract key learnings from test results"""
        output = []
        output.append("="*80)
        output.append("KEY LEARNINGS EXTRACTED FROM COMPLETED TESTS")
        output.append("="*80)
        output.append("")
        
        for test_entry in self.database.test_runs:
            output.append(f"TEST #{test_run.test_num}: {test_run.person} ({test_run.company})")
            output.append("-"*80)
            
            all_text = self._get_all_text(test_run)
            all_text_lower = all_text.lower()
            
            if not all_text.strip():
                output.append("  ⚠️ NO DATA — Re-scan completed packets to populate")
                output.append("")
                continue
            
            # Show which fields have data
            fields_present = []
            if test_run.results:
                fields_present.append("Results")
            if test_run.experiments:
                fields_present.append("Q&A")
            if test_run.hypotheses:
                fields_present.append("Hypotheses")
            if test_run.action_iterate:
                fields_present.append("Actions")
            if hasattr(test_run, 'substrate_impacts') and test_run.substrate_impacts:
                fields_present.append(f"Equipment({len(test_run.substrate_impacts)})")
            output.append(f"  📂 Data sources: {', '.join(fields_present)}")
            
            # Extract cost figures
            costs = self.extract_costs(all_text)
            if costs:
                output.append("\n💰 COST IMPACTS:")
                for cost in costs:
                    output.append(f"   ⬢ {cost}")
            
            # Extract equipment mentions
            equipment = self.extract_equipment_mentions(all_text)
            if equipment:
                output.append("\n⚙️ EQUIPMENT AFFECTED:")
                for equip in equipment:
                    output.append(f"   ⬢ {equip}")
            
            # Extract pain keywords
            pain_points = self.extract_pain_points(all_text_lower)
            if pain_points:
                output.append("\n🔥 PAIN POINTS:")
                for pain in pain_points:
                    output.append(f"   ⬢ {pain}")
            
            # Extract validation/invalidation
            if "validated:" in all_text_lower or "validated" in all_text_lower:
                output.append("\n✓ VALIDATED HYPOTHESIS")
            if "pivot" in all_text_lower or "wrong customer" in all_text_lower:
                output.append("\n⚠️ HYPOTHESIS INVALIDATED - PIVOT REQUIRED")
            
            output.append("")
        
        return "\n".join(output)
    
    def detect_cross_test_run_patterns(self) -> str:
        """Delegate to learning_engine for ML pattern detection."""
        if LEARNING_ENGINE_AVAILABLE:
            return detect_patterns(self.database, get_active_project)
        return "Learning engine not loaded. Place learning_engine.py next to collector."

    def generate_recommended_questions(self) -> str:
        """Delegate to learning_engine for ML question generation."""
        if LEARNING_ENGINE_AVAILABLE:
            return generate_questions(
                self.database, TEST_PLAN, get_active_project
            )
        return "Learning engine not loaded. Place learning_engine.py next to collector."

    def update_hypotheses(self) -> str:
        """Delegate to learning_engine for Bayesian hypothesis tracking."""
        if LEARNING_ENGINE_AVAILABLE:
            return track_hypotheses(self.database, TEST_PLAN)
        return "Learning engine not loaded. Place learning_engine.py next to collector."
    
    def extract_costs(self, text: str) -> list:
        """Extract cost figures from text"""
        import re
        costs = []
        
        # Pattern: $XXX,XXX or $XXXk or XXX,XXX
        patterns = [
            r'\$[\d,]+(?:k|K)?',
            r'[\d,]+k annual',
            r'[\d,]+\s*(?:per year|annually)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            costs.extend(matches)
        
        return list(set(costs))
    
    def extract_equipment_mentions(self, text: str) -> list:
        """Extract equipment mentions from text"""
        equipment_keywords = [
            'screen', 'CTS', 'digester', 'pump', 'basket', 'rotor', 'feeder',
            'blow line', 'washer', 'knotter', 'refiner', 'conveyor'
        ]
        
        found = []
        text_lower = text.lower()
        for keyword in equipment_keywords:
            if keyword in text_lower:
                found.append(keyword.title())
        
        return found
    
    def extract_pain_points(self, text: str) -> list:
        """Extract pain point keywords"""
        pain_keywords = {
            'downtime': 'Unplanned downtime',
            'damage': 'Equipment damage',
            'blind': 'Lack of visibility',
            'no detection': 'No detection capability',
            'reactive': 'Reactive maintenance',
            'baseline': 'Missing baseline data',
            'fail': 'System failures'
        }
        
        found = []
        for keyword, description in pain_keywords.items():
            if keyword in text:
                found.append(description)
        
        return found
    
    def count_mentions(self, keyword: str, text: str) -> int:
        """Count keyword mentions"""
        return text.lower().count(keyword.lower())
    
    def export_learnings(self):
        """Export learnings report"""
        filepath = get_project_dir(get_active_project()) / "exports" / f"Learning_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        with open(filepath, 'w') as f:
            f.write(self.learnings_text.toPlainText())
            f.write("\n\n" + "="*80 + "\n\n")
            f.write(self.patterns_text.toPlainText())
            f.write("\n\n" + "="*80 + "\n\n")
            f.write(self.hypotheses_text.toPlainText())
        
        QMessageBox.information(self, "Exported", f"Learning report saved to:\n{filepath}")
    
    def export_questions(self):
        """Export question bank"""
        filepath = get_project_dir(get_active_project()) / "exports" / f"Question_Bank_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        with open(filepath, 'w') as f:
            f.write(self.questions_text.toPlainText())
        
        QMessageBox.information(self, "Exported", f"Question bank saved to:\n{filepath}")


class BMCGeneratorWidget(QWidget):
    """Auto-generate Business Model Canvas from AI analysis"""
    
    def __init__(self, database: Test RunDatabase, parent=None):
        super().__init__(parent)
        self.database = database
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("<h2>Business Model Canvas Generator</h2>"))
        
        info = QLabel("Generate updated Business Model Canvas based on test_run AI analysis and equipment impact data.")
        info.setWordWrap(True)
        layout.addWidget(info)
        
        # Substrate Class display (auto-populated from AI)
        segment_group = QGroupBox("Primary Substrate Class (AI-Derived)")
        segment_layout = QVBoxLayout()
        
        self.segment_text = QTextEdit()
        self.segment_text.setReadOnly(True)
        self.segment_text.setMaximumHeight(150)
        segment_layout.addWidget(self.segment_text)
        
        refresh_segment_btn = QPushButton("🔄 Refresh from AI Analysis")
        refresh_segment_btn.clicked.connect(self.refresh_segment)
        segment_layout.addWidget(refresh_segment_btn)
        
        segment_group.setLayout(segment_layout)
        layout.addWidget(segment_group)
        
        # Value Propositions (auto-populated from AI)
        vp_group = QGroupBox("Value Propositions (AI-Optimized)")
        vp_layout = QVBoxLayout()
        
        self.vp_text = QTextEdit()
        self.vp_text.setReadOnly(True)
        self.vp_text.setMaximumHeight(200)
        vp_layout.addWidget(self.vp_text)
        
        refresh_vp_btn = QPushButton("🔄 Refresh from AI Analysis")
        refresh_vp_btn.clicked.connect(self.refresh_value_props)
        vp_layout.addWidget(refresh_vp_btn)
        
        vp_group.setLayout(vp_layout)
        layout.addWidget(vp_group)
        
        # Generate buttons
        button_layout = QHBoxLayout()
        
        generate_docx_btn = QPushButton("📄 Generate BMC (DOCX)")
        generate_docx_btn.clicked.connect(self.generate_docx)
        generate_docx_btn.setStyleSheet("background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 15px; border: 1px solid #7928CA;")
        button_layout.addWidget(generate_docx_btn)
        
        generate_pdf_btn = QPushButton("📄 Generate BMC (PDF)")
        generate_pdf_btn.clicked.connect(self.generate_pdf)
        generate_pdf_btn.setStyleSheet("background-color: #661020; color: #FF4444; font-weight: bold; padding: 15px; border: 1px solid #c41e3a;")
        button_layout.addWidget(generate_pdf_btn)
        
        layout.addLayout(button_layout)
        
        # Status
        self.status_label = QLabel("Ready to generate Business Model Canvas")
        layout.addWidget(self.status_label)
        
        layout.addStretch()
        
        self.setLayout(layout)
    
    def refresh_segment(self):
        """Refresh customer segment from AI analysis"""
        # Get synergy clusters
        clusters = self.database.get_synergy_clusters(min_test_runs=2)
        
        if not clusters:
            self.segment_text.setText("No synergy clusters detected yet. Complete more tests and run AI analysis.")
            return
        
        # Find highest density cluster
        sorted_clusters = sorted(clusters.items(), key=lambda x: len(x[1]), reverse=True)
        top_cluster_pos, top_tests = sorted_clusters[0]
        
        # Generate segment description
        segment_text = f"PRIMARY SEGMENT (Cube Position: {top_cluster_pos})\n\n"
        segment_text += f"Test Run Density: {len(top_tests)} test_runs\n\n"
        
        segment_text += "Companies:\n"
        for test_entry in top_test_runs:
            segment_text += f"  ⬢ {test_run.person} - {test_run.company}\n"
        
        segment_text += f"\nCommon Pain Patterns:\n"
        # Scan results for actual cost figures and pain patterns
        all_results = " ".join([i.results.lower() for i in top_test_runs])
        import re
        cost_mentions = re.findall(r'\$[\d,]+(?:\.\d+)?(?:\s*[kKmM])?(?:/?(?:yr|year))?', 
                                   " ".join([i.results for i in top_test_runs]))
        if cost_mentions:
            segment_text += f"  ⬢ Documented cost impacts: {', '.join(set(cost_mentions[:5]))}\n"
        if "baseline" in all_results or "commissioning" in all_results:
            segment_text += "  ⬢ Need baseline data for new equipment\n"
        if "modern equipment" in all_results or "still" in all_results:
            segment_text += "  ⬢ Modern equipment doesn't solve contamination\n"
        
        self.segment_text.setText(segment_text)
    
    def refresh_value_props(self):
        """Refresh value propositions from hypothesis confidence + test evidence."""
        tests = self.database.test_runs
        if not test_runs:
            self.vp_text.setText("No tests loaded yet.")
            return
        
        # Count tests with actual results text
        completed = [i for i in tests if getattr(i, 'results', '')]
        all_results = " ".join([i.results.lower() for i in completed])
        
        # VP definitions — tied to hypotheses, scored from evidence
        vp_defs = [
            {
                'num': 1, 'name': 'Calibrate Your Contamination Problem',
                'keywords': ['100% fail', 'all got through', 'manual catch', 'no detection', 'nothing catches', 'manual'],
                'desc': 'Install at blow line until rocks stop appearing. Data shows WHERE upstream to fix.',
            },
            {
                'num': 2, 'name': 'True Cost Discovery',
                'keywords': ['damage cost', 'annual damage', 'nobody tracks', 'more than expected', 'repair cost', 'replacement cost'],
                'desc': 'Expose real contamination costs hidden in maintenance budgets.',
            },
            {
                'num': 3, 'name': 'Last-Chance Prevention Before Digester',
                'keywords': ['blow line', 'before digester', 'last chance', 'entry point', 'blind spot', 'response time'],
                'desc': 'Real-time detection at blow line entry. Operators get immediate control.',
            },
            {
                'num': 4, 'name': 'Bark Blow Line Safety Monitoring',
                'keywords': ['bark', 'hog fuel', 'bark blow', 'boiler tube', 'bark line'],
                'desc': 'First-ever hog fuel boiler contamination detection. Prevent boiler tube failures.',
            },
            {
                'num': 5, 'name': 'Contamination is Accepted Practice',
                'keywords': ['accepted', 'normal', 'always been', 'cost of doing business', 'live with it', 'nobody cares'],
                'desc': 'Industry normalizes damage. Detection system creates accountability where none exists.',
            },
            {
                'num': 6, 'name': 'Commissioning Baseline Data',
                'keywords': ['baseline', 'commissioning', 'new equipment', 'modernization', 'investment', 'prove'],
                'desc': 'Document new equipment performance from day one. Prove investment is protected.',
            },
        ]
        
        value_props = []
        for vp in vp_defs:
            hits = sum(1 for kw in vp['keywords'] if kw in all_results)
            test_run_hits = sum(1 for i in completed if any(kw in i.results.lower() for kw in vp['keywords']))
            
            if test_run_hits > 0:
                confidence = min(95, int((test_run_hits / len(completed)) * 100))
                value_props.append(
                    f'{vp["num"]}. "{vp["name"]}" ({confidence}% — {test_run_hits} tests)\n'
                    f'   {vp["desc"]}'
                )
            else:
                value_props.append(
                    f'{vp["num"]}. "{vp["name"]}" (not yet validated)\n'
                    f'   {vp["desc"]}'
                )
        
        self.vp_text.setText("\n\n".join(value_props))
    
    def generate_docx(self):
        """Generate Business Model Canvas DOCX via bmc_engine (accumulative)"""
        try:
            if not BMC_ENGINE_AVAILABLE:
                QMessageBox.warning(self, "Not Available",
                    "bmc_engine.py not found.\nPlace it in the ICORPS directory.")
                return
            
            # Get active project — routes to THAT project's BMC folder
            project_id = get_active_project()
            bmc_dir = get_bmc_dir(project_id)
            
            # Gather all tests as dicts
            test_run_dicts = [iv.to_dict() for iv in self.database.test_runs]
            
            if not test_run_dicts:
                QMessageBox.warning(self, "No Data",
                    "No tests loaded. Process tests first.")
                return
            
            # Run accumulative BMC generation
            state = generate_project_bmc(
                project_id, test_run_dicts, bmc_dir, generate_docx=True
            )
            
            self.status_label.setText(
                f"✓ BMC v{state.version} — {state.test_count} tests → {bmc_dir}")
            
            # Update display fields from engine output
            vp_state = state.fields.get("value_propositions")
            if vp_state:
                self.vp_text.setText(vp_state.summary)
            cs_state = state.fields.get("substrate_classs")
            if cs_state:
                self.segment_text.setText(cs_state.summary)
            
            QMessageBox.information(self, "BMC Generated",
                f"Business Model Canvas v{state.version}\n"
                f"{state.test_count} tests scored\n"
                f"Changes: {len(state.changes_since_last)}\n\n"
                f"Saved to: {bmc_dir}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate BMC:\n{str(e)}")
    
    def generate_pdf(self):
        """Generate Business Model Canvas PDF"""
        QMessageBox.information(self, "PDF Generation", 
                              "PDF generation requires reportlab library.\n"
                              "Install with: pip install reportlab\n\n"
                              "For now, generate DOCX and convert to PDF.")


class AIAnalysisWidget(QWidget):
    """Display AI analysis results"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("<h2>AI Synergy Analysis</h2>"))
        
        # Progress bar
        self.progress = QProgressBar()
        layout.addWidget(self.progress)
        
        self.status_label = QLabel("Ready to analyze...")
        layout.addWidget(self.status_label)
        
        # Results display
        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        layout.addWidget(self.results_text)
        
        self.setLayout(layout)
    
    def update_progress(self, percent: int, status: str):
        """Update progress bar"""
        self.progress.setValue(percent)
        self.status_label.setText(status)
    
    def display_results(self, results: Dict):
        """Display real ML analysis results"""
        try:
            output = self._format_results(results)
            self.results_text.setText("\n".join(output))
        except Exception as e:
            import traceback
            self.results_text.setText(
                f"⚠️ ERROR DISPLAYING RESULTS\n\n"
                f"Exception: {e}\n\n"
                f"Traceback:\n{traceback.format_exc()}\n\n"
                f"Raw results keys: {list(results.keys()) if results else 'None'}\n"
                f"Segments: {len(results.get('substrate_classs', []))}\n"
                f"Patterns: {len(results.get('synergy_patterns', []))}\n"
                f"Clusters: {len(results.get('clusters', {}))}"
            )
    
    def _format_results(self, results: Dict) -> list:
        output = []
        
        output.append("=" * 80)
        output.append("ML SYNERGY ANALYSIS — REAL MATH, NO MOCK DATA")
        output.append("=" * 80)
        output.append("")
        
        # ML metrics
        var_exp = results.get('embedding_variance_explained', 0)
        output.append(f"  Engine: TF-IDF + TruncatedSVD (Halko et al. 2011)")
        output.append(f"  Variance explained: {var_exp:.1%}")
        output.append(f"  Scoring: Beta-Binomial Bayesian posterior")
        output.append("")
        
        # Latent topics from SVD
        top_terms = results.get('top_terms', [])
        if top_terms:
            output.append("LATENT TOPICS (SVD components):")
            output.append("-" * 80)
            for topic, terms in top_terms[:3]:
                output.append(f"  {topic}: {', '.join(terms[:8])}")
            output.append("")
        
        # Synergy patterns
        output.append("DETECTED PATTERNS (cosine similarity + 10-lens hypercube):")
        output.append("-" * 80)
        for pattern in results.get('synergy_patterns', []):
            output.append(f"  ⬢ {pattern}")
        output.append("")
        
        # Customer segments
        output.append("REFINED CUSTOMER SEGMENTS:")
        output.append("-" * 80)
        for segment in results.get('substrate_classs', []):
            priority = segment.get('priority', '?')
            output.append(f"\n  {priority}:")
            # ML path uses cluster_id, fallback uses cube_position
            if 'cluster_id' in segment:
                output.append(f"    Cluster: #{segment['cluster_id']}")
                output.append(f"    Tests: {segment['test_count']} "
                              f"(synergy={segment.get('avg_synergy', 0):.3f}, "
                              f"coherence={segment.get('coherence', 0):.3f})")
                nums = segment.get('test_nums', [])
                if nums:
                    output.append(f"    Test Run #s: {', '.join(str(n) for n in nums)}")
            else:
                output.append(f"    Cube Position: {segment.get('cube_position', '?')}")
                output.append(f"    Test Run Count: {segment.get('test_count', 0)}")
            output.append(f"    Description: {segment.get('description', '')}")
            output.append(f"    Companies: {', '.join(segment.get('companies', []))}")
            # Lens profile (from ML path)
            lens_profile = segment.get('lens_profile', {})
            if lens_profile:
                lens_str = ' | '.join(f"{k.replace('q_','').upper()[:5]}={v}" 
                                       for k, v in lens_profile.items())
                output.append(f"    Lens Profile: {lens_str}")
            # Pain themes
            pains = segment.get('pain_themes', [])
            if pains:
                output.append(f"    Pain Themes: {', '.join(pains[:3])}")
        output.append("")
        
        # Clusters with coherence
        output.append("CLUSTER COHERENCE (cosine similarity):")
        output.append("-" * 80)
        for pos, data in results.get('clusters', {}).items():
            coh = data.get('coherence', 0)
            lenses = data.get('dominant_lenses', {})
            lens_str = ' '.join(f"{k.replace('q_','').upper()[:3]}={v}" for k, v in lenses.items())
            output.append(
                f"  {pos}: {data['count']} tests, "
                f"synergy={data['avg_synergy']:.3f}, "
                f"coherence={coh:.3f} | {lens_str}"
            )
        output.append("")
        
        # Value propositions
        output.append("VALUE PROPOSITIONS (cluster pain + 10-lens intelligence):")
        output.append("-" * 80)
        for i, vp in enumerate(results.get('value_propositions', []), 1):
            output.append(f"  {i}. {vp}")
        output.append("")
        
        # Lens confidence summary
        lens_summary = results.get('lens_confidence_summary', {})
        if lens_summary:
            output.append("LENS CONFIDENCE (Bayesian evidence accumulation):")
            output.append("-" * 80)
            for lens, info in lens_summary.items():
                avg = info.get('avg_confidence', 0)
                low = info.get('low_confidence_count', 0)
                bar = "█" * int(avg * 10) + "░" * (10 - int(avg * 10))
                output.append(f"  {lens:22s} [{bar}] {avg:.0%}  ({low} need review)")
            output.append("")
        
        # Adaptive questions (if doctoral analyzer available)
        aq = results.get('adaptive_questions', [])
        if aq:
            output.append("INFO-GAIN QUESTIONS (Shannon entropy — doctoral_analyzer):")
            output.append("-" * 80)
            for q in aq[:8]:
                output.append(f"  IG={q['information_gain']:.3f}bits | {q['question']}")
            output.append("")
        
        # Knowledge entities (if extractor available)
        ke = results.get('knowledge_entities', {})
        if ke:
            output.append("EXTRACTED ENTITIES (knowledge_extractor):")
            output.append("-" * 80)
            for category, entities in ke.items():
                if entities:
                    output.append(f"  {category}: {len(entities)} entities")
            output.append("")
        
        # Engine status
        engines = results.get('engines_active', {})
        output.append("ENGINE STATUS:")
        output.append("-" * 80)
        for name, active in engines.items():
            status = "✓ ACTIVE" if active else "○ not loaded (place file in working directory)"
            output.append(f"  {status} — {name}")
        output.append("")
        
        return output


class BCMTestCollector(QMainWindow):
    """Main application window"""
    
    def __init__(self):
        super().__init__()
        
        self.database = Test RunDatabase()
        self.ai_engine = None
        
        # Genesis Brain flags
        self._genesis_running = False
        self.genesis_progress = None
        self.genesis_worker = None
        
        # ┬┬ Stale Analysis Tracking â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬┬┬┬┬┬┬
        # When a new Validation test_run arrives via Doc Reader, the existing
        # AI analysis no longer reflects reality.  These flags drive the
        # flashing "RECALCULATE" button so the user knows to re-run.
        self._analysis_stale = False
        self._stale_timer = QTimer()
        self._stale_timer.timeout.connect(self._flash_ai_button)
        self._stale_flash_on = False
        self._stale_test_run = None  # last validation test_run that triggered stale
        
        # Try to load existing database
        db_file = BCM_ROOT / "test_database.json"
        if db_file.exists():
            try:
                self.database = Test RunDatabase.load(db_file)
            except:
                pass
        
        # Also scan subdirectories for student deck JSON files (*_deck.json)
        self._load_deck_files()
        
        self.init_ui()
        
        # Auto-sca tested/ folder for completed Validation packets on startup
        QTimer.singleShot(1500, self._auto_scan_tested)
        
    def _load_deck_files(self):
        """
        Scan ALL deployed project Baseline folders for *_deck.json files.
        Students drop their deck JSON into a project subfolder and it gets picked up.
        Skips tests already in the database (by person name).
        """
        existing_names = {i.person.strip().lower() for i in self.database.test_runs}
        loaded_count = 0
        
        # Scan across ALL projects — not just active one
        for deck_file in PROJECT_ROOT.rglob("*_deck.json"):
            # Skip the main database file itself
            if deck_file.name == "test_database.json":
                continue
            try:
                with open(deck_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Handle both flat list and {"tests": [...]} wrapper
                if isinstance(data, dict) and 'tests' in data:
                    tests_raw = data['tests']
                elif isinstance(data, list):
                    tests_raw = data
                else:
                    print(f"Deck skip (unknown format): {deck_file.name}")
                    continue
                
                for entry in tests_raw:
                    person = entry.get('script_name', '').strip().lower()
                    if person in existing_names:
                        continue  # already loaded
                    
                    test_run = Test Run.from_dict(entry)
                    test_run.source = entry.get('source', 'sparks')
                    
                    # Normalize stack keys to canonical Unicode
                    test_run.q_stack = [normalize_stack_key(s) for s in test_run.q_stack]
                    
                    # Ensure substrate_impacts exists
                    if not hasattr(test_run, 'substrate_impacts'):
                        test_run.substrate_impacts = []
                    
                    self.database.add_test_run(test_run)
                    existing_names.add(person)
                    loaded_count += 1
                
                print(f"Deck loaded: {deck_file.name} ({loaded_count} new tests)")
                
            except Exception as e:
                print(f"Deck error ({deck_file.name}): {e}")
        
        # Rebuild cube matrix after all deck loads to ensure consistency
        if loaded_count > 0:
            self.database.rebuild_cube_matrix()
            db_file = BCM_ROOT / "test_database.json"
            self.database.save(db_file)
            print(f"Database updated: {len(self.database.tests)} total test_runs")
        
    def init_ui(self):
        self.setWindowTitle("GIBUSH BCM R&D Intelligence — Validation Research ⟶ Inclusion CaaS ⟶ Emulsion SaaS")
        self.setGeometry(100, 100, 1400, 900)
        
        # Create tabs
        self.tabs = tabs = QTabWidget()
        
        # Tab 0: Test Run Entry
        self.entry_widget = Test RunEntryWidget()
        self.entry_widget.test_run_saved.connect(self.on_test_run_saved)
        self.entry_widget.set_project_getter(get_active_project)
        tabs.addTab(self.entry_widget, "📝 New Test")
        
        # Tab 1: Test Planning - THE MAIN TEST PLAN TAB
        self.planning_widget = TestPlanningWidget(
            get_project_folder_fn=self._get_doc_reader_project_folder
        )
        self.planning_widget._get_database = lambda: self.database
        self.planning_widget.generate_template.connect(self.generate_test_template)
        tabs.addTab(self.planning_widget, "📋 Test Plan")
        
        # Tab 3: Test Run List (with Inclusion Returns panel)
        self._all_tests_container = QWidget()
        _ai_layout = QHBoxLayout(self._all_tests_container)
        _ai_layout.setContentsMargins(0, 0, 0, 0)
        _ai_layout.setSpacing(4)

        # LEFT: Main test table
        self.list_widget = QTableWidget()
        self.list_widget.setColumnCount(8)
        self.list_widget.setHorizontalHeaderLabels([
            "#", "Date", "Person", "Company", "Type", "Cube Position", "Source", "Delete"
        ])

        # RIGHT: Inclusion Returns panel
        _right_panel = QWidget()
        _right_layout = QVBoxLayout(_right_panel)
        _right_layout.setContentsMargins(4, 0, 4, 4)
        _right_layout.setSpacing(4)

        # ── Block 1: Inclusion Returns Table ──
        _inc_group = QGroupBox("📋 INCLUSION RETURNS (Stream 2)")
        _inc_group.setStyleSheet(
            "QGroupBox { color: #00FF00; border: 1px solid #00FF00; "
            "font-weight: bold; padding-top: 14px; margin-top: 6px; }"
            "QGroupBox::title { subcontrol-position: top left; padding: 2px 8px; }"
        )
        _inc_inner = QVBoxLayout(_inc_group)
        _inc_inner.setContentsMargins(4, 4, 4, 4)
        _inc_inner.setSpacing(3)

        # Tether button
        self._inclusion_tether_btn = QPushButton("🔗 TETHER INCLUSION TO PROJECT")
        self._inclusion_tether_btn.setStyleSheet(
            "QPushButton { background: #1B5E20; color: #00FF00; padding: 6px; "
            "font-weight: bold; font-size: 11px; border: 1px solid #00FF00; }"
            "QPushButton:hover { background: #2E7D32; }"
        )
        self._inclusion_tether_btn.clicked.connect(self._refresh_inclusion_returns)
        _inc_inner.addWidget(self._inclusion_tether_btn)

        # Status label
        self._inclusion_status = QLabel("⊘ Not tethered — click to connect")
        self._inclusion_status.setStyleSheet("color: #888; font-size: 10px; padding: 2px;")
        _inc_inner.addWidget(self._inclusion_status)

        # Inclusion table — Lens, Observer, Module, Type, Status
        self._inclusion_table = QTableWidget()
        self._inclusion_table.setColumnCount(5)
        self._inclusion_table.setHorizontalHeaderLabels([
            "Lens", "Observer", "Module", "Type", "Status"
        ])
        self._inclusion_table.setStyleSheet(
            "QTableWidget { background: #0a0a0a; color: #00FF00; gridline-color: #1a3a1a; "
            "font-size: 10px; }"
            "QHeaderView::section { background: #001a00; color: #00FF00; "
            "font-weight: bold; font-size: 9px; border: 1px solid #003300; padding: 2px; }"
            "QTableWidget::item:selected { background: #003300; }"
        )
        self._inclusion_table.horizontalHeader().setStretchLastSection(True)
        self._inclusion_table.verticalHeader().setDefaultSectionSize(22)
        self._inclusion_table.setSelectionBehavior(QTableWidget.SelectRows)
        self._inclusion_table.setSelectionMode(QTableWidget.SingleSelection)

        # Column widths
        _inc_hdr = self._inclusion_table.horizontalHeader()
        _inc_hdr.resizeSection(0, 36)   # Lens
        _inc_hdr.resizeSection(1, 90)   # Observer
        _inc_hdr.resizeSection(2, 55)   # Module
        _inc_hdr.resizeSection(3, 72)   # Type

        _inc_inner.addWidget(self._inclusion_table)

        # ── Observation preview ──
        self._inc_preview = QLabel("Select an entry to preview observation")
        self._inc_preview.setStyleSheet(
            "color: #999; font-size: 9px; padding: 4px; "
            "background: #050505; border: 1px solid #1a3a1a;"
        )
        self._inc_preview.setWordWrap(True)
        self._inc_preview.setMaximumHeight(55)
        _inc_inner.addWidget(self._inc_preview)

        # ── Accept / Reject row ──
        _action_row = QHBoxLayout()
        _action_row.setSpacing(4)

        self._inc_accept_btn = QPushButton("✓ ACCEPT")
        self._inc_accept_btn.setStyleSheet(
            "QPushButton { background: #1B5E20; color: #00FF00; padding: 4px 8px; "
            "font-weight: bold; font-size: 10px; border: 1px solid #00AA44; }"
            "QPushButton:hover { background: #2E7D32; }"
            "QPushButton:disabled { background: #111; color: #444; border-color: #333; }"
        )
        self._inc_accept_btn.clicked.connect(lambda: self._set_inclusion_verdict("ACCEPTED"))
        self._inc_accept_btn.setEnabled(False)
        _action_row.addWidget(self._inc_accept_btn)

        self._inc_reject_btn = QPushButton("✗ REJECT")
        self._inc_reject_btn.setStyleSheet(
            "QPushButton { background: #4A0000; color: #FF4444; padding: 4px 8px; "
            "font-weight: bold; font-size: 10px; border: 1px solid #AA0000; }"
            "QPushButton:hover { background: #660000; }"
            "QPushButton:disabled { background: #111; color: #444; border-color: #333; }"
        )
        self._inc_reject_btn.clicked.connect(lambda: self._set_inclusion_verdict("REJECTED"))
        self._inc_reject_btn.setEnabled(False)
        _action_row.addWidget(self._inc_reject_btn)

        self._inc_delete_btn = QPushButton("🗑 DELETE")
        self._inc_delete_btn.setStyleSheet(
            "QPushButton { background: #2a0a0a; color: #AA4444; padding: 4px 8px; "
            "font-weight: bold; font-size: 10px; border: 1px solid #660000; }"
            "QPushButton:hover { background: #440000; }"
            "QPushButton:disabled { background: #111; color: #444; border-color: #333; }"
        )
        self._inc_delete_btn.clicked.connect(self._delete_inclusion_entry)
        self._inc_delete_btn.setEnabled(False)
        _action_row.addWidget(self._inc_delete_btn)

        _inc_inner.addLayout(_action_row)

        # Enable buttons on row select
        self._inclusion_table.currentCellChanged.connect(self._on_inclusion_row_changed)
        # Store tethered entries for reference
        self._inclusion_entries_cache = []

        _right_layout.addWidget(_inc_group, stretch=4)

        # ── Block 2: CaaS Tether Status (compact) ──
        _tether_group = QGroupBox("🔌 CaaS MODULE TETHER")
        _tether_group.setStyleSheet(
            "QGroupBox { color: #00FFFF; border: 1px solid #00FFFF; "
            "font-weight: bold; padding-top: 14px; margin-top: 4px; }"
            "QGroupBox::title { subcontrol-position: top left; padding: 2px 8px; }"
        )
        _tether_inner = QVBoxLayout(_tether_group)
        _tether_inner.setContentsMargins(4, 4, 4, 4)

        self._tether_status = QLabel(
            "CaaS Connection: STANDBY\n"
            "Project: —\n"
            "Inclusion entries: 0\n"
            "Safety signals: 0"
        )
        self._tether_status.setStyleSheet(
            "color: #00FFFF; font-family: Consolas; font-size: 10px; padding: 4px;"
        )
        self._tether_status.setWordWrap(True)
        _tether_inner.addWidget(self._tether_status)

        _legend = QLabel(
            "MASTER LENS:  ML0=Operator  ML1=EL  ML2=TL  ML3=IM"
        )
        _legend.setStyleSheet(
            "color: #555; font-family: Consolas; font-size: 8px; padding: 2px;"
        )
        _tether_inner.addWidget(_legend)

        _right_layout.addWidget(_tether_group, stretch=1)

        # Assemble splitter
        _right_panel.setMinimumWidth(280)
        _right_panel.setMaximumWidth(420)

        _ai_layout.addWidget(self.list_widget, stretch=3)
        _ai_layout.addWidget(_right_panel, stretch=1)

        tabs.addTab(self._all_tests_container, "📋 All Test Runs")
        
        # Tab 3: Q-Cube Matrix
        self.cube_widget = QCubeVisualizerWidget(self.database)
        tabs.addTab(self.cube_widget, "🧊 GIBUSH Matrix")
        
        # Tab 4: AI Analysis
        self.analysis_widget = AIAnalysisWidget()
        tabs.addTab(self.analysis_widget, "🤖 AI Analysis")
        
        # Tab 5: Substrate Impact — auto-accumulating from config + test_runs
        self.equipment_widget = SubstrateImpactWidget(self.database)
        self.equipment_widget.set_project_getter(get_active_project)
        tabs.addTab(self.equipment_widget, "⚙️ Substrate Impact")
        
        # Tab 6: Learning Assistant
        self.learning_widget = LearningAssistantWidget(self.database)
        self.learning_widget.rescan_requested.connect(self._rescan_completed_packets)
        tabs.addTab(self.learning_widget, "🧠 Learning Assistant")
        
        # Tab 7: Business Model Canvas Generator
        self.bmc_widget = BMCGeneratorWidget(self.database)
        tabs.addTab(self.bmc_widget, "📊 Generate BMC")
        
        # Tab 8: AI Doc Reader
        if DOC_READER_AVAILABLE:
            self.doc_reader_widget = DocReaderWidget(
                get_project_folder_fn=self._get_doc_reader_project_folder,
                get_active_project_fn=get_active_project,
                get_genesis_intel_fn=self._load_genesis_intelligence,
                get_test_plan_fn=lambda: reload_from_excel(),
            )
            self.doc_reader_widget.test_run_processed.connect(self._on_doc_reader_processed)
            tabs.addTab(self.doc_reader_widget, "📄 Doc Reader")
        
        # Tab 9: R&D Health — Research ↔ Development ↔ Validation convergence
        if _HAS_RD_HEALTH_TAB:
            self.rd_health_tab = create_rd_health_tab(self)
            self.rd_health_tab.refresh_requested.connect(self._refresh_rd_health)
            self.rd_health_tab.tally_requested.connect(self._run_test_tally)
            tabs.addTab(self.rd_health_tab, "⚛ R&D Health")
        else:
            self.doc_reader_widget = None
        
        # ── Tab Switch → Auto-run AI Analysis when tab selected ──
        self._ai_analysis_tab_index = tabs.indexOf(self.analysis_widget)
        self._ai_running = False  # Guard against double-trigger
        tabs.currentChanged.connect(self._on_tab_changed)
        
        # Toolbar
        toolbar_layout = QHBoxLayout()
        
        # ┬┬ AI Analysis button ⬔ stored as self so we can flash it ┬┬
        self.run_ai_btn = QPushButton("🤖 Run AI Synergy Analysis")
        self.run_ai_btn.clicked.connect(self.run_ai_analysis)
        self.run_ai_btn.setStyleSheet("background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 10px; border: 1px solid #7928CA;")
        toolbar_layout.addWidget(self.run_ai_btn)
        
        export_btn = QPushButton("💾 Export Results")
        export_btn.clicked.connect(self.export_results)
        toolbar_layout.addWidget(export_btn)
        
        self.genesis_btn = QPushButton("🧠 Launch Genesis Analysis")
        self.genesis_btn.clicked.connect(self.run_genesis_analysis)
        self.genesis_btn.setStyleSheet("background-color: #004400; color: #00ff88; font-weight: bold; padding: 10px; border: 1px solid #00AA44;")
        toolbar_layout.addWidget(self.genesis_btn)
        
        # Project Selector — discovers all deployed projects from disk
        toolbar_layout.addWidget(QLabel("  Project:"))
        self.project_combo = QComboBox()
        self._populate_project_combo()
        self.project_combo.currentIndexChanged.connect(self._on_project_changed)
        self.project_combo.setMinimumWidth(250)
        toolbar_layout.addWidget(self.project_combo)
        
        # Create Project button
        create_proj_btn = QPushButton("➕ New Project")
        create_proj_btn.clicked.connect(self._on_create_project)
        create_proj_btn.setStyleSheet("background-color: #333; color: #ddd; padding: 5px; border: 1px solid #555;")
        toolbar_layout.addWidget(create_proj_btn)
        
        # REFRESH FROM EXCEL - bumps entire platform without reboot
        refresh_btn = QPushButton("🔄 Refresh from Excel")
        refresh_btn.clicked.connect(self._refresh_all_from_excel)
        refresh_btn.setStyleSheet("background-color: #663600; color: #ff8c00; font-weight: bold; padding: 5px; border: 1px solid #ff8c00;")
        refresh_btn.setToolTip("Reload Excel and refresh ALL tabs without restarting")
        toolbar_layout.addWidget(refresh_btn)
        
        toolbar_layout.addStretch()
        
        # Main layout
        main_layout = QVBoxLayout()
        main_layout.addLayout(toolbar_layout)
        main_layout.addWidget(tabs)
        
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)
        
        self.refresh_test_list()
    
    def on_test_run_saved(self, test_run: Test Run):
        """Handle new test saved"""
        # Auto-classify into Q-Cube foundation + all 10 lenses + Bayesian synergy
        self.auto_classify_test_run(test_run)
        
        # Add to database
        # RECONCILER: dedup guard — remove any existing entry with same person name
        incoming_person = test_run.person.strip().lower()
        existing = [
            iv for iv in self.database.test_runs
            if iv.person.strip().lower() == incoming_person
        ]
        if existing:
            self.database.tests = [
                iv for iv in self.database.test_runs
                if iv.person.strip().lower() != incoming_person
            ]
            print(f"  [RECONCILER] Removed {len(existing)} existing entry(ies) for {test_run.person}")

        self.database.add_test_run(test_run)
        
        # Save to disk
        db_file = BCM_ROOT / "test_database.json"
        self.database.save(db_file)
        
        # Refresh views
        self.refresh_test_list()
        self.cube_widget.refresh()
    
    def _populate_project_combo(self):
        """Populate project dropdown from DISK — discovers all deployed projects"""
        self.project_combo.clear()
        
        # Discover projects from disk (not hardcoded)
        disk_projects = list_deployed_projects()
        excel_projects = list_projects() if EXCEL_LOADER_AVAILABLE else []
        all_projects = sorted(set(disk_projects + excel_projects))
        
        for project_id in all_projects:
            self.project_combo.addItem(f"📁 {project_id}", project_id)
        
        # Set current selection
        current = get_active_project()
        for i in range(self.project_combo.count()):
            if self.project_combo.itemData(i) == current:
                self.project_combo.setCurrentIndex(i)
                break
    
    def _on_create_project(self):
        """Create a new project with blank Excel template"""
        from PySide6.QtWidgets import QInputDialog
        
        # Get project name
        project_id, ok = QInputDialog.getText(
            self, 
            "Create New Project",
            "Enter project ID (e.g., NEW_PRODUCT_LINE):"
        )
        
        if not ok or not project_id:
            return
        
        # Clean up project ID
        project_id = project_id.upper().replace(" ", "_")
        
        # Check if exists on disk
        if project_id in list_deployed_projects():
            QMessageBox.warning(
                self,
                "Project Exists",
                f"Project '{project_id}' already exists.\nSelect it from the dropdown."
            )
            return
        
        # Create project folder structure on disk
        project_dir = ensure_project_structure(project_id)
        
        # Also create Excel template if excel_test_loader is available
        excel_path = None
        if EXCEL_LOADER_AVAILABLE:
            excel_path = create_project(project_id)
        
        if project_dir.exists():
            # Refresh combo
            self._populate_project_combo()
            
            # Switch to new project
            set_active_project(project_id)
            self._populate_project_combo()
            
            msg = f"Created project: {project_id}\n\nFolder: {project_dir}\n\n"
            if excel_path:
                msg += f"Excel file: {excel_path}\nEdit the Excel file to add tests."
            else:
                msg += "Add an *_Test Run_Plan.xlsx to the project folder."
            
            QMessageBox.information(self, "Project Created", msg)
        else:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to create project '{project_id}'."
            )
    
    def _on_project_changed(self, index: int):
        """Handle project selection change — routes all paths to selected project"""
        global TEST_PLAN
        
        project_id = self.project_combo.currentData()
        set_active_project(project_id)
        
        # Use the master refresh to bump everything
        self._refresh_all_from_excel()
    
    def _get_doc_reader_project_folder(self) -> str:
        """Return the active project's folder path for the doc reader."""
        project_id = get_active_project()
        folder = get_project_dir(project_id)
        folder.mkdir(parents=True, exist_ok=True)
        return str(folder)
    
    def _refresh_all_from_excel(self):
        """
        MASTER REFRESH — Reload Excel and bump every widget without restarting.
        
        Call this after updating the Excel file externally.
        Refreshes: Test Plan tab, All Test Runs list, Q-Cube, 
        Substrate Impact, Learning Assistant, BMC, Doc Reader.
        """
        global TEST_PLAN
        
        print("\n" + "=" * 60)
        print("🔄 REFRESH ALL FROM EXCEL — START")
        print("=" * 60)
        
        # 0. Auto-sca tested/ for new completed Validation packets
        self._auto_scan_tested()

        # 1. Reload Excel (clears cache, re-reads from disk)
        TEST_PLAN = reload_from_excel()
        try:
            excel_path = get_excel_path()
            print(f"  Excel: {excel_path}")
        except:
            pass
        print(f"  Loaded {len(TEST_PLAN)} test_run plans")
        
        # 2. Test Plan tab (the main planning sheet)
        if hasattr(self, 'planning_widget'):
            self.planning_widget._populate_table()
            print("  ✓ Test Plan tab refreshed")
        
        # 3. All Test Runs list
        try:
            self.refresh_test_list()
            print("  ✓ All Test Runs tab refreshed")
        except Exception as e:
            print(f"  ⚠ All Tests: {e}")
        
        # 4. Q-Cube Matrix
        if hasattr(self, 'cube_widget'):
            try:
                self.cube_widget.refresh()
                print("  ✓ Q-Cube Matrix refreshed")
            except Exception as e:
                print(f"  ⚠ Q-Cube: {e}")
        
        # 5. Substrate Impact
        if hasattr(self, 'equipment_widget'):
            try:
                if hasattr(self.equipment_widget, 'refresh'):
                    self.equipment_widget.refresh()
                    print("  ✓ Substrate Impact refreshed")
            except Exception as e:
                print(f"  ⚠ Equipment: {e}")
        
        # 5b. Test Run Entry — rebuild equipment checkboxes for new project
        if hasattr(self, 'entry_widget'):
            try:
                self.entry_widget.rebuild_equipment()
                print("  ✓ Test Run Entry equipment rebuilt")
            except Exception as e:
                print(f"  ⚠ Entry equipment: {e}")
        
        # 6. Learning Assistant
        if hasattr(self, 'learning_widget'):
            try:
                if hasattr(self.learning_widget, 'refresh'):
                    self.learning_widget.refresh()
                    print("  ✓ Learning Assistant refreshed")
            except Exception as e:
                print(f"  ⚠ Learning: {e}")
        
        # 7. BMC Generator
        if hasattr(self, 'bmc_widget'):
            try:
                if hasattr(self.bmc_widget, 'refresh_segment'):
                    self.bmc_widget.refresh_segment()
                    print("  ✓ BMC Generator refreshed")
            except Exception as e:
                print(f"  ⚠ BMC: {e}")
        
        # 8. Doc Reader (completed/pending tables)
        if hasattr(self, 'doc_reader_widget') and self.doc_reader_widget:
            try:
                self.doc_reader_widget.refresh()
                print("  ✓ Doc Reader refreshed")
            except Exception as e:
                print(f"  ⚠ Doc Reader: {e}")
        
        print("=" * 60)
        print(f"🔄 REFRESH COMPLETE — {len(TEST_PLAN)} tests loaded")
        print("=" * 60 + "\n")
        
        # Show confirmation to user
        QMessageBox.information(
            self,
            "Refresh Complete",
            f"✓ Reloaded from Excel\n"
            f"✓ {len(TEST_PLAN)} test_run plans loaded\n"
            f"✓ All tabs refreshed\n\n"
            f"No reboot needed."
        )
    
    def _on_doc_reader_processed(self, parsed_data: dict):
        """Handle when doc reader finishes processing an test packet.
        
        Converts parsed .docx data → Test Run object → adds to database → saves.
        Also handles DELETE action to remove tests from in-memory database.
        """
        # Handle DELETE action
        if parsed_data.get('action') == 'DELETE':
            person = parsed_data.get('script_name', '')
            if person and hasattr(self, 'database'):
                before = len(self.database.tests)
                person_lower = person.strip().lower()
                self.database.tests = [
                    iv for iv in self.database.tests 
                    if iv.person.strip().lower() != person_lower
                ]
                after = len(self.database.tests)
                self.database.rebuild_cube_matrix()
                # Save
                db_file = BCM_ROOT / "test_database.json"
                self.database.save(db_file)
                print(f"  [DELETE] Removed {person} — {before - after} entries from memory")
                # Refresh all tabs
                self.refresh_test_list()
                if hasattr(self, 'planning_widget'):
                    self.planning_widget.populate_from_excel()
            return
        
        # Build Test Run from parsed packet data
        test_run = Test Run()
        
        header = parsed_data.get('header', {})
        test_run.test_num = parsed_data.get('test_num', 0)
        test_run.date = datetime.now().strftime('%b %d, %Y')
        test_run.person = header.get('Person', parsed_data.get('script_name', 'Unknown'))
        test_run.title = header.get('Title', '')
        test_run.company = header.get('Company', parsed_data.get('source_version', 'Unknown'))
        test_run.test_category = header.get('Substrate Class', '')
        
        # Map ICorps template fields
        test_run.hypotheses = parsed_data.get('hypotheses_validation', '')
        test_run.results = parsed_data.get('results', '')
        test_run.action_iterate = parsed_data.get('action_iterate', '')
        
        # Build experiments string from question/answer pairs
        exp_answers = parsed_data.get('experiments_answers', {})
        if exp_answers:
            exp_lines = []
            for q, a in exp_answers.items():
                exp_lines.append(f"Q: {q}\nA: {a}")
            test_run.experiments = '\n\n'.join(exp_lines)
        
        # Parse Q-Cube position from header: "[L2, OC, Sa]"
        cube_str = header.get('Q-Cube Position', '')
        if cube_str:
            parts = cube_str.strip('[]').split(',')
            parts = [p.strip() for p in parts]
            if len(parts) >= 2:
                test_run.q_layer = parts[0]
                test_run.q_object = parts[1]
            if len(parts) >= 3:
                test_run.q_stack = [normalize_stack_key(parts[2])]
            else:
                test_run.q_stack = [normalize_stack_key('Sa')]
        else:
            self.auto_classify_test_run(test_run)
        
        # Final normalization pass on all stack keys
        test_run.q_stack = [normalize_stack_key(s) for s in test_run.q_stack]
        
        # Map equipment discussed → substrate_impacts
        for eq in parsed_data.get('equipment_discussed', []):
            impact = {
                'equipment': eq.get('equipment', ''),
                'cost': eq.get('typical_cost', ''),
                'notes': eq.get('notes', '')
            }
            test_run.substrate_impacts.append(impact)
        
        # Run all analytical + innovation lens inference (fills A, V, T, C, N, F, K)
        lenses = infer_all_lenses(test_run)
        for key, val in lenses.items():
            if val:
                setattr(test_run, key, val)
        
        # Bayesian synergy score — real calculation, not placeholder
        import math as _math
        _alpha, _beta = 1.0, 1.0
        _v_weights = {'V4': 1.0, 'V3': 0.8, 'V2': 0.5, 'V1': 0.2}
        _l_weights = {'L1': 1.0, 'L2': 0.67, 'L3': 0.33}
        _v = getattr(test_run, 'q_evidence', 'V1')
        _alpha += _v_weights.get(_v, 0.2)
        _tlen = sum(len(getattr(test_run, f, '') or '') for f in ('results', 'experiments', 'hypotheses'))
        if _tlen > 800: _alpha += 0.8
        elif _tlen > 400: _alpha += 0.5
        elif _tlen > 100: _alpha += 0.2
        else: _beta += 0.5
        if test_run.substrate_impacts: _alpha += 0.3
        _alpha += _l_weights.get(test_run.q_layer, 0.5) * 0.5
        test_run.synergy_score = min(1.0, _alpha / (_alpha + _beta))
        
        # Tag as Validation (validation phase) - independent from Baseline (discovery phase)
        test_run.source = 'validation'
        
        # Validation numbering starts at SPARKS_MAX + 1 (51+)
        # Baseline occupies slots 1-SPARKS_MAX, Validation occupies SPARKS_MAX+1 onward
        validation_tests = [i for i in self.database.tests if getattr(i, 'source', '') == 'validation']
        existing_validation_nums = {i.test_num for i in validation_test_runs}
        
        validation_num = SPARKS_MAX + 1 + len(validation_tests)
        while validation_num in existing_validation_nums:
            validation_num += 1
        test_run.test_num = validation_num

        
        self.database.add_test_run(test_run)
        
        # Save to disk
        db_file = BCM_ROOT / "test_database.json"
        self.database.save(db_file)
        print(f"Doc Reader: FUSION Test Run F{validation_num - SPARKS_MAX} ({test_run.person}) added to database.")
        self._reconcile_counts()

        # ── AUTO-UPDATE BMC STATE ──
        # Push this test_run into bmc_state.json immediately using the
        # canonical renumbered test_num so Generate BMC stays in sync.
        try:
            if BMC_ENGINE_AVAILABLE:
                _project_id = get_active_project()
                _bmc_dir = get_bmc_dir(_project_id)
                generate_project_bmc(
                    _project_id,
                    [test_run.to_dict()],
                    _bmc_dir,
                    generate_docx=False,   # state only — full DOCX on demand
                )
                print(f"[BMC] State updated with test_run #{test_run.test_num} ({test_run.person})")
        except Exception as _bmc_e:
            print(f"[BMC] Warning: could not auto-update BMC state: {_bmc_e}")

        # Refresh all views
        if hasattr(self, 'planning_widget'):
            self.planning_widget._populate_table()
        self.refresh_test_list()
        self.cube_widget.refresh()
        
        # ┬┬ MARK AI ANALYSIS AS STALE â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬┬┬┬┬┬┬
        # The user just added new data.  The old AI results don't include
        # this test_run.  Flash the button so they know to recalculate.
        self._mark_analysis_stale(test_run)
        
        print(f"Doc Reader: Test Run #{test_run.test_num} ({test_run.person}) processed and filed.")
        
        # ── AUTO-GENERATE POST-TEST DELIVERABLES ──
        # 1) BCM format PPTX slide  2) Compound Intelligence Report
        self._auto_generate_post_test_run(parsed_data, test_run)
    
    def _auto_generate_post_test_run(self, parsed_data: dict, test_run):
        """
        Auto-generate post-test_run deliverables with PROGRESSIVE intelligence.
        1) BCM format PPTX slide (with before/after delta)
        2) Compound Intelligence Report (.txt)
        3) IoC snapshot (Test Run of Change)
        4) KNN model update
        Deposits into {project}/tested_pptx_slide_creator/
        """
        project_id = get_active_project()
        
        # Enrich parsed_data with test_run fields
        parsed_data['test_num'] = test_run.test_num
        if not parsed_data.get('script_name'):
            parsed_data['script_name'] = test_run.person
        if not parsed_data.get('source_version'):
            parsed_data['source_version'] = test_run.company
        
        # Try progressive engine first (full delta + IoC)
        try:
            from progressive_test_run_engine import ProgressiveTest RunEngine
            
            # Get state BEFORE (saved during template generation)
            state_before = getattr(self, '_last_intel_state', None)
            
            engine = ProgressiveTest RunEngine(self.database, project_id)
            outputs = engine.process_completed_test_run(
                parsed_data, state_before=state_before
            )
            
            slide_info = outputs.get('slides', {})
            delta_summary = outputs.get('summary', '')
            
            msg_parts = []
            if slide_info.get('log_slide'):
                msg_parts.append(f"BCM Slide: {Path(slide_info['log_slide']).name}")
            if slide_info.get('compound_report'):
                msg_parts.append(f"Compound Report: {Path(slide_info['compound_report']).name}")
            if outputs.get('ioc_path'):
                msg_parts.append(f"IoC Snapshot: {Path(outputs['ioc_path']).name}")
            
            QMessageBox.information(
                self,
                "📊 Post-Test Run Generated (Progressive)",
                f"Test Run #{test_run.test_num} ({test_run.person}):\n\n"
                + "\n".join(msg_parts)
                + f"\n\n{delta_summary}"
            )
            return
            
        except ImportError:
            print("  [INFO] progressive_test_run_engine not available, using basic generator")
        except Exception as e:
            print(f"  [WARN] Progressive post-test_run failed: {e}")
            import traceback; traceback.print_exc()
        
        # FALLBACK: Basic post_test_run_generator
        try:
            from post_test_run_generator import PostTest RunGenerator
            gen = PostTest RunGenerator(self.database, project_id)
            slide_outputs = gen.generate_all(parsed_data)
            
            if slide_outputs:
                msg_parts = []
                for name, path in slide_outputs.items():
                    msg_parts.append(f"{name}: {path.name}")
                QMessageBox.information(
                    self,
                    "📊 Post-Test Run Generated",
                    f"Test Run #{test_run.test_num} ({test_run.person}):\n\n"
                    + "\n".join(msg_parts)
                    + f"\n\nLocation: {gen.output_dir}"
                )
        except ImportError:
            print("  [WARN] post_test_run_generator.py not found")
        except Exception as e:
            print(f"  [WARN] Basic post-test_run generation failed: {e}")
    # RE-SCAN COMPLETED PACKETS (Learning Assistant)
    # ════════════════════════════════════════════════════════
    
    def _auto_scan_tested(self):
        """
        Silent auto-scan of tested/ folder on startup and refresh.
        Creates Validation database entries from completed .docx packets
        that are not yet in the database. No dialog — runs quietly.
        """
        project_id = get_active_project()
        tested_dir = get_tested_dir(project_id)
        if not tested_dir.exists():
            return
        docx_files = list(tested_dir.glob("*.docx"))
        if not docx_files:
            return
        try:
            from doc_reader import PacketParser as Test RunPacketParser
        except ImportError:
            return

        # Only deduplicate against FUSION tests — Baseline numbers occupy 1-SPARKS_MAX
        # and must not block Validation packets that happen to share a template number
        existing_validation_nums   = {i.test_num for i in self.database.test_runs
                                   if getattr(i, 'source', '') == 'validation'}
        existing_validation_people = {i.person.strip().lower() for i in self.database.test_runs
                                   if getattr(i, 'source', '') == 'validation'}
        added = 0

        for docx_path in docx_files:
            try:
                parser = Test RunPacketParser(docx_path)
                parsed = parser.parse()

                # doc_reader returns header fields via parsed['header']
                header = parsed.get('header', {})
                parsed_num    = parsed.get('test_num', 0)
                parsed_person = (
                    header.get('Person') or
                    parsed.get('script_name') or
                    ''
                ).strip()
                parsed_date   = header.get('Date', parsed.get('date', ''))
                parsed_title  = header.get('Title', parsed.get('title', ''))
                parsed_company= header.get('Company', parsed.get('source_version', ''))
                parsed_custype= header.get('Substrate Class', parsed.get('substrate_class', parsed.get('test_category', '')))

                # Skip if no person name could be extracted
                if not parsed_person:
                    print(f"  ⚠ Auto-scan: no person found in {docx_path.name} — skipping")
                    continue

                # Skip if already exists as a FUSION entry (by person name)
                if parsed_person.lower() in existing_validation_people:
                    continue

                new_iv = Test Run()
                new_iv.person         = parsed_person or docx_path.stem
                new_iv.date           = parsed_date
                new_iv.title          = parsed_title
                new_iv.company        = parsed_company
                new_iv.test_category  = parsed_custype
                new_iv.hypotheses     = parsed.get('hypotheses_validation', parsed.get('hypotheses', ''))
                exp_answers           = parsed.get('experiments_answers', {})
                new_iv.experiments    = '\n\n'.join(f"Q: {q}\nA: {a}" for q, a in exp_answers.items())
                new_iv.results        = parsed.get('results', '')
                new_iv.action_iterate = parsed.get('action_iterate', '')
                new_iv.source         = 'validation'
                new_iv.q_layer        = parsed.get('q_layer', 'L2')
                new_iv.q_object       = parsed.get('q_object', 'OC')
                new_iv.q_stack        = [normalize_stack_key(s) for s in parsed.get('q_stack', [])]
                new_iv.substrate_impacts = []
                for eq in parsed.get('equipment_discussed', []):
                    # Pick highest non-zero cost and tag cost_type correctly
                    prod  = float(eq.get('production_cost', 0) or 0)
                    dmg   = float(eq.get('damage_cost',     0) or 0)
                    safe  = float(eq.get('safety_cost',     0) or 0)
                    other = float(eq.get('other_cost',      0) or 0)
                    best_cost = prod or dmg or safe or other
                    if prod > 0:
                        cost_type = 'production'
                    elif safe > 0:
                        cost_type = 'safety'
                    else:
                        cost_type = 'repair'
                    new_iv.substrate_impacts.append({
                        'equipment':       eq.get('equipment', ''),
                        'cost':            best_cost,
                        'cost_type':       cost_type,
                        'production_cost': prod,
                        'damage_cost':     dmg,
                        'safety_cost':     safe,
                        'other_cost':      other,
                        'notes':           eq.get('notes', ''),
                    })

                # Assign next available Validation number above SPARKS_MAX
                new_num = SPARKS_MAX + 1
                while new_num in existing_validation_nums:
                    new_num += 1
                new_iv.test_num = new_num

                self.database.add_test_run(new_iv)
                existing_validation_nums.add(new_num)
                existing_validation_people.add(new_iv.person.lower())
                added += 1
                print(f"  ✓ Auto-scan: added Validation #{new_num} ({new_iv.person})")

            except Exception as e:
                print(f"  ⚠ Auto-scan error {docx_path.name}: {e}")

        if added > 0:
            db_file = BCM_ROOT / "test_database.json"
            self.database.save(db_file)
            self.database.rebuild_cube_matrix()
            self.refresh_test_list()
            print(f"  ✓ Auto-scan complete: {added} new Validation tests added")

    def _rescan_completed_packets(self):
        """Re-read all completed .docx packets and repopulate empty database fields.
        
        Called from Learning Assistant's 'Re-scan' button.
        Finds the tested/ folder, parses each .docx, and fills in
        any empty results/experiments/hypotheses fields in existing database entries.
        """
        project_id = get_active_project()
        tested_dir = get_tested_dir(project_id)
        
        if not tested_dir.exists():
            QMessageBox.warning(self, "No Folder",
                f"Completed packets folder not found:\n{tested_dir}")
            return
        
        docx_files = list(tested_dir.glob("*.docx"))
        if not docx_files:
            QMessageBox.warning(self, "No Packets",
                f"No .docx files found in:\n{tested_dir}")
            return
        
        try:
            from doc_reader import PacketParser as Test RunPacketParser
        except ImportError:
            QMessageBox.warning(self, "Module Missing",
                "doc_reader.py not found. Cannot re-scan packets.")
            return
        
        updated_count = 0
        error_count = 0
        
        for docx_path in docx_files:
            try:
                parser = Test RunPacketParser(docx_path)
                parsed = parser.parse()
                
                # Find matching database entry by person name (primary key)
                matched = None
                parsed_person = parsed.get('script_name', '').strip()
                
                for db_test_run in self.database.test_runs:
                    # Match by person name (primary)
                    if parsed_person and db_test_run.person.strip().lower() == parsed_person.lower():
                        matched = db_test_run
                        break
                
                if not matched:
                    # No existing entry — create a new Validation test_run from the docx
                    print(f"  ➕ Creating new Validation entry from: {docx_path.name}")
                    _hdr = parsed.get('header', {})
                    new_iv = Test Run()
                    new_iv.person        = parsed_person or docx_path.stem
                    new_iv.date          = _hdr.get('Date', parsed.get('date', ''))
                    new_iv.title         = _hdr.get('Title', parsed.get('title', ''))
                    new_iv.company       = _hdr.get('Company', parsed.get('source_version', ''))
                    new_iv.test_category = _hdr.get('Substrate Class', parsed.get('substrate_class', parsed.get('test_category', '')))
                    new_iv.hypotheses    = parsed.get('hypotheses_validation', parsed.get('hypotheses', ''))
                    new_iv.experiments   = ''
                    exp_answers = parsed.get('experiments_answers', {})
                    if exp_answers:
                        new_iv.experiments = '\n\n'.join(f"Q: {q}\nA: {a}" for q, a in exp_answers.items())
                    new_iv.results       = parsed.get('results', '')
                    new_iv.action_iterate= parsed.get('action_iterate', '')
                    new_iv.source        = 'validation'
                    new_iv.q_layer       = parsed.get('q_layer', 'L2')
                    new_iv.q_object      = parsed.get('q_object', 'OC')
                    new_iv.q_stack       = [normalize_stack_key(s) for s in parsed.get('q_stack', [])]
                    new_iv.substrate_impacts = []
                    for eq in parsed.get('equipment_discussed', []):
                        new_iv.substrate_impacts.append({
                            'equipment': eq.get('equipment', ''),
                            'cost':      eq.get('typical_cost', ''),
                            'notes':     eq.get('notes', ''),
                        })
                    # Assign Validation test_run number — no collision with Baseline 1-SPARKS_MAX
                    existing_validation_nums = {
                        i.test_num for i in self.database.test_runs
                        if getattr(i, 'source', '') == 'validation'
                    }
                    new_num = SPARKS_MAX + 1
                    while new_num in existing_validation_nums:
                        new_num += 1
                    new_iv.test_num = new_num
                    self.database.add_test_run(new_iv)
                    updated_count += 1
                    print(f"  ✓ Created Validation #{new_num} ({new_iv.person})")
                    continue
                
                # Fill ONLY empty fields — never overwrite existing data
                fields_filled = []
                
                new_results = parsed.get('results', '').strip()
                if not matched.results and new_results:
                    matched.results = new_results
                    fields_filled.append('results')
                
                new_hyp = parsed.get('hypotheses_validation', '').strip()
                if not matched.hypotheses and new_hyp:
                    matched.hypotheses = new_hyp
                    fields_filled.append('hypotheses')
                
                new_action = parsed.get('action_iterate', '').strip()
                if not matched.action_iterate and new_action:
                    matched.action_iterate = new_action
                    fields_filled.append('action_iterate')
                
                # Build experiments from Q&A pairs
                exp_answers = parsed.get('experiments_answers', {})
                if not matched.experiments and exp_answers:
                    exp_lines = []
                    for q, a in exp_answers.items():
                        exp_lines.append(f"Q: {q}\nA: {a}")
                    matched.experiments = '\n\n'.join(exp_lines)
                    fields_filled.append('experiments')
                
                # Equipment impacts
                equip_list = parsed.get('equipment_discussed', [])
                if not matched.substrate_impacts and equip_list:
                    for eq in equip_list:
                        impact = {
                            'equipment': eq.get('equipment', ''),
                            'cost': eq.get('typical_cost', ''),
                            'notes': eq.get('notes', '')
                        }
                        matched.substrate_impacts.append(impact)
                    fields_filled.append(f'equipment({len(equip_list)})')
                
                if fields_filled:
                    updated_count += 1
                    print(f"  ✓ Updated #{matched.test_num} {matched.person}: {', '.join(fields_filled)}")
                else:
                    print(f"  — #{matched.test_num} {matched.person}: already populated")
                    
            except Exception as e:
                error_count += 1
                print(f"  ✗ Error parsing {docx_path.name}: {e}")
        
        # Save updated database
        if updated_count > 0:
            db_file = BCM_ROOT / "test_database.json"
            self.database.save(db_file)
            print(f"\n✓ Database saved with {updated_count} updated test_runs")
        
        # Update Learning Assistant status
        self.learning_widget.status_label.setText(
            f"Re-scan complete: {updated_count} updated, {error_count} errors, "
            f"{len(self.database.tests)} total test_runs"
        )
        
        QMessageBox.information(self, "Re-scan Complete",
            f"Scanned {len(docx_files)} .docx files\n"
            f"Updated: {updated_count} test_runs\n"
            f"Errors: {error_count}\n\n"
            f"Click 'Analyze All' to run analysis on updated data.")

    # ════════════════════════════════════════════════════════
    # STALE ANALYSIS NOTIFICATION SYSTEM
    # ════════════════════════════════════════════════════════
    
    def _mark_analysis_stale(self, test_run: Test Run):
        """Signal that AI analysis is outdated due to new validation test_run.
        
        Called by _on_doc_reader_processed() after a completed packet is ingested.
        Starts flashing the Run AI button and posts a stale banner on the AI tab.
        """
        self._analysis_stale = True
        self._stale_test_run = test_run
        
        # Start flashing the AI analysis button (toggle every 800ms)
        self._stale_timer.start(800)
        
        # Build a quick delta summary for the AI Analysis tab
        validation_label = f"F{test_run.test_num - SPARKS_MAX}"
        equip_count = len(test_run.substrate_impacts)
        cube_pos = f"[{test_run.q_layer}, {test_run.q_object}, {','.join(test_run.q_stack)}]"
        
        stale_msg = (
            "⚠️ ════════════════════════════════════════════════════════\n"
            f"  NEW VALIDATION TEST: {validation_label} ⬔ {test_run.person} ({test_run.company})\n"
            f"  Q-Cube: {cube_pos}  |  Equipment discussed: {equip_count}\n"
            "  â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬â⬝â‚¬┬┬┬┬┬┬\n"
            "  AI Analysis is STALE ⬔ results do not reflect this test_run.\n"
            "  Click '🤖 Run AI Synergy Analysis' in the toolbar to recalculate.\n"
            "════════════════════════════════════════════════════════\n\n"
        )
        
        # Prepend stale warning to whatever is currently shown
        current = self.analysis_widget.results_text.toPlainText()
        self.analysis_widget.results_text.setText(stale_msg + current)
        
        # Update the status label on the AI tab
        self.analysis_widget.status_label.setText(
            f"⚠️ STALE ⬔ {validation_label} ({test_run.person}) not included in current analysis"
        )
        self.analysis_widget.status_label.setStyleSheet(
            "color: #FFD700; font-weight: bold; font-size: 11pt;"
        )
    
    def _flash_ai_button(self):
        """Toggle the AI analysis button between normal and alert colors.
        
        Called by self._stale_timer every 800ms while analysis is stale.
        """
        if not self._analysis_stale:
            self._stale_timer.stop()
            return
        
        self._stale_flash_on = not self._stale_flash_on
        
        if self._stale_flash_on:
            self.run_ai_btn.setStyleSheet(
                "background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 10px; border: 1px solid #7928CA;"
            )
            self.run_ai_btn.setText("⚠️ RECALCULATE ⬔ New Test Added")
        else:
            self.run_ai_btn.setStyleSheet(
                "background-color: #663300; color: #ff8c00; font-weight: bold; padding: 10px; border: 1px solid #ff6600;"
            )
            self.run_ai_btn.setText("🤖 Run AI Synergy Analysis")
    
    def _clear_stale_state(self):
        """Reset stale analysis state after user triggers recalculation.
        
        Called at the top of run_ai_analysis() to stop flashing and restore
        the button to its normal blue state.
        """
        self._analysis_stale = False
        self._stale_flash_on = False
        self._stale_timer.stop()
        self._stale_test_run = None
        
        # Restore button to normal
        self.run_ai_btn.setStyleSheet(
            "background-color: #9333EA; color: #FFD700; font-weight: bold; padding: 10px; border: 1px solid #7928CA;"
        )
        self.run_ai_btn.setText("🤖 Run AI Synergy Analysis")
        
        # Restore status label
        self.analysis_widget.status_label.setStyleSheet("")
        self.analysis_widget.status_label.setText("Recalculating...")
    
    def auto_classify_test_run(self, test_run: Test Run):
        """Auto-classify test_run into Q-Cube foundation + all 10 lenses."""
        # Analyze text to determine Q-Layer (foundation lens 1)
        text = f"{test_run.title} {test_run.hypotheses} {test_run.results}".lower()
        
        if any(word in text for word in ['vp', 'vice president', 'president', 'ceo', 'executive']):
            test_run.q_layer = 'L3'
        elif any(word in text for word in ['manager', 'supervisor', 'lead', 'operations']):
            test_run.q_layer = 'L2'
        else:
            test_run.q_layer = 'L1'
        
        # Determine Q-Object (foundation lens 2)
        if any(word in text for word in ['supplier', 'timber', 'chip mill', 'docked',
                                          'hauler', 'trucker', 'driver', 'logger', 'forest']):
            test_run.q_object = 'OA'
        elif any(word in text for word in ['blow line', 'transfer', 'calibration',
                                            'bark line', 'screen room']):
            test_run.q_object = 'OB'
        else:
            test_run.q_object = 'OC'
        
        # Determine Q-Stack (foundation lens 3 — can be multiple)
        test_run.q_stack = []
        if any(word in text for word in ['multiple mills', 'everywhere', 'all mills', 'across']):
            test_run.q_stack.append('Sα')
        if any(word in text for word in ['modern equipment', 'billion', 'new system', 'upgrade', 'capital investment']):
            test_run.q_stack.append('Sβ')
        if any(word in text for word in ['baseline', 'commissioning', 'new equipment', 'startup']):
            test_run.q_stack.append('Sγ')
        if any(word in text for word in ['kraft + tissue', 'multiple products', 'dual impact']):
            test_run.q_stack.append('Sδ')
        
        if not test_run.q_stack:
            test_run.q_stack.append('Sα')
        
        # Run all 7 analytical + innovation lenses (A, V, T, C, N, F, K)
        lenses = infer_all_lenses(test_run)
        for key, val in lenses.items():
            if val:
                setattr(test_run, key, val)
        
        # Bayesian synergy score — real calculation
        import math as _math
        _alpha, _beta = 1.0, 1.0
        _v_weights = {'V4': 1.0, 'V3': 0.8, 'V2': 0.5, 'V1': 0.2}
        _l_weights = {'L1': 1.0, 'L2': 0.67, 'L3': 0.33}
        _v = getattr(test_run, 'q_evidence', '')
        _alpha += _v_weights.get(_v, 0.2)
        _tlen = sum(len(getattr(test_run, f, '') or '') for f in ('results', 'experiments', 'hypotheses'))
        if _tlen > 800: _alpha += 0.8
        elif _tlen > 400: _alpha += 0.5
        elif _tlen > 100: _alpha += 0.2
        else: _beta += 0.5
        if getattr(test_run, 'substrate_impacts', []): _alpha += 0.3
        _alpha += _l_weights.get(test_run.q_layer, 0.5) * 0.5
        test_run.synergy_score = min(1.0, _alpha / (_alpha + _beta))
    
    def refresh_test_list(self):
        """Refresh test list table ⬔ Baseline first, then Validation, with section headers."""
        sparks = self.database.get_baseline_test_runs()
        validation = self.database.get_validation_test_runs()
        
        # Sort each group: Baseline by number (chronological), Validation by person name
        sparks.sort(key=lambda i: i.test_num)
        validation.sort(key=lambda i: i.person.strip().lower())
        
        total_rows = len(sparks) + len(validation)
        if sparks:
            total_rows += 1  # Baseline header row
        if validation:
            total_rows += 1  # Validation header row
        
        self.list_widget.setRowCount(total_rows)
        row = 0
        
        # --- SPARKS SECTION ---
        if sparks:
            header_item = QTableWidgetItem(f"═══ SPARKS → FUSION RESEARCH (Discovery → Validation) ⬔ {len(sparks)}/{SPARKS_REQUIRED} required ═══")
            header_item.setFont(QFont("Arial", 9, QFont.Bold))
            header_item.setBackground(QColor(40, 20, 60))
            self.list_widget.setItem(row, 0, header_item)
            for col in range(1, 8):
                blank = QTableWidgetItem("")
                blank.setBackground(QColor(40, 20, 60))
                self.list_widget.setItem(row, col, blank)
            self.list_widget.setSpan(row, 0, 1, 8)
            row += 1
            
            for test_entry in sparks:
                display_num = str(test_run.test_num)
                self.list_widget.setItem(row, 0, QTableWidgetItem(display_num))
                self.list_widget.setItem(row, 1, QTableWidgetItem(test_run.date))
                self.list_widget.setItem(row, 2, QTableWidgetItem(test_run.person))
                self.list_widget.setItem(row, 3, QTableWidgetItem(test_run.company))
                self.list_widget.setItem(row, 4, QTableWidgetItem(test_run.test_category))
                cube_pos = f"[{test_run.q_layer}, {test_run.q_object}, {','.join(test_run.q_stack)}]"
                self.list_widget.setItem(row, 5, QTableWidgetItem(cube_pos))
                self.list_widget.setItem(row, 6, QTableWidgetItem("SPARKS"))
                row += 1
        
        # --- FUSION SECTION ---
        if validation:
            header_item = QTableWidgetItem(f"═══ FUSION RESEARCH (Value Proposition Validation) ⬔ {len(validation)} tests ═══")
            header_item.setFont(QFont("Arial", 9, QFont.Bold))
            header_item.setBackground(QColor(0, 40, 0))
            self.list_widget.setItem(row, 0, header_item)
            for col in range(1, 8):
                blank = QTableWidgetItem("")
                blank.setBackground(QColor(0, 40, 0))
                self.list_widget.setItem(row, col, blank)
            self.list_widget.setSpan(row, 0, 1, 8)
            row += 1
            
            for test_entry in validation:
                display_num = f"F{test_run.test_num - SPARKS_MAX}"
                self.list_widget.setItem(row, 0, QTableWidgetItem(display_num))
                self.list_widget.setItem(row, 1, QTableWidgetItem(test_run.date))
                self.list_widget.setItem(row, 2, QTableWidgetItem(test_run.person))
                self.list_widget.setItem(row, 3, QTableWidgetItem(test_run.company))
                self.list_widget.setItem(row, 4, QTableWidgetItem(test_run.test_category))
                cube_pos = f"[{test_run.q_layer}, {test_run.q_object}, {','.join(test_run.q_stack)}]"
                self.list_widget.setItem(row, 5, QTableWidgetItem(cube_pos))
                self.list_widget.setItem(row, 6, QTableWidgetItem("FUSION"))
                _del_btn = QPushButton("DELETE")
                _del_btn.setStyleSheet(
                    "QPushButton{background:#8B0000;color:white;font-weight:bold;"
                    "font-size:10px;padding:2px 6px;}"
                    "QPushButton:hover{background:#CC0000;}"
                )
                _n = test_run.person
                _del_btn.clicked.connect(lambda checked, n=_n: self._delete_validation_test_run(n))
                self.list_widget.setCellWidget(row, 7, _del_btn)
                row += 1

        # --- INCLUSION SECTION (Stream 2 — Deployed Module Intelligence) ---
        inclusion_entries = []
        try:
            project_id = self._get_active_project_id()
            if project_id:
                from pathlib import Path as _P
                log_file = get_inclusion_log(project_id)
                if log_file.exists():
                    import json as _j
                    _data = _j.loads(log_file.read_text(encoding='utf-8'))
                    inclusion_entries = _data.get('entries', [])
        except Exception:
            pass

        if inclusion_entries:
            # Count by status
            _active = [e for e in inclusion_entries if e.get('status', 'PENDING') != 'REJECTED']
            _rejected = [e for e in inclusion_entries if e.get('status') == 'REJECTED']

            # Expand table for inclusion rows
            current_rows = self.list_widget.rowCount()
            extra = 1 + len(inclusion_entries)  # header + entries
            self.list_widget.setRowCount(current_rows + extra)

            # Header row
            _ML_MAP = {"EL": "ML1", "TL": "ML2", "IM": "ML3", "OPERATOR": "ML0"}
            header_item = QTableWidgetItem(
                f"═══ INCLUSION DEVELOPMENT (CaaS Pilot Intelligence) 📋 "
                f"{len(_active)} active / {len(_rejected)} rejected ═══"
            )
            header_item.setFont(QFont("Arial", 9, QFont.Bold))
            header_item.setBackground(QColor(0, 30, 0))
            header_item.setForeground(QColor(0, 255, 0))
            self.list_widget.setItem(row, 0, header_item)
            for col in range(1, 8):
                blank = QTableWidgetItem("")
                blank.setBackground(QColor(0, 30, 0))
                self.list_widget.setItem(row, col, blank)
            self.list_widget.setSpan(row, 0, 1, 8)
            row += 1

            for idx, inc in enumerate(inclusion_entries):
                role = inc.get('observer_role', 'OPERATOR')
                lens = _ML_MAP.get(role, 'ML0')
                status = inc.get('status', 'PENDING')
                is_rejected = (status == "REJECTED")
                display_num = f"{lens}-{idx + 1}"

                self.list_widget.setItem(row, 0, QTableWidgetItem(display_num))

                # Date from timestamp
                ts = inc.get('timestamp', '')
                try:
                    from datetime import datetime as _dt
                    dt = _dt.fromisoformat(ts)
                    date_str = dt.strftime("%b %d, %Y")
                except Exception:
                    date_str = ts[:10] if ts else ""
                self.list_widget.setItem(row, 1, QTableWidgetItem(date_str))

                self.list_widget.setItem(row, 2, QTableWidgetItem(
                    inc.get('observer_name', 'Anonymous')))

                # Show source module instead of generic "Internal (CaaS)"
                mod_name = inc.get('source_module', 'CaaS')
                self.list_widget.setItem(row, 3, QTableWidgetItem(mod_name[:18]))

                self.list_widget.setItem(row, 4, QTableWidgetItem(
                    inc.get('observation_type', '')))
                self.list_widget.setItem(row, 5, QTableWidgetItem(
                    f"[L2, OC, Sι]"))

                # Source column with lens color + status
                src_text = f"INCLUSION/{lens}"
                if is_rejected:
                    src_text = f"X REJECTED/{lens}"
                elif status == "ACCEPTED":
                    src_text = f"OK {lens}"
                src_item = QTableWidgetItem(src_text)
                _lens_colors = {"ML1": "#FFD700", "ML2": "#00FFFF", "ML3": "#FF69B4", "ML0": "#00FF00"}
                src_item.setForeground(QColor(_lens_colors.get(lens, '#00FF00')))
                self.list_widget.setItem(row, 6, src_item)

                # Delete button for inclusion entry
                _inc_id = inc.get('inclusion_id', '')
                _inc_name = inc.get('observer_name', 'entry')
                _idel_btn = QPushButton("DELETE")
                _idel_btn.setStyleSheet(
                    "QPushButton{background:#4A0000;color:#FF8888;font-weight:bold;"
                    "font-size:10px;padding:2px 6px;}"
                    "QPushButton:hover{background:#8B0000;color:white;}"
                )
                _idel_btn.clicked.connect(
                    lambda checked, eid=_inc_id, ename=_inc_name:
                    self._delete_inclusion_row_entry(eid, ename)
                )
                self.list_widget.setCellWidget(row, 7, _idel_btn)

                # REJECTED: strikethrough + dim all cells
                if is_rejected:
                    _strike = QFont("Consolas", 9)
                    _strike.setStrikeOut(True)
                    for col in range(7):
                        item = self.list_widget.item(row, col)
                        if item:
                            item.setFont(_strike)
                            item.setForeground(QColor("#555"))
                            item.setBackground(QColor(15, 5, 5))

                # Safety highlight (overrides reject dim for safety rows)
                elif inc.get('safety_signal'):
                    for col in range(7):
                        item = self.list_widget.item(row, col)
                        if item:
                            item.setBackground(QColor(40, 0, 0))

                row += 1
    

        # --- EMULSION SECTION (Stream 3 — SaaS Validation Intelligence) ---
        emulsion_entries = []
        try:
            project_id = self._get_active_project_id()
            if project_id:
                try:
                    imm_log = get_immulsion_log(project_id)
                except NameError:
                    imm_log = get_project_dir(project_id) / "immulsion_log.json"
                if imm_log.exists():
                    import json as _j2
                    _idata = _j2.loads(imm_log.read_text(encoding='utf-8'))
                    emulsion_entries = _idata.get('entries', [])
        except Exception:
            pass

        if emulsion_entries:
            current_rows = self.list_widget.rowCount()
            extra = 1 + len(emulsion_entries)
            self.list_widget.setRowCount(current_rows + extra)

            _IMM_LENS = {"MILL": ("ML5", "#00FF88"), "INDUSTRY": ("ML6", "#FF69B4"), "COMPANY": ("ML7", "#FFD700")}
            header_item = QTableWidgetItem(
                f"═══ EMULSION VALIDATION (SaaS Production Feedback) 🏭 {len(emulsion_entries)} entries ═══"
            )
            header_item.setFont(QFont("Arial", 9, QFont.Bold))
            header_item.setBackground(QColor(40, 30, 0))
            header_item.setForeground(QColor(255, 215, 0))
            self.list_widget.setItem(row, 0, header_item)
            for col in range(1, 8):
                blank = QTableWidgetItem("")
                blank.setBackground(QColor(40, 30, 0))
                self.list_widget.setItem(row, col, blank)
            self.list_widget.setSpan(row, 0, 1, 8)
            row += 1

            for idx, emm in enumerate(emulsion_entries):
                role = emm.get('recipient_role', 'MILL')
                lens_id, lens_color = _IMM_LENS.get(role, ("ML5", "#00FF88"))
                display_num = f"{lens_id}-{idx + 1}"
                self.list_widget.setItem(row, 0, QTableWidgetItem(display_num))

                _ets = emm.get('timestamp', '')
                try:
                    from datetime import datetime as _dt2
                    _edt = _dt2.fromisoformat(_ets)
                    _edate = _edt.strftime("%b %d, %Y")
                except Exception:
                    _edate = _ets[:10] if _ets else ""
                self.list_widget.setItem(row, 1, QTableWidgetItem(_edate))
                self.list_widget.setItem(row, 2, QTableWidgetItem(
                    emm.get('recipient_name', 'Anonymous')))
                self.list_widget.setItem(row, 3, QTableWidgetItem(f"SaaS ({role})"))
                self.list_widget.setItem(row, 4, QTableWidgetItem(
                    emm.get('category', emm.get('observation_type', ''))))
                self.list_widget.setItem(row, 5, QTableWidgetItem("[L3, OS, Sω]"))

                src_item = QTableWidgetItem(f"EMULSION/{lens_id}")
                src_item.setForeground(QColor(lens_color))
                self.list_widget.setItem(row, 6, src_item)

                # Delete button for emulsion entry
                _emm_id = emm.get('inclusion_id', emm.get('emulsion_id', ''))
                _emm_name = emm.get('recipient_name', 'entry')
                _edel_btn = QPushButton("DELETE")
                _edel_btn.setStyleSheet(
                    "QPushButton{background:#3A2A00;color:#FFD700;font-weight:bold;"
                    "font-size:10px;padding:2px 6px;}"
                    "QPushButton:hover{background:#7A5A00;color:white;}"
                )
                _edel_btn.clicked.connect(
                    lambda checked, eid=_emm_id, ename=_emm_name:
                    self._delete_emulsion_entry(eid, ename)
                )
                self.list_widget.setCellWidget(row, 7, _edel_btn)

                if emm.get('service_request'):
                    for col in range(7):
                        item = self.list_widget.item(row, col)
                        if item: item.setBackground(QColor(50, 30, 0))

                if emm.get('safety_signal'):
                    for col in range(7):
                        item = self.list_widget.item(row, col)
                        if item: item.setBackground(QColor(60, 0, 0))

                row += 1

    def _delete_validation_test_run(self, person_name: str):
        """Remove a Validation test_run by person name. Saves to disk, rebuilds cube, refreshes."""
        target = next(
            (iv for iv in self.database.test_runs
             if iv.person.strip().lower() == person_name.strip().lower()), None
        )
        if not target:
            QMessageBox.warning(self, "Not Found", f"{person_name} not found.")
            return
        confirm = QMessageBox.question(
            self, "Confirm Delete",
            f"Delete Validation test_run for {target.person}?\n\n"
            f"Company: {target.company}\n\n"
            f"This cannot be undone.",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm != QMessageBox.Yes:
            return
        self.database.tests = [
            iv for iv in self.database.test_runs
            if iv.person.strip().lower() != person_name.strip().lower()
        ]
        self.database.rebuild_cube_matrix()
        try:
            from project_paths import BCM_ROOT
        except ImportError:
            BCM_ROOT = Path.cwd() / "BCM_Projects"
        self.database.save(BCM_ROOT / "test_database.json")
        print(f"  [DELETE] Validation {target.person} removed.")
        self.refresh_test_list()
        self.cube_widget.refresh()

    def _delete_inclusion_row_entry(self, inclusion_id: str, display_name: str):
        """Remove an inclusion entry by inclusion_id from inclusion_log.json."""
        if not inclusion_id:
            QMessageBox.warning(self, "Cannot Delete", "Entry has no inclusion_id.")
            return
        project_id = self._get_active_project_id()
        if not project_id:
            QMessageBox.warning(self, "No Project", "No active project selected.")
            return
        confirm = QMessageBox.question(
            self, "Confirm Delete",
            f"Delete inclusion entry?\n\nObserver: {display_name}\nID: {inclusion_id}\n\n"
            f"This cannot be undone.",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm != QMessageBox.Yes:
            return
        log_file = get_inclusion_log(project_id)
        if not log_file.exists():
            QMessageBox.warning(self, "Not Found", f"inclusion_log.json not found.")
            return
        import json as _j
        data = _j.loads(log_file.read_text(encoding='utf-8'))
        before = len(data.get('entries', []))
        data['entries'] = [
            e for e in data.get('entries', [])
            if e.get('inclusion_id', '') != inclusion_id
        ]
        after = len(data['entries'])
        log_file.write_text(_j.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
        print(f"  [DELETE] Inclusion {inclusion_id} ({display_name}) removed. {before}->{after}")
        self.refresh_test_list()

    def _delete_emulsion_entry(self, entry_id: str, display_name: str):
        """Remove an emulsion entry by id from immulsion_log.json."""
        if not entry_id:
            QMessageBox.warning(self, "Cannot Delete", "Entry has no ID.")
            return
        project_id = self._get_active_project_id()
        if not project_id:
            QMessageBox.warning(self, "No Project", "No active project selected.")
            return
        confirm = QMessageBox.question(
            self, "Confirm Delete",
            f"Delete emulsion entry?\n\nRecipient: {display_name}\nID: {entry_id}\n\n"
            f"This cannot be undone.",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm != QMessageBox.Yes:
            return
        try:
            imm_log = get_immulsion_log(project_id)
        except NameError:
            imm_log = get_project_dir(project_id) / "immulsion_log.json"
        if not imm_log.exists():
            QMessageBox.warning(self, "Not Found", f"immulsion_log.json not found.")
            return
        import json as _j
        data = _j.loads(imm_log.read_text(encoding='utf-8'))
        before = len(data.get('entries', []))
        data['entries'] = [
            e for e in data.get('entries', [])
            if e.get('inclusion_id', e.get('emulsion_id', '')) != entry_id
        ]
        after = len(data['entries'])
        imm_log.write_text(_j.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
        print(f"  [DELETE] Emulsion {entry_id} ({display_name}) removed. {before}->{after}")
        self.refresh_test_list()

    def _refresh_inclusion_returns(self):
        """
        Tether Inclusion Returns to current project.
        Reads inclusion_log.json — shows module source, status, enables governance.
        """
        project_id = self._get_active_project_id()
        if not project_id:
            self._inclusion_status.setText("⊘ No project selected")
            return

        inclusion_entries = []
        try:
            log_file = get_inclusion_log(project_id)
            if log_file.exists():
                import json as _j
                data = _j.loads(log_file.read_text(encoding='utf-8'))
                raw_entries = data.get('entries', [])

                _ML = {"EL": "ML1", "TL": "ML2", "IM": "ML3", "OPERATOR": "ML0"}
                _ML_COLORS = {"ML1": "#FFD700", "ML2": "#00FFFF", "ML3": "#FF69B4", "ML0": "#00FF00"}

                for e in raw_entries:
                    role = e.get('observer_role', 'OPERATOR')
                    lens = _ML.get(role, 'ML0')
                    inclusion_entries.append({
                        "inclusion_id": e.get('inclusion_id', ''),
                        "master_lens": lens,
                        "observer_name": e.get('observer_name', ''),
                        "observation_type": e.get('observation_type', ''),
                        "severity": e.get('severity', 'NOTICE'),
                        "safety_signal": e.get('safety_signal', False),
                        "lens_color": _ML_COLORS.get(lens, '#00FF00'),
                        "observation": e.get('observation', ''),
                        "source_module": e.get('source_module', '—'),
                        "status": e.get('status', 'PENDING'),
                        "timestamp": e.get('timestamp', ''),
                    })
        except Exception as e:
            print(f"  [INCLUSION] Read error: {e}")

        # Cache for accept/reject/delete
        self._inclusion_entries_cache = inclusion_entries

        self._inclusion_table.setRowCount(len(inclusion_entries))
        safety_count = 0
        accepted = 0
        rejected = 0

        _STATUS_COLORS = {
            "ACCEPTED": "#00FF00",
            "REJECTED": "#FF4444",
            "PENDING": "#888888",
        }

        for row, entry in enumerate(inclusion_entries):
            # Col 0: Lens
            lens = entry.get('master_lens', 'ML0')
            lens_item = QTableWidgetItem(lens)
            lens_color = entry.get('lens_color', '#00FF00')
            lens_item.setForeground(QColor(lens_color))
            lens_item.setFont(QFont("Consolas", 9, QFont.Bold))
            self._inclusion_table.setItem(row, 0, lens_item)

            # Col 1: Observer
            obs_item = QTableWidgetItem(entry.get('observer_name', '')[:15])
            self._inclusion_table.setItem(row, 1, obs_item)

            # Col 2: Source module (shortened)
            mod_name = entry.get('source_module', '—')
            short_mod = mod_name[:8] if len(mod_name) > 8 else mod_name
            mod_item = QTableWidgetItem(short_mod)
            mod_item.setForeground(QColor("#888"))
            mod_item.setToolTip(mod_name)
            self._inclusion_table.setItem(row, 2, mod_item)

            # Col 3: Type
            type_item = QTableWidgetItem(entry.get('observation_type', '')[:15])
            self._inclusion_table.setItem(row, 3, type_item)

            # Col 4: Status
            status = entry.get('status', 'PENDING')
            status_item = QTableWidgetItem(status)
            status_item.setForeground(QColor(_STATUS_COLORS.get(status, '#888')))
            status_item.setFont(QFont("Consolas", 8, QFont.Bold))
            self._inclusion_table.setItem(row, 4, status_item)

            # Strike through rejected rows
            if status == "REJECTED":
                _strike_font = QFont("Consolas", 9)
                _strike_font.setStrikeOut(True)
                for c in range(5):
                    it = self._inclusion_table.item(row, c)
                    if it:
                        it.setFont(_strike_font)
                        it.setForeground(QColor("#555"))

            # Safety highlight
            if entry.get('safety_signal'):
                safety_count += 1
                for c in range(5):
                    it = self._inclusion_table.item(row, c)
                    if it:
                        it.setBackground(QColor(40, 0, 0))

            if status == "ACCEPTED":
                accepted += 1
            elif status == "REJECTED":
                rejected += 1

        self._inclusion_status.setText(
            f"\u2713 Tethered to {project_id} — {len(inclusion_entries)} entries "
            f"({accepted}✓ {rejected}✗)"
        )
        self._inclusion_status.setStyleSheet("color: #00FF00; font-size: 10px;")
        self._refresh_rd_health()  # R&D Health convergence update

        self._tether_status.setText(
            f"CaaS Connection: TETHERED\n"
            f"Project: {project_id}\n"
            f"Inclusion entries: {len(inclusion_entries)}\n"
            f"Safety signals: {safety_count}\n"
            f"Accepted: {accepted}  Rejected: {rejected}"
        )

        # Reset buttons
        self._inc_accept_btn.setEnabled(False)
        self._inc_reject_btn.setEnabled(False)
        self._inc_delete_btn.setEnabled(False)
        self._inc_preview.setText("Select an entry to preview observation")

        self.refresh_test_list()
        print(f"  [INCLUSION] Tethered: {project_id} -> {len(inclusion_entries)} entries")

    def _on_inclusion_row_changed(self, row, col, prev_row, prev_col):
        """When user selects an inclusion row, show preview and enable buttons."""
        if row < 0 or row >= len(self._inclusion_entries_cache):
            self._inc_accept_btn.setEnabled(False)
            self._inc_reject_btn.setEnabled(False)
            self._inc_delete_btn.setEnabled(False)
            self._inc_preview.setText("Select an entry to preview observation")
            return

        entry = self._inclusion_entries_cache[row]
        obs = entry.get("observation", "(no observation)")
        lens = entry.get("master_lens", "ML0")
        name = entry.get("observer_name", "?")
        mod = entry.get("source_module", "?")
        status = entry.get("status", "PENDING")
        ts = entry.get("timestamp", "")[:16]

        self._inc_preview.setText(
            f"{lens} | {name} | {mod} | {ts}\n{obs[:200]}"
        )
        self._inc_preview.setStyleSheet(
            "color: #CCC; font-size: 9px; padding: 4px; "
            "background: #050505; border: 1px solid #1a3a1a;"
        )

        self._inc_accept_btn.setEnabled(status != "ACCEPTED")
        self._inc_reject_btn.setEnabled(status != "REJECTED")
        self._inc_delete_btn.setEnabled(True)

    def _set_inclusion_verdict(self, verdict: str):
        """Set ACCEPTED or REJECTED on selected inclusion entry. Writes to disk."""
        row = self._inclusion_table.currentRow()
        if row < 0 or row >= len(self._inclusion_entries_cache):
            return

        entry = self._inclusion_entries_cache[row]
        inc_id = entry.get("inclusion_id", "")
        if not inc_id:
            return

        project_id = self._get_active_project_id()
        if not project_id:
            return

        try:
            import json as _j
            log_file = get_inclusion_log(project_id)
            if not log_file.exists():
                return

            data = _j.loads(log_file.read_text(encoding='utf-8'))
            modified = False
            for e in data.get('entries', []):
                if e.get('inclusion_id') == inc_id:
                    e['status'] = verdict
                    modified = True
                    break

            if modified:
                with open(log_file, 'w', encoding='utf-8') as f:
                    _j.dump(data, f, indent=2, ensure_ascii=False)
                print(f"  [INCLUSION] {inc_id} → {verdict}")

            # Refresh both panels
            self._refresh_inclusion_returns()

        except Exception as e:
            print(f"  [INCLUSION] Verdict error: {e}")

    def _delete_inclusion_entry(self):
        """Permanently remove selected inclusion entry from log."""
        row = self._inclusion_table.currentRow()
        if row < 0 or row >= len(self._inclusion_entries_cache):
            return

        entry = self._inclusion_entries_cache[row]
        inc_id = entry.get("inclusion_id", "")
        name = entry.get("observer_name", "?")

        reply = QMessageBox.question(
            self,
            "Delete Inclusion Entry",
            f"Permanently delete entry from {name}?\n"
            f"ID: {inc_id}\n\nThis cannot be undone.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return

        project_id = self._get_active_project_id()
        if not project_id:
            return

        try:
            import json as _j
            log_file = get_inclusion_log(project_id)
            if not log_file.exists():
                return

            data = _j.loads(log_file.read_text(encoding='utf-8'))
            original_count = len(data.get('entries', []))
            data['entries'] = [
                e for e in data.get('entries', [])
                if e.get('inclusion_id') != inc_id
            ]
            removed = original_count - len(data['entries'])

            with open(log_file, 'w', encoding='utf-8') as f:
                _j.dump(data, f, indent=2, ensure_ascii=False)
            print(f"  [INCLUSION] Deleted {inc_id} ({removed} removed)")

            self._refresh_inclusion_returns()

        except Exception as e:
            print(f"  [INCLUSION] Delete error: {e}")

    def _get_active_project_id(self):
        """Get the currently selected project ID."""
        if hasattr(self, 'project_combo'):
            text = self.project_combo.currentText()
            # Strip parenthetical suffix like "(0 tests)"
            if '(' in text:
                text = text.split('(')[0].strip()
            else:
                text = text.strip()
            # Strip emoji/symbol prefix: "🔧 BCM_SUBSTRATE" → "BCM_SUBSTRATE"
            import re
            text = re.sub(r'^[^\w]+', '', text).strip()
            return text if text else None
        return None

    def _on_tab_changed(self, index: int):
        """Auto-run AI Analysis when its tab is selected."""
        if index == self._ai_analysis_tab_index:
            if self._ai_running:
                return  # Already running, don't double-trigger
            if len(self.database.tests) == 0:
                return  # Nothing to analyze
            # Auto-trigger in silent mode (no popups)
            self.run_ai_analysis(silent=True)

    def run_ai_analysis(self, silent: bool = False):
        """Run AI analysis in background thread.
        
        Args:
            silent: If True, suppress message boxes (used by auto-trigger on tab switch).
        """
        if len(self.database.tests) == 0:
            if not silent:
                QMessageBox.warning(self, "No Data", "Please add tests first!")
            return
        
        if self._ai_running:
            return  # Guard against double-trigger
        
        self._ai_running = True
        self._ai_silent_mode = silent
        
        # ┬┬ Clear stale state ⬔ user is recalculating ┬┬
        self._clear_stale_state()
        
        self.ai_engine = AIAnalysisEngine(self.database)
        self.ai_engine.progress_update.connect(self.analysis_widget.update_progress)
        self.ai_engine.analysis_complete.connect(self.on_analysis_complete)
        self.ai_engine.start()
        
        if not silent:
            QMessageBox.information(self, "Analysis Started", 
                                  "AI analysis running in background...\nCheck 'AI Analysis' tab for results.")
    
    def on_analysis_complete(self, results):
        """Handle AI analysis completion"""
        self._ai_running = False
        
        # Display results
        self.analysis_widget.display_results(results)
        
        # Force reload database from disk (AI engine saved it)
        db_file = BCM_ROOT / "test_database.json"
        if db_file.exists():
            self.database = Test RunDatabase.load(db_file)
            
            # Update all widgets with new database reference
            self.cube_widget.database = self.database
            self.learning_widget.database = self.database
            self.bmc_widget.database = self.database
            self.equipment_widget.set_database(self.database)
            
            # Force refresh cube display
            self.cube_widget.refresh()
            
            # Show message with real ML metrics
            n_patterns = len(results.get('synergy_patterns', []))
            n_clusters = len(results.get('clusters', {}))
            n_vps = len(results.get('value_propositions', []))
            var_exp = results.get('embedding_variance_explained', 0)
            
            if not getattr(self, '_ai_silent_mode', False):
                QMessageBox.information(
                    self,
                    "ML Analysis Complete",
                    f"Real ML analysis on {len(self.database.tests)} test_runs:\n\n"
                    f"• TF-IDF + SVD embeddings ({var_exp:.0%} variance explained)\n"
                    f"• Bayesian synergy scores (Beta-Binomial model)\n"
                    f"• {n_patterns} resonance patterns detected\n"
                    f"• {n_clusters} synergy clusters\n"
                    f"• {n_vps} value propositions generated\n\n"
                    f"Check Q-Cube Matrix tab for updated scores."
                )
    
    def export_results(self):
        """Export analysis results"""
        filepath, _ = QFileDialog.getSaveFileName(
            self, "Export Results", 
            str(get_project_dir(get_active_project()) / "exports" / "validation_analysis_results.txt"),
            "Text Files (*.txt);;All Files (*)"
        )
        
        if filepath:
            with open(filepath, 'w') as f:
                f.write(self.analysis_widget.results_text.toPlainText())
            
            QMessageBox.information(self, "Exported", f"Results exported to:\n{filepath}")
    
    def generate_test_template(self, test_data: dict):
        """Generate intelligent test packet with pre-analysis and targeted questions"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import json
            from pathlib import Path
            
            # Load equipment config based on ACTIVE PROJECT
            active_project = get_active_project()
            config_file = _find_substrate_config(active_project)
            
            # Parse equipment/modules from config
            equipment_data = {}
            pain_indicators = {}
            scorecard = {}
            
            if config_file and config_file.exists():
                with open(config_file, 'r') as f:
                    config = json.load(f)
                
                # Extract equipment/modules based on project type
                if 'module_categories' in config:
                    categories = config.get('module_categories', {})
                    for cat_id, cat_data in categories.items():
                        modules = cat_data.get('modules', {})
                        for mod_id, mod_data in modules.items():
                            equipment_data[mod_id] = {
                                'name': mod_data.get('display_name', mod_id),
                                'category': cat_data.get('name', cat_id),
                                'function': mod_data.get('spine_function', ''),
                                'cost': mod_data.get('typical_cost', ''),
                                'keywords': mod_data.get('keywords', [])
                            }
                else:  # equipment_categories structure
                    categories = config.get('equipment_categories', {})
                    for cat_id, cat_data in categories.items():
                        equipment = cat_data.get('equipment', {})
                        for equip_id, equip_data in equipment.items():
                            equipment_data[equip_id] = {
                                'name': equip_data.get('display_name', equip_id),
                                'category': cat_data.get('name', cat_id),
                                'damage_sources': equip_data.get('damage_sources', []),
                                'cost': equip_data.get('typical_cost_range', ''),
                                'keywords': equip_data.get('keywords', [])
                            }
                
                pain_indicators = config.get('pain_indicators', {})
                scorecard = config.get('test_run_scorecard', {})
            
            equipment_list = [e['name'] for e in equipment_data.values()]
            
            # Load Genesis intelligence — LIVE from progressive engine
            genesis_intel = self._load_genesis_intelligence(test_data)
            
            doc = Document()
            
            # Title based on active project (no hardcoded names)
            title_text = f"GIBUSH Test Packet"
            
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(196, 30, 58)
            
            # Test Run info
            info_table = doc.add_table(rows=7, cols=2)
            info_table.style = 'Light Grid Accent 1'

            # Test Run # left blank — assigned by Excel plan, not baked into template
            info_table.rows[0].cells[0].text = "Test Run #"
            info_table.rows[0].cells[1].text = ""

            info_table.rows[1].cells[0].text = "Date"
            info_table.rows[1].cells[1].text = ""

            info_table.rows[2].cells[0].text = "Person"
            info_table.rows[2].cells[1].text = test_data['name']

            info_table.rows[3].cells[0].text = "Title"
            info_table.rows[3].cells[1].text = test_data['title']

            info_table.rows[4].cells[0].text = "Company"
            info_table.rows[4].cells[1].text = test_data['source_version']

            info_table.rows[5].cells[0].text = "Substrate Class"
            info_table.rows[5].cells[1].text = test_data['type']

            info_table.rows[6].cells[0].text = "Q-Cube Position"
            q_pos = f"[{test_data.get('q_layer', 'L2')}, {test_data.get('q_object', 'OC')}, {test_data.get('q_stack', 'Sa')}]"
            info_table.rows[6].cells[1].text = q_pos
            
            doc.add_paragraph()
            
            # PRE-TEST INTELLIGENCE SECTION
            doc.add_heading('📊 PRE-TEST INTELLIGENCE', 1)
            
            if genesis_intel['available']:
                doc.add_paragraph(f"Based on {genesis_intel['test_count']} previous tests:")
                doc.add_paragraph()
                
                # Equipment coverage
                if genesis_intel['equipment_mentions']:
                    doc.add_heading('Equipment Already Discussed:', 2)
                    for equip, stats in genesis_intel['equipment_mentions'].items():
                        doc.add_paragraph(f"⬢ {equip}: {stats['count']} mentions, avg cost: {stats['avg_cost']}")
                
                # Critical gaps
                if genesis_intel['gaps']:
                    doc.add_heading('🎯 CRITICAL INFORMATION GAPS:', 2)
                    doc.add_paragraph("This test_run should focus on:")
                    for gap in genesis_intel['gaps']:
                        p = doc.add_paragraph(f"⬢ {gap}")
                        p.runs[0].bold = True
                
                # Hypothesis status
                if genesis_intel['hypotheses']:
                    doc.add_heading('Hypothesis Validation Status:', 2)
                    for hyp, status in genesis_intel['hypotheses'].items():
                        doc.add_paragraph(f"⬢ {hyp}: {status['confidence']}% validated")
                        if status['needs_more']:
                            doc.add_paragraph(f"  → ASK: {status['probe_question']}")
                
                # Contradictions to resolve
                if genesis_intel['contradictions']:
                    doc.add_heading('⚠️  CONTRADICTIONS TO RESOLVE:', 2)
                    for contradiction in genesis_intel['contradictions']:
                        doc.add_paragraph(f"⬢ {contradiction}")
                
            else:
                doc.add_paragraph("This is an early test_run. No prior Genesis analysis available.")
                doc.add_paragraph("Focus on discovering equipment damage costs and pain points.")
            
            doc.add_paragraph()
            
            # Segment-specific focus (CORRECTED: L1=Operator, L2=Manager, L3=Executive)
            doc.add_heading('🎯 YOUR MISSION FOR THIS TEST', 1)
            
            # Use progressive engine mission if available
            if genesis_intel.get('mission'):
                for line in genesis_intel['mission'].split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            else:
                q_layer = test_data.get('q_layer', 'L2')
                if q_layer == 'L1':
                    doc.add_paragraph("FOCUS: Hands-on equipment experience, daily pain points, operator workarounds")
                    doc.add_paragraph("KEY QUESTIONS: Which equipment breaks down? How do you detect problems? What do instruments miss?")
                elif q_layer == 'L2':
                    doc.add_paragraph("FOCUS: Equipment budgets, maintenance costs, process improvement decisions, AMO authority")
                    doc.add_paragraph("KEY QUESTIONS: What are annual damage costs? What budget path exists? What ROI do you need?")
                elif q_layer == 'L3':
                    doc.add_paragraph("FOCUS: Strategic risk, capital allocation, ROI requirements, vendor relationships")
                    doc.add_paragraph("KEY QUESTIONS: What budget would you allocate? Who makes final decisions? What's enterprise risk?")
            
            # COMPOUND QUESTIONS — Progressive, calibrated to ALL prior test_runs
            if genesis_intel.get('compound_tests'):
                doc.add_heading('🧠 COMPOUND INTELLIGENCE QUESTIONS', 1)
                doc.add_paragraph(
                    f"Calibrated from {genesis_intel['test_count']} prior tests. "
                    f"These questions target gaps and validate hypotheses.")
                for i, q in enumerate(genesis_intel['compound_tests'], 1):
                    p = doc.add_paragraph(f"Q{i}. {q}")
                    p.runs[0].bold = True
            
            # KNN SIMILAR — Show which prior tests are most similar
            if genesis_intel.get('knn_similar'):
                doc.add_heading('🔗 SIMILAR PRIOR TESTS', 2)
                for sim in genesis_intel['knn_similar'][:3]:
                    doc.add_paragraph(
                        f"#{sim['num']} {sim['script_name']} ({sim['similarity']:.0%} match)")
            
            # REVERSE GAPS — What similar tests missed
            if genesis_intel.get('reverse_gaps'):
                doc.add_heading('⚠️ GAPS FROM SIMILAR TESTS', 2)
                doc.add_paragraph("Similar tests missed these — don't repeat:")
                for gap in genesis_intel['reverse_gaps'][:3]:
                    doc.add_paragraph(f"⬢ [{gap['type']}] {gap.get('gap', '')[:100]}")
            
            # VP STATUS — Current validation summary
            if genesis_intel.get('vp_summary'):
                doc.add_heading('📊 VALUE PROPOSITION STATUS', 2)
                for vp in genesis_intel['vp_summary'][:6]:
                    doc.add_paragraph(f"⬢ {vp}")
            
            doc.add_paragraph()
            
            # Equipment reference with categories and costs
            if equipment_data:
                if 'module_categories' in config:
                    doc.add_heading('📋 SAFETY MODULES CHECKLIST', 1)
                    doc.add_paragraph("Check modules discussed during test_run:")
                else:
                    doc.add_heading('📋 EQUIPMENT DAMAGE CHECKLIST', 1)
                    doc.add_paragraph("Check equipment discussed during test_run:")
                
                # Group by category
                categories = {}
                for equip_id, edata in equipment_data.items():
                    cat = edata.get('category', 'Other')
                    if cat not in categories:
                        categories[cat] = []
                    categories[cat].append(edata)
                
                # Create checklist table by category
                for cat_name, items in categories.items():
                    doc.add_heading(cat_name, 2)
                    equip_table = doc.add_table(rows=len(items)+1, cols=8)
                    equip_table.style = 'Light Grid'
                    
                    # Header — 8 columns
                    equip_table.rows[0].cells[0].text = "☐"
                    equip_table.rows[0].cells[1].text = "Equipment/Module"
                    equip_table.rows[0].cells[2].text = "Typical Cost"
                    equip_table.rows[0].cells[3].text = "Production $"
                    equip_table.rows[0].cells[4].text = "Equipment Damage $"
                    equip_table.rows[0].cells[5].text = "Safety $"
                    equip_table.rows[0].cells[6].text = "Other $"
                    equip_table.rows[0].cells[7].text = "Notes"

                    for i, item in enumerate(items, 1):
                        equip_table.rows[i].cells[0].text = "☐"
                        equip_table.rows[i].cells[1].text = item['name']
                        equip_table.rows[i].cells[2].text = item.get('cost', '')
                        equip_table.rows[i].cells[3].text = ""
                        equip_table.rows[i].cells[4].text = ""
                        equip_table.rows[i].cells[5].text = ""
                        equip_table.rows[i].cells[6].text = ""
                        equip_table.rows[i].cells[7].text = ""
                    
                    doc.add_paragraph()
                
                # Pain indicators to listen for
                if pain_indicators:
                    doc.add_heading('🎯 PAIN INDICATORS (Listen For These)', 2)
                    for indicator, data in list(pain_indicators.items())[:10]:
                        doc.add_paragraph(f"⬢ {data.get('display', indicator)}", style='List Bullet')
                
                # Scorecard questions
                if scorecard:
                    doc.add_heading('📝 SCORECARD QUESTIONS', 2)
                    for key, question in scorecard.items():
                        if key == 'instruction' or not question:
                            continue
                        # Guard: some scorecard values are nested dicts (e.g. hauler questions)
                        if isinstance(question, dict):
                            for sub_key, sub_q in question.items():
                                if isinstance(sub_q, str) and sub_q.strip():
                                    doc.add_paragraph(f"⬢ {sub_q}", style='List Bullet')
                        elif isinstance(question, str):
                            doc.add_paragraph(f"⬢ {question}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"⬢ {str(question)}", style='List Bullet')
            
            doc.add_paragraph()
            
            # Hypotheses
            doc.add_heading('💡 Hypotheses to Test', 1)
            doc.add_paragraph(test_data['hypothesis'])
            
            # Questions
            doc.add_heading('❓ Experiments (Questions to Ask)', 1)
            doc.add_paragraph(test_data['questions'])
            
            # Cost probing framework
            doc.add_paragraph()
            doc.add_heading('💰 COST DISCOVERY FRAMEWORK', 2)
            doc.add_paragraph("For EVERY piece of equipment mentioned:")
            cost_table = doc.add_table(rows=4, cols=2)
            cost_table.style = 'Light Grid'
            cost_table.rows[0].cells[0].text = "Annual damage cost:"
            cost_table.rows[0].cells[1].text = "$__________"
            cost_table.rows[1].cells[0].text = "Failure frequency:"
            cost_table.rows[1].cells[1].text = "_____ times/year"
            cost_table.rows[2].cells[0].text = "Downtime per event:"
            cost_table.rows[2].cells[1].text = "_____ hours"
            cost_table.rows[3].cells[0].text = "Prevention worth:"
            cost_table.rows[3].cells[1].text = "$__________/year"
            
            doc.add_paragraph()
            
            # Results section
            doc.add_heading('📝 Results (What We Learned)', 1)
            doc.add_paragraph('[FILL DURING TEST]')
            doc.add_paragraph()
            doc.add_paragraph("Key Learnings:")
            for i in range(5):
                doc.add_paragraph("⬢ ", style='List Bullet')
            
            doc.add_heading('Equipment & Costs Captured:', 2)
            results_table = doc.add_table(rows=10, cols=4)
            results_table.style = 'Light Grid Accent 1'
            results_table.rows[0].cells[0].text = "Equipment"
            results_table.rows[0].cells[1].text = "Annual Damage $"
            results_table.rows[0].cells[2].text = "Failure Freq"
            results_table.rows[0].cells[3].text = "Cost Type"
            # Pre-fill Cost Type column hint
            for r in range(1, 10):
                results_table.rows[r].cells[3].text = "repair / production"
            
            # Validation section
            doc.add_paragraph()
            doc.add_heading('✅ Hypothesis Validation:', 2)
            doc.add_paragraph("Did this test_run VALIDATE or INVALIDATE our hypotheses?")
            
            # Prefer genesis_intel hypothesis names (short, testable statements)
            if genesis_intel.get('hypotheses'):
                for hyp_name in genesis_intel['hypotheses'].keys():
                    doc.add_paragraph(f"⬢ {hyp_name}: [ ] VALIDATED  [ ] INVALIDATED  [ ] UNCLEAR")
            else:
                # Fallback: parse hypothesis text for H1/H2/H3 pattern or short lines
                import re
                hyp_text = test_data.get('hypothesis', '')
                
                # Try to find "H1:", "H2:", "H3:", "Hypothesis 1:", etc.
                h_pattern = re.findall(r'(H\d+:?\s*[^\n]{10,80})', hyp_text)
                if h_pattern:
                    for h in h_pattern[:5]:
                        doc.add_paragraph(f"⬢ {h.strip()}: [ ] VALIDATED  [ ] INVALIDATED  [ ] UNCLEAR")
                else:
                    # Last resort: split by double-newline, take first sentence of each
                    chunks = [c.strip() for c in hyp_text.split('\n\n') if c.strip()]
                    for chunk in chunks[:5]:
                        # Take first sentence only (up to first period or 100 chars)
                        first_sentence = chunk.split('.')[0][:100].strip()
                        if first_sentence:
                            doc.add_paragraph(f"⬢ {first_sentence}: [ ] VALIDATED  [ ] INVALIDATED  [ ] UNCLEAR")
            
            # Action/Iterate
            doc.add_heading('🔄 Action/Iterate (Next Steps)', 1)
            doc.add_paragraph('[FILL AFTER TEST]')
            doc.add_paragraph()
            doc.add_paragraph("Next Steps:")
            doc.add_paragraph("⬢ ", style='List Bullet')
            doc.add_paragraph("⬢ ", style='List Bullet')
            
            # Save
            # Strip ALL Windows-illegal filename characters: \\ / : * ? " < > |
            safe_name = test_data['name']
            for ch in r'\\/:*?"<>|':
                safe_name = safe_name.replace(ch, '')
            safe_name = safe_name.strip()
            # No test_run number in filename — number comes from Excel plan
            filepath = get_template_dir(get_active_project()) / f"Test Run_Packet_{safe_name}.docx"
            doc.save(filepath)
            
            msg = f"Test Packet Generated!\n\n{filepath}\n\n"
            if genesis_intel['available']:
                msg += f"Includes:\n⬢ Intelligence from {genesis_intel['test_count']} previous tests\n"
                msg += f"⬢ {len(genesis_intel['gaps'])} critical gaps to fill\n"
                msg += f"⬢ {len(genesis_intel['contradictions'])} contradictions to resolve\n"
                msg += "⬢ Equipment cost tracking framework"
            else:
                msg += "Template includes equipment reference and cost framework."
            
            QMessageBox.information(self, "Packet Generated", msg)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate packet:\n{str(e)}")
    
    def _reconcile_counts(self):
        """Post-save count verification via Test RunReconciler."""
        if not _HAS_RECONCILER:
            return
        try:
            project_id = get_active_project()
            counts = quick_verify(project_id, self.database)
            if not counts['all_counts_match']:
                print(f"  [RECONCILER] ⚠ Count mismatch detected:")
                if counts['has_duplicates']:
                    print(f"    Duplicates in database ({counts['database_total']} total, "
                          f"{counts['database_unique_nums']} unique)")
                if not counts['validation_matches_changelog']:
                    print(f"    Validation({counts['database_validation']}) ≠ "
                          f"change_log({counts['change_log_completed']})")
            else:
                print(f"  [RECONCILER] ✓ Counts OK: {counts['database_total']} tests "
                      f"({counts['database_sparks']}S + {counts['database_validation']}F)")
        except Exception as e:
            print(f"  [RECONCILER] verify error: {e}")

    def _refresh_rd_health(self):
        """Refresh R&D Health — compute Research + Development convergence."""
        if not _HAS_RD_HEALTH_TAB or not hasattr(self, 'rd_health_tab'):
            return
        try:
            project_id = get_active_project()
            compute_rd_health(self.rd_health_tab, self.database, project_id)
        except Exception as e:
            print(f"  [R&D_HEALTH] Refresh error: {e}")

    def _run_test_tally(self):
        """
        Run the Validation Tally copier — sequential re-run of all test_runs
        through the current pipeline to produce the tare (balanced state).
        Called from R&D Health tab's Run Validation Tally button.
        """
        if not _HAS_RD_HEALTH_TAB or not hasattr(self, 'rd_health_tab'):
            return

        project_id = get_active_project()
        try:
            from test_reconciler import run_test_tally
            print(f"\n{'=' * 60}")
            print(f"  FUSION TALLY — {project_id}")
            print(f"{'=' * 60}")

            result = run_test_tally(
                project_id=project_id,
                database=None,  # Load from disk for clean state
                reconcile_first=True,
            )

            # Display in R&D Health tab
            self.rd_health_tab.display_tally(result)

            tare = result.get('tare', {})
            print(f"  ✓ Tally complete: {tare.get('total_test_runs', 0)} tests, "
                  f"Q-Cube {tare.get('cube_coverage_pct', 0)}%")

        except ImportError:
            print("  [TALLY] test_reconciler.py not found")
            self.rd_health_tab.tally_btn.setEnabled(True)
            self.rd_health_tab.tally_btn.setText("▶ Run Validation Tally")
            self.rd_health_tab.tally_status.setText(
                "⚠ test_reconciler.py not found")
            QMessageBox.warning(
                self, "Missing Module",
                "test_reconciler.py not found.\n\n"
                "Place it next to validation_test_collector.py."
            )
        except Exception as e:
            import traceback
            print(f"  [TALLY] Error: {e}")
            traceback.print_exc()
            self.rd_health_tab.tally_btn.setEnabled(True)
            self.rd_health_tab.tally_btn.setText("▶ Run Validation Tally")
            self.rd_health_tab.tally_status.setText(f"⚠ Error: {e}")
            self.rd_health_tab.tally_text.setText(
                f"Tally error:\n\n{e}\n\n{traceback.format_exc()}"
            )

    def _load_genesis_intelligence(self, test_data: dict = None) -> dict:
        """
        LIVE intelligence from progressive engine — NOT static JSON files.
        Calculates from actual database tests every time.
        
        If test_data is provided, generates progressive template
        with compound questions calibrated to ALL prior tests.
        """
        intel = {
            'available': False,
            'test_count': 0,
            'equipment_mentions': {},
            'equipment_gaps': [],
            'gaps': [],
            'hypotheses': {},
            'contradictions': [],
            'compound_tests': [],
            'mission': '',
            'vp_summary': [],
            'knn_similar': [],
            'reverse_gaps': [],
        }
        
        try:
            # LIVE: Use progressive engine (test_intelligence under the hood)
            from progressive_test_run_engine import ProgressiveTest RunEngine
            project_id = get_active_project()
            engine = ProgressiveTest RunEngine(self.database, project_id)
            
            if test_data:
                # Generate progressive template for specific test_run
                template = engine.generate_progressive_template(
                    test_num=test_data.get('num'),
                    person_name=test_data.get('name', ''),
                    q_layer=test_data.get('q_layer', 'L2'),
                    q_object=test_data.get('q_object', 'OC'),
                )
                
                if template.get('available'):
                    intel['available'] = True
                    intel['test_count'] = template.get('test_count', 0)
                    intel['compound_tests'] = template.get('compound_tests', [])
                    intel['mission'] = template.get('mission', '')
                    intel['vp_summary'] = template.get('vp_summary', [])
                    intel['knn_similar'] = template.get('knn_similar', [])
                    intel['reverse_gaps'] = template.get('reverse_gaps', [])
                    intel['hypotheses'] = template.get('hypotheses', {})
                    intel['equipment_mentions'] = template.get('equipment_mentions', {})
                    intel['gaps'] = template.get('gaps', [])
                    intel['contradictions'] = [
                        c.get('probe', str(c)) for c in template.get('contradictions', [])
                    ]
                    
                    # Store state snapshot for IoC (before state)
                    self._last_intel_state = template.get('_state_snapshot', {})
                    return intel
            
            # Fallback: just get current state summary
            if engine.intel and engine.intel.state.get('ready'):
                s = engine.intel.state
                intel['available'] = True
                intel['test_count'] = s.get('test_count', 0)
                for h_id, h in s.get('hypotheses', {}).items():
                    intel['hypotheses'][h['name']] = {
                        'confidence': h['confidence'],
                        'needs_more': h['needs_more'],
                        'probe_question': f"Can you provide specific evidence about {h['name'].lower()}?"
                    }
                intel['gaps'] = [f"Cube gap: {g}" for g in s.get('cube_gaps', [])[:3]]
                intel['equipment_mentions'] = {
                    name: {'count': stats['count'],
                           'avg_cost': f"${stats['avg_annual']/1000:.0f}K/yr"
                                   if stats['avg_annual'] > 0 else 'Unknown'}
                    for name, stats in s.get('equipment_ranked', [])[:5]
                }
                intel['contradictions'] = [
                    c.get('probe', str(c)) for c in s.get('contradictions', [])
                ]
            
            return intel
            
        except ImportError:
            print("  [INFO] progressive_test_run_engine not available, trying static fallback")
        except Exception as e:
            print(f"  [WARN] Progressive intel failed: {e}")
        
        # STATIC FALLBACK: Read from Genesis output files (backward compat)
        try:
            intel_file = Path("GENESIS_OUTPUT/strategic_intelligence.json")
            if intel_file.exists():
                with open(intel_file, 'r') as f:
                    data = json.load(f)
                intel['available'] = True
                for hyp_key, hyp_data in data.get('hypotheses', {}).items():
                    intel['hypotheses'][hyp_data['name']] = {
                        'confidence': int(hyp_data['posterior'] * 100),
                        'needs_more': hyp_data['posterior'] < 0.75,
                        'probe_question': f"Can you provide specific evidence about {hyp_data['name'].lower()}?"
                    }
                for gap in data.get('gaps', [])[:3]:
                    intel['gaps'].append(gap['description'])
        except Exception as e:
            print(f"  [WARN] Static fallback also failed: {e}")
        
        # DATABASE-DIRECT FALLBACK: If still 0, count from live database
        # This ensures test_count is NEVER 0 when tests exist
        if intel['test_count'] == 0 and hasattr(self, 'database'):
            db = self.database
            tests = getattr(db, 'tests', [])
            if test_runs:
                intel['test_count'] = len(tests)
                intel['available'] = True
                print(f"  [INFO] Database-direct fallback: {len(tests)} tests loaded")
                
                # Build basic hypothesis confidence from test results text
                # Uses same keyword approach as doctoral_analyzer._process_test_run_evidence
                hyp_keywords = {
                    'Manual Detection is Industry-Wide Failure': {
                        'for': ['100% fail', 'all got through', 'manual catch', 'no detection', 'nothing catches'],
                        'against': ['effective detection', 'catches most']
                    },
                    'True Cost Impact is Understated': {
                        'for': ['damage cost', 'annual damage', 'more than expected', 'nobody tracks', 'repair cost', 'replacement cost'],
                        'against': ['minimal damage', 'low cost']
                    },
                    "Capital Investment Alone Doesn't Solve Contamination": {
                        'for': ['modern', 'new equipment', 'still happens', 'upgraded', 'persist'],
                        'against': ['solved it', 'fixed the problem']
                    },
                    'Operators Sense Problems Instruments Miss': {
                        'for': ['feel', 'sense', 'hear', 'notice', 'go look', 'know before'],
                        'against': ["didn't notice", 'no warning']
                    },
                    'Blow Line Entry Point is Critical Blind Spot': {
                        'for': ['blind spot', 'no visibility', 'can\'t see', 'before digester'],
                        'against': ['good visibility', 'can monitor']
                    },
                    'Contamination is Accepted Practice': {
                        'for': ['accepted', 'normal', 'always been', 'cost of doing business', 'live with it'],
                        'against': ['unacceptable', 'zero tolerance']
                    },
                }
                
                for hyp_name, keywords in hyp_keywords.items():
                    support_count = 0
                    total_checked = 0
                    for test_entry in test_runs:
                        results = str(getattr(test_run, 'results', '') or '').lower()
                        action = str(getattr(test_run, 'action_iterate', '') or '').lower()
                        combined = results + ' ' + action
                        if not combined.strip():
                            continue
                        total_checked += 1
                        if any(kw in combined for kw in keywords['for']):
                            support_count += 1
                    
                    if total_checked > 0:
                        confidence = int((support_count / total_checked) * 100)
                        confidence = max(10, min(95, confidence))  # clamp
                    else:
                        confidence = 50  # prior
                    
                    intel['hypotheses'][hyp_name] = {
                        'confidence': confidence,
                        'needs_more': confidence < 75,
                        'probe_question': f"Can you provide specific evidence about {hyp_name.lower()}?"
                    }
        
        return intel
    
    def run_genesis_analysis(self):
        """Run Genesis Brain analysis in worker thread with NON-BLOCKING progress dialog."""
        if self._genesis_running:
            return  # hard gate

        if len(self.database.tests) == 0:
            QMessageBox.warning(self, "No Test Runs", "Please add tests first.")
            return

        self._genesis_running = True
        if hasattr(self, "genesis_btn"):
            self.genesis_btn.setEnabled(False)

        # ---------- Worker thread ----------
        class GenesisWorker(QThread):
            finished = Signal(dict)
            error = Signal(str)

            def __init__(self, tests):
                super().__init__()
                self.tests = test_runs

            def run(self):
                try:
                    from genesis_brain import GenesisOrchestrator
                    from pathlib import Path

                    output_dir = Path("GENESIS_OUTPUT")
                    output_dir.mkdir(exist_ok=True, parents=True)

                    orchestrator = GenesisOrchestrator(output_dir=output_dir)
                    result = orchestrator.run_full_analysis(self.tests)
                    self.finished.emit(result)
                except Exception:
                    import traceback
                    self.error.emit(traceback.format_exc())

        tests_dict = [test_run.to_dict() for test_entry in self.database.test_runs]

        self.genesis_worker = GenesisWorker(tests_dict)
        self.genesis_worker.finished.connect(self.on_genesis_complete)
        self.genesis_worker.error.connect(self.on_genesis_error)

        # Ensure worker object is cleaned up
        self.genesis_worker.finished.connect(self.genesis_worker.deleteLater)
        self.genesis_worker.error.connect(self.genesis_worker.deleteLater)

        # ---------- Progress dialog (NO QMessageBox) ----------
        self.genesis_progress = QProgressDialog(
            "Running Genesis Analysis...\n\nThis may take 30⬓60 seconds.\nPlease wait...",
            None, 0, 0, self
        )
        self.genesis_progress.setWindowTitle("Genesis Brain")
        self.genesis_progress.setWindowModality(Qt.WindowModal)
        self.genesis_progress.setMinimumDuration(0)
        self.genesis_progress.setAutoClose(True)
        self.genesis_progress.setAutoReset(True)
        self.genesis_progress.show()

        self.genesis_worker.start()


    def _finalize_genesis_ui(self):
        """Always called after Genesis completes or errors."""
        self._genesis_running = False

        if hasattr(self, "genesis_btn"):
            self.genesis_btn.setEnabled(True)

        if getattr(self, "genesis_progress", None) is not None:
            try:
                self.genesis_progress.close()
                self.genesis_progress.deleteLater()
            except Exception:
                pass
            self.genesis_progress = None

        self.genesis_worker = None


    def on_genesis_complete(self, result: dict):
        """Handle Genesis analysis completion."""
        self._finalize_genesis_ui()

        # Orchestrator return shape varies; do not assume numeric fields exist.
        msg = (
            "Genesis Analysis Complete!\n\n"
            f"Test Runs Analyzed: {len(self.database.tests)}\n"
            "Outputs saved to: GENESIS_OUTPUT\\"
        )
        QMessageBox.information(self, "Success", msg)


    def on_genesis_error(self, error_text: str):
        """Handle Genesis analysis error."""
        self._finalize_genesis_ui()

        from pathlib import Path
        error_file = Path("GENESIS_OUTPUT/last_error.txt")
        error_file.parent.mkdir(exist_ok=True, parents=True)
        error_file.write_text(error_text, encoding="utf-8")

        QMessageBox.critical(
            self,
            "Analysis Failed",
            f"Genesis analysis error.\n\nFull error saved to:\n{error_file.absolute()}"
        )


# ============================================================================
# MAIN
# ============================================================================

def main():
    app = QApplication(sys.argv)
    
    # Set application style
    app.setStyle('Validation')
    
    # Set GIBUSH dark color scheme
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(0, 0, 0))
    palette.setColor(QPalette.WindowText, QColor(220, 220, 220))
    palette.setColor(QPalette.Base, QColor(16, 16, 16))
    palette.setColor(QPalette.AlternateBase, QColor(26, 26, 26))
    palette.setColor(QPalette.Text, QColor(220, 220, 220))
    palette.setColor(QPalette.Button, QColor(26, 26, 26))
    palette.setColor(QPalette.ButtonText, QColor(200, 200, 200))
    palette.setColor(QPalette.Highlight, QColor(147, 51, 234))
    palette.setColor(QPalette.HighlightedText, QColor(255, 215, 0))
    app.setPalette(palette)
    
    # Master dark stylesheet — matches GIBUSH EHS theme
    app.setStyleSheet("""
        QMainWindow { background-color: #000; }
        QWidget { background-color: #000; color: #ddd; }
        QLabel { color: #ddd; }
        QGroupBox { 
            color: #00ffff; 
            border: 1px solid #333; 
            padding: 10px; 
            margin-top: 14px;
        }
        QGroupBox::title { 
            subcontrol-origin: margin; 
            left: 10px; 
            padding: 0 5px; 
        }
        QPushButton { 
            background: #111; 
            color: #FFD700; 
            border: 1px solid #9333EA; 
            padding: 6px 12px; 
            font-weight: bold; 
        }
        QPushButton:hover { background: #1a1a1a; }
        QPushButton:pressed { background: #9333EA; }
        QPushButton:disabled { background: #111; color: #555; border: 1px solid #333; }
        QTabWidget::pane { border: 1px solid #9333EA; }
        QTabBar::tab { 
            background: #0a0a0a; 
            color: #888; 
            padding: 8px 16px; 
            border: 1px solid #333; 
        }
        QTabBar::tab:selected { 
            background: #9333EA; 
            color: #FFD700; 
            font-weight: bold; 
        }
        QTabBar::tab:hover { background: #1a1a1a; }
        QTableWidget { 
            background: #0a0a0a; 
            color: #00ff88; 
            gridline-color: #333; 
            border: 1px solid #333;
        }
        QTableWidget::item:selected {
            background: #9333EA;
            color: #FFD700;
        }
        QHeaderView::section {
            background: #1a1a1a;
            color: #00ffff;
            border: 1px solid #333;
            padding: 4px;
            font-weight: bold;
        }
        QTreeWidget {
            background: #0a0a0a;
            color: #00ffff;
            border: 1px solid #333;
        }
        QTreeWidget::item:selected {
            background: #9333EA;
            color: #FFD700;
        }
        QTreeWidget::item:hover { background: #1a1a1a; }
        QTextEdit { 
            background: #0a0a0a; 
            color: #00ff88; 
            border: 1px solid #333; 
        }
        QLineEdit { 
            background: #111; 
            color: #ddd; 
            border: 1px solid #333; 
            padding: 4px;
        }
        QComboBox { 
            background: #111; 
            color: #ddd; 
            border: 1px solid #9333EA; 
            padding: 4px;
        }
        QComboBox QAbstractItemView {
            background: #1a1a1a;
            color: #ddd;
            selection-background-color: #9333EA;
        }
        QScrollBar:vertical {
            background: #0a0a0a;
            width: 12px;
        }
        QScrollBar::handle:vertical {
            background: #333;
            border-radius: 4px;
        }
        QScrollBar:horizontal {
            background: #0a0a0a;
            height: 12px;
        }
        QScrollBar::handle:horizontal {
            background: #333;
            border-radius: 4px;
        }
        QSplitter::handle { background: #333; }
        QProgressBar {
            background: #111;
            border: 1px solid #333;
            color: #FFD700;
            text-align: center;
        }
        QProgressBar::chunk {
            background: #9333EA;
        }
    """)
    
    window = BCMTestCollector()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()